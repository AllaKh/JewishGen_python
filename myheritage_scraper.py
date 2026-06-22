#!/usr/bin/env python3
"""
myheritage_scraper.py  —  v6
─────────────────────────────
Logic:
1. Open the URL chosen in GUI (e.g. https://www.myheritage.co.il/  or
   https://www.myheritage.com/?lang=RU  etc.)
2. If the page is .co.il → accept cookie banner ("לקבל הכל").
3. Click "Вход" / "כניסה" / "Log in" link to open the login form.
4. Fill email + password, click submit ("היכנס" / "Авторизация" / "Log in").
5. If a 2-FA code dialog appears → GUI shows a modal asking user to type
   the code; scraper waits for it via an asyncio.Event, then enters the code.
6. After login → navigate to the search URL for the chosen domain, fill
   the search form, collect results, open each qualifying record, save.

IMPORTANT: a PERSISTENT browser profile (.mh_profile) is used so cookies and
session survive between runs — after the first successful login + email 2FA,
later runs reuse the session and MyHeritage stops flagging logins as suspicious.
"""

import asyncio, difflib, imaplib, email as _email_lib
import io, os, re, sys, time, hashlib, urllib.parse
from pathlib import Path

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    import browser_util
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

from docx_util import add_page_numbers

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

try:
    from PIL import Image
    _PIL_OK = True
except ImportError:
    _PIL_OK = False


def _to_png(data: bytes):
    """Return PNG bytes for any image (WEBP/JPEG/…) so python-docx can embed
    it. MyHeritage serves photos as WEBP which docx can't read directly."""
    if not data:
        return None
    # already a docx-friendly format? (JPEG/PNG/GIF/BMP magic)
    if (data[:3] == b"\xff\xd8\xff" or data[:8] == b"\x89PNG\r\n\x1a\n"
            or data[:4] == b"GIF8" or data[:2] == b"BM"):
        return data
    if _PIL_OK:
        try:
            im = Image.open(io.BytesIO(data)).convert("RGB")
            out = io.BytesIO()
            im.save(out, format="PNG")
            return out.getvalue()
        except Exception:
            return None
    return None

# ── Constants ─────────────────────────────────────────────────────────────── #
MIN_MATCH_PCT = 80
HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"

# Persistent Chromium profile — keeps MyHeritage cookies/session between runs
# so the anti-bot stops flagging the login as suspicious.
from paths_util import user_data_dir
MH_PROFILE_DIR = user_data_dir() / ".mh_profile"   # writable even in a packaged install

# Site presets — key → (login_url, search_url, has_cookie_banner)
SITE_PRESETS = {
    "Israel (.co.il)":        ("https://www.myheritage.co.il/login",
                               "https://www.myheritage.co.il/research/search/all/all",
                               True),
    "English (.com EN)":      ("https://www.myheritage.com/login?lang=EN",
                               "https://www.myheritage.com/research/search/all/all?lang=EN",
                               True),
    "Russian (.com RU)":      ("https://www.myheritage.com/login?lang=RU",
                               "https://www.myheritage.com/research/search/all/all?lang=RU",
                               True),
    "Hebrew (.com HE)":       ("https://www.myheritage.com/login?lang=HE",
                               "https://www.myheritage.com/research/search/all/all?lang=HE",
                               True),
    "French (.com FR)":       ("https://www.myheritage.com/login?lang=FR",
                               "https://www.myheritage.com/research/search/all/all?lang=FR",
                               True),
    "German (.com DE)":       ("https://www.myheritage.com/login?lang=DE",
                               "https://www.myheritage.com/research/search/all/all?lang=DE",
                               True),
    "Spanish (.com ES)":      ("https://www.myheritage.com/login?lang=ES",
                               "https://www.myheritage.com/research/search/all/all?lang=ES",
                               True),
    "Portuguese (.com PT)":   ("https://www.myheritage.com/login?lang=PT",
                               "https://www.myheritage.com/research/search/all/all?lang=PT",
                               True),
}

FILTER_OPTIONS = ["All Records", "Historical Records", "Family Trees"]

# ── Multi-language UI labels ─────────────────────────────────────────────────
# MyHeritage renders the research form in the SITE language (the lang= in the
# URL, picked by the GUI "Site / Language" selector). Every place we locate a
# pill / section / button BY ITS VISIBLE TEXT must therefore use the label in
# the active site language. The GUI itself is always English and sends canonical
# English values; the scraper translates them here.
#
# We pass ALL supported-language variants to the text matchers: only the variant
# matching the loaded site exists on the page, and the data-automations verify-id
# confirms the correct popup actually opened — so the same code drives every
# site. Supported: en, ru, he, fr, de, es, pt. (Non-EN/RU labels are best-effort
# — but the data-automations IDs do the real field/apply work and are language-
# independent, so a wrong word only affects finding a pill by text, never input.)
#
# NOTE: the user's MyHeritage account UI language may override the URL lang after
# the page loads (e.g. it reverts to Russian). That's fine: _ui_labels returns
# EVERY language's variant, so whatever language the page ends up in, its label
# is in the list and still matches.
def _site_lang(site):
    """Map a Site/Language preset to a language code (en/ru/he/fr/de/es/pt)."""
    s = (site or "").lower()
    if "russ" in s or "(.com ru" in s or " ru)" in s:
        return "ru"
    if "hebrew" in s or "co.il" in s or "israel" in s or "(.com he" in s:
        return "he"
    if "french" in s or "(.com fr" in s:
        return "fr"
    if "german" in s or "(.com de" in s:
        return "de"
    if "spanish" in s or "(.com es" in s:
        return "es"
    if "portug" in s or "(.com pt" in s:
        return "pt"
    return "en"

_UI_I18N = {
    "search":      {"en": "Search",      "ru": "Поиск",            "he": "חיפוש",
                    "fr": "Rechercher",  "de": "Suchen",           "es": "Buscar",
                    "pt": "Pesquisar"},
    "more":        {"en": "More",        "ru": "Больше",           "he": "עוד",
                    "fr": "Plus",        "de": "Mehr",             "es": "Más",
                    "pt": "Mais"},
    "father":      {"en": "Father",      "ru": "Отец",             "he": "אב",
                    "fr": "Père",        "de": "Vater",            "es": "Padre",
                    "pt": "Pai"},
    "mother":      {"en": "Mother",      "ru": "Мать",             "he": "אם",
                    "fr": "Mère",        "de": "Mutter",           "es": "Madre",
                    "pt": "Mãe"},
    "spouse":      {"en": "Spouse",      "ru": ["Супруг(-а)", "Супруг"],
                    "he": ["בן/בת זוג", "בן זוג"],
                    "fr": ["Conjoint(e)", "Conjoint"], "de": ["Ehepartner", "Partner"],
                    "es": "Cónyuge",     "pt": "Cônjuge"},
    "death":       {"en": "Death",       "ru": "Смерть",           "he": "פטירה",
                    "fr": "Décès",       "de": "Tod",              "es": ["Defunción", "Fallecimiento"],
                    "pt": ["Falecimento", "Morte"]},
    "residence":   {"en": "Residence",   "ru": "Местожительство",  "he": "מגורים",
                    "fr": "Résidence",   "de": ["Wohnsitz", "Wohnort"], "es": "Residencia",
                    "pt": "Residência"},
    "military":    {"en": "Military",    "ru": "Вооруженные силы", "he": "צבא",
                    "fr": ["Service militaire", "Militaire"], "de": ["Militär", "Militärdienst"],
                    "es": ["Servicio militar", "Militar"], "pt": ["Serviço militar", "Militar"]},
    "immigration": {"en": "Immigration", "ru": "Иммиграция",       "he": "הגירה",
                    "fr": "Immigration", "de": "Einwanderung",     "es": "Inmigración",
                    "pt": "Imigração"},
    "keywords":    {"en": "Keywords",    "ru": "Ключевые слова",   "he": "מילות מפתח",
                    "fr": "Mots-clés",   "de": ["Schlüsselwörter", "Stichwörter"],
                    "es": "Palabras clave", "pt": "Palavras-chave"},
    "gender":      {"en": "Gender",      "ru": "Пол",              "he": "מין",
                    "fr": "Sexe",        "de": "Geschlecht",       "es": ["Sexo", "Género"],
                    "pt": ["Sexo", "Gênero"]},
    "apply":       {"en": "Apply",       "ru": "Применить",        "he": "החל",
                    "fr": "Appliquer",   "de": ["Anwenden", "Übernehmen"], "es": "Aplicar",
                    "pt": "Aplicar"},
    "male":        {"en": "Male",        "ru": "Мужчина",          "he": "זכר",
                    "fr": ["Homme", "Masculin"], "de": ["Männlich", "Mann"],
                    "es": ["Hombre", "Masculino"], "pt": ["Masculino", "Homem"]},
    "female":      {"en": "Female",      "ru": "Женщина",          "he": "נקבה",
                    "fr": ["Femme", "Féminin"], "de": ["Weiblich", "Frau"],
                    "es": ["Mujer", "Femenino"], "pt": ["Feminino", "Mulher"]},
}

# Stable language order used when no specific site-language is given.
_LANGS = ["en", "ru", "he", "fr", "de", "es", "pt"]

def _ui_labels(concept, lang=None):
    """All language variants of a concept's UI label, site-lang first when known.
    Accepts dict values that are a single string or a list of strings. Iterates
    every language present, so adding a language to _UI_I18N just works."""
    d = _UI_I18N.get(concept, {})
    out, seen = [], set()
    for k in ([lang] if lang else []) + _LANGS + list(d.keys()):
        vals = d.get(k) or []
        if isinstance(vals, str):
            vals = [vals]
        for v in vals:
            if v and v not in seen:
                seen.add(v); out.append(v)
    return out

def _search_words(lang=None):
    """Every language's word for the search-submit button (site-lang first)."""
    return _ui_labels("search", lang)

# ── Helpers ───────────────────────────────────────────────────────────────── #
def _name_sim(a, b):
    a, b = a.strip().lower(), b.strip().lower()
    return difflib.SequenceMatcher(None, a, b).ratio() * 100 if a and b else 0.0


# Names come in mixed scripts («Александр-Вольф Сандерс», «Alexandre-Wolf Sanders»)
# — fold everything onto one coarse Latin form before comparing.
_CYR2LAT = {"а":"a","б":"b","в":"v","г":"g","ґ":"g","д":"d","е":"e","ё":"e","є":"e",
            "ж":"zh","з":"z","и":"i","і":"i","ї":"i","й":"i","к":"k","л":"l","м":"m",
            "н":"n","о":"o","п":"p","р":"r","с":"s","т":"t","у":"u","ф":"f","х":"h",
            "ц":"c","ч":"ch","ш":"sh","щ":"sch","ъ":"","ы":"i","ь":"","э":"e","ю":"iu","я":"ia"}
_NAME_SEP = re.compile(r"[\s.,()/\[\]«»\"'\-—–]+")


def _fold_name_token(t):
    t = "".join(_CYR2LAT.get(ch, ch) for ch in (t or "").lower())
    for a, b in (("shch","sch"),("kh","h"),("ts","c"),("ph","f"),("x","ks"),
                 ("w","v"),("y","i"),("j","i")):
        t = t.replace(a, b)
    return re.sub(r"[^a-z0-9]+", "", t)


def _gname_tokens(s):
    return [t for t in (_fold_name_token(x) for x in _NAME_SEP.split(s or "")) if t]


def _stem_match(a, b):
    if a == b:
        return True
    if len(a) >= 5 and len(b) >= 5 and a[:5] == b[:5]:
        return True
    return difflib.SequenceMatcher(None, a, b).ratio() >= 0.72


def _tok_in(q, rt):
    """A query token is present among record tokens (exact / stem / initial)."""
    return any(q == r or _stem_match(q, r)
               or (len(q) == 1 and r.startswith(q))
               or (len(r) == 1 and q.startswith(r)) for r in rt)


def _name_relevant(record_name, first_q, last_q):
    """Keep records that match the searched name fully: the FIRST given name is
    positional (record must START with «Alexander»-variant, so «Walter Alexander»
    is dropped), and EVERY other given token — incl. the middle initial «W» —
    must be present (so «Alexander Alfred/Lorn/Bertram Sanders» are dropped, but
    «Alexander-Wolf»/«Александр-Вольф» pass: «W» matches «Wolf»/«Вольф»). The
    surname must be present. Lenient on transliteration / variants."""
    rt = _gname_tokens(record_name)
    if not rt:
        return True
    qf = _gname_tokens(first_q)
    if qf:
        if len(qf[0]) >= 2 and not _stem_match(qf[0], rt[0]):
            return False                              # first token ≠ «Alexander»
        for q in qf[1:]:                              # middle initials / names («W»)
            if not _tok_in(q, rt):
                return False
    for lt in _gname_tokens(last_q):
        if len(lt) >= 2 and not _tok_in(lt, rt):
            return False                              # surname missing
    return True


def _year_ok(want_year, rec_year, tol=5):
    """True unless BOTH years are known and differ by more than `tol`. Drops the
    «Alexander W. (~1848)» people when the search asked for ~1897."""
    try:
        wy = int(re.sub(r"\D", "", str(want_year))[:4] or 0)
        ry = int(re.sub(r"\D", "", str(rec_year))[:4] or 0)
    except Exception:
        return True
    if not wy or not ry:
        return True
    return abs(wy - ry) <= tol


def _name_year(name):
    """Birth year embedded in a result name, e.g. «Alexander W. (~1848) Sanders»
    or «… (1848-1920)» → «1848». MyHeritage shows the life span in parentheses
    for family-tree people, so this catches the wrong «~1848» Alexanders that
    carry no «Рождение» line on the card."""
    m = re.search(r"\(\s*~?\s*(1[5-9]\d\d|20\d\d)\b", name or "")
    return m.group(1) if m else ""


def _detail_birth_year(rec):
    """A birth-year estimate from the OPENED record, to drop wrong-era people the
    card stage couldn't (no birth line on the card). Priority:
      1. the person's own «Рождение»/«Birth» field;
      2. the SPOUSE's year as a same-generation proxy — «Alexander married Cynthia
         (b.1848)» reveals an 1840s man even with no birth field. Parents are a
         generation older, so they are NOT used (would wrongly drop the 1897 man
         whose parents were born 1858)."""
    td = rec.get("table_data", {}) or {}
    # 1) the person's OWN birth field — short-circuits (so correct people whose
    #    parents/siblings were born earlier are never judged by those years).
    for k, v in td.items():
        if re.search(r"рожд|birth|geboren|naiss|nacim|geb\.", k, re.I):
            m = re.search(r"\b(1[5-9]\d\d|20\d\d)\b", str(v))
            if m:
                return m.group(1)
    # 2) no own birth → earliest SPOUSE/CHILD year (same-or-later generation),
    #    from a structured field OR inlined in a blob («Источник (дерево)» often
    #    glues «… Жена Cynthia (born West) 1848 …»). NOT parents/siblings (older).
    _SK = re.compile(r"\b(жена|муж|супруг|spouse|wife|husband|дети|ребён|ребен|"
                     r"сын|доч|child|children|son|daughter)\b", re.I)
    # NB: allow ANY chars in the gap (not [^0-9]) so a leading day breaks nothing:
    # «Супруг(-а): Sarah … 10 мар 1864» — the «10» must not block reaching 1864.
    _SB = re.compile(r"(?:жена|муж|супруг|жен[ыа]|spouse|wife|husband|дети|ребён|"
                     r"ребен|child|children).{0,80}?(1[5-9]\d\d|20\d\d)", re.I)
    years = []
    for k, v in td.items():
        v = str(v)
        if _SK.search(k):
            years += [int(y) for y in re.findall(r"\b(1[5-9]\d\d|20\d\d)\b", v)]
        years += [int(y) for y in _SB.findall(v)]
    return str(min(years)) if years else ""


def _year_not_earlier(want, got, tol=5):
    """The user's rule: DROP records whose year is EARLIER than the searched birth
    year, KEEP equal or later. Returns False (→ drop) only when `got` is more than
    `tol` years BEFORE `want`. Unknown years are kept (lenient)."""
    try:
        wy = int(re.sub(r"\D", "", str(want))[:4] or 0)
        gy = int(re.sub(r"\D", "", str(got))[:4] or 0)
    except Exception:
        return True
    if not wy or not gy:
        return True
    return gy >= wy - tol


def _type_ok(rec, filt):
    """Honour «Уточнить по типу записи» by filtering on the record's category:
    «Семейные деревья» keeps only trees, «Исторические записи» keeps only the
    rest. «All Records» keeps everything."""
    if filt not in ("Historical Records", "Family Trees"):
        return True
    blob = (str(rec.get("category", "")) + " " +
            str(rec.get("source_text", ""))).lower()
    url = str(rec.get("url", "")).lower()
    is_tree = bool(re.search(r"дерев|family\s*tree|tree|geni", blob)) \
        or "/collection-1/" in url
    return is_tree if filt == "Family Trees" else (not is_tree)

def safe_fn(s, n=80):
    return re.sub(r'\s+', '_', re.sub(r'[\\/*?:"<>|]', '_', s.strip()))[:n] or "result"

# ── Cookie banner ─────────────────────────────────────────────────────────── #
# Specific «accept ALL» variants only — NOT bare «Принять»/«Accept»/«OK», which
# could match «Принять только необходимые» / «Accept only necessary».
_COOKIE_TEXTS = ["Принять все", "Принять всё", "Разрешить все", "Разрешить всё",
                 "Accept all", "Accept All", "Allow all", "Allow All",
                 "Соглашаюсь со всеми", "לקבל הכל", "קבל הכל"]
# ONLY accept-specific controls — NEVER a generic «first button in a cookie box»
# (that grabbed a footer «Cookie policy» / «Decline» button and BROKE the login).
_COOKIE_SELS = ["#onetrust-accept-btn-handler", ".onetrust-accept-btn-handler",
                'button[id*="accept-all" i]', 'button[id*="acceptall" i]',
                'button[aria-label*="Accept all" i]',
                'button[aria-label*="Принять все" i]',
                '[data-testid*="accept-all" i]', '[data-testid*="acceptall" i]']


async def _accept_cookies(page, log):
    """Accept the cookie / consent banner wherever it shows — MAIN frame OR an
    iframe — retrying a few times (it can load late). It ONLY clicks an explicit
    ACCEPT control: pass 1 = a button whose TEXT is «Принять все»/«Accept all»
    (across ALL frames first, so a stray cookie-link in the main frame can't be
    clicked before the real banner in an iframe); pass 2 = accept-specific
    selectors (OneTrust / accept-all). Always safe; returns True if dismissed."""
    for _ in range(3):                                # banner can load late
        await asyncio.sleep(1.0)
        # PASS 1 — explicit accept BUTTON TEXT, all frames before any selector
        for fr in page.frames:
            for text in _COOKIE_TEXTS:
                try:
                    btn = fr.get_by_role(
                        "button", name=re.compile(re.escape(text), re.I))
                    if await btn.count():
                        await btn.first.click(timeout=3000)
                        log(f"  ✓ Cookie banner accepted ('{text}')")
                        await asyncio.sleep(0.6)
                        return True
                except Exception:
                    pass
        # PASS 2 — accept-SPECIFIC selectors only
        for fr in page.frames:
            for sel in _COOKIE_SELS:
                try:
                    el = fr.locator(sel).first
                    if await el.count() and await el.is_visible():
                        await el.click(timeout=2500)
                        log(f"  ✓ Cookie banner accepted ({sel})")
                        await asyncio.sleep(0.6)
                        return True
                except Exception:
                    pass
    return False

# ── Fill input field ─────────────────────────────────────────────────────── #
async def _fill(page, selectors, value, label, log):
    """
    Fill a form field reliably.

    MyHeritage uses  type="text" autocomplete="off"  for the email field and
    a normal type="password" for the password field.  Playwright's .fill()
    sometimes gets cleared by the site's JS on fields with autocomplete="off".
    The safest approach is:
      1. Click the element to give it real focus
      2. Select-all + Delete to clear any existing value
      3. Use page.keyboard.type() — simulates real keystrokes, never blocked
      4. Verify the value stuck; if not, try JS direct assignment as last resort
    """
    if not value:
        return True
    for sel in selectors:
        try:
            el = page.locator(sel).first
            if not await el.count():
                continue
            await el.scroll_into_view_if_needed(timeout=4000)

            # Give the field real browser focus via JS (more reliable than click)
            await page.evaluate(
                "sel => { const el = document.querySelector(sel); "
                "if (el) { el.focus(); } }",
                sel,
            )
            await asyncio.sleep(0.15)

            # Clear existing value
            await el.click(timeout=3000)
            await page.keyboard.press("Control+a")
            await page.keyboard.press("Delete")
            await asyncio.sleep(0.1)

            # Type character by character — works with autocomplete="off"
            await page.keyboard.type(value, delay=40)
            await asyncio.sleep(0.2)

            typed = await el.input_value(timeout=2000)
            if typed.strip():
                log(f"  ✓ {label}: OK ({len(typed)} chars)")
                return True

            # Last resort: JS direct value set + events
            await page.evaluate("""([s, v]) => {
                const el = document.querySelector(s);
                if (!el) return;
                const setter = Object.getOwnPropertyDescriptor(
                    window.HTMLInputElement.prototype, 'value').set;
                setter.call(el, v);
                el.dispatchEvent(new Event('input',  {bubbles: true}));
                el.dispatchEvent(new Event('change', {bubbles: true}));
                el.dispatchEvent(new KeyboardEvent('keyup', {bubbles: true}));
            }""", [sel, value])
            typed = await el.input_value(timeout=2000)
            if typed.strip():
                log(f"  ✓ {label}: JS OK ({len(typed)} chars)")
                return True

        except Exception as exc:
            log(f"  (selector {sel} failed: {exc})")
            continue

    log(f"  !! Field not found: {label}")
    return False

# ── Open login form ───────────────────────────────────────────────────────── #
async def _open_login_form(page, log):
    """
    Ensure the email/password form is visible.

    IMPORTANT: the login_url we navigate to is already MyHeritage's /login
    page, so the form is normally present immediately. We must NOT click any
    random "Log in" / "Войти" text — on the home/footer that can match a
    Facebook social link and navigate the tab to facebook.com (the bug Alla
    hit). So:
      1. If we're already on a /login or /signin URL → do nothing.
      2. If the email field is already present → do nothing.
      3. Only as a last resort click a login control that is a real
         MyHeritage button/link (never a social link).
    """
    # 1) Already on the login page → form is there.
    if "/login" in page.url.lower() or "/signin" in page.url.lower():
        return True

    # 2) Email field already on the page → nothing to open.
    for sel in ('#registrationEmail', 'input[name="registrationEmail"]',
                'input[type="email"]'):
        try:
            if await page.locator(sel).first.count():
                return True
        except Exception:
            pass

    # 3) Last resort: click a real MyHeritage login control, scoped to
    #    same-site hrefs only (skip facebook/google/twitter/etc).
    LOGIN_TEXTS = ["כניסה", "Вход", "Войти", "Авторизация",
                   "Log in", "Login", "Sign in"]
    for text in LOGIN_TEXTS:
        try:
            el = page.get_by_role(
                "link", name=re.compile(rf"^{re.escape(text)}$", re.I)).first
            if await el.count():
                href = (await el.get_attribute("href") or "").lower()
                # Skip social / external links
                if any(bad in href for bad in
                       ("facebook", "google", "twitter", "apple",
                        "instagram", "linkedin")):
                    continue
                await el.click(timeout=5000)
                await asyncio.sleep(1.5)
                log(f"  ✓ Opened login form (clicked '{text}')")
                return True
        except Exception:
            pass
    log("  (login form already visible / no nav link needed)")
    return True

# ── YANDEX 2FA CODE READER ────────────────────────────────────────────────── #

def _imap_server(email_addr: str) -> tuple:
    """Return (host, port) for IMAP based on the email domain."""
    domain = email_addr.split("@")[-1].lower() if "@" in email_addr else ""
    if "gmail" in domain:
        return ("imap.gmail.com", 993)
    if any(x in domain for x in ("yandex", "ya.ru")):
        return ("imap.yandex.ru", 993)
    if any(x in domain for x in ("outlook", "hotmail", "live", "msn")):
        return ("outlook.office365.com", 993)
    if "mail.ru" in domain:
        return ("imap.mail.ru", 993)
    if "yahoo" in domain:
        return ("imap.mail.yahoo.com", 993)
    # generic fallback
    return (f"imap.{domain}", 993)


def _imap_read_mh_code(imap_email: str, imap_password: str,
                       max_wait_sec: int = 90) -> str | None:
    """
    Read MyHeritage verification code from the user's mailbox via IMAP.
    Follows the exact pattern from the project's test suite:
      - Connect via IMAP4_SSL
      - Login with full email + password
      - Search FROM "@myheritage.com"
      - Fetch last matching message
      - Extract 6-digit code with regex \\b(\\d{6})\\b
      - Poll every 5 seconds up to max_wait_sec
    """
    host, port = _imap_server(imap_email)
    deadline = time.time() + max_wait_sec
    while time.time() < deadline:
        try:
            mail = imaplib.IMAP4_SSL(host, port)
            mail.login(imap_email, imap_password)
            mail.select("INBOX")
            _, msg_ids = mail.search(None, '(FROM "@myheritage.com")')
            ids = msg_ids[0].split() if msg_ids[0] else []
            if not ids:
                # Also try broader search
                _, msg_ids2 = mail.search(None, '(FROM "myheritage")')
                ids = msg_ids2[0].split() if msg_ids2[0] else []
            if ids:
                _, msg_data = mail.fetch(ids[-1], "(RFC822)")
                msg = _email_lib.message_from_bytes(msg_data[0][1])
                body = ""
                if msg.is_multipart():
                    for part in msg.walk():
                        if part.get_content_type() in ("text/plain", "text/html"):
                            body = part.get_payload(decode=True).decode(
                                "utf-8", errors="ignore")
                            break
                else:
                    body = msg.get_payload(decode=True).decode(
                        "utf-8", errors="ignore")
                mail.close(); mail.logout()
                m = re.search(r"\b(\d{6})\b", body)
                if m:
                    return m.group(1)
            else:
                mail.close(); mail.logout()
        except Exception as exc:
            pass  # retry after sleep
        time.sleep(5)
    return None


async def _browser_read_yandex_code(ctx, mail_email: str, mail_password: str,
                                    log, timeout: int = 120) -> str | None:
    """
    Open a NEW TAB at https://mail.yandex.ru/, log in with password, open the
    topmost email and read the 6-digit MyHeritage code. Then close the tab.
    Uses the EXACT selectors from the live Yandex login flow.
    """
    async def _click_any(sels, label, wait=8000):
        """Wait for and click the first matching selector. Returns True/False."""
        for sel in sels:
            try:
                el = page.locator(sel).first
                await el.wait_for(state="visible", timeout=wait)
                await el.click(timeout=5000)
                log(f"  2FA: {label}  ({sel})")
                return True
            except Exception:
                continue
        log(f"  2FA: !! не нашёл — {label}")
        return False

    async def _fill_any(sels, value, label, wait=8000):
        for sel in sels:
            try:
                el = page.locator(sel).first
                await el.wait_for(state="visible", timeout=wait)
                await el.click(timeout=4000)
                await el.fill(value)
                await asyncio.sleep(0.4)
                if (await el.input_value()).strip():
                    log(f"  2FA: {label} заполнено  ({sel})")
                    return True
            except Exception:
                continue
        log(f"  2FA: !! поле не найдено — {label}")
        return False

    # Protect the mail tab from the junk-tab auto-closer (_on_new_page): without
    # this it closes the tab after ~3s for "not being MyHeritage", which killed
    # the 2FA flow on fresh logins (notably the .co.il site, whose cookies don't
    # carry over from .com so it always needs a real login + email code).
    try:
        ctx._mh_pause_autoclose = True
    except Exception:
        pass
    page = await ctx.new_page()
    try:
        page._mh_protected = True
    except Exception:
        pass
    try:
        await page.bring_to_front()
        log("  2FA: открываю новую вкладку mail.yandex.ru …")
        await page.goto("https://mail.yandex.ru/", wait_until="domcontentloaded",
                        timeout=30000)
        await asyncio.sleep(2)

        # The PERSISTENT profile usually keeps the Yandex session, so the inbox
        # opens straight away. In that case there is NO login form — running the
        # login dance clicks a stray "Войти", navigates AWAY from the inbox and
        # breaks the tab. So detect "already logged in" and skip login entirely.
        async def _mail_logged_in():
            try:
                return await page.evaluate(r"""() => {
                    if (document.querySelector('#header-login-button')) return false;
                    const txt = e => (e.textContent || '').trim();
                    const compose = Array.from(
                        document.querySelectorAll('a,button,span,div'))
                        .some(e => /^(Написать|Compose|Написати|כתוב)$/i.test(txt(e)));
                    const rows = document.querySelectorAll(
                        '[class*="MessageSnippet"], .mail-MessageSnippet, '
                        + 'li[data-id]').length;
                    return compose || rows > 0
                           || /[#/](inbox|message)/i.test(location.href);
                }""")
            except Exception:
                return False

        # Poll up to ~10s for the logged-in state — the persistent session can
        # briefly show the landing page before redirecting into the inbox.
        logged_in_mail = False
        for _ in range(10):
            if await _mail_logged_in():
                logged_in_mail = True
                break
            await asyncio.sleep(1)

        if logged_in_mail:
            log("  2FA: уже залогинен в Яндекс (профиль) — вход пропускаю")
        else:
            # 1) Click "Войти" — ONLY the specific header button (a loose
            #    text match grabs a stray «Войти» ad/footer link → breaks flow).
            await _click_any(['#header-login-button'], "нажал «Войти»")
            await asyncio.sleep(2)

            # 2) Enter email/username (wait for passport page to render)
            await _fill_any(
                ['input[data-testid="text-field-input"][autocomplete="username"]',
                 'input[autocomplete="username"]',
                 'input[name="login"]', '#passp-field-login'],
                mail_email, f"логин {mail_email}", wait=15000)

            # 3) Click "Next"
            await _click_any(['button[data-testid="add-user-next"]',
                              'button:has-text("Next")', 'button:has-text("Далее")'],
                             "нажал Next")
            await asyncio.sleep(2)

            # 4) Choose "Log in with your password" (skip the one-time-code step)
            await _click_any(['span:has-text("Log in with your password")',
                              'button:has-text("Log in with your password")',
                              'span:has-text("Войти с паролем")',
                              'button:has-text("Войти с паролем")'],
                             "выбрал «Log in with your password»")
            await asyncio.sleep(1.5)

            # 5) Enter password
            filled_pw = await _fill_any(
                ['input[data-testid="text-field-input"][autocomplete="current-password"]',
                 'input[autocomplete="current-password"]',
                 'input[type="password"]', '#passp-field-passwd'],
                mail_password, "пароль почты", wait=15000)
            if filled_pw:
                await page.keyboard.press("Enter")
                await asyncio.sleep(1)
                await _click_any(['button[data-testid="add-user-next"]',
                                  'button:has-text("Sign in")',
                                  'button:has-text("Войти")'],
                                 "подтвердил вход", wait=4000)

            # 6) Wait for the inbox to load after login
            try:
                await page.wait_for_url(
                    lambda u: "mail.yandex" in u and "passport" not in u,
                    timeout=25000)
            except Exception:
                await asyncio.sleep(5)
        log("  2FA: в почте, ищу письма с кодами …")

        # MyHeritage sends TWO emails with DIFFERENT codes:
        #   B) "Ваш код подтверждения для входа в MyHeritage" / "verification
        #      code" — this is the code the «Введите верификационный код» field
        #      actually wants (priority).
        #   A) "Confirm a login attempt on MyHeritage" — the anti-fraud
        #      challenge ("flagged as suspicious"); its code is the fallback.
        # We read BOTH from the inbox snippets and return them in priority
        # order so the caller can try B first, then A.
        async def _scan_codes():
            # Read the code STRAIGHT FROM THE INBOX LIST (the email's row header
            # + preview), WITHOUT opening the message — that is how it worked.
            # CRUCIAL: match WHOLE message rows, not leaf elements. The 6-digit
            # code sits in the PREVIEW while the «код подтверждения для входа»
            # phrase is in the SUBJECT; with broad 'a,li,div' selectors they land
            # in DIFFERENT elements, so neither matches both → nothing found.
            # The MessageSnippet/li[data-id] container holds the whole row.
            return await page.evaluate(r"""() => {
                const rows = Array.from(document.querySelectorAll(
                    '[class*="MessageSnippet" i], .mail-MessageSnippet, '
                    + '[data-test-id="message-snippet"], li[data-id], '
                    + 'div[role="listitem"], a[href*="#message"]'));
                let verify = '', attempt = '';
                for (const r of rows) {
                    const t = (r.textContent || '');
                    if (!/myheritage/i.test(t)) continue;
                    const m = t.match(/\b(\d{6})\b/);
                    if (!m) continue;                 // row with no 6-digit → skip
                    const code = m[1];
                    if (!verify && /(подтверждения для входа|верификацион|verification code|confirmation code)/i.test(t))
                        verify = code;                // B — the one the field wants
                    else if (!attempt && /(login attempt|confirm a login|попытк)/i.test(t))
                        attempt = code;               // A — anti-fraud challenge
                    else if (!verify)
                        verify = code;                // MyHeritage + code, no clear phrase → treat as B
                }
                return {verify, attempt};
            }""")

        async def _open_and_read():
            """Snippet scan can miss the code (truncated/changed DOM) — so OPEN
            the newest MyHeritage *verification* email and read the code from its
            body. This is the step that used to work and was lost in a refactor.
            We match MyHeritage + a verification phrase, so the hh.ru ad at the
            very top is skipped (the lesson's «не открывать верхнее письмо»)."""
            try:
                marked = await page.evaluate(r"""() => {
                    const norm = s => (s || '').toLowerCase();
                    document.querySelectorAll('[data-pw-mail]').forEach(
                        e => e.removeAttribute('data-pw-mail'));
                    const rows = Array.from(document.querySelectorAll(
                        '[class*="MessageSnippet"], .mail-MessageSnippet, '
                        + '[data-test-id="message-snippet"], li[data-id], a'));
                    for (const r of rows) {                 // DOM order = newest first
                        const t = norm(r.textContent);
                        if (t.includes('myheritage') &&
                            /(код подтверждения для входа|верификацион|verification code|confirmation code)/.test(t)) {
                            r.setAttribute('data-pw-mail', '1');
                            return true;
                        }
                    }
                    return false;
                }""")
                if not marked:
                    return ""
                log("  2FA: открываю письмо MyHeritage с кодом …")
                await page.locator('[data-pw-mail="1"]').first.click(timeout=6000)
                await asyncio.sleep(2.5)
                # Prefer the OPENED message's body container (so we don't pick an
                # older code still shown in the inbox list); fall back to body.
                body = await page.evaluate(r"""() => {
                    const el = document.querySelector(
                        '[class*="MessageBody" i], .mail-Message-Body, '
                        + '[class*="message-body" i], [class*="msgBody" i], '
                        + '[class*="MessageViewer" i]');
                    return (el && el.innerText) || document.body.innerText || '';
                }""")
                m = (re.search(
                        r"(?:странице входа в MyHeritage|login screen|"
                        r"verification code|код подтверждения)[^\d]{0,40}(\d{6})",
                        body or "", re.I)
                     or re.search(r"\b(\d{6})\b", body or ""))
                code = m.group(1) if m else ""
                if code:
                    log(f"  2FA: код из ОТКРЫТОГО письма: {code}")
                # Return to the inbox so the next reload/scan keeps working.
                try:
                    await page.goto("https://mail.yandex.ru/",
                                    wait_until="domcontentloaded", timeout=15000)
                except Exception:
                    pass
                return code
            except Exception as e:
                log(f"  2FA: !! не смог открыть письмо: {e}")
                return ""

        deadline = asyncio.get_event_loop().time() + timeout
        last = {"verify": "", "attempt": ""}
        while asyncio.get_event_loop().time() < deadline:
            # Stop if the mail tab was closed (don't spin on a dead page).
            if page.is_closed():
                log("  2FA: вкладка почты закрыта — прекращаю чтение")
                break
            # Let the freshly-sent emails arrive, then reload to get them on top
            await asyncio.sleep(5)
            try:
                await page.reload(wait_until="domcontentloaded", timeout=15000)
                await asyncio.sleep(2)
            except Exception:
                pass
            try:
                last = await _scan_codes()
            except Exception:
                last = {"verify": "", "attempt": ""}
            # Snippet scan missed the verification code (B) → OPEN the email and
            # read the code from its body (the reliable path that used to work).
            if not last.get("verify"):
                opened = await _open_and_read()
                if opened:
                    last["verify"] = opened
            # We have at least the verification code (B) — good to go
            if last.get("verify") or last.get("attempt"):
                codes = []
                if last.get("verify"):
                    codes.append(last["verify"])
                    log(f"  2FA: код «подтверждения для входа» (B): {last['verify']}")
                if last.get("attempt") and last["attempt"] not in codes:
                    codes.append(last["attempt"])
                    log(f"  2FA: код «login attempt» (A): {last['attempt']}")
                # Give the second email a couple seconds in case only one is in yet
                if not last.get("verify"):
                    await asyncio.sleep(4)
                    try:
                        again = await _scan_codes()
                        if again.get("verify") and again["verify"] not in codes:
                            codes.insert(0, again["verify"])
                            log(f"  2FA: код «подтверждения для входа» (B): {again['verify']}")
                    except Exception:
                        pass
                return codes

        log("  2FA: коды в почте не найдены за отведённое время")
        return []
    except Exception as e:
        log(f"  2FA browser error: {e}")
        return []
    finally:
        # Close the mail tab regardless of outcome, then re-enable the auto-closer
        try:
            await page.close()
        except Exception:
            pass
        try:
            ctx._mh_pause_autoclose = False
        except Exception:
            pass


async def _get_2fa_codes(ctx, mail_email: str, mail_password: str,
                         log, ask_2fa_code=None) -> list:
    """
    Return a PRIORITY LIST of candidate 2FA codes (MyHeritage sends two
    different emails). The caller tries them in order.
      1. Browser: open a new tab to Yandex mail, read both codes.
      2. IMAP fallback for non-Yandex providers.
      3. GUI dialog as a last resort.
    `mail_email` is the SAME email used for MyHeritage login.
    """
    if mail_email and mail_password:
        codes = await _browser_read_yandex_code(ctx, mail_email, mail_password, log)
        if codes:
            return codes
        log("  2FA: пробую IMAP как запасной вариант …")
        code = await asyncio.to_thread(
            _imap_read_mh_code, mail_email, mail_password, 60)
        if code:
            log("  2FA: код получен по IMAP ✓")
            return [code]

    if ask_2fa_code:
        log("  2FA: запрашиваю код у пользователя …")
        c = ask_2fa_code()
        return [c] if c else []
    return []


# ── LOGIN ─────────────────────────────────────────────────────────────────── #
async def _login(page, login_url, has_cookies,
                 email, password, log,
                 ask_2fa_code=None,
                 imap_password=None) -> bool:
    """
    Full login flow:
    1. Navigate to login_url
    2. Accept cookie banner if has_cookies
    3. Open login form if needed
    4. Fill email + password
    5. Handle 2-FA code dialog if it appears
    """
    log(f"  → Navigating to {login_url} …")
    try:
        await page.goto(login_url, wait_until="domcontentloaded", timeout=35000)
    except Exception as exc:
        log(f"  !! Cannot load login page: {exc}")
        return False

    await asyncio.sleep(2)
    log(f"     Landed: {page.url}")

    # Step 1: cookies (always — the consent banner can show on any locale)
    await _accept_cookies(page, log)
    await asyncio.sleep(0.5)

    # Step 0: ALREADY logged in? With the persistent profile MyHeritage
    # redirects away from /login (e.g. to the family-site home). If we are not
    # on a login/signin page, there's nothing to do — skip the whole form.
    cur = page.url.lower()
    if "login" not in cur and "signin" not in cur:
        log("  ✓ Уже авторизован (сессия из профиля) — логин не нужен")
        # Close any social popup tab that may have opened
        await _close_social_tabs(page.context, page, log)
        return True

    # Step 2: open login form (click nav link if needed)
    await _open_login_form(page, log)

    # Step 3: wait for email field
    # Exact selectors from real page HTML (RU site, 2025):
    #   <input type="text" name="registrationEmail" id="registrationEmail" autocomplete="off">
    #   <input type="password" name="registrationLoginPassword" id="registrationLoginPassword">
    EMAIL_SELS = [
        '#registrationEmail',
        'input[name="registrationEmail"]',
        'input[data-automations="login_email_block_input"]',
        '#email',
        'input[name="email"]',
        'input[type="email"]',
        'input[placeholder*="эл. почты" i]',
        'input[placeholder*="email" i]',
        'input[placeholder*="mail" i]',
        'input[placeholder*="דוא" i]',
        'input[placeholder*="почт" i]',
    ]
    PASS_SELS = [
        '#registrationLoginPassword',
        'input[name="registrationLoginPassword"]',
        '#password',
        'input[name="password"]',
        'input[type="password"]',
        'input[placeholder*="Пароль" i]',
        'input[placeholder*="пароль" i]',
        'input[placeholder*="assword" i]',
        'input[placeholder*="סיסמ" i]',
    ]
    # Wait up to 10s for the email field to appear
    email_visible = False
    for sel in EMAIL_SELS:
        try:
            await page.locator(sel).first.wait_for(state="visible", timeout=6000)
            email_visible = True
            log(f"  ✓ Email field visible: {sel}")
            break
        except Exception:
            continue

    if not email_visible:
        log("  !! Email field not found — saving screenshot for debugging…")
        try:
            ss = Path(__file__).resolve().parent / "debug_login.png"
            await page.screenshot(path=str(ss), full_page=True)
            log(f"     Screenshot saved: {ss}")
        except Exception:
            pass
        return False

    # Step 4: fill form
    ok_e = await _fill(page, EMAIL_SELS, email,    "Email",    log)
    ok_p = await _fill(page, PASS_SELS,  password, "Password", log)
    if not ok_e or not ok_p:
        return False

    # Step 5: submit
    SUBMIT_SELS = [
        'button:has-text("היכנס")',        # Hebrew submit
        'button:has-text("Авторизация")',  # Russian submit (seen in screenshot)
        'button:has-text("Войти")',
        'button:has-text("Log in")',
        'button:has-text("Sign in")',
        'button:has-text("Login")',
        'button[type="submit"]',
        'input[type="submit"]',
    ]
    submitted = False
    for sel in SUBMIT_SELS:
        try:
            el = page.locator(sel).first
            if await el.count():
                await el.click(timeout=6000)
                submitted = True
                log(f"  ✓ Clicked submit ({sel})")
                break
        except Exception:
            continue
    if not submitted:
        await page.keyboard.press("Enter")

    # ── Step 6: wait for EITHER login success OR a 2FA code field ──────── #
    # MyHeritage often shows a "verify it's you" code screen on the SAME
    # /login URL after submit, and it can take >5s to appear. So we POLL.
    TFA_SELS = [
        'input[placeholder*="код" i]',
        'input[placeholder*="code" i]',
        'input[placeholder*="verification" i]',
        'input[placeholder*="верификац" i]',
        'input[type="number"][maxlength]',
        'input[maxlength="6"]',
        'input[autocomplete="one-time-code"]',
        'input[autocomplete*="one-time"]',
        'input[name*="code"]',
        'input[name*="otp"]',
        'input[name*="token"]',
        'input[inputmode="numeric"]',
    ]

    async def _logged_in() -> bool:
        u = page.url.lower()
        if any(x in u for x in ("login", "signin", "verify", "auth")):
            return False
        # Guard: if a password or code field is still visible we are NOT in yet
        for sel in ('input[type="password"]', 'input[autocomplete="one-time-code"]',
                    'input[maxlength="6"]'):
            try:
                if await page.locator(sel).first.is_visible(timeout=500):
                    return False
            except Exception:
                pass
        return True

    async def _find_tfa():
        for sel in TFA_SELS:
            try:
                el = page.locator(sel).first
                if await el.count() and await el.is_visible():
                    return sel
            except Exception:
                pass
        return None

    tfa_sel = None
    for _ in range(30):                       # poll up to ~30s
        await asyncio.sleep(1)
        # Check for the 2FA code field FIRST — MH may sit on a non-/login URL
        # while showing the verification screen, so a premature "logged in"
        # check would wrongly skip it.
        tfa_sel = await _find_tfa()
        if tfa_sel:
            log(f"  ⚠  2FA code field detected ({tfa_sel})")
            break
        if await _logged_in():
            log(f"  ✓ Logged in (no 2FA). URL: {page.url}")
            return True

    if tfa_sel:
        # MyHeritage sends two emails with two different codes — get both.
        codes = await _get_2fa_codes(page.context, email, imap_password,
                                     log, ask_2fa_code)
        if not codes:
            log("  !! 2FA: код не получен — прерываю.")
            return False

        async def _digits_in_field() -> str:
            try:
                return await page.evaluate("""() => {
                    const els = document.querySelectorAll(
                        'input[inputmode="numeric"], input[autocomplete="one-time-code"], '
                        + 'input[maxlength="6"], input[name*="code"], input[type="number"]');
                    let s = '';
                    els.forEach(e => { if (e.offsetParent !== null) s += (e.value || ''); });
                    return s.replace(/\\D/g, '');
                }""")
            except Exception:
                return ""

        async def _clear_code():
            """Clear all visible code boxes before entering a new code."""
            try:
                boxes = page.locator(
                    'input[inputmode="numeric"], input[autocomplete="one-time-code"], '
                    'input[maxlength="6"], input[name*="code"]')
                n = await boxes.count()
                for i in range(max(n, 1)):
                    b = boxes.nth(i) if n else page.locator(tfa_sel).first
                    try:
                        await b.click(timeout=1500)
                        await page.keyboard.press("Control+a")
                        await page.keyboard.press("Delete")
                    except Exception:
                        pass
            except Exception:
                pass

        async def _enter_code(code: str) -> bool:
            for _ in range(3):
                try:
                    fld = page.locator(tfa_sel).first
                    await fld.click(timeout=4000)
                    await page.keyboard.press("Control+a")
                    await page.keyboard.press("Delete")
                    await page.keyboard.type(code, delay=140)
                    await asyncio.sleep(0.6)
                except Exception:
                    pass
                if code in (await _digits_in_field()):
                    return True
                # segmented field — one digit per box
                try:
                    boxes = page.locator(
                        'input[inputmode="numeric"], input[autocomplete="one-time-code"]')
                    if await boxes.count() >= len(code):
                        for i, ch in enumerate(code):
                            b = boxes.nth(i)
                            await b.click(timeout=2000); await b.fill(ch)
                            await asyncio.sleep(0.1)
                        await asyncio.sleep(0.5)
                except Exception:
                    pass
                if code in (await _digits_in_field()):
                    return True
                await asyncio.sleep(0.4)
            return False

        async def _submit_once() -> bool:
            # The confirm button lives INSIDE the verification MODAL (next to the
            # code field). On the page behind it there is ALSO an "Авторизация"
            # button (the login form) — we must NOT click that one. So: find the
            # code field, climb to its modal/dialog/form container, locate the
            # action button inside it, tag it, and click THAT with a real click.
            tagged = await page.evaluate(r"""() => {
                document.querySelectorAll('[data-pw-2fa]').forEach(
                    e => e.removeAttribute('data-pw-2fa'));
                const norm = s => (s || '').replace(/\s+/g, ' ').trim();
                const fld = document.querySelector(
                    'input[inputmode="numeric"], input[autocomplete="one-time-code"], '
                    + 'input[maxlength="6"], input[name*="code"]');
                if (!fld) return false;
                // climb to a reasonable container (dialog/modal/form)
                let box = fld.closest('[role=dialog], .modal, form') || fld.parentElement;
                for (let i = 0; i < 5 && box && box.parentElement; i++) {
                    const btns = box.querySelectorAll('button, [role=button]');
                    // a submit-ish button whose text is Отправить/Авторизация/Continue
                    const re = /^(Отправить|Авторизация|Continue|Продолжить|Verify|Подтвердить|Submit|Send)$/i;
                    let hit = Array.from(btns).find(b => re.test(norm(b.textContent)));
                    if (!hit && btns.length) hit = btns[btns.length - 1];
                    if (hit) { hit.setAttribute('data-pw-2fa', '1'); return true; }
                    box = box.parentElement;
                }
                return false;
            }""")
            if tagged:
                try:
                    btn = page.locator('[data-pw-2fa="1"]').first
                    await btn.scroll_into_view_if_needed(timeout=3000)
                    await btn.click(timeout=5000)
                    log("  ✓ 2FA: нажал кнопку подтверждения в окне кода")
                    return True
                except Exception as e:
                    log(f"  2FA: реальный клик не вышел ({e}); пробую JS-клик")
                    try:
                        await page.evaluate(
                            "() => { const b=document.querySelector('[data-pw-2fa=\"1\"]');"
                            " if (b) b.click(); }")
                        log("  ✓ 2FA: нажал кнопку (JS)")
                        return True
                    except Exception:
                        pass
            # Last resort: a single Enter
            await page.keyboard.press("Enter")
            log("  2FA: Enter в поле кода")
            return False

        # Bring MH tab back to front (mail tab was focused) so input lands here
        try:
            await page.bring_to_front()
            await asyncio.sleep(0.4)
        except Exception:
            pass

        # Try each candidate code: enter → submit ONCE → wait for success.
        for ci, code in enumerate(codes, 1):
            log(f"  → Пробую код {ci}/{len(codes)}: {code}")
            await _clear_code()
            ok = await _enter_code(code)
            log("  ✓ 2FA: код введён" if ok
                else "  !! 2FA: код не подтвердился в поле — всё равно отправляю")
            await asyncio.sleep(0.4)
            await _submit_once()
            log("  2FA: отправил, жду результат…")
            for _ in range(18):
                await asyncio.sleep(1)
                if await _logged_in():
                    log(f"  ✓ Logged in after 2FA. URL: {page.url}")
                    return True
            log(f"  !! 2FA: код {code} не подошёл, пробую следующий…")
            await asyncio.sleep(1)

        log("  !! 2FA: ни один код не подошёл.")

    # Neither success nor 2FA → report any error message / anti-bot block
    cur = page.url
    try:
        body = (await page.evaluate(
            "() => (document.body.innerText || '').slice(0, 600)") or "").lower()
        if any(k in body for k in (
                "robot", "captcha", "recaptcha", "подозрительн", "too many",
                "слишком много", "заблокир", "blocked", "unusual activity",
                "необычн", "verify you", "докажите")):
            log("  !! Похоже, MyHeritage временно ОГРАНИЧИЛ вход (анти-бот/капча). "
                "Это не баг кода — частые попытки входа усугубляют. Подожди "
                "10–15 мин и войди в браузере профиля вручную один раз.")
        err = (await page.locator('[class*="error" i],[role="alert"]')
               .first.text_content(timeout=2000) or "").strip()
        if err:
            log(f"  !! Login error: {err!r}")
    except Exception:
        pass
    log(f"  !! Login failed. URL: {cur}")
    return False

# ── SELECT FAMILY SITE (FP/select-site.php) ───────────────────────────────── #
async def _handle_select_site(page, family_site, log):
    """
    Determine the research search URL for the chosen family site.
      - On FP/select-site.php → click the chosen site (by name, else first).
      - Already on a family-sites/<slug>/<ID> page → extract <ID> from the URL.
    Returns the research URL (research?s=<id>) or None.
    """
    cur = page.url
    # Case B: we're already inside a family site → derive id from the URL
    m = re.search(r"/family-sites/[^/]+/([A-Z0-9]+)", cur)
    if m and "select-site" not in cur.lower():
        sid = m.group(1)
        lang = "RU"
        lm = re.search(r"[?&]lang=([A-Za-z]+)", cur)
        if lm:
            lang = lm.group(1)
        log(f"  → Уже на семейном сайте (id {sid}) — иду в поиск")
        return f"https://www.myheritage.com/research?s={sid}&lang={lang}"

    if "select-site" not in cur.lower():
        return None
    log("  → Страница выбора семейного сайта")
    await asyncio.sleep(1.5)

    # Each site link:  <a ... onclick="goToSiteClicked('SITEID',
    #   'https://www.myheritage.com/family-sites/<slug>/<SITEID>?lang=..')">Name</a>
    target = await page.evaluate(r"""(wanted) => {
        const links = Array.from(document.querySelectorAll('a[onclick*="goToSiteClicked"]'));
        const parse = (a) => {
            const m = (a.getAttribute('onclick') || '')
                .match(/goToSiteClicked\(\s*'([^']+)'\s*,\s*'([^']+)'/);
            return m ? {id: m[1], url: m[2], name: (a.textContent||'').trim()} : null;
        };
        const all = links.map(parse).filter(Boolean);
        if (!all.length) return null;
        if (wanted) {
            const hit = all.find(s => s.name.toLowerCase()
                                       .includes(wanted.toLowerCase()));
            if (hit) return hit;
        }
        return all[0];   // default: first (usually the admin's own site)
    }""", family_site or "")

    if not target:
        log("  !! Сайты не найдены на странице выбора")
        return None
    log(f"  → Перехожу на сайт: {target['name']}")
    await page.goto(target["url"], wait_until="domcontentloaded", timeout=35000)
    await asyncio.sleep(2)
    # Build the research search URL for this site: /research?s=<id>&lang=..
    lang = "RU"
    m = re.search(r"[?&]lang=([A-Za-z]+)", target["url"])
    if m:
        lang = m.group(1)
    return f"https://www.myheritage.com/research?s={target['id']}&lang={lang}"


# ── close stray social popup tabs (Facebook/Google) ───────────────────────── #
async def _close_social_tabs(ctx, keep, log):
    for p in list(ctx.pages):
        if p is keep:
            continue
        try:
            u = (p.url or "").lower()
            if any(s in u for s in ("facebook.com", "accounts.google", "apple.com")):
                await p.close()
                log("  ✓ Закрыл лишнюю вкладку соцсети")
        except Exception:
            pass


# ── SEARCH FORM ───────────────────────────────────────────────────────────── #
FIRST_NAME_SELS = [
    'input[data-automations="research-family_first_name"]',
    'input[placeholder="Имя и отчество"]',
    'input[placeholder*="Имя" i]', 'input[placeholder*="first" i]',
]
LAST_NAME_SELS = [
    'input[data-automations="research-family_last_name"]',
    'input[placeholder="Фамилия"]',
    'input[placeholder*="Фамилия" i]', 'input[placeholder*="last" i]',
]


async def _find_form_root(page, sels, log, secs=25):
    """
    Return the frame (main or iframe) that contains the research search form,
    polling up to `secs`. Detect by the first-name field OR the «Поиск» submit
    button (the form can be in the main frame or an iframe).
    """
    _sw = _search_words()
    btn_sels = ([f'button:has(span.button_content:has-text("{w}"))' for w in _sw]
                + [f'button:has-text("{w}")' for w in _sw])
    for _t in range(secs):
        await asyncio.sleep(1)
        for fr in page.frames:
            for sel in list(sels) + btn_sels:
                try:
                    if await fr.locator(sel).first.count():
                        log(f"  ✓ Форма поиска готова ({_t+1}с"
                            f"{', iframe' if fr is not page.main_frame else ''})")
                        return fr
                except Exception:
                    continue
    return None


async def _set_name_match(root, page, params, log):
    """MyHeritage has name-matching options in a dropdown under the «Имя» field
    («Искать совпадения строго по имени» / «Варианты написания» / «Совпадение
    инициалов» / «Начинается с…»). They decide whether the search surfaces
    spelling / initial variants — e.g. «Alexander-Wolf Sanders (Shenderovich)»
    for «Alexander W Sanders». The desired state comes from the GUI checkboxes.
    Checkboxes are role=checkbox spans (data-automations=check_box_control_label)."""
    DESIRED = [
        ("строго по имени",   bool(params.get("name_strict"))),
        ("Варианты написани", bool(params.get("name_variants"))),
        ("Совпадение инициал", bool(params.get("name_initials"))),
        ("Начинается с",      bool(params.get("name_startswith"))),
    ]
    # HARD-WON facts about these checkboxes (verified on the live DOM):
    #  • JS .click() WORKS (React applies it); a REAL Playwright click does NOT —
    #    the option spans are 0×0 overlays → click() times out → nothing changed.
    #  • EVERY option is DUPLICATED (mobile + desktop renders) → must click ALL
    #    matches, not just the first (the first was a hidden dup → «не передавалось»).
    #  • «Искать совпадения строго по имени» is shared by the NAME and the SURNAME
    #    popup → for «строго» touch ONLY the copies whose popup also has «Варианты
    #    написания» (that's the NAME popup), never the surname's.
    try:
        fn = root.locator('[data-pw-rf="first"]').first
        if await fn.count():
            await fn.click(timeout=2500)          # focus → opens the name popup
            await asyncio.sleep(0.7)
        res = await root.evaluate(r"""(desired) => {
            const norm = s => (s || '').replace(/\s+/g, ' ').trim();
            const all = [...document.querySelectorAll(
                '[role=checkbox][data-automations=check_box_control_label]')];
            // NAME popup = «строго» whose CLOSE ancestor (≤3 levels — they sit as
            // siblings under one «relative_layout») also holds «Варианты». 3, NOT
            // 8: 8 reached the shared «search_form» and wrongly flagged surname's.
            const inNamePopup = (b) => {
                let n = b;
                for (let i = 0; i < 3 && n; i++) {
                    if ([...n.querySelectorAll(
                          '[role=checkbox][data-automations=check_box_control_label]')]
                          .some(x => norm(x.textContent).includes('Варианты написани')))
                        return true;
                    n = n.parentElement;
                }
                return false;
            };
            const out = [];
            for (const [frag, want] of desired) {
                let cbs = all.filter(x => norm(x.textContent).includes(frag));
                if (frag.includes('строго')) cbs = cbs.filter(inNamePopup);
                let did = false;
                for (const b of cbs)
                    if ((b.getAttribute('aria-checked') === 'true') !== want) {
                        b.click(); did = true;       // click ALL duplicates
                    }
                if (did) out.push(frag + (want ? ' вкл' : ' выкл'));
            }
            return out;
        }""", DESIRED)
        if res:
            log(f"  → Имя, опции совпадения: {', '.join(res)}")
    except Exception as e:
        log(f"  !! опции имени: {type(e).__name__}: {e}")


async def _set_field_checkbox(root, page, field_tag, frag, want, label, log,
                              exclude_name=False):
    """Set a field's match checkbox via JS .click() (the only thing that works —
    the spans are 0×0 overlays; a real Playwright click / visibility filter fails).
    Clicks ALL duplicates. `exclude_name`: for the surname «строго» (text shared
    with the name popup) — skip copies whose CLOSE ancestor (≤3) holds «Варианты»
    (those belong to the name popup), so we touch only the surname's own copy."""
    try:
        fld = root.locator(f'[data-pw-rf="{field_tag}"]').first
        if await fld.count():
            await fld.click(timeout=2500)             # focus → its popup opens
            await asyncio.sleep(0.6)
        res = await root.evaluate(r"""(args) => {
            const [frag, want, excludeName] = args;
            const norm = s => (s || '').replace(/\s+/g, ' ').trim();
            const inNamePopup = (b) => {
                let n = b;
                for (let i = 0; i < 3 && n; i++) {
                    if ([...n.querySelectorAll(
                          '[role=checkbox][data-automations=check_box_control_label]')]
                          .some(x => norm(x.textContent).includes('Варианты написани')))
                        return true;
                    n = n.parentElement;
                }
                return false;
            };
            let cbs = [...document.querySelectorAll(
                '[role=checkbox][data-automations=check_box_control_label]')]
                .filter(x => norm(x.textContent).toLowerCase()
                              .includes(frag.toLowerCase()));
            if (excludeName) cbs = cbs.filter(b => !inNamePopup(b));
            let did = false;
            for (const b of cbs)
                if ((b.getAttribute('aria-checked') === 'true') !== want) {
                    b.click(); did = true;
                }
            return did ? (want ? 'вкл' : 'выкл') : '';
        }""", [frag, want, exclude_name])
        if res:
            log(f"  → {label}: {res}")
    except Exception as e:
        log(f"  !! {label}: {type(e).__name__}")


async def _set_year_match(root, page, params, log):
    """Birth-year match. VERIFIED on the live DOM: the ± tolerance radios render
    ONLY AFTER «Точное совпадение года» is TICKED («сначала точное, потом разброс»).
    So: tick exact → the radios appear → click «match_plus_minus_{N}» (N=0 is «В
    этом году» = exact). `year_match`: «exact»/«0» → ±0; «1/2/5/10/20» → that radio."""
    ym = str(params.get("year_match") or "").strip().lower()
    if not ym or not params.get("birth_year"):
        return
    try:
        fld = root.locator('[data-pw-rf="year"]').first
        if await fld.count():
            await fld.click(timeout=2500)
            await asyncio.sleep(0.5)
        # ONE tick of the exact-match checkbox (reveals the radios), then ONE click
        # on the tolerance radio — the VISIBLE copy only, no «дёрганья» over dups.
        # Located by container/locator (language-independent).
        await root.evaluate(r"""() => {
            const ex = [...document.querySelectorAll(
                '[data-automations="birth_year_open_field_filter_matching_options_dropdown"]'
                + ' [role=checkbox][data-automations=check_box_control_label]')]
                .filter(b => b.getBoundingClientRect().width > 0);
            const b = ex[0] || document.querySelector(
                '[data-automations="birth_year_open_field_filter_matching_options_dropdown"]'
                + ' [role=checkbox][data-automations=check_box_control_label]');
            if (b && b.getAttribute('aria-checked') !== 'true') b.click();
        }""")
        await asyncio.sleep(0.5)                        # let the radios render
        n = "0" if re.search(r"exact|точн|^0$", ym) else (re.sub(r"\D", "", ym) or "5")
        res = await root.evaluate(r"""(n) => {
            const rs = [...document.querySelectorAll(
                '[data-automations="birth_year_open_field_filter_match_plus_minus_' + n + '"]')];
            const vis = rs.filter(r => r.getBoundingClientRect().width > 0);
            const r = vis[0] || rs[0];                  // ONE click, prefer visible
            if (r) { r.click(); return '±' + n; }
            return 'нет радио ±' + n;
        }""", n)
        log(f"  → Совпадение года: {'точно (В этом году)' if n == '0' and res.startswith('±') else res}")
    except Exception as e:
        log(f"  !! год match: {type(e).__name__}")


async def _fill_research_basic(root, page, params, log) -> bool:
    """
    Robustly fill the basic research fields. Selectors for the inputs vary, but
    the «Поиск» button is reliable — so we walk up from it to the form
    container and TAG its text inputs (by data-automations / placeholder /
    order), then fill the tagged fields. Works in main frame or iframe.
    """
    tagged = await root.evaluate(r"""(searchWords) => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
        const vis = el => el && el.offsetParent !== null;
        const isSearch = t => searchWords.some(w => norm(t).toLowerCase() === w.toLowerCase());
        // find the search submit button (in the active site language)
        let btn = Array.from(document.querySelectorAll('button, [role=button], span'))
            .find(b => isSearch(b.textContent));
        // climb to a container that holds >= 2 visible text inputs (the form)
        let form = null, n = btn;
        for (let i = 0; i < 9 && n; i++) {
            const ins = Array.from(n.querySelectorAll('input'))
                .filter(x => vis(x) && /^(text|search|number|)$/i.test(x.type || ''));
            if (ins.length >= 2) { form = n; break; }
            n = n.parentElement;
        }
        if (!form) form = document.querySelector('form') || document.body;
        const inputs = Array.from(form.querySelectorAll('input'))
            .filter(x => vis(x) && /^(text|search|number|)$/i.test(x.type || ''));
        document.querySelectorAll('[data-pw-rf]').forEach(e => e.removeAttribute('data-pw-rf'));
        const byAuto = sub => inputs.find(i =>
            ((i.getAttribute('data-automations') || '').toLowerCase()).includes(sub));
        const byPh = re => inputs.find(i => re.test(i.getAttribute('placeholder') || ''));
        const first = byAuto('first_name') || byPh(/Имя|first/i) || inputs[0];
        const last  = byAuto('last_name')  || byPh(/Фамилия|last|surname/i) || inputs[1];
        const year  = byAuto('birth')      || byPh(/Год рождения|birth\s*year/i);
        const place = byAuto('place')      || byPh(/Насел|place/i);
        const tag = (el, name) => { if (el) el.setAttribute('data-pw-rf', name); };
        tag(first, 'first'); tag(last, 'last'); tag(year, 'year'); tag(place, 'place');
        return {first: !!first, last: !!last, year: !!year, place: !!place,
                count: inputs.length};
    }""", _search_words(params.get("lang")))
    log(f"  → Поля формы: {tagged}")

    async def _type(name, value, label):
        if not value:
            return
        try:
            el = root.locator(f'[data-pw-rf="{name}"]').first
            if not await el.count():
                log(f"  !! Field not found: {label}")
                return
            await el.scroll_into_view_if_needed(timeout=4000)
            await el.click(timeout=3000)
            await page.keyboard.press("Control+a")
            await page.keyboard.press("Delete")
            await page.keyboard.type(value, delay=40)
            await asyncio.sleep(0.2)
            log(f"  ✓ {label}: OK")
        except Exception as e:
            log(f"  !! {label}: {e}")

    await _type("first", params.get("first_name", ""), "First name")
    # set each field's match options while THAT field is focused and its options
    # popup is open — BEFORE moving on. Name: 4 checkboxes; surname: «строго по
    # имени»; year: exact / ±N; place: «должно соответствовать».
    await _set_name_match(root, page, params, log)
    await _type("last",  params.get("surname", ""),    "Surname")
    # Mirror the GUI state to the site (check OR uncheck), so it's visible there.
    # exclude_name=True → don't touch the name popup's «строго» (shared text).
    if params.get("surname"):
        await _set_field_checkbox(root, page, "last", "строго",
                                  bool(params.get("surname_strict")),
                                  "Совпадение фамилии (строго)", log,
                                  exclude_name=True)
    await _type("year",  params.get("birth_year", ""), "Birth year")
    await _set_year_match(root, page, params, log)
    await _type("place", params.get("birth_place", ""), "Birth place")
    if params.get("birth_place"):
        await _set_field_checkbox(root, page, "place", "должно соответств",
                                  bool(params.get("place_match")),
                                  "Совпадение места", log)
    return bool(tagged and (tagged.get("first") or tagged.get("last")))


async def _fill_root(root, page, sels, value, label, log):
    """Fill a field inside `root` (page or frame); type via page.keyboard so
    it works whether the field is in the main frame or an iframe."""
    if not value:
        return True
    for sel in sels:
        try:
            el = root.locator(sel).first
            if not await el.count():
                continue
            await el.scroll_into_view_if_needed(timeout=4000)
            await el.click(timeout=3000)
            await page.keyboard.press("Control+a")
            await page.keyboard.press("Delete")
            await page.keyboard.type(value, delay=40)
            await asyncio.sleep(0.2)
            if (await el.input_value(timeout=2000)).strip():
                log(f"  ✓ {label}: OK")
                return True
        except Exception:
            continue
    log(f"  !! Field not found: {label}")
    return False


# ── ADVANCED SEARCH (pills → popup with data-automations → «Применить») ─────── #
async def _fill_advanced(root, page, params, log):
    """
    Each advanced detail is a tag/pill (div.tag_label «Отец»/«Мать»/«Супруг(-а)»/
    «Смерть»). Clicking it opens a popup whose inputs have exact data-automations
    ids; after filling we MUST click the «Применить» (apply) button.
    "+ Больше" opens extra event/relative/other tags + «Точное совпадение всех
    параметров».

    On the EN/HE sites these pills carry English/Hebrew text — the site language
    comes from params["lang"] and drives _ui_labels(concept, lang).
    """
    lang = params.get("lang")
    async def _open_pill(texts, verify_auto_id):
        """REAL mouse-click the visible chip (JS .click() does NOT open the
        React popup), then confirm the popup opened by waiting for its input
        (verify_auto_id) to become visible. Returns True only if it opened."""
        for fr in page.frames:
            for t in texts:
                # tag via JS so we can click the right element via Playwright
                try:
                    tagged = await fr.evaluate(r"""(label) => {
                        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
                        document.querySelectorAll('[data-pw-pill]').forEach(
                            e => e.removeAttribute('data-pw-pill'));
                        const els = Array.from(document.querySelectorAll(
                            'div, span, button, [role=button], [class*="tag" i], [class*="chip" i]'));
                        for (const e of els) {
                            const txt = norm(e.textContent);
                            if (!txt || txt.length > 40 || e.offsetParent === null) continue;
                            if (txt === label || txt.startsWith(label + ':')
                                || txt.startsWith(label + ' ')) {
                                e.setAttribute('data-pw-pill', '1');
                                return true;
                            }
                        }
                        return false;
                    }""", t)
                    if not tagged:
                        continue
                    chip = fr.locator('[data-pw-pill="1"]').first
                    # real mouse click (bubbles → React handler fires)
                    await chip.scroll_into_view_if_needed(timeout=3000)
                    await chip.click(timeout=4000)
                    if not verify_auto_id:
                        await asyncio.sleep(0.8)
                        return True
                    # confirm the popup's input is now visible
                    for _ in range(8):
                        await asyncio.sleep(0.4)
                        v_fr, v_el = await _find_auto(verify_auto_id, want_visible=True,
                                                      wait_secs=1)
                        if v_el:
                            return True
                    # popup didn't open — try clicking the closest clickable parent
                    try:
                        await chip.evaluate(
                            "e => { const b = e.closest('button,[role=button]'); "
                            "if (b && b!==e) b.click(); }")
                        for _ in range(6):
                            await asyncio.sleep(0.4)
                            v_fr, v_el = await _find_auto(verify_auto_id,
                                                          want_visible=True, wait_secs=1)
                            if v_el:
                                return True
                    except Exception:
                        pass
                except Exception:
                    continue
        return False

    async def _find_auto(auto_id, want_visible=True, wait_secs=5):
        """Return (frame, locator) for the VISIBLE data-automations element.
        MyHeritage keeps hidden template duplicates, so we must skip those and
        wait for the popup's real (visible) one to appear."""
        for _ in range(wait_secs * 2):
            for fr in page.frames:
                try:
                    loc = fr.locator(f'[data-automations="{auto_id}"]')
                    n = await loc.count()
                    for i in range(n):
                        el = loc.nth(i)
                        if not want_visible or await el.is_visible():
                            return fr, el
                except Exception:
                    continue
            if not want_visible:
                break
            await asyncio.sleep(0.5)
        return None, None

    async def _fill_auto(auto_id, value):
        if not value:
            return
        fr, el = await _find_auto(auto_id, want_visible=True)
        if not el:
            log(f"  !! {auto_id} не виден")
            return
        # Try a normal click+type first; fall back to a JS value-set with the
        # React native setter + input/change events (handles hidden/animating).
        try:
            await el.click(timeout=3000)
            await page.keyboard.press("Control+a")
            await page.keyboard.press("Delete")
            await page.keyboard.type(str(value), delay=40)
            await asyncio.sleep(0.2)
            if (await el.input_value()).strip():
                log(f"  ✓ {auto_id} = {value!r}")
                return
        except Exception:
            pass
        try:
            await fr.evaluate(r"""([id, val]) => {
                const el = Array.from(document.querySelectorAll('[data-automations="'+id+'"]'))
                    .find(e => e.offsetParent !== null) ||
                    document.querySelector('[data-automations="'+id+'"]');
                if (!el) return;
                const setter = Object.getOwnPropertyDescriptor(
                    window.HTMLInputElement.prototype, 'value').set;
                el.focus(); setter.call(el, val);
                el.dispatchEvent(new Event('input',  {bubbles: true}));
                el.dispatchEvent(new Event('change', {bubbles: true}));
            }""", [auto_id, str(value)])
            log(f"  ✓ {auto_id} = {value!r} (JS)")
        except Exception as e:
            log(f"  !! {auto_id}: {e}")

    async def _apply(auto_id):
        """Click the popup's «Применить» button — exact id first, else generic."""
        fr, el = await _find_auto(auto_id, want_visible=True, wait_secs=2)
        if el:
            try:
                await el.click(timeout=4000)
                await asyncio.sleep(0.6)
                return True
            except Exception:
                pass
        for fr in page.frames:
            for sel in ('button.search_form_apply_button',
                        'button:has-text("Применить")', 'button:has-text("Apply")',
                        'button:has-text("החל")'):
                try:
                    b = fr.locator(sel).first
                    if await b.count() and await b.is_visible():
                        await b.click(timeout=4000)
                        await asyncio.sleep(0.6)
                        return True
                except Exception:
                    continue
        return False

    # ── Family members: pill → two name inputs → apply ──────────────────────
    fam = [
        ("father", _ui_labels("father", lang), "father_first_name_field",
         "father_last_name_field", "father_filter_apply_button"),
        ("mother", _ui_labels("mother", lang), "mother_first_name_field",
         "mother_last_name_field", "mother_filter_apply_button"),
        ("spouse", _ui_labels("spouse", lang), "spouse_first_name_field",
         "spouse_last_name_field", "spouse_filter_apply_button"),
    ]
    for key, tags, fa, la, apply_id in fam:
        fv = params.get(key, "")
        lv = params.get(f"{key}_last", "")
        if not fv and not lv:
            continue
        if await _open_pill(tags, fa):
            log(f"  → открыл «{tags[0]}»")
            await _fill_auto(fa, fv)
            await _fill_auto(la, lv)
            await _apply(apply_id)
        else:
            log(f"  !! «{tags[0]}»: попап не открылся")

    # ── Death: pill → year + place → apply ──────────────────────────────────
    if params.get("death_year") or params.get("death_place"):
        if await _open_pill(_ui_labels("death", lang), "death_year_field"):
            log("  → открыл «Смерть»")
            await _fill_auto("death_year_field", params.get("death_year", ""))
            await _fill_auto("death_place_field", params.get("death_place", ""))
            await _apply("death_filter_apply_button")
        else:
            log("  !! «Смерть»: попап не открылся")

    # ── "+ Больше": extra events / keywords / gender / exact-all ─────────────
    need_more = (any(params.get(f) for f in
                     ("residence", "military", "immigration", "keywords"))
                 or params.get("gender", "Any") not in ("Any", "Любой")
                 or params.get("exact_match"))
    if need_more:
        # Open the "+ Больше" panel and VERIFY it opened by waiting for one of
        # its hallmark texts to appear (the exact-match checkbox / event tags).
        async def _more_open():
            # Hallmark texts of the open "+ More" panel in RU / EN / HE.
            hallmark = re.compile(
                r"Точное совпадение всех параметров|событие из жизни|"
                r"Семейное положение|"
                r"Exact match for all|life event|Marital status|"
                r"התאמה מדויקת|אירוע|מצב משפחתי|"
                r"Correspondance exacte|État civil|"          # fr
                r"Genaue Übereinstimmung|Familienstand|"      # de
                r"Coincidencia exacta|Estado civil|"          # es
                r"Correspondência exata", re.I)               # pt
            for fr in page.frames:
                try:
                    el = fr.get_by_text(hallmark).first
                    if await el.count() and await el.is_visible():
                        return True
                except Exception:
                    continue
            return False

        # "+ Больше" is a TOGGLE — clicking it twice closes it again. So check
        # whether it's already open BEFORE each click, and click only if not.
        opened_more = False
        for _attempt in range(4):
            if await _more_open():
                opened_more = True
                break
            await _open_pill(_ui_labels("more", lang), "")
            await asyncio.sleep(1.0)
        log("  → панель «+ Больше» открыта" if opened_more
            else "  !! панель «+ Больше» не открылась")

        for key, tags, place_auto, apply_id in [
            ("residence", _ui_labels("residence", lang),
             "residence_place_field", "residence_filter_apply_button"),
            ("military", _ui_labels("military", lang),
             "military_place_field", "military_filter_apply_button"),
            ("immigration", _ui_labels("immigration", lang),
             "immigration_place_field", "immigration_filter_apply_button"),
        ]:
            if params.get(key):
                if await _open_pill(tags, place_auto):
                    await _fill_auto(place_auto, params[key])
                    await _apply(apply_id)

        if params.get("keywords"):
            if await _open_pill(_ui_labels("keywords", lang), "keywords_field"):
                await _fill_auto("keywords_field", params["keywords"])
                await _apply("keywords_filter_apply_button")

        g = params.get("gender", "Any")
        if g not in ("Any", "Любой"):
            if await _open_pill(_ui_labels("gender", lang), ""):
                # Click the gender value in the site language (try every variant;
                # only the one rendered on the loaded site exists).
                concept = "female" if g in ("Female", "Женщина") else "male"
                for val in _ui_labels(concept, lang):
                    try:
                        opt = root.get_by_text(
                            re.compile(rf"^{re.escape(val)}$", re.I)).first
                        if await opt.count():
                            await opt.click(timeout=3000)
                            break
                    except Exception:
                        continue
                await _apply("gender_filter_apply_button")

        # «Точное совпадение всех параметров». aria-checked never flips by
        # plain clicks → the click is intercepted by the visual checkbox box
        # and/or aria-checked is decorative. So: tag the label; FORCE-click
        # several candidate targets (the visible checkbox box, the parent
        # control, the label) at real coordinates; verify "checked" via EITHER
        # aria-checked OR a nearby input[type=checkbox].checked, re-read fresh.
        if params.get("exact_match"):
            async def _exact_state():
                """Tag the control; return (frame, checked_bool) or (None,None)."""
                for fr in page.frames:
                    try:
                        st = await fr.evaluate(r"""() => {
                            const norm = s => (s || '').replace(/\s+/g, ' ').trim();
                            document.querySelectorAll('[data-pw-ex],[data-pw-exbox]')
                                .forEach(e => { e.removeAttribute('data-pw-ex');
                                                e.removeAttribute('data-pw-exbox'); });
                            for (const e of document.querySelectorAll('span,label,[role=checkbox]')) {
                                const t = norm(e.textContent);
                                if (t.length > 60) continue;
                                if (!/Точное совпадение всех параметров/i.test(t)) continue;
                                if (e.offsetParent === null) continue;
                                // control = nearest clickable wrapper
                                const ctrl = e.closest('label, [role=checkbox]')
                                          || e.parentElement || e;
                                ctrl.setAttribute('data-pw-ex', '1');
                                // the visual checkbox box (first small sibling/child)
                                const box = ctrl.querySelector(
                                    '[class*="checkbox" i]:not([class*="label" i]), '
                                    + 'svg, [class*="box" i], [class*="control" i]');
                                if (box) box.setAttribute('data-pw-exbox', '1');
                                // state from aria-checked or a real input
                                const inp = ctrl.querySelector('input[type=checkbox]');
                                const aria = (e.closest('[role=checkbox]')
                                    || ctrl).getAttribute('aria-checked');
                                const checked = (inp && inp.checked) || aria === 'true';
                                return checked;
                            }
                            return 'none';
                        }""")
                        if st != 'none':
                            return fr, bool(st)
                    except Exception:
                        continue
                return None, None

            done = False
            for _attempt in range(10):
                fr, checked = await _exact_state()
                if fr is None:
                    await asyncio.sleep(0.5)
                    continue
                if checked:
                    log("  ✓ Точное совпадение всех параметров")
                    done = True
                    break
                for sel in ('[data-pw-exbox="1"]', '[data-pw-ex="1"]'):
                    try:
                        tgt = fr.locator(sel).first
                        if not await tgt.count():
                            continue
                        await tgt.scroll_into_view_if_needed(timeout=1500)
                        await tgt.click(timeout=2500, force=True)
                        await asyncio.sleep(1.0)
                        _, now = await _exact_state()
                        if now:
                            done = True
                            break
                    except Exception:
                        continue
                if not done:
                    try:
                        await fr.locator('[data-pw-ex="1"]').first.focus()
                        await page.keyboard.press("Space")
                        await asyncio.sleep(1.0)
                        _, now = await _exact_state()
                        if now:
                            done = True
                    except Exception:
                        pass
                if done:
                    log("  ✓ Точное совпадение всех параметров")
                    break
                await asyncio.sleep(0.4)
            params["_exact_ok"] = done
            if not done:
                # MyHeritage's checkbox is decorative — its state never flips by
                # click. Not a problem: we use the normal (fuzzy) search and filter
                # results ourselves by имя/год/тип (exact mode would drop the
                # «Александр-Вольф (Shenderovich)» variants). Just a calm note.
                log("  → «Точное совпадение всех параметров» сайт не переключает — "
                    "не нужно, фильтрую сам по имени/году/типу")


async def _apply_record_type_filter(page, record_filter, log):
    """Click the RESULTS-page «Уточнить по типу записи» refine radio so MyHeritage
    filters server-side (instant) instead of us reading every record and dropping
    the wrong type. The «refine» control lives on the results page (NOT the search
    form — clicking the form's category nav-link went to /category-5000/)."""
    if record_filter not in ("Historical Records", "Family Trees"):
        return False
    wanted = {
        "Historical Records": ["Исторические записи", "Historical records",
                               "Historical Records", "Documents historiques",
                               "Historische Aufzeichnungen", "Registros históricos",
                               "Registos históricos", "רשומות היסטוריות"],
        "Family Trees":       ["Семейные деревья", "Family trees", "Family Trees",
                               "Arbres généalogiques", "Stammbäume",
                               "Árboles genealógicos", "Árvores genealógicas",
                               "עצי משפחה"],
    }[record_filter]
    # JS finder/clicker — robust to the hashed CSS-module class. Finds the refine
    # radio label by text and clicks it + its radio ancestor.
    CLICK_JS = r"""(labels) => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
        const nodes = [...document.querySelectorAll(
            '[class*="styled_radio_label_paragraph"], [class*="styled_radio_label"], '
            + '[class*="radio_label"], [class*="RadioLabel"], label, [role="radio"]')];
        for (const lab of labels) {
            let el = nodes.find(e => norm(e.textContent) === lab)
                  || nodes.find(e => norm(e.textContent).includes(lab));
            if (el) {
                const tgt = el.closest('label, [class*="radio" i], [role="radio"]') || el;
                try { tgt.click(); } catch (e) {}
                try { el.click(); } catch (e) {}
                return norm(el.textContent).slice(0, 40);
            }
        }
        return '';
    }"""
    # wait for the result list to render first
    for _ in range(20):
        try:
            if await page.evaluate(
                    "() => document.querySelectorAll('a[href*=\"showRecord\"]').length"):
                break
        except Exception:
            pass
        await asyncio.sleep(0.5)
    try:
        before = await page.evaluate(
            "() => (document.querySelector('a[href*=\"showRecord\"]')||{}).href || ''")
    except Exception:
        before = ""
    # poll for the refine radio (it renders after the cards), across frames,
    # scrolling the sidebar into view
    clicked = ""
    for _ in range(16):                              # ~16s
        for fr in page.frames:
            try:
                clicked = await fr.evaluate(CLICK_JS, wanted)
            except Exception:
                clicked = ""
            if clicked:
                break
        if clicked:
            break
        try:
            await page.evaluate("() => window.scrollBy(0, 350)")
        except Exception:
            pass
        await asyncio.sleep(1.0)
    if not clicked:
        log(f"  !! радио «{record_filter}» на странице результатов не найдено — "
            f"оставляю свой пост-фильтр")
        return False
    # wait for the results to refresh (first record link changes)
    changed = False
    for _ in range(24):                              # ~12s
        await asyncio.sleep(0.5)
        try:
            now = await page.evaluate(
                "() => (document.querySelector('a[href*=\"showRecord\"]')||{}).href || ''")
        except Exception:
            now = before
        if now != before:
            changed = True
            break
    log(f"  ✓ Сайт отфильтровал по типу записи: {record_filter} (клик: «{clicked}»)"
        + ("" if changed else " — набор не сменился"))
    return True


# English (canonical, as shown in the GUI) → the label MyHeritage shows in each site
# language. The GUI is ALWAYS English; the scraper translates here so the filter works
# on any-language site. ru solid (the user's account); he/fr/de/es/pt best-effort.
_CAT_I18N = {
    "All Collections":                   ["Все коллекции", "כל האוספים", "Toutes les collections", "Alle Sammlungen", "Todas las colecciones", "Todas as coleções"],
    "Schools & Universities":            ["Школы и университеты", "בתי ספר ואוניברסיטאות", "Écoles et universités", "Schulen & Universitäten", "Escuelas y universidades", "Escolas e universidades"],
    "Census & Voter Lists":              ["Перепись и списки избирателей", "מפקדי אוכלוסין ורשימות מצביעים", "Recensements et listes électorales", "Volkszählungen & Wählerlisten", "Censos y listas de votantes", "Censos e listas de eleitores"],
    "Directories, Guides & References":   ["Справочники, путеводители и ссылки", "מדריכים והפניות", "Annuaires, guides et références", "Verzeichnisse, Anleitungen & Referenzen", "Directorios, guías y referencias", "Diretórios, guias e referências"],
    "Histories, Memories & Biographies":  ["Истории, мемуары и биографии", "היסטוריות, זיכרונות וביוגרפיות", "Histoires, mémoires et biographies", "Geschichten, Erinnerungen & Biografien", "Historias, memorias y biografías", "Histórias, memórias e biografias"],
    "Maps":                              ["Карты", "מפות", "Cartes", "Karten", "Mapas", "Mapas"],
    "Books & Publications":              ["Книги и публикации", "ספרים ופרסומים", "Livres et publications", "Bücher & Veröffentlichungen", "Libros y publicaciones", "Livros e publicações"],
    "Birth, Marriage & Death":           ["Реестры рождения, браков и смерти", "לידה, נישואין ופטירה", "Naissance, mariage et décès", "Geburt, Heirat & Tod", "Nacimiento, matrimonio y defunción", "Nascimento, casamento e óbito"],
    "Immigration & Travel":              ["Иммиграция и путешествия", "הגירה ונסיעות", "Immigration et voyages", "Einwanderung & Reisen", "Inmigración y viajes", "Imigração e viagens"],
    "Public Records":                    ["Публичные отчёты", "רשומות ציבוריות", "Archives publiques", "Öffentliche Aufzeichnungen", "Registros públicos", "Registros públicos"],
    "Military":                          ["Вооруженные силы", "צבא", "Militaire", "Militär", "Militar", "Militar"],
    "Photos":                            ["Фото", "תמונות", "Photos", "Fotos", "Fotos", "Fotos"],
    "Government, Land, Court & Wills":    ["Правительство, земля, суды и завещания", "ממשל, קרקעות, בתי משפט וצוואות", "Gouvernement, terres, tribunaux et testaments", "Regierung, Land, Gericht & Testamente", "Gobierno, tierras, tribunales y testamentos", "Governo, terras, tribunais e testamentos"],
    "Family Trees":                      ["Семейные деревья", "אילנות יוחסין", "Arbres généalogiques", "Stammbäume", "Árboles genealógicos", "Árvores genealógicas"],
    "Newspapers":                        ["Газеты", "עיתונים", "Journaux", "Zeitungen", "Periódicos", "Jornais"],
}


async def _expand_category_facet(page, log):
    """Click any «Show more / Показать ещё» link inside the «Narrow down by category»
    facet so every category is visible before we pick one. MH renders these long lists
    lazily and needs ~2 seconds to draw them fully, so we pause after each expand."""
    MORE_JS = r"""(words) => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim().toLowerCase();
        let n = 0;
        for (const e of document.querySelectorAll('a, span, div, button')) {
            if (e.children.length) continue;            // leaf only
            if (words.includes(norm(e.textContent))) {
                try { e.click(); n++; } catch (_) {}
            }
        }
        return n;
    }"""
    words = ["show more", "see more", "show all", "показать ещё", "показать еще",
             "показать больше", "показать все", "показать всё", "voir plus",
             "mehr anzeigen", "ver más", "ver mas", "mostrar más"]
    for _ in range(5):
        try:
            n = await page.evaluate(MORE_JS, words)
        except Exception:
            n = 0
        if not n:
            break
        log(f"  → «Показать ещё» в категориях: +{n}")
        await asyncio.sleep(2)                           # long list draws ~2s


async def _apply_category_filter(page, label, log):
    """Click the results-page «Narrow down by category» row whose .name matches the
    English `label` OR any of its language variants (the site shows the label in its own
    language) so MyHeritage narrows server-side. Best-effort — the facet selectors are
    inferred from the live HTML; validate/extend the tree with myheritage_filter_crawler.py."""
    labels = [label] + _CAT_I18N.get(label, [])
    CLICK_JS = r"""(labels) => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
        const want = labels.map(norm);
        const sel = '[class*="narrow_down_link"], [data-automations="action_text"]';
        for (const n of document.querySelectorAll(sel)) {
            const nm = n.querySelector('.name');
            if (want.includes(norm(nm ? nm.textContent : n.textContent))) {
                try { n.click(); } catch (e) {}
                try { (nm || n).click(); } catch (e) {}
                return true;
            }
        }
        return false;
    }"""
    for _ in range(20):                              # wait for the result cards
        try:
            if await page.evaluate(
                    "() => document.querySelectorAll('a[href*=\"showRecord\"]').length"):
                break
        except Exception:
            pass
        await asyncio.sleep(0.5)
    await asyncio.sleep(2)                           # the long category list draws ~2s
    await _expand_category_facet(page, log)          # reveal «Show more» categories
    try:
        before = await page.evaluate(
            "() => (document.querySelector('a[href*=\"showRecord\"]')||{}).href || ''")
    except Exception:
        before = ""
    clicked = False
    for _ in range(16):
        for fr in page.frames:
            try:
                clicked = await fr.evaluate(CLICK_JS, labels)
            except Exception:
                clicked = False
            if clicked:
                break
        if clicked:
            break
        try: await page.evaluate("() => window.scrollBy(0, 350)")
        except Exception: pass
        await asyncio.sleep(1.0)
    if not clicked:
        log(f"  !! категория «{label}» на странице результатов не найдена")
        return False
    for _ in range(24):                             # wait for the narrowed refresh
        await asyncio.sleep(0.5)
        try:
            now = await page.evaluate(
                "() => (document.querySelector('a[href*=\"showRecord\"]')||{}).href || ''")
        except Exception:
            now = before
        if now != before:
            break
    log(f"  ✓ Применена категория: {label}")
    return True


async def _search(page, search_url, params, has_cookies, log):
    log(f"  → Navigating to search: {search_url}")
    try:
        await page.goto(search_url, wait_until="domcontentloaded", timeout=35000)
    except Exception as exc:
        log(f"  !! Cannot open search page: {exc}")
        return False
    # ALWAYS try (not just .co.il): the consent banner can reappear on any locale
    # and overlays the form, blocking the field clicks.
    await _accept_cookies(page, log)
    await _close_social_tabs(page.context, page, log)

    # Wait for the research form (searches the main frame AND any iframe)
    log("  → Жду форму поиска…")
    root = await _find_form_root(page, FIRST_NAME_SELS, log, secs=25)
    if root is None:
        log("  !! Форма поиска не появилась — пробую главный фрейм")
        root = page.main_frame
    # The banner sometimes loads only after the form — sweep once more so it
    # doesn't sit on top of the inputs while we fill them.
    await _accept_cookies(page, log)

    # Fill basic fields (JS-tag the form's inputs, then type) — robust to the
    # varying selectors; works in main frame or iframe.
    await _fill_research_basic(root, page, params, log)

    # ── Advanced search: pills → popup (data-automations) → «Применить» ──────
    try:
        await _fill_advanced(root, page, params, log)
    except Exception as _exc:
        log(f"  !! advanced search error (продолжаю поиск): {_exc}")

    # NB: record-type («Уточнить по типу записи»: Исторические записи / Семейные
    # деревья) is NOT clicked on the form. Clicking the label matched a CATEGORY
    # NAV LINK («Исторические записи» in the menu) → navigated to /category-5000/
    # → killed the search (0 results, name re-filled twice). Instead we search ALL
    # types and FILTER the collected records by their category afterwards (see
    # «_type_ok» in the read loop) — which is what the filter is meant to do.

    # Submit search — results usually open in a NEW TAB
    ctx = page.context
    results_page = None
    _sw = _search_words(params.get("lang"))
    submit_sels = ([f'span.button_content:has-text("{w}")' for w in _sw]
                   + [f'button:has-text("{w}")' for w in _sw]
                   + ['button[type="submit"]', 'input[type="submit"]'])
    pages_before = set(ctx.pages)
    search_url_now = page.url
    clicked = False
    for sel in submit_sels:
        try:
            el = root.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=6000)
                clicked = True
                log(f"  ✓ Нажал «Поиск» ({sel})")
                break
        except Exception:
            pass
    if not clicked:
        await page.keyboard.press("Enter")
        log("  → Поиск: Enter")

    # Results open in a NEW TAB (a MyHeritage results page). Poll up to 25s for
    # a new MH page, or for the current page to navigate to a results URL.
    results_page = None
    for _ in range(25):
        await asyncio.sleep(1)
        new_pages = [p for p in ctx.pages if p not in pages_before]
        for np in new_pages:
            try:
                u = (np.url or "").lower()
                if "myheritage" in u and "research" in u:
                    results_page = np
                    break
            except Exception:
                pass
        if results_page:
            break
        # same-tab navigation to a results page?
        if page.url != search_url_now and "research" in page.url.lower():
            results_page = page
            break
    if results_page is None:
        # fall back to the newest MH tab, else the search page
        mh = [p for p in ctx.pages if "myheritage" in (p.url or "").lower()]
        results_page = mh[-1] if mh else page

    try:
        await results_page.bring_to_front()
        await results_page.wait_for_load_state("networkidle", timeout=20000)
    except Exception:
        await asyncio.sleep(3)
    log(f"  ✓ Страница результатов: {results_page.url[:90]}")

    # Use the form's NATURAL (fuzzy) search — exactSearch= empty — which is what
    # the user's own search does. MyHeritage's exactSearch=1 (strict) is WRONG
    # here: for «Alexander W Sanders» it DROPS the relevant variants
    # «Alexander-Wolf Sanders (Shenderovich)» / «Александр-Вольф Сандерс
    # (Шендерович)» (hyphenated first name ≠ exact «Alexander») while letting
    # through unrelated «Alexander Sanders». The result set stays bounded by the
    # «expanded criteria» divider (everything up to «Walter Alexander …») and the
    # page/40-record caps. So we NEVER force exactSearch=1 — strip it if present.
    u = results_page.url
    if re.search(r"[?&]exactSearch=1\b", u):
        new_u = re.sub(r"([?&]exactSearch=)1\b", r"\g<1>", u)
        try:
            log("  → Нечёткий поиск формы (exactSearch пустой)…")
            await results_page.goto(new_u, wait_until="domcontentloaded", timeout=30000)
            got = 0
            for _ in range(30):
                await asyncio.sleep(1)
                try:
                    got = await results_page.evaluate(
                        "() => document.querySelectorAll("
                        "'a[href*=\"showRecord\"], a[href*=\"recordTitle\"]').length")
                except Exception:
                    got = 0
                if got > 0:
                    break
            log(f"  ✓ Результаты формы ({got} ссылок на странице): "
                f"{results_page.url[:80]}")
        except Exception as _e:
            log(f"  !! не вышло переключить на нечёткий поиск: {_e}")

    # Let the SITE filter by record type (server-side, instant). Our «_type_ok»
    # post-filter stays as a safety net if the refine control isn't found.
    try:
        await _apply_record_type_filter(
            results_page, params.get("record_filter", "All Records"), log)
    except Exception as _e:
        log(f"  !! фильтр по типу записи на сайте не применился: {_e}")

    return results_page

async def _collect_one_page(page):
    """
    Extract record cards from the CURRENT results page.
    Returns {rows, stop}: rows are the EXACT matches (those appearing BEFORE
    the «расширения критериев / expanded search criteria» divider); stop=True
    once that divider is seen so pagination halts (the records below it are
    the relaxed, irrelevant matches).
    """
    return await page.evaluate(r"""() => {
        const out = [];
        const seen = new Set();
        // Find the "expanded criteria" divider — everything after it is relaxed.
        // The phrase may be split across child spans, so search the SMALLEST
        // element whose text contains both key words (not just leaf nodes).
        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
        let divider = null, dlen = 1e9;
        for (const e of document.querySelectorAll('div, p, span, section, h1, h2, h3')) {
            const t = norm(e.textContent);
            if (t.length > 300 || t.length >= dlen) continue;
            const ru = /расширени/i.test(t) && /критери/i.test(t);
            const en = /expand/i.test(t) && /criteri/i.test(t);
            const he = /הרחב/.test(t) && /קריטריון|תנאי/.test(t);
            const fr = /élarg/i.test(t)  && /critère/i.test(t);
            const de = /erweiter/i.test(t) && /kriterien/i.test(t);
            const es = /ampli/i.test(t)  && /criterio/i.test(t);
            const pt = /expand|amplia/i.test(t) && /critério/i.test(t);
            if (ru || en || he || fr || de || es || pt) { divider = e; dlen = t.length; }
        }
        const beforeDivider = (el) => {
            if (!divider) return true;
            // el comes before divider in document order?
            return !!(divider.compareDocumentPosition(el)
                      & Node.DOCUMENT_POSITION_PRECEDING);
        };
        // Collect ALL records (incl. those after the «expanded criteria» divider —
        // name-variant matches like «Alexander-Wolf Sanders (Shenderovich)» live
        // there). The Python side keeps the relevant ones by first-name match.
        const SKIPIMG = /avatar|silhouette|placeholder|default|no_?photo|no_?image|sprite|icon|\.svg|blank|spacer|stock/i;
        const anchors = Array.from(document.querySelectorAll(
            'a[href*="showRecord"], a[href*="recordTitle"]'));
        for (const a of anchors) {
            const href = a.href || '';
            if (!href || seen.has(href)) continue;
            seen.add(href);
            const card = a.closest('[class*="result" i], li, article, div') || a;
            let name = '';
            const h = card.querySelector('h1,h2,h3,[class*="recordTitle" i],[class*="name" i]');
            if (h) name = norm(h.textContent);
            if (!name) {
                const m = href.match(/recordTitle=([^&]+)/);
                if (m) { try { name = decodeURIComponent(m[1].replace(/\+/g, ' ')); } catch(e){} }
            }
            let score = -1;
            const sm = (card.textContent || '').match(/(\d{1,3})\s*%/);
            if (sm) score = parseInt(sm[1], 10);
            // result-card portrait thumbnail — used as a photo fallback when the
            // record's own detail page yields none (family-tree people).
            let thumb = '';
            for (const im of card.querySelectorAll('img')) {
                let s = im.getAttribute('src') || im.getAttribute('data-src') || '';
                const ss = im.getAttribute('srcset') || '';
                if (ss) s = ss.split(',').pop().trim().split(/\s+/)[0];
                if (!s || SKIPIMG.test(s)) continue;
                if (s.startsWith('//')) s = 'https:' + s;
                if (/^https?:/.test(s)) { thumb = s; break; }
            }
            // BIRTH year from the card — taken right after «Рождение»/«Birth»,
            // NOT the first year in the card (which is the COLLECTION year, e.g.
            // «…США, 1936-2007» / «…1950 года» — that wrongly dropped the real
            // 1897 records). Empty year ⇒ keep (lenient).
            let year = '';
            const ctxt = card.innerText || '';
            const ym = ctxt.match(/(?:Рождение|Рожд|Birth|Born|Né[e]?|Geboren|Nacim)[\s\S]{0,25}?(1[6-9]\d\d|20[0-2]\d)/i);
            if (ym) year = ym[1];
            out.push({url: href, name_text: name.slice(0, 200), score, thumb, year});
        }
        return {rows: out, stop: false, hasDivider: !!divider};
    }""")


# Visible-element test shared by the per-page + pagination helpers. MyHeritage renders
# DUPLICATE controls (mobile + desktop) with the SAME data-automations; Playwright's
# .first kept grabbing a HIDDEN twin, so the click was a silent no-op — that is why the
# page stayed at 20 and never advanced. We always act on the VISIBLE one.
_VIS_JS = ("const V = e => { if (!e) return false; const r = e.getClientRects();"
           " if (!r.length) return false; const s = getComputedStyle(e);"
           " return s.visibility !== 'hidden' && s.display !== 'none' && s.opacity !== '0'; };")


_SEL_HDR = ('[data-automations="selector_header"], '
            '[data-automations="selector_header_container"], .selector_header')


async def _set_results_per_page(page, log, want="50"):
    """Set «Results per page» to `want` (50). The selector
    (`data-automations="selector_header"`, the span showing «20») is rendered LAZILY near
    the bottom pagination and may sit in a sub-frame — a live run showed `found:0` because
    we checked too early / only the top frame. So we SCROLL down and search EVERY frame
    before acting on the VISIBLE control (class fallback `.selector_header` too)."""
    frame = None
    for _ in range(6):
        for fr in page.frames:
            try:
                c = await fr.evaluate("(s) => document.querySelectorAll(s).length", _SEL_HDR)
            except Exception:
                c = 0
            if c:
                frame = fr
                break
        if frame:
            break
        # selector is lazy / below the fold → scroll everything to the bottom and retry
        for fr in page.frames:
            try:
                await fr.evaluate("() => window.scrollTo(0, document.body.scrollHeight)")
            except Exception:
                pass
        await asyncio.sleep(1.2)
    if frame is None:
        log("  !! селектор «результатов на странице» не найден (после скролла/фреймов)")
        return
    try:
        info = await frame.evaluate("(s) => {" + _VIS_JS + r"""
            const hs = [...document.querySelectorAll(s)];
            const h = hs.find(V) || hs[0];
            return h ? {found:true, total:hs.length, visible:hs.filter(V).length,
                       cur:(h.textContent||'').trim()}
                     : {found:false, total:hs.length};
        }""", _SEL_HDR)
        log(f"  per-page selector: {info}")
        if not info.get("found"):
            return
        if info.get("cur") == want:
            log(f"  → Результатов на странице уже {want}")
            return
        # open the dropdown on the VISIBLE header (real click, JS fallback)
        try:
            await frame.locator(
                '[data-automations="selector_header"]:visible, .selector_header:visible'
            ).first.click(timeout=4000)
        except Exception:
            await frame.evaluate("(s) => {" + _VIS_JS + """
                const h = [...document.querySelectorAll(s)].find(V);
                if (h) h.click();
            }""", _SEL_HDR)
        await asyncio.sleep(2)                          # option list draws ~2s
        # click the VISIBLE «50» option (leaf with exact text)
        res = await frame.evaluate("(w) => {" + _VIS_JS + r"""
            const cands = [...document.querySelectorAll(
                'li, span, a, div, button, [role="option"], [data-automations]')];
            const hit = cands.filter(e => e.children.length === 0
                        && (e.textContent || '').trim() === w && V(e));
            if (hit.length) { hit[hit.length-1].click(); return {clicked:true, n:hit.length}; }
            return {clicked:false};
        }""", want)
        log(f"  per-page option «{want}»: {res}")
        await asyncio.sleep(2)                          # results reload with `want`/page
        now = await frame.evaluate(
            "(s) => { const h = document.querySelector(s); return h ? (h.textContent||'').trim() : ''; }",
            _SEL_HDR)
        log(f"  → Результатов на странице теперь: {now}")
    except Exception as e:
        log(f"  → per-page: {e}")


async def _goto_next_results(page, log):
    """Advance to the next results page via the VISIBLE next_icon
    (`data-automations="next_icon"`). Real click on the visible twin + JS fallback +
    mid-wait retry; returns True only when the first record link actually changes."""
    try:
        before = await page.evaluate(
            """() => { const a = document.querySelector('a[href*=\"showRecord\"]');
                       return a ? a.href : ''; }""")
        info = await page.evaluate("() => {" + _VIS_JS + r"""
            const all = [...document.querySelectorAll('[data-automations="next_icon"]')];
            const v = all.find(V) || all[0];
            if (!v) return {found:false, total:all.length};
            const cls = ((v.className || '') + '').toLowerCase();
            const dis = /disabled/.test(cls) || v.getAttribute('aria-disabled') === 'true';
            return {found:true, total:all.length, visible:all.filter(V).length, disabled:dis};
        }""")
        log(f"  next_icon: {info}")
        if not info.get("found"):
            log("  → «Далее»: next_icon нет — последняя страница")
            return False
        if info.get("disabled"):
            log("  → «Далее»: кнопка неактивна — последняя страница")
            return False
        try:
            await page.locator(
                '[data-automations="next_icon"]:visible').first.click(timeout=5000)
        except Exception:
            await page.evaluate("() => {" + _VIS_JS + """
                const a = [...document.querySelectorAll('[data-automations="next_icon"]')].find(V)
                          || document.querySelector('[data-automations="next_icon"]');
                if (a) a.click();
            }""")
        for k in range(15):                             # wait for the AJAX page swap
            await asyncio.sleep(1)
            after = await page.evaluate(
                """() => { const a = document.querySelector('a[href*=\"showRecord\"]');
                           return a ? a.href : ''; }""")
            if after and after != before:
                log("  → Перешёл на следующую страницу")
                return True
            if k == 6:                                  # halfway: retry the click
                await page.evaluate("() => {" + _VIS_JS + """
                    const a = [...document.querySelectorAll('[data-automations="next_icon"]')].find(V);
                    if (a) a.click();
                }""")
        log("  → «Далее»: страница не сменилась")
        return False
    except Exception as e:
        log(f"  → «Далее»: {e}")
        return False


# ── COLLECT RESULTS (exact matches across pages) ──────────────────────────── #
async def _diag_results(page, log):
    """One-shot dump of what the results page actually contains, so we can fix
    a 0-results run precisely instead of guessing (the live site can't be
    inspected directly)."""
    try:
        info = await page.evaluate(r"""() => {
            const q = s => { try { return document.querySelectorAll(s).length; }
                             catch(e){ return -1; } };
            const sample = Array.from(document.querySelectorAll('a[href*="record" i]'))
                .slice(0, 6).map(a => (a.getAttribute('href') || '').slice(0, 130));
            const bt = (document.body && document.body.innerText) || '';
            return {
                url: location.href,
                title: document.title,
                showRecord:  q('a[href*="showRecord"]'),
                recordTitle: q('a[href*="recordTitle"]'),
                anyRecord:   q('a[href*="record" i]'),
                resultCards: q('[class*="result" i]'),
                hasForm:     q('input[name*="first" i], input[name*="last" i]'),
                divider: ((/расширени|expand/i.test(bt) && /критери|criteri/i.test(bt))
                          || (/הרחב/.test(bt) && /קריטריון|תנאי/.test(bt))),
                noResults: /(ничего не найдено|no results|0 результатов|не дал|nothing|לא נמצאו)/i.test(bt),
                bodyLen: bt.length,
                sample
            };
        }""")
        log("  🔎 ДИАГНОСТИКА страницы результатов:")
        log(f"     url:   {info.get('url','')[:160]}")
        log(f"     title: {info.get('title','')}")
        log(f"     showRecord={info.get('showRecord')} "
            f"recordTitle={info.get('recordTitle')} "
            f"anyRecord={info.get('anyRecord')} "
            f"resultCards={info.get('resultCards')} "
            f"formInputs={info.get('hasForm')}")
        log(f"     divider≈{info.get('divider')} "
            f"noResultsText≈{info.get('noResults')} "
            f"bodyLen={info.get('bodyLen')}")
        for s in (info.get("sample") or []):
            log(f"     ↪ {s}")
    except Exception as e:
        log(f"  🔎 диагностика не вышла: {e}")


async def _collect(page, log, max_pages=25, want_first="", want_last="", want_year=""):
    # Wait for result cards to render
    results, seen_urls = [], set()
    for _t in range(25):
        await asyncio.sleep(1)
        try:
            res = await _collect_one_page(page)
        except Exception:
            res = {"rows": [], "stop": False}
        if res.get("rows"):
            break

    # Show 50 per page to reduce the number of page turns (MH reloads ~2s)
    await _set_results_per_page(page, log, "50")
    await asyncio.sleep(2)

    page_no = 1
    while page_no <= max_pages:
        try:
            res = await _collect_one_page(page)
        except Exception:
            res = {"rows": [], "stop": False}
        rows = res.get("rows", [])
        new = new_rel = 0
        for r in rows:
            if r["url"] not in seen_urls:
                seen_urls.add(r["url"])
                results.append(r)
                new += 1
                rel = ((not want_first and not want_last)
                       or _name_relevant(r.get("name_text", ""), want_first, want_last))
                ry = r.get("year", "") or _name_year(r.get("name_text", ""))
                if rel and _year_not_earlier(want_year, ry):
                    new_rel += 1
        log(f"  → Страница {page_no}: записей {len(rows)} "
            f"(новых {new}, релевантных {new_rel}), всего {len(results)}")
        # Walk EVERY result page. Only halt when MyHeritage's «expanded criteria»
        # divider is reached (everything after it is irrelevant) or a page brings
        # no new records at all — NOT merely because a page's names stopped fuzzy-
        # matching (that wrongly cut pagination short on deeper pages).
        if res.get("stop"):
            log("  → достигнут разделитель «расширения критериев» — стоп")
            break
        if results and new == 0:
            log("  → новых записей на странице нет — последняя страница")
            break
        try:
            advanced = await _goto_next_results(page, log)
        except Exception:
            advanced = False
        if not advanced:
            break
        page_no += 1
        await asyncio.sleep(2.5)
        try:
            await page.wait_for_load_state("networkidle", timeout=15000)
        except Exception:
            await asyncio.sleep(2)
    log(f"  → Всего точных записей: {len(results)}")
    # Nothing collected → dump the page so the log tells us why (right page? any
    # record links at all? a "no results" message? still showing the form?).
    if not results:
        await _diag_results(page, log)
    return results

# ── DETAIL PAGE ───────────────────────────────────────────────────────────── #
async def _detail(page, url, has_cookies, log, card_thumb=""):
    d = {"url": url, "full_name": "", "category": "", "table_data": {},
         "profile_url": "", "source_text": "", "thumb_bytes": None,
         "household": "", "is_historical": False, "paywall": False,
         "doc_bytes": None, "doc_ext": ""}
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=35000)
        await _accept_cookies(page, log)          # always — banner can reappear
        await asyncio.sleep(1.2)
        # PAYWALL: a free account gets a LIMITED number of full record views per
        # run; once exhausted MyHeritage redirects record links to
        # «/FP/search-plans.php» (the «Купить подписку» wall). Reading that page
        # yields an empty record — detect it, retry once (sometimes transient),
        # then flag it so the caller skips it and reports a clean summary.
        for _ in range(2):
            if "search-plans" not in (page.url or "").lower():
                break
            await asyncio.sleep(1.0)
            try:
                await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                await asyncio.sleep(1.5)
            except Exception:
                break
        if "search-plans" in (page.url or "").lower():
            d["paywall"] = True
            return d
        # BOT-CHECK: too-frequent access makes MyHeritage redirect to
        # «recaptcha-challenge.php» / show «…возможно, Вы — робот». Reading that
        # yields the warning text as a fake record — detect & skip (treat like a
        # wall). We never solve captchas; the user must slow down / re-login.
        cur = (page.url or "").lower()
        is_bot = "recaptcha" in cur or "captcha-challenge" in cur
        if not is_bot:
            try:
                # match the actual warning, not a stray «recaptcha» script name
                is_bot = await page.evaluate(
                    "() => { const t = (document.body.innerText || '').slice(0, 600);"
                    " return /(возможно|полага|похоже|система).{0,30}робот|"
                    "вы\\s*[-—]\\s*робот|you are (a )?robot|"
                    "подтвердите[^]{0,30}не\\s*робот/i.test(t); }")
            except Exception:
                is_bot = False
        if is_bot:
            d["paywall"] = True            # same handling: skip + report
            d["botcheck"] = True
            return d
        # The record body (esp. family-tree «Члены семьи»: Родители / Родные
        # брат-сестра / Жена / Дети) and the document scan load LAZILY — scroll
        # through the whole page, then wait until the field-row count stops
        # growing. Reading too early was why one tree showed only «Родители».
        try:
            prev, stable = -1, 0
            for _ in range(12):                       # ~9s cap
                await page.evaluate(
                    "() => window.scrollTo(0, document.body.scrollHeight)")
                await asyncio.sleep(0.45)
                cnt = await page.evaluate(
                    "() => document.querySelectorAll('tr.recordFieldsRow').length")
                if cnt and cnt == prev:
                    stable += 1
                    if stable >= 2:                   # two equal reads → settled
                        break
                else:
                    stable = 0
                prev = cnt
            await page.evaluate("() => window.scrollTo(0, 0)")
            await asyncio.sleep(0.4)
        except Exception:
            pass
    except Exception as exc:
        log(f"    !! {exc}")
        return d

    # Extract record data by WHITELISTED genealogy labels (avoids the sidebar
    # "more records"/Geni garbage). Labels are unique enough to scan the page.
    info = await page.evaluate(r"""() => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
        const res = {name: '', category: '', fields: [], household: '', profile: '', photo: ''};

        // category line
        let catEl = Array.from(document.querySelectorAll('*')).find(e =>
            !e.children.length &&
            /^(В категории|In category)/i.test(norm(e.textContent)));
        if (catEl) res.category = norm(catEl.textContent)
            .replace(/^В категории:?\s*/i, '').replace(/^In category:?\s*/i, '');

        // record title = the h1 nearest the category line (NOT the account name)
        const h1s = Array.from(document.querySelectorAll('h1, h2'));
        if (catEl) {
            // the title h1 sits just above the category line
            let best = null, bestDist = 1e9;
            const cy = catEl.getBoundingClientRect().top;
            for (const h of h1s) {
                const t = norm(h.textContent);
                if (!t || t.length > 120) continue;
                const dy = cy - h.getBoundingClientRect().top;
                if (dy >= 0 && dy < bestDist) { bestDist = dy; best = h; }
            }
            if (best) res.name = norm(best.textContent);
        }
        if (!res.name && h1s.length) res.name = norm(h1s[0].textContent);

        // ── Capture EVERY field, using MyHeritage's REAL markup (read from the
        // live DOM, not guessed):
        //   • main fields  = «tr.recordFieldsRow» → «.recordFieldLabel» / «.recordFieldValue»
        //   • sub-sections («Перепись» …) = a nested «table.multi_table» (2 columns)
        //   • household («Домочадцы»)      = «table.groupTable» (Родство | Имя | Возраст)
        // Values are read with <script>/<style> stripped — that map-callout
        // «<script>jQuery(function(){…})» inside the value cell was the junk.
        // newline-preserving normaliser (collapses spaces but keeps line breaks)
        const normNL = s => (s || '').replace(/[ \t ]+/g, ' ')
            .replace(/ *\n+ */g, '\n').replace(/\n{2,}/g, '\n').trim();
        const valText = (cell) => {
            const c = cell.cloneNode(true);
            c.querySelectorAll('script, style').forEach(s => s.remove());
            // Family-member cells list several people, each wrapped in an <a>
            // (Родители / Родные брат-сестра / Дети). Put EACH on its own line —
            // otherwise «Имя1 годы Имя2 годы …» runs together («слипшееся»).
            const links = c.querySelectorAll('a');
            if (links.length >= 2) {
                links.forEach((a, idx) => {
                    if (idx > 0)
                        a.parentNode.insertBefore(
                            document.createTextNode('\n'), a);
                });
                return normNL(c.textContent);
            }
            return norm(c.textContent);
        };
        const seen = new Set();
        const push = (k, v) => {
            k = norm(k); v = normNL(v);
            // NB: «Родные брат/сестра» can list 12+ people (700+ chars). The old
            // 600-char cap silently DROPPED that whole field for trees with many
            // siblings (why one tree had «всё» and another lost the siblings).
            // A single record-field cell never legitimately exceeds a few KB.
            if (!k || !v || k === v || v.length > 4000) return;
            const key = k + '=' + v;
            if (!seen.has(key)) { seen.add(key); res.fields.push([k, v]); }
        };
        // main fields (skip rows whose value holds a sub-table — handled below)
        document.querySelectorAll('tr.recordFieldsRow').forEach(tr => {
            const lab = tr.querySelector('.recordFieldLabel');
            const val = tr.querySelector('.recordFieldValue');
            if (!val) return;
            if (val.querySelector('table.multi_table, table.groupTable')) return;
            push(lab ? norm(lab.textContent) : '', valText(val));
        });
        // 2-column sub-tables (Перепись: Город|Detroit  Выпуски|T628 …)
        document.querySelectorAll('table.multi_table').forEach(mt => {
            mt.querySelectorAll('tr').forEach(row => {
                const cells = Array.from(row.children);
                for (let i = 0; i + 1 < cells.length; i += 2)
                    push(cells[i].textContent, cells[i + 1].textContent);
            });
        });
        // household («Домочадцы») = groupTable rows «Родство | Имя | Возраст».
        // Keep cell POSITIONS (don't drop empty cells) so the Word table columns
        // stay aligned when a row has no «Родство»; trim trailing empties only.
        const hh = [];
        document.querySelectorAll('table.groupTable').forEach(gt => {
            gt.querySelectorAll('tr').forEach(row => {
                let c = Array.from(row.children).map(x => norm(x.textContent));
                while (c.length && !c[c.length - 1]) c.pop();   // trim trailing
                if (c.some(x => x)) hh.push(c.join('  |  '));
            });
        });
        res.household = hh.join('\n');

        // "Посмотреть полный профиль на этом сайте" — search whole page
        const prof = Array.from(document.querySelectorAll('a[href]')).find(a =>
            /полный профиль|full profile|profil complet/i.test(norm(a.textContent)));
        if (prof) res.profile = prof.href;

        // record photo — pick the LARGEST real image (portrait / document scan),
        // not the collection cover (flag/crown) in the right «искать в этой
        // коллекции» sidebar. Photos are lazy-loaded → use src/data-src/srcset.
        const sidebar = Array.from(document.querySelectorAll('*')).find(e =>
            /искать в этой коллекции|search this collection/i.test(e.innerText || '')
            && (e.innerText || '').length < 1500);
        const SKIP = /avatar|icon|sprite|placeholder|logo|geni|brand|\.svg|badge|flag|blank|spacer|loading|pixel|1x1|default|silhouette|no_?photo|no_?image|tombstone|headstone|gravestone|stock|gender|unknown/i;
        const bestSrc = (im) => {
            const ss = im.getAttribute('srcset') || im.getAttribute('data-srcset') || '';
            if (ss) {
                const parts = ss.split(',').map(s => s.trim().split(' ')[0]).filter(Boolean);
                if (parts.length) return parts[parts.length - 1];
            }
            return im.getAttribute('src') || im.getAttribute('data-src')
                || im.getAttribute('data-original') || im.getAttribute('data-lazy') || '';
        };
        const candidates = [];
        for (const im of document.querySelectorAll('img')) {
            if (sidebar && sidebar.contains(im)) continue;   // skip the collection cover
            const s = bestSrc(im);
            if (!s || !/^https?:|^\/\//.test(s)) continue;
            if (SKIP.test(s) || SKIP.test(im.getAttribute('alt') || '')) continue;
            const r = im.getBoundingClientRect();
            const w = Math.max(r.width || 0, im.naturalWidth || 0,
                               parseInt(im.getAttribute('width') || '0', 10) || 0);
            const h = Math.max(r.height || 0, im.naturalHeight || 0,
                               parseInt(im.getAttribute('height') || '0', 10) || 0);
            const area = w * h;
            if (area < 50 * 50) continue;           // skip tiny pixels/icons
            candidates.push({src: s.startsWith('//') ? 'https:' + s : s, area});
        }
        // also CSS background-image photos (some cards use them)
        for (const el of document.querySelectorAll('[style*="background-image" i], [class*="photo" i], [class*="thumbnail" i]')) {
            if (sidebar && sidebar.contains(el)) continue;   // skip the collection cover
            const bg = (el.style && el.style.backgroundImage) || '';
            const m = bg.match(/url\((['"]?)(.*?)\1\)/i);
            if (!m) continue;
            let s = m[2];
            if (!s || SKIP.test(s)) continue;
            const r = el.getBoundingClientRect();
            const area = (r.width || 0) * (r.height || 0);
            if (area < 50 * 50) continue;
            candidates.push({src: s.startsWith('//') ? 'https:' + s : s, area});
        }
        candidates.sort((a, b) => b.area - a.area);
        res.photo = candidates.length ? candidates[0].src : '';

        // source line ("Семейные деревья MyHeritage", "Geni ...") as TEXT
        let src = Array.from(document.querySelectorAll('*')).find(e =>
            !e.children.length &&
            /^(Источник|Source|В категории|In category)/i.test(norm(e.textContent)));
        res.source = src ? norm(src.textContent) : '';

        // ── «Источник» SECTION (family-tree submitter block): who submitted the
        // tree, profile/photo counts, «Обновлено …». The user explicitly wants
        // this captured. Find the «recordSectionTitle» = «Источник», climb to its
        // section wrapper, take the text, strip the heading + action links.
        let submitter = '';
        const stitle = Array.from(
            document.querySelectorAll('.recordSectionTitle, [class*="SectionTitle" i]'))
            .find(e => /^(Источник|Source)/i.test(norm(e.textContent)));
        if (stitle) {
            let box = stitle;
            for (let i = 0; i < 4; i++) {
                if (box.parentElement &&
                    norm(box.parentElement.textContent).length < 450)
                    box = box.parentElement; else break;
            }
            // textContent glued phrases («Michael NeymanОбновлено…») because the
            // элементы стоят без пробела. Clone and put a line break before each
            // block-ish child so every phrase is on its own line.
            const c = box.cloneNode(true);
            c.querySelectorAll('script, style').forEach(s => s.remove());
            c.querySelectorAll('div, p, li, a, br, h1, h2, h3, h4, h5, span')
                .forEach(e => e.parentNode &&
                    e.parentNode.insertBefore(document.createTextNode('\n'), e));
            submitter = normNL(c.textContent)
                .replace(/^\s*Источник[:\s]*/i, '').replace(/^\s*Source[:\s]*/i, '')
                .replace(/\n?\s*Посмотреть полный профиль[^]*?(?:сайте|site)/i, '')
                .replace(/\n?\s*Связаться с[^]*$/i, '')
                .replace(/\n?\s*Contact\b[^]*$/i, '')
                .replace(/\n{2,}/g, '\n').trim().slice(0, 400);
        }
        res.submitter = submitter;

        res.historical = /историческ|historical/i.test(res.category || '');
        return res;
    }""")

    d["full_name"]  = (info.get("name") or "").strip()
    d["category"]   = (info.get("category") or "").strip()
    d["profile_url"] = info.get("profile") or ""
    # Source = collection name from the URL; the DOM «Источник» element often was
    # just the bare heading «Источники» (junk) — drop that as a fallback.
    dom_src = (info.get("source") or "").strip()
    if re.fullmatch(r"(источник[аи]?|source[s]?)\W*", dom_src, re.I):
        dom_src = ""
    d["source_text"] = _collection_name(url) or dom_src
    d["is_historical"] = bool(info.get("historical"))
    td = {}
    for pair in info.get("fields", []):
        if isinstance(pair, list) and len(pair) == 2:
            td[str(pair[0])] = str(pair[1])
    # «Источник» section (family-tree submitter: «Michael Neyman · 2 819 профилей
    # в 4 деревьях · 806 фото · Обновлено …») — add it as a table field too.
    submitter = (info.get("submitter") or "").strip()
    if submitter and len(submitter) > 4 and not re.fullmatch(
            r"(источник[аи]?|source[s]?)\W*", submitter, re.I):
        td.setdefault("Источник (дерево)", submitter)
    d["table_data"] = td
    d["household"] = (info.get("household") or "").strip()

    # ── Photo: click the magnifier (лупа) to open the FULL-size photo ───────
    # The record card shows a small thumbnail with a zoom control
    # <div class="... main_record_image_zoom">. Clicking it opens a popup with
    # the large image. We download the FULL photo (saved to disk by the caller)
    # and embed a SMALL copy in Word.
    full_url = ""
    try:
        zoom = page.locator(
            '.main_record_image_zoom, .imageZoom, [class*="image_zoom" i], '
            '[class*="record_image" i] [class*="zoom" i]').first
        if await zoom.count():
            await zoom.scroll_into_view_if_needed(timeout=3000)
            await zoom.click(timeout=4000)
            await asyncio.sleep(1.5)
            full_url = await page.evaluate(r"""() => {
                // largest image inside the opened lightbox/modal
                const SKIP = /sprite|icon|logo|geni|brand|\.svg|avatar|badge|default|silhouette|no_?photo|no_?image|tombstone|headstone|gravestone|stock|gender|unknown/i;
                let best = '', area = 0;
                const scope = document.querySelector(
                    '[class*="modal" i], [role="dialog"], [class*="lightbox" i], '
                    + '[class*="overlay" i], [class*="popup" i]') || document;
                for (const im of scope.querySelectorAll('img')) {
                    const s = im.src || im.getAttribute('data-src') || '';
                    if (!s || SKIP.test(s)) continue;
                    const r = im.getBoundingClientRect();
                    const a = (r.width || im.naturalWidth || 0) *
                              (r.height || im.naturalHeight || 0);
                    if (a > area) { area = a; best = s; }
                }
                return best;
            }""")
            # close the popup
            try:
                closer = page.locator(
                    '[class*="modal" i] [class*="close" i], [role="dialog"] '
                    '[aria-label*="Close" i], [class*="lightbox" i] [class*="close" i]'
                ).first
                if await closer.count():
                    await closer.click(timeout=2000)
                else:
                    await page.keyboard.press("Escape")
            except Exception:
                await page.keyboard.press("Escape")
            await asyncio.sleep(0.4)
    except Exception:
        pass

    # ── Document download (historical records: census / passenger lists). The
    # download control lives in the FULLSCREEN viewer, which opens by clicking the
    # overlay button ON the scan («.fullscreen_overlay_button», revealed on hover).
    # Only THERE does «Загрузить документ» → documentViewer.downloadSource() give
    # the FULL original. We open it, click download inside expect_download, then
    # leave fullscreen. Family-tree people have a portrait (full_url) and NO
    # document viewer — skip them. (This is the gwar/FamilySearch download pattern.)
    is_doc = False
    if not full_url:
        try:
            is_doc = await page.evaluate(
                "() => !!document.querySelector('.fullscreen_overlay_button, "
                "img.document_viewer_image, .mediaItemDownload, "
                "[onclick*=\"downloadSource\"]')")
        except Exception:
            is_doc = False
    if is_doc:
        try:
            # 1) hover the scan to reveal the overlay, then click it → fullscreen
            opened = False
            try:
                img = page.locator(
                    'img.document_viewer_image, [class*="recordImage" i] img, '
                    '[class*="record_image" i]').first
                if await img.count():
                    await img.scroll_into_view_if_needed(timeout=2500)
                    await img.hover(timeout=2500)
                    await asyncio.sleep(0.4)
            except Exception:
                pass
            fsb = page.locator(
                '.fullscreen_overlay_button, [class*="fullscreen_overlay_button" i], '
                '[class*="fullscreen" i][class*="button" i]').first
            if await fsb.count():
                try:
                    await fsb.click(timeout=4000, force=True)
                    opened = True
                except Exception:
                    opened = False
            if not opened:                            # fall back to the hash route
                try:
                    await page.evaluate("() => { location.hash = 'fullscreen'; }")
                    await asyncio.sleep(1.0)
                    opened = True
                except Exception:
                    opened = False
            # 2) inside fullscreen: «Загрузить документ» → catch the download
            if opened:
                await asyncio.sleep(1.2)              # let the viewer initialise
                try:
                    # short timeout: downloadSource() usually doesn't fire a
                    # Playwright download here — fail fast and use the full-screen
                    # image (_fs_img) below, which IS the full-res scan.
                    async with page.expect_download(timeout=9000) as dl_info:
                        clicked = False
                        try:
                            btn = page.locator(
                                'span.mediaItemDownload, [onclick*="downloadSource" i], '
                                '[title*="агрузить документ" i]').first
                            if await btn.count():
                                await btn.click(timeout=5000, force=True)
                                clicked = True
                        except Exception:
                            clicked = False
                        if not clicked:
                            await page.evaluate(
                                "() => { try { documentViewer.downloadSource(); } "
                                "catch (e) {} }")
                    dl = await dl_info.value
                    name = (dl.suggested_filename or "").lower()
                    p = await dl.path()
                    if p:
                        b = Path(p).read_bytes()
                        if b and len(b) > 3000:
                            if   b[:4] == b"%PDF":                  ext = ".pdf"
                            elif b[:3] == b"\xff\xd8\xff":           ext = ".jpg"
                            elif b[:8] == b"\x89PNG\r\n\x1a\n":      ext = ".png"
                            elif b[:4] == b"RIFF" and b[8:12] == b"WEBP": ext = ".webp"
                            elif b[:4] in (b"II*\x00", b"MM\x00*"): ext = ".tif"
                            elif name.endswith((".jpg", ".jpeg")):  ext = ".jpg"
                            elif name.endswith(".png"):             ext = ".png"
                            elif name.endswith(".pdf"):             ext = ".pdf"
                            elif name.endswith((".tif", ".tiff")):  ext = ".tif"
                            else:                                   ext = ".jpg"
                            d["doc_bytes"] = b
                            d["doc_ext"]   = ext
                            log(f"    📄 документ скачан {len(b)//1024}KB ({ext})")
                            if ext != ".pdf":
                                d["thumb_bytes"] = b
                except Exception:
                    pass    # download button didn't fire — full-res scan via _fs_img
                # Belt-and-suspenders: the fullscreen viewer has the FULL image
                # loaded now — grab its URL so the fallback can fetch it if the
                # download event never fired (e.g. opened inline instead).
                if not d.get("doc_bytes"):
                    try:
                        fsimg = await page.evaluate(r"""() => {
                            let best = '', area = 0;
                            for (const im of document.querySelectorAll('img')) {
                                const a = (im.naturalWidth||0)*(im.naturalHeight||0);
                                const s = im.currentSrc||im.src||'';
                                if (a > area && a > 500*500 &&
                                    /myheritage|myheritageimages/i.test(s)) {
                                    area = a; best = s;
                                }
                            }
                            return best;
                        }""")
                        if fsimg:
                            d["_fs_img"] = fsimg
                    except Exception:
                        pass
            try:                                      # leave fullscreen
                await page.keyboard.press("Escape")
                await asyncio.sleep(0.3)
            except Exception:
                pass
        except Exception:
            pass

    # Fallbacks if no document downloaded: the viewer scan image (wait for it to
    # decode — the cached 1950 census worked, lazy 1930/Ellis didn't), then the
    # portrait, preview, or card thumbnail.
    if not d.get("doc_bytes") and not d.get("thumb_bytes"):
        scan2 = ""
        if is_doc:
            try:
                for _ in range(14):                   # ~7s for the lazy decode
                    got = await page.evaluate(r"""() => {
                        let best = '', area = 0;
                        for (const im of document.querySelectorAll(
                            'img.document_viewer_image, img[class*="document_viewer" i]')) {
                            const a = (im.naturalWidth||0) * (im.naturalHeight||0);
                            if (a > area) { area = a; best = im.currentSrc||im.src||''; }
                        }
                        return {src: best, area};
                    }""")
                    if got and got.get("area", 0) > 500 * 500:
                        scan2 = got["src"]
                        break
                    await page.evaluate("() => window.scrollBy(0, 400)")
                    await asyncio.sleep(0.5)
            except Exception:
                pass
        elif not full_url:
            # Non-document portrait (Geni / other tree). info.photo can miss it:
            # (a) it's lazy/undecoded, (b) its URL contains «geni» which the main
            # SKIP wrongly excluded (that was «нет фото» on the Geni record). Scroll
            # and re-grab the largest real image, NOT skipping «geni».
            try:
                for _ in range(8):                    # ~4s
                    got = await page.evaluate(r"""() => {
                        const SKIP = /avatar|sprite|placeholder|\/logo|brand|\.svg|badge|flag|blank|spacer|loading|pixel|1x1|default|silhouette|no_?photo|no_?image|gender|unknown|facebook|sprite/i;
                        let best = '', area = 0;
                        for (const im of document.querySelectorAll('img')) {
                            const s = im.currentSrc || im.src || im.getAttribute('data-src') || '';
                            if (!s || !/^https?:|^\/\//.test(s) || SKIP.test(s)) continue;
                            const a = (im.naturalWidth||0) * (im.naturalHeight||0);
                            if (a > area) { area = a; best = s; }
                        }
                        return {src: best.startsWith('//') ? 'https:'+best : best, area};
                    }""")
                    if got and got.get("area", 0) > 140 * 140:
                        scan2 = got["src"]
                        break
                    await page.evaluate("() => window.scrollBy(0, 400)")
                    await asyncio.sleep(0.5)
                await page.evaluate("() => window.scrollTo(0, 0)")
            except Exception:
                pass
        fs_img = d.pop("_fs_img", "")
        photo = full_url or fs_img or scan2 or info.get("photo") or card_thumb
        src_lbl = (" (из лупы)" if full_url
                   else " (скан, полноэкранный)" if fs_img
                   else " (скан)" if scan2
                   else " (превью)" if info.get("photo") else " (из карточки)")
        if photo:
            try:
                r = await page.request.get(photo, timeout=15000)
                if r.ok:
                    body = await r.body()
                    if len(body) > 1000:
                        d["thumb_bytes"] = body
                        log(f"    📷 фото {len(body)//1024}KB{src_lbl}")
                    else:
                        log("    📷 фото слишком маленькое — пропуск")
                else:
                    log(f"    !! фото HTTP {r.status}")
            except Exception as e:
                log(f"    !! фото не скачалось: {e}")
        else:
            log("    (фото на странице не найдено)")
    elif d.get("doc_ext") == ".pdf" and not d.get("thumb_bytes"):
        # PDF document downloaded — get the preview tile for the Word thumbnail
        prev = info.get("photo") or ""
        if prev:
            try:
                r = await page.request.get(prev, timeout=15000)
                if r.ok:
                    bb = await r.body()
                    if len(bb) > 1000:
                        d["thumb_bytes"] = bb
            except Exception:
                pass
    return d

# ── OUTPUT ───────────────────────────────────────────────────────────────── #
def _hyperlink(para, text, url):
    part = para.part
    rid  = part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl   = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run  = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c    = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u    = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t    = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


# The collection / source name is carried in the record URL, e.g.
#   /research/collection-10970/новая-зеландия-индекс-записей-о-рождении-1840-1901?…
# → «новая зеландия индекс записей о рождении 1840 1901». Far better than the
# DOM «Источник» element, which often just grabbed the bare heading «Источники».
def _collection_name(url: str) -> str:
    m = re.search(r"/research/collection-\d+/([^/?#]+)", url or "")
    if not m:
        return ""
    name = urllib.parse.unquote(m.group(1)).replace("-", " ").strip()
    return re.sub(r"\s{2,}", " ", name)


_YEAR_RE = re.compile(r"\b(1[5-9]\d\d|20\d\d)\b")


def _record_year(rec: dict) -> str:
    """A representative year for the document — the census/collection year if the
    source carries one, otherwise a year taken from the date fields."""
    yrs = _YEAR_RE.findall(rec.get("source_text", "") or "")
    if yrs:
        return f"{yrs[0]}–{yrs[1]}" if len(yrs) >= 2 else yrs[0]
    td = rec.get("table_data", {}) or {}
    for k, v in td.items():
        if re.search(r"рожд|смерт|брак|прожив|погреб|захорон|крещен|"
                     r"birth|death|marriage|residence|burial|baptism", k, re.I):
            m = _YEAR_RE.search(str(v))
            if m:
                return m.group(1)
    for v in td.values():
        m = _YEAR_RE.search(str(v))
        if m:
            return m.group(1)
    return ""


def _docx_add_record(doc, i, rec):
    """Render ONE record into an open Document (shared by fresh + append)."""
    # Exactly ONE blank line between cards (never two) — put it BEFORE each card
    # except the first, and add nothing trailing.
    if i > 1:
        doc.add_paragraph("")
    doc.add_heading(f"{i}. {rec.get('full_name','—')}", level=2)
    p = doc.add_paragraph()
    p.add_run("Category: ").bold = True
    p.add_run(rec.get("category", "—"))
    p2 = doc.add_paragraph()
    p2.add_run("Match: ").bold = True
    p2.add_run(f"{rec.get('score','?')}%")
    # Year of the document (census / record year) — always show it.
    py = doc.add_paragraph()
    py.add_run("Год: ").bold = True
    py.add_run(_record_year(rec) or "—")
    # Source = the collection name (e.g. "Семейные деревья MyHeritage") as TEXT
    if rec.get("source_text"):
        ps = doc.add_paragraph()
        ps.add_run("Источник: ").bold = True
        ps.add_run(rec["source_text"])
    if rec.get("url"):
        p3 = doc.add_paragraph()
        p3.add_run("Ссылка: ").bold = True
        _hyperlink(p3, "Открыть запись на MyHeritage", rec["url"])
    td = rec.get("table_data", {})
    if td:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Field"; hdr[1].text = "Value"
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for f, v in td.items():
            row = tbl.add_row().cells
            row[0].text = str(f)
            # multi-line value (family members — one person per line)
            lines = str(v).split("\n")
            row[1].text = lines[0] if lines else ""
            for extra in lines[1:]:
                if extra.strip():
                    row[1].add_paragraph(extra)
    # Household members («Домочадцы») — render as a real table (Родство | Имя |
    # Возраст …), the user wanted it tabular, not a text blob. Blank line BEFORE
    # the «Домочадцы:» label (separate it from the fields table), not after.
    hh = rec.get("household")
    if hh:
        doc.add_paragraph("")
        doc.add_paragraph().add_run("Домочадцы:").bold = True
        rows = [[c.strip() for c in ln.split("|")]
                for ln in hh.split("\n") if ln.strip()]
        ncol = max((len(r) for r in rows), default=0)
        if ncol >= 2:
            htbl = doc.add_table(rows=0, cols=ncol)
            htbl.style = "Table Grid"
            for r in rows:
                cells = htbl.add_row().cells
                for j in range(ncol):
                    cells[j].text = r[j] if j < len(r) else ""
        else:                                    # single-column fallback
            para = doc.add_paragraph()
            for k, ln in enumerate(x.strip() for x in hh.split("\n") if x.strip()):
                if k:
                    para.add_run().add_break()
                para.add_run(ln)
    # Photo thumbnail (small in Word; full-size saved to disk separately).
    # Convert WEBP→PNG so python-docx can embed it.
    tb = rec.get("thumb_bytes")
    if tb:
        png = _to_png(tb)
        if png:
            try:
                doc.add_picture(io.BytesIO(png), width=Inches(2.2))
            except Exception:
                pass
    dp = rec.get("doc_path")
    if dp and Path(dp).exists():
        p = doc.add_paragraph(); p.add_run("Файл: ").bold = True
        p.add_run(str(Path(dp).resolve()))           # exact path where it was saved
    # "View full profile on this site" link
    if rec.get("profile_url"):
        pp = doc.add_paragraph()
        pp.add_run("Полный профиль: ").bold = True
        _hyperlink(pp, "Посмотреть полный профиль на этом сайте",
                   rec["profile_url"])
    # NB: no trailing blank — the single inter-card blank is added at the TOP of
    # the next card (see the «if i > 1» guard) so there are never two in a row.


def write_docx(path, records, qlines, append=False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx not installed")
    existing = append and Path(path).exists()
    if existing:
        # Open the existing document and append a clearly-marked new batch.
        doc = Document(str(path))
        doc.add_page_break()
        sep = doc.add_heading(
            f"➕ Appended {len(records)} more record(s) — "
            f"{time.strftime('%Y-%m-%d %H:%M')}", level=1)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Search parameters:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph("")
    else:
        doc = Document()
        sec = doc.sections[0]
        sec.page_width = Mm(297); sec.page_height = Mm(210)
        sec.left_margin = sec.right_margin = Mm(18)
        sec.top_margin  = sec.bottom_margin = Mm(15)
        h = doc.add_heading("MyHeritage Search Results", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Search parameters:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph(f"Records saved: {len(records)}  (match ≥ {MIN_MATCH_PCT}%)")
        doc.add_paragraph("")
    for i, rec in enumerate(records, 1):
        _docx_add_record(doc, i, rec)
    add_page_numbers(doc)
    doc.save(path)

def write_xlsx(path, records, qlines, append=False):
    if not _OPENPYXL_OK:
        raise RuntimeError("openpyxl not installed")
    HF = PatternFill("solid", fgColor="2A4A7F")
    HN = Font(bold=True, color="FFFFFF", size=11)
    LINKF = Font(color="0563C1", underline="single")
    TS = Side(style="thin", color="B0B8C8")
    T  = Border(left=TS, right=TS, top=TS, bottom=TS)

    def _put(c, name, val):
        """Write a cell; the URL column becomes a hidden hyperlink («Открыть»)."""
        if name == "URL" and val:
            c.value = "Открыть"; c.hyperlink = val; c.font = LINKF
        else:
            c.value = val
        c.border = T
        c.alignment = Alignment(wrap_text=True, vertical="top")

    # Genealogy fields present across THESE records.
    aff = []
    for rec in records:
        for k in rec.get("table_data", {}):
            if k not in aff:
                aff.append(k)
    base_cols = ["#", "Full Name", "Category", "Match %", "Год", "Файл", "URL"]

    def _doc_path(rec):
        dp = rec.get("doc_path")
        return str(Path(dp).resolve()) if dp and Path(dp).exists() else ""

    existing = append and Path(path).exists()
    if existing:
        # Append rows to the existing sheet, continuing the # numbering and
        # adding columns for any genealogy fields not already present.
        wb = load_workbook(str(path))
        ws = wb.active
        header = [ws.cell(row=1, column=c).value
                  for c in range(1, ws.max_column + 1)]
        for name in base_cols + aff:
            if name not in header:
                header.append(name)
                c = ws.cell(row=1, column=len(header), value=name)
                c.font = HN; c.fill = HF; c.border = T
                c.alignment = Alignment(horizontal="center",
                                        vertical="center", wrap_text=True)
        col_idx = {name: i + 1 for i, name in enumerate(header)}
        last_n = 0
        for r in range(2, ws.max_row + 1):
            try:
                last_n = max(last_n, int(ws.cell(row=r, column=1).value))
            except Exception:
                pass
        start_row = ws.max_row + 1
        for off, rec in enumerate(records):
            td = rec.get("table_data", {})
            rowdata = {"#": last_n + off + 1,
                       "Full Name": rec.get("full_name", ""),
                       "Category":  rec.get("category", ""),
                       "Match %":   rec.get("score", ""),
                       "Год":       _record_year(rec),
                       "Файл":      _doc_path(rec),
                       "URL":       rec.get("url", "")}
            for f in aff:
                rowdata[f] = td.get(f, "")
            for name, val in rowdata.items():
                ci = col_idx.get(name)
                if ci:
                    _put(ws.cell(row=start_row + off, column=ci), name, val)
        cols = header
    else:
        wb = Workbook(); ws = wb.active; ws.title = "MyHeritage"
        cols = base_cols + aff
        for ci, cn in enumerate(cols, 1):
            c = ws.cell(row=1, column=ci, value=cn)
            c.font = HN; c.fill = HF; c.border = T
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for ri, rec in enumerate(records, 2):
            td   = rec.get("table_data", {})
            vals = [ri-1, rec.get("full_name",""), rec.get("category",""),
                    rec.get("score",""), _record_year(rec), _doc_path(rec),
                    rec.get("url","")] \
                + [td.get(f,"") for f in aff]
            for ci, val in enumerate(vals, 1):
                _put(ws.cell(row=ri, column=ci), cols[ci-1], val)
    # Auto-size every column.
    for ci in range(1, len(cols) + 1):
        letter = get_column_letter(ci)
        mw = max(8, *(len(str(ws.cell(row=r, column=ci).value or ""))
                      for r in range(1, ws.max_row + 1)))
        ws.column_dimensions[letter].width = min(mw + 4, 60)
    wb.save(path)

# Anti-detect init script (webdriver/plugins/languages spoof + Facebook kill) —
# the SAME fingerprint hiding that stops MyHeritage flagging «ты скрипт».
_ANTI_DETECT_JS = r"""
    Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
    Object.defineProperty(navigator, 'plugins',   {get: () => [1,2,3,4,5]});
    Object.defineProperty(navigator, 'languages', {get: () => ['en-US','en']});
    window.chrome = { runtime: {} };
    const _open = window.open;
    window.open = function(u, ...rest) {
        try { if (u && /facebook|fbcdn|accounts\.google|apple\.com/i.test(u)) return null; }
        catch (e) {}
        return _open ? _open.call(window, u, ...rest) : null;
    };
    const killFB = () => {
        try {
            document.querySelectorAll(
                'a[href*="facebook.com"], a[href*="facebook.net"], '
                + 'iframe[src*="facebook"], [class*="facebook" i], '
                + '[data-href*="facebook"], .fb-page, .fb-like, .fb_iframe_widget'
            ).forEach(e => { try { e.remove(); } catch (x) {} });
        } catch (x) {}
    };
    try { setInterval(killFB, 400); } catch (e) {}
    if (document.addEventListener) document.addEventListener('DOMContentLoaded', killFB);
"""


async def make_browser_context(pw):
    """Launch the SAME anti-detect persistent MyHeritage context the main scraper
    uses (the .mh_profile + webdriver/plugins/languages spoof + Facebook blocking),
    so other tools (the category crawler) get the identical, non-bot-flagged session
    instead of rolling their own bare context. Returns (ctx, page) with one clean page.
    The caller still does cookie-accept / login via _accept_cookies / _login."""
    for lk in ("SingletonLock", "SingletonCookie", "SingletonSocket"):
        try:
            (MH_PROFILE_DIR / lk).unlink()
        except Exception:
            pass
    MH_PROFILE_DIR.mkdir(parents=True, exist_ok=True)
    ctx = await browser_util.launch_persistent(pw, 
        str(MH_PROFILE_DIR),
        headless=False, accept_downloads=True, no_viewport=True,
        args=["--start-maximized", "--disable-blink-features=AutomationControlled",
              "--disable-infobars", "--disable-dev-shm-usage"],
    )
    await ctx.add_init_script(_ANTI_DETECT_JS)
    for _pat in ("**/*facebook.com/**", "**/*facebook.net/**",
                 "**/connect.facebook.net/**", "**/*fbcdn.net/**"):
        try:
            await ctx.route(_pat, lambda r: r.abort())
        except Exception:
            pass
    page = ctx.pages[0] if ctx.pages else await ctx.new_page()
    for extra in list(ctx.pages):
        if extra is not page:
            try:
                await extra.close()
            except Exception:
                pass
    return ctx, page


# ── MAIN ENTRY POINT ─────────────────────────────────────────────────────── #
async def run_scraper(*,
    site_preset    = "Israel (.co.il)",
    first_name     = "", surname       = "",
    name_strict    = False, name_variants = True, name_initials = True, name_startswith = False,
    surname_strict = False, # «Искать совпадения строго по имени» (фамилия)
    year_match     = True, year_exact = True,  year_1 = False, year_2 = False, year_5 = False, year_10 = False, year_20 = False, # «exact» / «1» / «2» / «5» / «10» / «20»
    place_match    = False, # «Местоположение должно соответствовать»
    birth_year     = "", birth_place   = "",
    father         = "", father_last   = "",
    mother         = "", mother_last   = "",
    spouse         = "", spouse_last   = "",
    death_year     = "",
    death_place    = "", residence     = "",
    military       = "", immigration   = "",
    # new life-event dates (GUI redesign) + relatives — accepted; used best-effort
    marriage_year  = "", marriage_place = "",
    military_year  = "", military_place = "",
    immigration_year = "", immigration_place = "",
    any_year       = "", any_place      = "",
    residence_year = "",
    child          = "", child_last     = "",
    sibling        = "", sibling_last   = "",
    dates          = None, relatives    = None,
    keywords       = "", gender        = "Any",
    exact_match    = False,
    record_filter  = "All Records",
    record_type    = "All records",
    category       = "All collections",
    categories     = None,   # «Ограничить поиск по категории» — list of ticked labels
    output_format  = "both",
    output_folder  = Path("."),
    email          = None, password    = None,
    log            = print,
    progress       = None,
    cancel_event   = None,
    ask_2fa_code   = None,   # callable() → str, fallback if IMAP auto-read fails
    imap_password  = None,   # mail password for MH account email (to auto-read 2FA code)
    family_site    = "",     # which family site to enter on select-site.php (name substring)
    ask_file_conflict = None, # callable(list[str]) → "overwrite"/"append"/"skip"
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    login_url, search_url, has_cookies = SITE_PRESETS.get(
        site_preset, SITE_PRESETS["Israel (.co.il)"]
    )

    want_docx = output_format in ("docx", "both")
    want_xlsx = output_format in ("xlsx", "both")
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)

    qname  = " ".join(p for p in (first_name, surname) if p)
    # Map the GUI record-type onto the internal record_filter values. Accept the
    # site-language label the GUI actually sends — Russian now; other languages
    # to be added tomorrow. (English aliases kept so nothing regresses.)
    _RT_MAP = {"Все записи": "All Records",
               "Исторические записи": "Historical Records",
               "Семейные деревья": "Family Trees",
               "All records": "All Records",
               "Historical records": "Historical Records",
               "Family trees": "Family Trees"}
    if record_type in _RT_MAP:
        record_filter = _RT_MAP[record_type]
    params = dict(
        first_name=first_name, surname=surname,
        name_strict=name_strict, name_variants=name_variants,
        name_initials=name_initials, name_startswith=name_startswith,
        surname_strict=surname_strict, year_match=year_match,
        place_match=place_match,
        birth_year=birth_year, birth_place=birth_place,
        father=father, father_last=father_last,
        mother=mother, mother_last=mother_last,
        spouse=spouse, spouse_last=spouse_last,
        death_year=death_year, death_place=death_place,
        residence=residence, military=military,
        immigration=immigration, keywords=keywords,
        gender=gender, exact_match=exact_match,
        record_filter=record_filter, category=category,
        # Site language (ru/en/he) — drives which language's UI text the advanced
        # search matches when locating pills/sections/buttons on the page.
        lang=_site_lang(site_preset),
    )
    qlines = [f"{k}: {v}" for k, v in params.items()
              if k != "lang" and v and v not in ("Any", False, "All Records")]
    summary = {"ok": False}

    _prog(0, "Launching browser…")

    async with async_playwright() as pw:
        # PERSISTENT profile — cookies + history live in MH_PROFILE_DIR between
        # runs. Once you've logged in (and passed the email 2FA) once, the next
        # runs reuse the session and MyHeritage's anti-bot stops flagging the
        # login as suspicious (no more "Confirm a login attempt" challenge).
        MH_PROFILE_DIR.mkdir(parents=True, exist_ok=True)
        # NB: do NOT hard-code user_agent. It used to be pinned to «Chrome/124»,
        # but Playwright's bundled Chromium has since moved on (141+). A UA that
        # claims 124 while the engine exposes 141's APIs is a textbook bot tell —
        # that mismatch is exactly why MyHeritage started flagging «ты скрипт».
        # Headed Chromium already reports a correct, matching «Chrome/<ver>» UA.
        ctx = await browser_util.launch_persistent(pw, 
            str(MH_PROFILE_DIR),
            headless=False,
            accept_downloads=True,
            no_viewport=True,
            args=["--start-maximized",
                  "--disable-blink-features=AutomationControlled",
                  "--disable-infobars",
                  "--disable-dev-shm-usage"],
        )
        # Hide automation fingerprint — same technique that lets
        # familysearch_scraper.py bypass bot detection
        await ctx.add_init_script("""
            Object.defineProperty(navigator, 'webdriver', {get: () => undefined});
            Object.defineProperty(navigator, 'plugins',   {get: () => [1,2,3,4,5]});
            Object.defineProperty(navigator, 'languages', {get: () => ['en-US','en']});
            window.chrome = { runtime: {} };
            // ── Nuke Facebook entirely so it can never open a tab ──────────
            const _open = window.open;
            window.open = function(u, ...rest) {
                try { if (u && /facebook|fbcdn|accounts\\.google|apple\\.com/i.test(u)) return null; }
                catch (e) {}
                return _open ? _open.call(window, u, ...rest) : null;
            };
            const killFB = () => {
                try {
                    document.querySelectorAll(
                        'a[href*=\"facebook.com\"], a[href*=\"facebook.net\"], '
                        + 'iframe[src*=\"facebook\"], [class*=\"facebook\" i], '
                        + '[data-href*=\"facebook\"], .fb-page, .fb-like, .fb_iframe_widget'
                    ).forEach(e => { try { e.remove(); } catch (x) {} });
                } catch (x) {}
            };
            try { setInterval(killFB, 400); } catch (e) {}
            if (document.addEventListener) {
                document.addEventListener('DOMContentLoaded', killFB);
            }
        """)
        # Abort the Facebook SDK + plugin requests at the network level so the
        # social widget never loads and never spawns popup tabs.
        for _pat in ("**/*facebook.com/**", "**/*facebook.net/**",
                     "**/connect.facebook.net/**", "**/*fbcdn.net/**"):
            try:
                await ctx.route(_pat, lambda r: r.abort())
            except Exception:
                pass

        # Auto-close any JUNK tab the instant it opens (Facebook/Google/blank)
        # so it never steals focus. MyHeritage result tabs (myheritage.com) are
        # kept.
        def _on_new_page(p):
            async def _maybe_close():
                try:
                    # Poll the tab's URL for up to 3s: close the instant it is
                    # (or becomes) Facebook/blank-junk; keep MyHeritage tabs.
                    for _ in range(10):
                        # NEVER close a tab the scraper opened on purpose (the
                        # 2FA mail-reading tab) — guarded by a context flag and a
                        # per-page mark set in _browser_read_yandex_code.
                        if (getattr(ctx, "_mh_pause_autoclose", False)
                                or getattr(p, "_mh_protected", False)):
                            return
                        u = (p.url or "").lower()
                        # Keep mail / login-passport tabs used to read the code.
                        if any(s in u for s in ("yandex", "passport", "mail.",
                                                "/mail", "outlook.", "office365")):
                            return
                        if any(s in u for s in ("facebook", "fbcdn",
                                                "accounts.google", "apple.com")):
                            await p.close()
                            return
                        if "myheritage" in u:
                            return                      # real results tab — keep
                        await asyncio.sleep(0.3)
                    # Still blank after 3s and not MyHeritage → junk, close it
                    if ("myheritage" not in (p.url or "").lower()
                            and not getattr(ctx, "_mh_pause_autoclose", False)
                            and not getattr(p, "_mh_protected", False)):
                        await p.close()
                except Exception:
                    pass
            try:
                asyncio.create_task(_maybe_close())
            except Exception:
                pass
        ctx.on("page", _on_new_page)

        # The persistent profile RESTORES whatever tabs were open last time
        # (e.g. old Facebook/blank tabs). Keep ONE page, close the rest.
        page = ctx.pages[0] if ctx.pages else await ctx.new_page()
        for extra in list(ctx.pages):
            if extra is not page:
                try:
                    await extra.close()
                except Exception:
                    pass

        try:
            if email and password:
                _prog(5, "Logging in…")
                logged_in = await _login(
                    page, login_url, has_cookies,
                    email, password, log,
                    ask_2fa_code=ask_2fa_code,
                    imap_password=imap_password,
                )
                if not logged_in:
                    summary["error"]   = "login_failed"
                    summary["message"] = "Login failed — check credentials or 2FA code."
                    return summary
            else:
                log("  No credentials — guest mode.")
                try:
                    await page.goto(login_url, wait_until="domcontentloaded", timeout=20000)
                    if has_cookies:
                        await _accept_cookies(page, log)
                except Exception:
                    pass

            if _done():
                return summary

            # If MyHeritage shows the family-site chooser, pick a site and get
            # the research URL for it (overrides the generic search_url).
            site_research_url = await _handle_select_site(page, family_site, log)
            if site_research_url:
                search_url = site_research_url

            _prog(15, "Opening search page…")
            results_page = await _search(page, search_url, params, has_cookies, log)
            if not results_page:
                summary["error"]   = "search_failed"
                summary["message"] = "Could not reach search results page."
                return summary
            # All result collection happens on the (possibly new) results tab
            page = results_page

            if _done():
                return summary

            # «Ограничить поиск по категории» — narrow server-side by each ticked
            # category before collecting (best-effort, like the record-type refine).
            _cats = [c for c in (categories or []) if c]
            if _cats:
                _prog(28, "Ограничение по категории…")
                for _c in _cats:
                    await _apply_category_filter(page, _c, log)

            _prog(30, "Collecting results…")
            raw = await _collect(page, log,
                                 want_first=params.get("first_name", ""),
                                 want_last=params.get("surname", ""),
                                 want_year=params.get("birth_year", ""))
            log(f"  Candidates: {len(raw)}")

            # Fallback: exactSearch=1 narrows on ALL fields at once, so a
            # many-field query (father + spouse + death year …) can legitimately
            # return zero exact matches. Rather than report "nothing found",
            # drop back to the plain results (still bounded by the «expanded
            # criteria» divider) so the user always gets the relevant records.
            if (params.get("exact_match") and params.get("_exact_ok")
                    and not raw and params.get("_noexact_url")):
                log("  !! Точный поиск (exactSearch=1) дал 0 записей — "
                    "слишком строго; откатываюсь к обычным результатам")
                try:
                    await page.goto(params["_noexact_url"],
                                    wait_until="domcontentloaded", timeout=30000)
                    await page.wait_for_load_state("networkidle", timeout=15000)
                    params["_exact_ok"] = False
                    # Only the first page — MyHeritage sorts by relevance, so the
                    # top results are the right person; deeper pages are noise.
                    raw = await _collect(page, log, max_pages=1)
                    log(f"  Candidates (без exactSearch, 1 стр.): {len(raw)}")
                except Exception as _e:
                    log(f"  !! откат к обычным результатам не вышел: {_e}")

            # Results collected are MyHeritage's matches (everything BEFORE the
            # «expanded criteria» divider). Keep them all: use the card's match
            # % when shown, otherwise treat as an exact MyHeritage match (100%).
            # Do NOT re-filter by name similarity — that wrongly drops correct
            # multi-word names (e.g. "Тамара Хананновна Рогинская (Рубина)").
            wf, wl = params.get("first_name", ""), params.get("surname", "")
            wy = params.get("birth_year", "")
            qualified, dropped_name = [], 0
            for r in raw:
                s = r["score"]
                if s < 0:
                    s = 100.0
                if s < MIN_MATCH_PCT:
                    continue
                # Keep only the right person: first name «Alexander» + middle «W»
                # (variants «Alexander-Wolf»/«Александр-Вольф» pass) + surname, and
                # a birth year near the searched one (drops «Alexander W. ~1848»).
                if (wf or wl) and not _name_relevant(r.get("name_text", ""), wf, wl):
                    dropped_name += 1
                    continue
                # year: card «Рождение …» OR the «(~1848)» span in the name; drop
                # only people EARLIER than the searched birth year (keep later).
                ry = r.get("year", "") or _name_year(r.get("name_text", ""))
                if wy and not _year_not_earlier(wy, ry):
                    dropped_name += 1
                    continue
                r["score"] = round(s, 1)
                qualified.append(r)
            if dropped_name:
                log(f"  → отброшено (не тот человек / год): {dropped_name}")
            # Safety: never let the name filter wipe EVERYTHING — fall back to
            # MyHeritage's own set if it did.
            if not qualified and dropped_name:
                log("  !! фильтр имени отсёк всё — оставляю результаты как есть")
                for r in raw:
                    s = r["score"] if r["score"] >= 0 else 100.0
                    if s >= MIN_MATCH_PCT:
                        r["score"] = round(s, 1); qualified.append(r)
            log(f"  Подходящих записей: {len(qualified)}")

            # Safety: if the «expanded criteria» divider wasn't found, the site
            # returns a huge ranked set — process only the top 40 (most relevant,
            # which is where the name-variant matches sit). Normally the divider
            # bounds it well below 40 and this never triggers.
            if len(qualified) > 40:
                log(f"  !! слишком много результатов ({len(qualified)}) — беру первые "
                    f"40 (самые релевантные сверху)")
                qualified = qualified[:40]

            if not qualified:
                _prog(100, "No results above threshold.")
                summary.update({"ok": True, "docx_count": 0, "xlsx_path": None,
                                "n_records": 0,
                                "message": f"No records with match ≥{MIN_MATCH_PCT}%."})
                return summary

            # Folder for the full-size photos
            images_dir = output_folder / "images" / (safe_fn(qname) or "myheritage")

            records = []
            paywalled = []                       # records hidden behind the wall
            dropped_year = dropped_type = 0
            _rf = params.get("record_filter", "All Records")
            if _rf != "All Records":
                log(f"  → Фильтр результатов по типу записи: {_rf}")
            n = len(qualified)
            for i, r in enumerate(qualified, 1):
                if _done():
                    break
                _prog(35 + int(50 * i / n), f"[{i}/{n}] Reading record…")
                dp = None
                try:
                    dp = await ctx.new_page()
                    det = await _detail(dp, r["url"], has_cookies, log,
                                        card_thumb=r.get("thumb", ""))
                    det["score"] = r["score"]
                    if not det.get("full_name"):
                        det["full_name"] = r["name_text"]
                    # Paywalled → don't pollute the document with an empty record;
                    # collect the name so we can report it at the end.
                    if det.get("paywall"):
                        paywalled.append(det["full_name"])
                        if det.get("botcheck"):
                            log("    ⚠ MyHeritage показал проверку «робот» "
                                "(слишком частые запросы) — пропуск")
                        else:
                            log("    ⚠ за пейволлом MyHeritage (бесплатный лимит "
                                "просмотров исчерпан) — пропуск")
                        continue
                    # Post-detail filters (the card stage couldn't see these):
                    #  • year now visible on the record page is EARLIER than the
                    #    searched birth year → wrong (earlier) person (1848/1864).
                    by = _detail_birth_year(det)
                    if wy and by and not _year_not_earlier(wy, by):
                        log(f"    ⤫ {det['full_name']} — год {by} раньше {wy}, "
                            f"пропуск")
                        dropped_year += 1
                        continue
                    #  • «Уточнить по типу записи» (historical / family trees)
                    if not _type_ok(det, _rf):
                        log(f"    ⤫ {det['full_name']} — не тот тип записи "
                            f"({_rf}), пропуск")
                        dropped_type += 1
                        continue
                    records.append(det)
                    log(f"    ✓ {det['full_name']} — {det['score']}%")
                except Exception as _exc:
                    log(f"    !! запись пропущена: {_exc}")
                    if "closed" in str(_exc).lower():
                        break          # context/browser gone — stop gracefully
                finally:
                    if dp is not None:
                        try:
                            await dp.close()
                        except Exception:
                            pass
                await asyncio.sleep(0.7)

            if dropped_year:
                log(f"  → отброшено по году (раньше {wy}): {dropped_year}")
            if dropped_type:
                log(f"  → отброшено по типу записи ({_rf}): {dropped_type}")
            log(f"  ✓ В документ войдёт записей: {len(records)}")
            if paywalled:
                log(f"  ⚠ {len(paywalled)} запис(ь/и) за пейволлом MyHeritage "
                    f"(нужна платная подписка / исчерпан бесплатный лимит "
                    f"просмотров): " + "; ".join(paywalled[:12]))

            # Drop MyHeritage placeholder images (gray silhouette, tombstone,
            # stock collection covers) — they repeat byte-for-byte. A hash is a
            # placeholder when it is either TINY and repeated (the silhouette) OR
            # it shows up under ≥3 DIFFERENT names (a generic image reused for
            # unrelated people). A real photo repeated for ONE person is kept.
            phash = {}
            for det in records:
                tb = det.get("thumb_bytes")
                if not tb:
                    continue
                h = hashlib.md5(tb).hexdigest()
                det["_phash"] = h
                e = phash.setdefault(h, {"n": 0, "size": len(tb), "names": set()})
                e["n"] += 1
                e["names"].add((det.get("full_name") or "").strip().lower())
            placeholders = {h for h, e in phash.items()
                            if (e["size"] < 12000 and e["n"] >= 2)   # tiny silhouette
                            or len(e["names"]) >= 3}                  # generic across people
            dropped = 0
            for det in records:
                if det.pop("_phash", None) in placeholders:
                    det["thumb_bytes"] = None
                    dropped += 1
            if dropped:
                log(f"  → отброшено картинок-заглушек (повторяются): {dropped}")
            # Save the surviving photos to disk. MyHeritage serves WEBP, which
            # Windows won't open as «.jpg» (that was «нет фотографий в директории»)
            # — convert to PNG (same as the Word copy) and save as «.png».
            saved_imgs = 0
            for i, det in enumerate(records, 1):
                # Prefer the FULL downloaded document (census/passenger scan, may
                # be a .pdf); fall back to the portrait/preview image.
                doc = det.get("doc_bytes")
                tb  = det.get("thumb_bytes")
                if doc:
                    data = doc
                    ext  = det.get("doc_ext") or ".jpg"
                elif tb:
                    data = _to_png(tb) or tb         # WEBP → PNG if PIL can; else raw
                    if data[:3] == b"\xff\xd8\xff":
                        ext = ".jpg"
                    elif data[:8] == b"\x89PNG\r\n\x1a\n":
                        ext = ".png"
                    elif data[:4] == b"RIFF" and data[8:12] == b"WEBP":
                        ext = ".webp"                # PIL couldn't convert → keep WEBP
                    elif data[:4] == b"GIF8":
                        ext = ".gif"
                    else:
                        ext = ".png"
                else:
                    continue
                try:
                    images_dir.mkdir(parents=True, exist_ok=True)
                    parts = [det.get("full_name") or f"record_{i}"]
                    cat = (det.get("category") or "").strip()
                    yr  = _record_year(det)
                    if cat:
                        parts.append(cat)
                    if yr:
                        parts.append(yr)
                    # «_{i}» keeps the name unique (3 records of one person no
                    # longer overwrite each other to a single file).
                    fn = safe_fn(" — ".join(parts)) + f"_{i}" + ext
                    (images_dir / fn).write_bytes(data)
                    det["doc_path"] = str((images_dir / fn).resolve())   # path for Word/Excel
                    saved_imgs += 1
                except Exception as _e:
                    log(f"    !! файл не сохранён на диск: {type(_e).__name__}: {_e}")
            if saved_imgs:
                log(f"  → сохранено файлов на диск: {saved_imgs} → {images_dir}")

            _prog(88, "Saving files…")
            base   = safe_fn(f"myheritage_{qname}") or "myheritage_results"
            docx_p = output_folder / f"{base}.docx"
            xlsx_p = output_folder / f"{base}.xlsx"

            # If output files already exist, ask the user what to do: overwrite,
            # append the new records, or skip saving. Default (no callback / no
            # conflict) is a plain overwrite — the previous behaviour.
            existing_names = [p.name for p, want in
                              ((docx_p, want_docx), (xlsx_p, want_xlsx))
                              if want and records and p.exists()]
            decision = "overwrite"
            if existing_names and ask_file_conflict:
                try:
                    decision = (ask_file_conflict(existing_names)
                                or "overwrite").lower()
                except Exception as _e:
                    log(f"  !! file-conflict dialog error: {_e}")
                    decision = "overwrite"
                log(f"  → Файл(ы) уже существуют {existing_names} → выбор: {decision}")
            append = (decision == "append")

            sd = sx = False
            if decision == "skip":
                log("  → Сохранение пропущено по выбору пользователя "
                    "(существующие файлы не тронуты).")
            else:
                if want_docx and records:
                    write_docx(docx_p, records, qlines, append=append)
                    sd = True
                    log(f"  → Word: {docx_p.name}"
                        f"{' (дополнен)' if append and docx_p.name in existing_names else ''}")
                if want_xlsx and records:
                    write_xlsx(xlsx_p, records, qlines, append=append)
                    sx = True
                    log(f"  → Excel: {xlsx_p.name}"
                        f"{' (дополнен)' if append and xlsx_p.name in existing_names else ''}")

            _prog(100, f"Done — {len(records)} record(s).")
            summary.update({
                "ok":            True,
                "docx_count":    1 if sd else 0,
                "xlsx_path":     str(xlsx_p) if sx else None,
                "output_folder": str(output_folder),
                "n_records":     len(records),
            })

        except Exception as exc:
            summary.update({"error": "exception",
                            "message": f"{type(exc).__name__}: {exc}"})
            log(f"  !! {exc}")
        finally:
            # Persistent context — closing it saves cookies/session for next run
            try:
                await ctx.close()
            except Exception:
                pass

    return summary
