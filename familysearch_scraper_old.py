#!/usr/bin/env python3
"""
familysearch_scraper.py — v13
================================
EXACT FLOW (as required):

1.  Open home page → fill First/Last Names + Place + Birth Year → click SEARCH.
2.  Wait 5 s for results to render.
3.  Sign in HERE via the nav "Sign In" link (BEFORE clicking HR tab).
      - Click [data-testid="no-loggedin-sign-in-button"]  (or any "SIGN IN" nav link)
      - Login page: fill #userName and #password with credentials from GUI
      - Click button#login  ("SIGN IN" submit button)
      - Wait for redirect back to familysearch.org
      *** Sign in ONCE and ONLY ONCE — never sign in again in this session ***
4.  Click Historical Records tab [data-testid="hr-tab"].
      Wait for URL to contain "tab=records" (NOT networkidle — hangs forever).
5.  IF advanced fields provided: click Advanced Search button
      [data-testid="advanced-search-form-button"], fill modal, click Search.
6.  Check/set results per page = 60.
7.  Collect tbody tr rows.  Table structure confirmed by browser inspection:
      cells[0] = "More" link  (href=/ark:...)   ← NOT the person's name!
      cells[1] = thumbnail/placeholder image
      cells[2] = <strong>Name</strong> + collection (2nd line)
      cells[3] = Events
      cells[4] = Relationships
8.  Filter rows: name similarity >= 80 %.
9.  For each qualifying record:
      a. Open record in a new tab (already logged-in context).
      b. Scrape all text data (dl/dt/dd or table rows).
      c. Get thumbnail image bytes  → embed in Word as small picture.
      d. Download full-resolution JPG via viewer toolbar:
           click image → viewer → download arrow → JPG Only
           → DOWNLOAD [data-testid="full-text-confirm-download"]
      e. Close tab.
10. Save  {FirstNames} {LastNames}.docx  — tables + small image per record.
    Save  {FirstNames} {LastNames}.xlsx  — text only.
"""

import asyncio, difflib, io, json, os, re, shutil, sys
from pathlib import Path

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

# ── Config ───────────────────────────────────────────────────────────────────── #
_HERE     = Path(__file__).resolve().parent
_CFG_PATH = _HERE / "config" / "familysearch.json"
try:
    _CFG = json.loads(_CFG_PATH.read_text(encoding="utf-8"))
except Exception:
    _CFG = {}

HOME_URL      = _CFG.get("home_url", "https://www.familysearch.org/en/global")
MIN_MATCH     = int(_CFG.get("min_match", 80))
BAD_PATHS     = _CFG.get("bad_paths", [
    "/records/images", "/search/linker", "/linker",
    "/en/tree/", "/tree/person/", "/tree/",
    "/catalog", "/wiki", "/books", "/films",
])
_dl           = _CFG.get("downloads_dir", "")
DOWNLOADS_DIR = Path(_dl) if _dl else Path.home() / "Downloads"
HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)
FS_BASE = "https://www.familysearch.org"
_IMG_SKIP = ("icon", "logo", "sprite", "avatar", "pixel", "button",
             "badge", "placeHolder", "placeholder", ".svg", "fscdn.org")


# ── Helpers ──────────────────────────────────────────────────────────────────── #

def _sim(a: str, b: str) -> float:
    a = re.sub(r"\s+", " ", a.strip().lower())
    b = re.sub(r"\s+", " ", b.strip().lower())
    if not a or not b:
        return 0.0
    wa, wb = set(a.split()), set(b.split())
    return max(len(wa & wb) / max(len(wa), 1),
               difflib.SequenceMatcher(None, a, b).ratio()) * 100


def safe_fn(s: str, n: int = 100) -> str:
    s = re.sub(r'[\\/*?:"<>|]', "_", s.strip())
    return re.sub(r"\s+", " ", s)[:n].strip() or "document"


def _is_record(href: str) -> bool:
    if not href:
        return False
    for bad in BAD_PATHS:
        if bad in href:
            return False
    return "/ark:" in href


def _abs_url(href: str) -> str:
    if not href:
        return ""
    return FS_BASE + href if href.startswith("/") else href


# ── Type into a form field ───────────────────────────────────────────────────── #

async def _type(page, sels: list, val: str, label: str, log) -> bool:
    if not val:
        return True
    for sel in sels:
        try:
            el = page.locator(sel).first
            if not await el.count():
                continue
            await el.scroll_into_view_if_needed(timeout=4000)
            await el.click(timeout=3000)
            await page.keyboard.press("Control+a")
            await page.keyboard.press("Delete")
            await asyncio.sleep(0.1)
            await page.keyboard.type(val, delay=40)
            await asyncio.sleep(0.2)
            if (await el.input_value(timeout=2000)).strip():
                log(f"  OK  {label}")
                return True
        except Exception:
            continue
    log(f"  !!  field not found: {label}")
    return False


# ── Step 1: Search ───────────────────────────────────────────────────────────── #

async def _search_from_home(page, first_names, last_names,
                             place_lived, birth_year, log) -> None:
    log(f"  Opening: {HOME_URL}")
    await page.goto(HOME_URL, wait_until="domcontentloaded", timeout=30000)
    await asyncio.sleep(3)

    for sel in ['input[placeholder*="First and Middle" i]',
                'input[id*="givenName" i]']:
        try:
            await page.locator(sel).first.wait_for(state="visible", timeout=10000)
            break
        except Exception:
            continue

    await _type(page,
        ['input[placeholder*="First and Middle" i]',
         'input[id*="givenName" i]', 'input[name*="givenName" i]'],
        first_names, "First Names", log)

    await _type(page,
        ['input[placeholder*="Last or Maiden" i]',
         'input[id*="surname" i]', 'input[name*="surname" i]'],
        last_names, "Last Names", log)

    if place_lived:
        await _type(page,
            ['input[placeholder*="City, County, State" i]'],
            place_lived, "Place Lived", log)

    if birth_year:
        await _type(page,
            ['input[placeholder*="Birth Year" i]', 'input[id*="birthYear" i]'],
            birth_year, "Birth Year", log)

    for sel in ['button:has-text("SEARCH")', 'button:has-text("Search")',
                'button[type="submit"]']:
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=6000)
                log("  SEARCH clicked")
                break
        except Exception:
            continue

    try:
        await page.wait_for_url(lambda u: "discovery/results" in u, timeout=20000)
    except Exception:
        pass
    # Wait 5 s for results to render (required!)
    await asyncio.sleep(5)
    log(f"  Results page: {page.url}")


# ── Step 2: Sign in via nav button ───────────────────────────────────────────── #

async def _sign_in_via_nav(page, username: str, password: str, log) -> bool:
    """
    Click the nav Sign In link on the results page.
    This opens ident.familysearch.org/en/identity/login/.
    Fill #userName + #password, click button#login.
    Called ONCE before clicking HR tab.
    """
    log("  Looking for Sign In nav link...")

    # Wait up to 8 s for the not-logged-in Sign In button
    sign_in_el = None
    for sel in [
        '[data-testid="no-loggedin-sign-in-button"]',
        'a[data-testid="no-loggedin-sign-in-button"]',
        # Fallback: any nav-level "SIGN IN" link
        'header a:has-text("SIGN IN")',
        'nav a:has-text("SIGN IN")',
        'a:has-text("SIGN IN")',
    ]:
        try:
            el = page.locator(sel).first
            await el.wait_for(state="visible", timeout=3000)
            sign_in_el = el
            log(f"  Found Sign In link: {sel}")
            break
        except Exception:
            continue

    if sign_in_el is None:
        log("  No Sign In nav link found — assuming already logged in")
        return True

    await sign_in_el.click(timeout=5000)
    # Navigate to login page
    try:
        await page.wait_for_url(lambda u: "login" in u, timeout=10000)
    except Exception:
        pass
    await asyncio.sleep(2)
    log(f"  On login page: {page.url}")

    return await _fill_login_form(page, username, password, log)


async def _fill_login_form(page, username: str, password: str, log) -> bool:
    """
    Fill #userName and #password, then click button#login.
    Uses page.fill() — most reliable for React controlled inputs.
    """
    log("  Waiting for #userName field...")
    try:
        await page.locator('#userName').wait_for(state="visible", timeout=15000)
    except Exception:
        log(f"  !! #userName not found. URL: {page.url}")
        return False

    await asyncio.sleep(1)

    # Fill with Playwright's fill() — correctly handles React inputs
    try:
        await page.locator('#userName').fill(username)
        log(f"  Filled #userName")
        await asyncio.sleep(0.3)
        await page.locator('#password').fill(password)
        log(f"  Filled #password")
        await asyncio.sleep(0.3)
    except Exception as exc:
        log(f"  fill() failed: {exc} — trying keyboard...")
        for sel, val, lbl in [('#userName', username, 'username'),
                               ('#password', password, 'password')]:
            try:
                await page.locator(sel).click(timeout=3000)
                await page.keyboard.press("Control+a")
                await page.keyboard.press("Delete")
                await page.keyboard.type(val, delay=60)
                await asyncio.sleep(0.3)
                log(f"  Typed {lbl}")
            except Exception as e2:
                log(f"  !! Cannot fill {lbl}: {e2}")
                return False

    # Verify fields are not empty
    u_val = await page.locator('#userName').input_value()
    p_val = await page.locator('#password').input_value()
    log(f"  Fields: user={'✓' if u_val else '✗'}, pass={'✓' if p_val else '✗'}")

    if not u_val.strip() or not p_val.strip():
        log("  !! Fields still empty — login will fail")
        return False

    # Click the Sign In submit button — ONLY button#login, nothing else
    clicked = False
    for sel in ['button#login', '#login', 'button[id="login"]']:
        try:
            btn = page.locator(sel).first
            if await btn.count() and await btn.is_visible():
                await btn.click(timeout=5000)
                log(f"  Clicked Sign In button ({sel})")
                clicked = True
                break
        except Exception:
            continue

    if not clicked:
        # Last resort: press Enter (won't click Church Account link)
        await page.keyboard.press("Enter")
        log("  Sign In via Enter key")

    # Wait for redirect away from login
    try:
        await page.wait_for_url(
            lambda u: "familysearch.org" in u and "login" not in u,
            timeout=30000)
    except Exception:
        await asyncio.sleep(6)

    cur = page.url
    ok  = "familysearch.org" in cur and "login" not in cur
    log(f"  Sign-in {'OK ✓' if ok else 'FAILED ✗'}  URL: {cur}")
    return ok


# ── Step 3: Click Historical Records tab ─────────────────────────────────────── #

async def _click_historical_tab(page, log) -> None:
    """
    Click HR tab. URL changes to include 'tab=records'.
    Wait for that URL change — NOT networkidle (hangs forever).
    """
    for sel in ['[data-testid="hr-tab"]']:
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                try:
                    await page.wait_for_url(
                        lambda u: "tab=records" in u, timeout=10000)
                except Exception:
                    await asyncio.sleep(2)
                await asyncio.sleep(1)
                log(f"  HR tab clicked — {page.url}")
                return
        except Exception:
            pass

    # Fallback: text-based
    for label in ("Historical Records", "Historical records"):
        try:
            el = page.get_by_role("tab", name=re.compile(re.escape(label), re.I)).first
            if not await el.count():
                el = page.locator(f'[role="tab"]:has-text("{label}")').first
            if await el.count():
                await el.click(timeout=5000)
                try:
                    await page.wait_for_url(lambda u: "tab=records" in u, timeout=10000)
                except Exception:
                    await asyncio.sleep(2)
                await asyncio.sleep(1)
                log(f"  HR tab clicked (text fallback) — {page.url}")
                return
        except Exception:
            pass
    log("  (HR tab not found)")


# ── Advanced Search modal ────────────────────────────────────────────────────── #

async def _advanced_search(page, adv: dict, log) -> None:
    log("  Opening Advanced Search...")
    opened = False
    for sel in [
        '[data-testid="advanced-search-form-button"]',
        'button:has-text("Advanced Search")',
        'button:has-text("ADVANCED SEARCH")',
    ]:
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                await asyncio.sleep(1.5)
                opened = True
                log(f"  Opened ({sel[:50]})")
                break
        except Exception:
            continue

    if not opened:
        log("  !! Advanced Search button not found")
        return

    event_map = {
        "birth_place":     ("BIRTH",     'input[placeholder*="City, County, State, Province"]'),
        "birth_year":      ("BIRTH",     'input[placeholder="Year"]'),
        "death_place":     ("DEATH",     'input[placeholder*="City, County, State, Province"]'),
        "death_year":      ("DEATH",     'input[placeholder="Year"]'),
        "marriage_place":  ("MARRIAGE",  'input[placeholder*="City, County, State, Province"]'),
        "marriage_year":   ("MARRIAGE",  'input[placeholder="Year"]'),
        "residence_place": ("RESIDENCE", 'input[placeholder*="City, County, State, Province"]'),
        "residence_year":  ("RESIDENCE", 'input[placeholder="Year"]'),
        "any_place":       ("ANY",       'input[placeholder*="City, County, State, Province"]'),
        "any_year":        ("ANY",       'input[placeholder="Year"]'),
    }
    opened_tabs: set = set()
    for key, (tab_label, input_sel) in event_map.items():
        val = adv.get(key, "")
        if not val:
            continue
        if tab_label not in opened_tabs:
            try:
                t = page.get_by_text(re.compile(rf"^{re.escape(tab_label)}$", re.I)).first
                if await t.count():
                    await t.click(timeout=3000)
                    await asyncio.sleep(0.7)
                    opened_tabs.add(tab_label)
            except Exception:
                pass
        await _type(page, [input_sel], val, key, log)

    fam_map = {
        "spouse": ("SPOUSE",       "Spouse's First Names",       "Spouse's Last Names"),
        "father": ("FATHER",       "Father's First Names",       "Father's Last Names"),
        "mother": ("MOTHER",       "Mother's First Names",       "Mother's Last Names"),
        "other":  ("OTHER PERSON", "Other Person's First Names", "Other Person's Last Names"),
    }
    for key, (tab_label, first_ph, last_ph) in fam_map.items():
        fv = adv.get(f"{key}_first", "")
        lv = adv.get(f"{key}_last",  "")
        if not fv and not lv:
            continue
        try:
            t = page.get_by_text(re.compile(rf"^{re.escape(tab_label)}$", re.I)).first
            if await t.count():
                await t.click(timeout=3000)
                await asyncio.sleep(0.7)
        except Exception:
            pass
        if fv:
            await _type(page, [f'input[placeholder="{first_ph}"]'], fv, f"{key}_first", log)
        if lv:
            await _type(page, [f'input[placeholder="{last_ph}"]'], lv, f"{key}_last", log)

    if adv.get("country"):
        try:
            t = page.get_by_text(re.compile(r"^LOCATION$", re.I)).first
            if await t.count():
                await t.click(timeout=3000)
                await asyncio.sleep(0.5)
        except Exception:
            pass
        await _type(page, ['input[placeholder="Country or Location"]'],
                   adv["country"], "country", log)
    if adv.get("state"):
        await _type(page, ['input[placeholder="State or Province"]'],
                   adv["state"], "state", log)
    if adv.get("keywords"):
        await _type(page, ['input[placeholder*="keyword" i]'],
                   adv["keywords"], "keywords", log)

    for sel in ['button:has-text("SEARCH")', 'button:has-text("Search")']:
        try:
            el = page.locator(sel).last
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                break
        except Exception:
            continue

    try:
        await page.wait_for_url(lambda u: "tab=records" in u, timeout=15000)
    except Exception:
        await asyncio.sleep(3)
    log("  Advanced Search submitted.")


# ── Set 60 per page ──────────────────────────────────────────────────────────── #

async def _set_60(page, log) -> None:
    for sel in ['select[aria-label*="result" i]', 'select[name*="result" i]',
                'select[id*="result" i]']:
        try:
            el = page.locator(sel).last
            if not await el.count():
                continue
            opts = await el.evaluate("e => Array.from(e.options).map(o => o.value)")
            if "60" in opts:
                await el.select_option(value="60")
                await asyncio.sleep(1)
                log("  Results per page → 60")
                return
        except Exception:
            continue


# ── Collect result rows ──────────────────────────────────────────────────────── #

async def _collect(page, qname: str, log) -> list:
    """
    Table structure (browser-confirmed):
      cells[0] = "More" link  (href=/ark:...) — NOT the person's name
      cells[1] = image
      cells[2] = <strong>Name</strong> + collection on 2nd line
      cells[3] = Events
      cells[4] = Relationships
    """
    await asyncio.sleep(2)
    results, seen = [], set()
    rows = await page.query_selector_all("tbody tr")
    log(f"  Rows: {len(rows)}")

    for row in rows:
        try:
            cells = await row.query_selector_all("td")
            if not cells:
                continue

            # Get record URL from ark link
            lnk_href = None
            for a in await row.query_selector_all("a[href]"):
                href = (await a.get_attribute("href") or "").strip()
                if _is_record(href):
                    lnk_href = href
                    break
            if not lnk_href:
                continue

            url = _abs_url(lnk_href)
            if url in seen:
                continue
            seen.add(url)

            # Name from cells[2] <strong> tag — NOT link text
            name = ""
            coll = ""
            name_idx = 2 if len(cells) > 3 else max(0, len(cells) - 3)
            if len(cells) > name_idx:
                nc = cells[name_idx]
                try:
                    bold = await nc.query_selector("strong, b")
                    if bold:
                        name = (await bold.text_content() or "").strip()
                except Exception:
                    pass
                if not name:
                    full = (await nc.text_content() or "").strip()
                    lines = [ln.strip() for ln in full.splitlines() if ln.strip()]
                    name = lines[0] if lines else ""
                    coll = lines[1] if len(lines) > 1 else ""
                else:
                    try:
                        full = (await nc.text_content() or "").strip()
                        lines = [ln.strip() for ln in full.splitlines() if ln.strip()]
                        coll = lines[1] if len(lines) > 1 else ""
                    except Exception:
                        pass

            if not name:
                continue

            evts = (await cells[name_idx + 1].text_content() or "").strip() if len(cells) > name_idx + 1 else ""
            rels = (await cells[name_idx + 2].text_content() or "").strip() if len(cells) > name_idx + 2 else ""

            score = round(_sim(qname, name), 1)
            results.append({"url": url, "name": name, "coll": coll,
                            "evts": evts, "rels": rels, "score": score})
            log(f"    {score:5.1f}%  {name}")
        except Exception:
            continue

    log(f"  Qualified candidates: {len(results)}")
    return results


# ── Fetch image bytes from a URL ─────────────────────────────────────────────── #

async def _fetch_image(ctx, src: str) -> bytes | None:
    """Download an image URL and return bytes, or None if < 5 KB."""
    if not src or not src.startswith("http"):
        return None
    ip = await ctx.new_page()
    try:
        r = await ip.goto(src, timeout=15000)
        if r and r.ok:
            body = await r.body()
            if len(body) > 5000:
                return body
    except Exception:
        pass
    finally:
        await ip.close()
    return None


# ── Find best image src on a page ────────────────────────────────────────────── #

async def _find_image_src(page) -> str:
    """Return the src of the best candidate image on the record page."""
    candidates = []
    for el in await page.query_selector_all("img[src]"):
        try:
            src = (await el.get_attribute("src") or "").strip()
            if not src.startswith("http"):
                continue
            if any(b.lower() in src.lower() for b in _IMG_SKIP):
                continue
            w = await el.evaluate("e => e.naturalWidth")
            h = await el.evaluate("e => e.naturalHeight")
            if w > 50 or h > 50:
                candidates.append((w * h, src))
        except Exception:
            continue
    if candidates:
        candidates.sort(reverse=True)
        return candidates[0][1]
    return ""


# ── Download full-res JPG from viewer ────────────────────────────────────────── #

async def _download_full_image(ctx, page, dest_dir: Path,
                                title: str, log) -> str | None:
    """
    Click the image → viewer opens → click download arrow →
    select JPG Only → click DOWNLOAD [data-testid="full-text-confirm-download"].
    """
    dest_dir.mkdir(parents=True, exist_ok=True)
    fname  = safe_fn(title) + ".jpg"
    dest   = dest_dir / fname
    before = set(DOWNLOADS_DIR.glob("*.jpg")) | set(DOWNLOADS_DIR.glob("*.jpeg"))
    pages_before = set(ctx.pages)

    # Click the image to open viewer
    thumb_clicked = False
    img_src = await _find_image_src(page)
    if img_src:
        try:
            el = page.locator(f'img[src="{img_src}"]').first
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                await asyncio.sleep(3)
                log(f"    Image clicked")
                thumb_clicked = True
        except Exception:
            pass

    if not thumb_clicked:
        # Try generic image selectors
        for sel in [
            'img[src*="dz/v1"]', 'img[src*="apiv2"]',
            'img[alt="Thumbnail"]', 'img[class*="imageThumb" i]',
            '[class*="image" i] img', 'main img',
        ]:
            try:
                el = page.locator(sel).first
                if not await el.count():
                    continue
                src = await el.get_attribute("src") or ""
                if any(b.lower() in src.lower() for b in _IMG_SKIP):
                    continue
                if await el.is_visible():
                    await el.click(timeout=5000)
                    await asyncio.sleep(3)
                    log(f"    Image clicked ({sel})")
                    thumb_clicked = True
                    break
            except Exception:
                continue

    if not thumb_clicked:
        log("    No image found on this record page")
        return None

    # Viewer may open in a new tab
    viewer = page
    await asyncio.sleep(1)
    new_pg = set(ctx.pages) - pages_before
    if new_pg:
        viewer = list(new_pg)[0]
        try:
            await viewer.wait_for_load_state("domcontentloaded", timeout=10000)
        except Exception:
            pass
        await asyncio.sleep(2)
        log("    Viewer in new tab")

    # Click download arrow button in viewer toolbar
    dl_clicked = False
    for sel in [
        'button[aria-label*="Download" i]',
        'button[title*="Download" i]',
        '[data-testid*="download" i]:not([data-testid="full-text-confirm-download"])',
        'button[class*="download" i]',
        '[class*="toolbar"] button:nth-last-child(3)',
        '[class*="toolbar"] button:nth-last-child(2)',
        '[class*="tools"] button:nth-last-child(2)',
    ]:
        try:
            el = viewer.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=4000)
                await asyncio.sleep(1.5)
                dl_clicked = True
                log(f"    Download button clicked ({sel})")
                break
        except Exception:
            continue

    if not dl_clicked:
        log("    Download button not found — saving thumbnail preview")
        if viewer is not page:
            try: await viewer.close()
            except Exception: pass
        # Fallback: save thumbnail bytes
        body = await _fetch_image(ctx, img_src) if img_src else None
        if body:
            dest_dir.mkdir(parents=True, exist_ok=True)
            fp = dest_dir / (safe_fn(title) + "_preview.jpg")
            fp.write_bytes(body)
            log(f"    Thumbnail saved: {fp.name} ({len(body)//1024}KB)")
            return str(fp)
        return None

    # Select JPG Only in popup
    for lbl in ("JPG Only", "JPG only", "JPG"):
        try:
            el = viewer.get_by_text(lbl, exact=True).first
            if await el.count():
                await el.click(timeout=3000)
                await asyncio.sleep(0.5)
                log("    JPG Only selected")
                break
        except Exception:
            continue

    # Click DOWNLOAD — exact data-testid first
    downloaded = None
    try:
        async with viewer.expect_download(timeout=30000) as dl_info:
            for sel in [
                '[data-testid="full-text-confirm-download"]',
                'button:has-text("DOWNLOAD")',
                'button:has-text("Download")',
            ]:
                try:
                    btn = viewer.locator(sel).first
                    if await btn.count() and await btn.is_visible():
                        await btn.click(timeout=5000)
                        log(f"    DOWNLOAD clicked ({sel})")
                        break
                except Exception:
                    continue
        dl = await dl_info.value
        await dl.save_as(str(dest))
        downloaded = str(dest)
        log(f"    Full-res saved: {fname} ({dest.stat().st_size // 1024}KB)")
    except Exception as exc:
        log(f"    expect_download failed ({exc}) — watching Downloads folder...")
        for _ in range(15):
            await asyncio.sleep(1)
            after = set(DOWNLOADS_DIR.glob("*.jpg")) | set(DOWNLOADS_DIR.glob("*.jpeg"))
            new_f = after - before
            if new_f:
                src_f = max(new_f, key=lambda p: p.stat().st_mtime)
                shutil.move(str(src_f), str(dest))
                downloaded = str(dest)
                log(f"    Moved: {fname}")
                break

    # Close stray tabs
    await asyncio.sleep(0.5)
    for pg in list(ctx.pages):
        if pg not in (page, viewer) and "familysearch" not in pg.url:
            try: await pg.close()
            except Exception: pass
    if viewer is not page:
        try: await viewer.close()
        except Exception: pass

    if downloaded:
        return downloaded

    # Absolute fallback: save thumbnail
    if img_src:
        body = await _fetch_image(ctx, img_src)
        if body:
            dest_dir.mkdir(parents=True, exist_ok=True)
            fp = dest_dir / (safe_fn(title) + "_preview.jpg")
            fp.write_bytes(body)
            log(f"    Fallback preview: {fp.name}")
            return str(fp)
    return None


# ── Scrape one record page ───────────────────────────────────────────────────── #

async def _scrape_record(ctx, url: str, name_hint: str,
                          images_root: Path, log) -> dict:
    for bad in BAD_PATHS:
        if bad in url:
            return _empty(url, name_hint)

    rec  = _empty(url, name_hint)
    page = await ctx.new_page()
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=30000)

        # Should not redirect to login — we're already signed in
        if "login" in page.url:
            log(f"    !! Unexpected login redirect — session may have expired")
            return rec

        await asyncio.sleep(3)  # wait for lazy-loaded content

        # Page title
        for sel in ["h1", '[class*="title" i]', "h2"]:
            try:
                t = (await page.locator(sel).first
                     .text_content(timeout=3000) or "").strip()
                if 3 < len(t) < 300:
                    rec["title"] = t
                    break
            except Exception:
                pass

        # Structured text: dl/dt/dd
        td: dict = {}
        try:
            dts = await page.query_selector_all("dl dt")
            dds = await page.query_selector_all("dl dd")
            for dt, dd in zip(dts, dds):
                k = (await dt.text_content() or "").strip().rstrip(":")
                v = (await dd.text_content() or "").strip()
                if k and v:
                    td[k] = v
        except Exception:
            pass

        # Fallback: table rows
        if not td:
            try:
                for row in await page.query_selector_all("table tr"):
                    cells = await row.query_selector_all("td, th")
                    if len(cells) >= 2:
                        k = (await cells[0].text_content() or "").strip().rstrip(":")
                        v = (await cells[1].text_content() or "").strip()
                        if k and v:
                            td[k] = v
            except Exception:
                pass

        rec["table_data"] = td

        # Image label for folder/file naming
        img_label = name_hint
        if td:
            parts = [name_hint]
            for k in ("Event Type", "Type", "Event"):
                if td.get(k):
                    parts.append(td[k]); break
            for k in ("Event Date", "Date", "Death Date", "Birth Date",
                      "Marriage Date", "Naturalization Date"):
                if td.get(k):
                    parts.append(td[k]); break
            if len(parts) > 1:
                img_label = " — ".join(parts)

        img_dir = images_root / safe_fn(img_label)

        # Get thumbnail bytes for Word (small image)
        img_src = await _find_image_src(page)
        if img_src:
            rec["thumb_bytes"] = await _fetch_image(ctx, img_src)
            if rec["thumb_bytes"]:
                log(f"    Thumbnail: {len(rec['thumb_bytes'])//1024}KB")

        # Download full-res image
        img_path = await _download_full_image(ctx, page, img_dir, img_label, log)
        rec["images"] = [img_path] if img_path else []

    except Exception as exc:
        log(f"    !! Record error: {exc}")
    finally:
        await page.close()
    return rec


def _empty(url, name):
    return {"url": url, "title": name, "name": name,
            "table_data": {}, "links": [], "images": [],
            "thumb_bytes": None,
            "events": "", "relationships": "", "collection": ""}


# ── Word output ──────────────────────────────────────────────────────────────── #

def _add_hyperlink(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


def write_docx(path: Path, records: list, qlines: list) -> None:
    if not _DOCX_OK:
        raise RuntimeError("python-docx not installed")
    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Mm(297); sec.page_height = Mm(210)
    sec.left_margin = sec.right_margin = Mm(15)
    sec.top_margin  = sec.bottom_margin = Mm(15)

    h = doc.add_heading("FamilySearch Search Results", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Search parameters:")
    for ln in qlines:
        doc.add_paragraph(ln, style="List Bullet")
    doc.add_paragraph(f"Records found: {len(records)}  (match ≥ {MIN_MATCH}%)")
    doc.add_paragraph("")

    for i, rec in enumerate(records, 1):
        title = rec.get("title") or rec.get("name", "—")
        doc.add_heading(f"{i}. {title}", level=2)

        # Source hyperlink
        if rec.get("url"):
            pp = doc.add_paragraph()
            pp.add_run("Source: ").bold = True
            _add_hyperlink(pp, rec["url"], rec["url"])

        p = doc.add_paragraph()
        p.add_run("Match: ").bold = True
        p.add_run(f"{rec.get('score', '?')}%")

        # Data table — mirrors the site record layout
        td = rec.get("table_data", {})
        rows_data: list[tuple] = []
        if rec.get("collection"):
            rows_data.append(("Collection", rec["collection"]))
        if rec.get("events"):
            rows_data.append(("Events", rec["events"]))
        if rec.get("relationships"):
            rows_data.append(("Relationships", rec["relationships"]))
        for field, value in td.items():
            rows_data.append((str(field), str(value)))

        if rows_data:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Field"; hdr[1].text = "Value"
            for cell in hdr:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for field, value in rows_data:
                row = tbl.add_row().cells
                row[0].text = field
                row[1].text = value

        doc.add_paragraph("")

        # Image: full-res if downloaded, else thumbnail bytes
        imgs = rec.get("images", [])
        tb   = rec.get("thumb_bytes")

        if imgs and Path(imgs[0]).exists():
            doc.add_paragraph("Document image:").runs[0].bold = True
            try:
                doc.add_picture(imgs[0], width=Inches(4))
            except Exception:
                doc.add_paragraph(f"  [{Path(imgs[0]).name}]")
        elif tb:
            doc.add_paragraph("Document image (thumbnail):").runs[0].bold = True
            try:
                doc.add_picture(io.BytesIO(tb), width=Inches(4))
            except Exception:
                doc.add_paragraph("  [image could not be embedded]")
        else:
            doc.add_paragraph("  [No image for this record]")

        doc.add_paragraph("")
    doc.save(path)


# ── Excel output ─────────────────────────────────────────────────────────────── #

def write_xlsx(path: Path, records: list, qlines: list) -> None:
    if not _OPENPYXL_OK:
        raise RuntimeError("openpyxl not installed")
    wb = Workbook(); ws = wb.active; ws.title = "FamilySearch"
    HF = PatternFill("solid", fgColor="006B6B")
    HN = Font(bold=True, color="FFFFFF", size=11)
    TS = Side(style="thin", color="B0C8C8")
    T  = Border(left=TS, right=TS, top=TS, bottom=TS)

    aff: list = []
    for rec in records:
        for k in rec.get("table_data", {}):
            if k not in aff:
                aff.append(k)

    cols = ["#", "Title", "Match %", "Collection", "Events",
            "Relationships", "Image file", "URL"] + aff

    for ci, cn in enumerate(cols, 1):
        c = ws.cell(row=1, column=ci, value=cn)
        c.font = HN; c.fill = HF; c.border = T
        c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    for ri, rec in enumerate(records, 2):
        td   = rec.get("table_data", {})
        imgs = "\n".join(Path(p).name for p in rec.get("images", []))
        vals = [
            ri - 1,
            rec.get("title", rec.get("name", "")),
            rec.get("score", ""),
            rec.get("collection", ""),
            rec.get("events", ""),
            rec.get("relationships", ""),
            imgs,
            rec.get("url", ""),
        ] + [td.get(f, "") for f in aff]
        for ci, val in enumerate(vals, 1):
            c = ws.cell(row=ri, column=ci, value=val)
            c.border = T
            c.alignment = Alignment(wrap_text=True, vertical="top")

    for ci in range(1, len(cols) + 1):
        letter = get_column_letter(ci)
        mw = max(
            len(str(cols[ci - 1])),
            *(len(str(ws.cell(row=r, column=ci).value or "").split("\n")[0])
              for r in range(2, ws.max_row + 1)),
            8,
        )
        ws.column_dimensions[letter].width = min(mw + 4, 60)
    wb.save(path)


# ── Main entry point ─────────────────────────────────────────────────────────── #

async def run_scraper(
    *,
    first_names:   str       = "",
    last_names:    str       = "",
    place_lived:   str       = "",
    birth_year:    str       = "",
    tab:           str       = "Historical Records",
    advanced:      dict|None = None,
    output_format: str       = "both",
    output_folder            = Path("."),
    email:         str|None  = None,
    password:      str|None  = None,
    log                      = print,
    progress                 = None,
    cancel_event             = None,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    adv       = advanced or {}
    has_adv   = any(v for v in adv.values() if v and v not in (False, "Unspecified"))
    want_docx = output_format in ("docx", "both")
    want_xlsx = output_format in ("xlsx", "both")
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    images_root = output_folder / "images"

    qname     = " ".join(p for p in (first_names, last_names) if p)
    qlines    = [ln for ln in [
        f"First Names: {first_names}",
        f"Last Names: {last_names}",
        f"Place Lived: {place_lived}",
        f"Birth Year: {birth_year}",
    ] if not ln.endswith(": ")]
    summary   = {"ok": False}
    file_base = safe_fn(qname) if qname else "familysearch_results"

    _prog(0, "Launching browser...")

    async with async_playwright() as pw:
        browser = await pw.chromium.launch(
            headless=False,
            args=["--start-maximized",
                  "--disable-blink-features=AutomationControlled"],
        )
        ctx  = await browser.new_context(no_viewport=True, accept_downloads=True)
        page = await ctx.new_page()

        try:
            # ── 1. Search ────────────────────────────────────────────────── #
            _prog(5, "Searching on home page...")
            await _search_from_home(page, first_names, last_names,
                                    place_lived, birth_year, log)
            if _done(): return summary

            # ── 2. Sign in ONCE via nav Sign In link ─────────────────────── #
            if email and password:
                _prog(12, "Signing in...")
                ok = await _sign_in_via_nav(page, email, password, log)
                if not ok:
                    summary.update({"error":   "login_failed",
                                    "message": "Sign-in failed — check credentials."})
                    return summary
                # If redirected away from results, navigate back
                if "discovery/results" not in page.url:
                    results_url = (
                        f"{FS_BASE}/en/search/discovery/results"
                        f"?q.givenName={first_names.replace(' ', '+')}"
                        f"&q.surname={last_names.replace(' ', '+')}"
                    )
                    if place_lived:
                        results_url += f"&q.anyPlace={place_lived.replace(' ', '+')}"
                    if birth_year:
                        results_url += f"&q.birthLikeDate.from={birth_year}&q.birthLikeDate.to={birth_year}"
                    _prog(15, "Navigating back to results after sign-in...")
                    await page.goto(results_url, wait_until="domcontentloaded", timeout=20000)
                    await asyncio.sleep(3)
            else:
                log("  No credentials provided — skipping sign-in")

            if _done(): return summary

            # ── 3. Historical Records tab ────────────────────────────────── #
            _prog(18, "Clicking Historical Records tab...")
            await _click_historical_tab(page, log)
            if _done(): return summary

            # ── 4. Advanced Search (if requested) ────────────────────────── #
            if has_adv:
                _prog(22, "Opening Advanced Search...")
                await _advanced_search(page, adv, log)
                if _done(): return summary

            # ── 5. Set 60 per page ───────────────────────────────────────── #
            _prog(26, "Setting results per page...")
            await _set_60(page, log)

            # ── 6. Collect results ───────────────────────────────────────── #
            _prog(30, "Collecting results...")
            raw       = await _collect(page, qname, log)
            qualified = [r for r in raw if r["score"] >= MIN_MATCH]
            log(f"  Qualified (≥{MIN_MATCH}%): {len(qualified)}")

            if not qualified:
                _prog(100, f"No results with match ≥ {MIN_MATCH}%.")
                summary.update({"ok": True, "n_records": 0,
                                "message": f"No records ≥ {MIN_MATCH}%."})
                return summary

            # ── 7. Scrape each record ────────────────────────────────────── #
            records: list = []
            n = len(qualified)
            for i, r in enumerate(qualified, 1):
                if _done():
                    break
                _prog(30 + int(55 * i / n),
                      f"[{i}/{n}] {r['name'][:60]}...")
                det = await _scrape_record(ctx, r["url"], r["name"],
                                           images_root, log)
                det["score"]         = r["score"]
                det["collection"]    = r.get("coll", "")
                det["events"]        = r.get("evts", "")
                det["relationships"] = r.get("rels", "")
                records.append(det)
                log(f"    ✓  {det['title'][:70]}  ({r['score']}%)")
                await asyncio.sleep(0.5)

            # ── 8. Save files ────────────────────────────────────────────── #
            _prog(88, "Saving output files...")
            docx_p = output_folder / f"{file_base}.docx"
            xlsx_p = output_folder / f"{file_base}.xlsx"
            sd = sx = False
            if want_docx and records:
                write_docx(docx_p, records, qlines)
                sd = True
                log(f"  Word: {docx_p}")
            if want_xlsx and records:
                write_xlsx(xlsx_p, records, qlines)
                sx = True
                log(f"  