"""
wikisource_scraper.py — Вікіджерела (Ukrainian Wikisource, uk.wikisource.org).

Archive scans of Jewish records (Ревізькі казки, метричні книги, рекрутські
списки …) live on uk.wikisource.org as an organised tree under
«Архів:Єврейське містечко/<губернія>», where each document is a link to a PDF
on Wikimedia Commons.

This scraper is **API-based** (MediaWiki action API + direct Commons download) —
no browser is needed. Commons full-resolution URLs are built straight from the
file name (the canonical upload-path md5 algorithm), so we skip the whole
«Download all sizes → Full resolution / click the cover twice → wait for the
huge PDF to render» dance a human has to do.

Two search modes:
  • locality — walk «Архів:Єврейське містечко/<губернія>», keep documents under
    the requested place (повіт / місто / містечко), filtered by document type.
  • code     — search by an archive reference (e.g. «ДАКрО/185/1/49») and grab
    the matching files.

Then every matching PDF is streamed to the output folder and a Word manifest of
what was downloaded is written.
"""

import urllib.request
import urllib.parse
import urllib.error
import json
import hashlib
import re
import time
import asyncio
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False

# ── Constants ───────────────────────────────────────────────────────────────
UA = "JewishGenSearch/1.0 (genealogy research)"
WS_API = "https://uk.wikisource.org/w/api.php"
COMMONS_API = "https://commons.wikimedia.org/w/api.php"
UPLOAD = "https://upload.wikimedia.org/wikipedia/commons"
ROOT = "Архів:Єврейське містечко"

# Search namespaces on uk.wikisource: 0 main, 116 «Архів», 252 «Індекс».
SR_NS = "0|116|252"

# Regions (губернії / воєводства / …) as they appear on the root page → their
# English / Russian display names. The GUI shows «English / Русский»; the search
# always uses the Ukrainian key (nobody types Ukrainian). «Примітки» (footnotes)
# is intentionally excluded. The live list is also re-read at run time so newly
# added regions still work in «all regions» mode.
REGIONS = {
    "Брацлавське воєводство":              ("Bratslav Voivodeship", "Брацлавское воеводство"),
    "Волинська губернія":                  ("Volhynia Governorate", "Волынская губерния"),
    "Герцогство Буковина":                 ("Duchy of Bukovina", "Герцогство Буковина"),
    "Закарпаття":                          ("Transcarpathia", "Закарпатье"),
    "Катеринославська губернія":           ("Yekaterinoslav Governorate", "Екатеринославская губерния"),
    "Київська губернія":                   ("Kiev Governorate", "Киевская губерния"),
    "Королівство Галичини та Володимирії": ("Kingdom of Galicia and Lodomeria", "Королевство Галиции и Лодомерии"),
    "Подільська губернія":                 ("Podolia Governorate", "Подольская губерния"),
    "Полтавська губернія":                 ("Poltava Governorate", "Полтавская губерния"),
    "Таврійська губернія":                 ("Taurida Governorate", "Таврическая губерния"),
    "Харківська губернія":                 ("Kharkov Governorate", "Харьковская губерния"),
    "Херсонська губернія":                 ("Kherson Governorate", "Херсонская губерния"),
    "Чернігівська губернія":               ("Chernigov Governorate", "Черниговская губерния"),
    "Переписи різних місцевостей":         ("Censuses of various localities", "Переписи разных местностей"),
}
GUBERNIAS = list(REGIONS.keys())
_SKIP_REGIONS = {"Примітки"}


def _region_ru(region: str) -> str:
    """Russian display name for a region (the Ukrainian page title stays for the API)."""
    return REGIONS.get(region, ("", ""))[1] or region


def _region_en(region: str) -> str:
    """English display name for a region (used in the English-GUI status line)."""
    return REGIONS.get(region, ("", ""))[0] or region

# Document type → keyword list (matched, lower-cased, against the link label).
DOC_TYPES = {
    "Ревізькі казки":               ["ревізьк", "ревизьк", "ревизск"],
    "Додаткові ревізькі казки":     ["додаткові ревізьк", "додаткова ревізьк", "дополнительн"],
    "Метричні книги":               ["метричн", "метрическ", "метрика", "книга запис",
                                     "книга народжен", "книга шлюб", "книга помер",
                                     "записів народжених", "записів шлюб", "записів помер",
                                     "запис актів", "народжен", "шлюб", "помер"],
    "Рекрутські / призовні списки": ["рекрут", "призов", "набор", "призивн",
                                     "черговий список", "повинн"],
    "Інші документи про євреїв":    ["євре", "єврей", "рабин", "рабинськ", "рабинів",
                                     "юдей", "іудей", "синагог", "кагал", "евре"],
}

# Document type → (English, Russian) for the GUI; keys match DOC_TYPES.
DOC_TYPE_TR = {
    "Ревізькі казки":               ("Revision lists (census)", "Ревизские сказки"),
    "Додаткові ревізькі казки":     ("Additional revision lists", "Дополнительные ревизские сказки"),
    "Метричні книги":               ("Vital records (metrical books)", "Метрические книги"),
    "Рекрутські / призовні списки": ("Recruitment / conscription lists", "Рекрутские / призывные списки"),
    "Інші документи про євреїв":    ("Other Jewish records", "Прочие еврейские документы"),
}

# Any Jewish-related token (used by the «Інші про євреїв» / default filter).
_JEWISH = ["євре", "єврей", "рабин", "юдей", "іудей", "синагог", "кагал"]


def safe_fn(s: str, maxlen: int = 120) -> str:
    s = re.sub(r"[\\/:*?\"<>|\n\r\t]+", " ", s or "").strip()
    s = re.sub(r"\s{2,}", " ", s)
    return (s[:maxlen]).strip() or "document"


# ── MediaWiki API ───────────────────────────────────────────────────────────
def _api(base: str, params: dict, retries: int = 3, log=print) -> dict:
    params = {**params, "format": "json", "formatversion": "2"}
    url = base + "?" + urllib.parse.urlencode(params, quote_via=urllib.parse.quote)
    last = None
    for a in range(retries):
        try:
            req = urllib.request.Request(url, headers={"User-Agent": UA})
            with urllib.request.urlopen(req, timeout=45) as r:
                return json.loads(r.read().decode("utf-8", "replace"))
        except Exception as e:
            last = e
            time.sleep(2 + 2 * a)
    log(f"   !! API error: {type(last).__name__}: {last}")
    return {}


def _wikitext(title: str, log=print) -> str:
    d = _api(WS_API, {"action": "parse", "page": title,
                      "prop": "wikitext", "redirects": 1}, log=log)
    return ((d.get("parse") or {}).get("wikitext") or "") if d else ""


def _search(query: str, limit: int, log=print) -> list:
    out, off = [], 0
    while len(out) < limit:
        d = _api(WS_API, {"action": "query", "list": "search", "srsearch": query,
                          "srlimit": min(500, limit - len(out)), "sroffset": off,
                          "srnamespace": SR_NS, "srwhat": "text"}, log=log)
        hits = (d.get("query") or {}).get("search") or []
        out += [h["title"] for h in hits]
        cont = (d.get("continue") or {}).get("sroffset")
        if not hits or cont is None:
            break
        off = cont
    return out


_H2_RE = re.compile(r"(?m)^==\s*([^=].*?)\s*==\s*$")


def _live_regions(log=print) -> list:
    """The current top-level region list from the root page («Примітки» dropped)."""
    wt = _wikitext(ROOT, log)
    regs = [h.strip() for h in _H2_RE.findall(wt) if h.strip() not in _SKIP_REGIONS]
    return regs or GUBERNIAS


def _section_slice(wikitext: str, heading: str) -> str:
    """Wikitext between «== heading ==» and the next «== … ==» (for regions that
    have no subpage and live directly on the root page)."""
    out, grab = [], False
    for ln in (wikitext or "").splitlines():
        m = _H2_RE.match(ln)
        if m:
            grab = (m.group(1).strip() == heading.strip())
            continue
        if grab:
            out.append(ln)
    return "\n".join(out)


def _region_wikitext(region: str, root_wt: str = "", log=print) -> str:
    """Wikitext for a region: its subpage if it exists, else its root-page slice."""
    wt = _wikitext(f"{ROOT}/{region}", log)
    if wt.strip():
        return wt
    root_wt = root_wt or _wikitext(ROOT, log)
    return _section_slice(root_wt, region)


# ── Commons file → download URL ─────────────────────────────────────────────
def _md5_shard(filename: str) -> str:
    fn = filename.replace(" ", "_")
    h = hashlib.md5(fn.encode("utf-8")).hexdigest()
    return f"{h[0]}/{h[:2]}/{urllib.parse.quote(fn)}"


def _file_dl_url(filename: str) -> str:
    """Canonical Commons full-resolution URL (no API call needed)."""
    return f"{UPLOAD}/{_md5_shard(filename)}"


def _file_preview_url(filename: str, width: int = 1024) -> str:
    """First-page JPG preview of a PDF on Commons."""
    fn = filename.replace(" ", "_")
    h = hashlib.md5(fn.encode("utf-8")).hexdigest()
    return (f"{UPLOAD}/thumb/{h[0]}/{h[:2]}/{urllib.parse.quote(fn)}"
            f"/page1-{width}px-thumbnail.pdf.jpg")


def _file_page_url(filename: str) -> str:
    return "https://commons.wikimedia.org/wiki/File:" + urllib.parse.quote(
        filename.replace(" ", "_"))


def _imageinfo_url(filename: str, log=print) -> str:
    """Fallback: ask Commons for the real URL (handles odd shards / redirects)."""
    d = _api(COMMONS_API, {"action": "query", "titles": "File:" + filename,
                           "prop": "imageinfo", "iiprop": "url", "redirects": 1},
             log=log)
    for p in (d.get("query") or {}).get("pages") or []:
        ii = p.get("imageinfo") or []
        if ii and ii[0].get("url"):
            return ii[0]["url"]
    return ""


def _imageinfo_thumb(filename: str, width: int = 1024, log=print) -> str:
    """A valid first-page JPG thumbnail URL (Commons renders any width on demand)."""
    d = _api(COMMONS_API, {"action": "query", "titles": "File:" + filename,
                           "prop": "imageinfo", "iiprop": "url",
                           "iiurlwidth": width, "redirects": 1}, log=log)
    for p in (d.get("query") or {}).get("pages") or []:
        ii = p.get("imageinfo") or []
        if ii and ii[0].get("thumburl"):
            return ii[0]["thumburl"]
    return ""


# ── Wikitext tree parsing ───────────────────────────────────────────────────
_FILE_RE = re.compile(r"\[\[\s*(?:c|commons|:c|:commons):File:([^|\]]+?)\s*(?:\|([^\]]*))?\]\]", re.I)
# A document linked through a Wikisource «Архів:CODE» index page (NOT the
# «Єврейське містечко» navigation tree itself) — resolved to its Commons file.
_INDEX_RE = re.compile(r"\[\[\s*(Архів:(?!Єврейське містечко)[^|\]#]+?)\s*(?:\|([^\]]*))?\]\]")
# A Commons file inside an «Архів:CODE» index page — either a wikilink
# («[[File:…]]» / «[[Файл:…]]») or a template parameter («| link_commons = File:…»
# in {{Архіви/справа}}).
_ANYFILE_RE = re.compile(r"\[\[\s*(?:c:|commons:|:)?(?:File|Файл|Image|Зображення)\s*:\s*([^|\]]+?)\s*(?:\|[^\]]*)?\]\]", re.I)
_PARAM_FILE_RE = re.compile(r"(?:link_commons|commons|файл|file|link)\s*=\s*(?:\[\[)?\s*(?:c:|commons:|:)?(?:File|Файл)\s*:\s*([^|\]\n}]+)", re.I)
_HEAD_RE = re.compile(r"^(={2,6})\s*(.+?)\s*\1\s*$")
_TERM_RE = re.compile(r"^;\s*(.+?)\s*$")        # «; Місто Житомир»


def _index_to_file(title: str, log=print) -> str:
    """An «Архів:CODE» index page → the Commons file it links to."""
    wt = _wikitext(title, log)
    m = _PARAM_FILE_RE.search(wt) or _ANYFILE_RE.search(wt)
    if not m:
        return ""
    return re.sub(r"\s+", " ", m.group(1).replace("_", " ")).strip()


def _moved_target(filename: str, log=print) -> str:
    """If a file was moved without a redirect, the move log holds the new name."""
    d = _api(COMMONS_API, {"action": "query", "list": "logevents", "letype": "move",
                           "letitle": "File:" + filename, "lelimit": 3}, log=log)
    for ev in (d.get("query") or {}).get("logevents") or []:
        tt = (ev.get("params") or {}).get("target_title")
        if tt:
            return tt[5:].replace("_", " ").strip() if tt.lower().startswith("file:") else tt
    return ""


def _clean_wiki(s: str) -> str:
    s = s or ""
    s = re.sub(r"\[\[[^|\]]*\|([^\]]*)\]\]", r"\1", s)   # [[link|text]] → text
    s = re.sub(r"\[\[([^\]]*)\]\]", r"\1", s)            # [[text]] → text
    s = re.sub(r"'{2,}", "", s)                          # bold/italic
    s = re.sub(r"<[^>]+>", "", s)                        # html
    s = re.sub(r"\s{2,}", " ", s)
    return s.strip()


def _parse_tree(wikitext: str) -> list:
    """→ [{file, label, section}] — section = «повіт / місто / містечко» chain."""
    items, ctx, term = [], [], ""
    for raw in (wikitext or "").splitlines():
        line = raw.strip()
        hm = _HEAD_RE.match(line)
        if hm:
            lvl, head = len(hm.group(1)), _clean_wiki(hm.group(2))
            ctx = [(l, h) for (l, h) in ctx if l < lvl] + [(lvl, head)]
            term = ""
            continue
        tm = _TERM_RE.match(line)
        if tm:
            term = _clean_wiki(tm.group(1))
            continue
        sec = " / ".join(h for _, h in ctx)
        if term:
            sec = (sec + " / " + term) if sec else term
        for m in _FILE_RE.finditer(raw):          # direct Commons file
            fn = re.sub(r"\s+", " ", m.group(1).replace("_", " ")).strip()
            items.append({"file": fn, "index": "",
                          "label": _clean_wiki(m.group(2) or fn), "section": sec})
        for m in _INDEX_RE.finditer(raw):          # «Архів:CODE» index page
            title = m.group(1).strip()
            items.append({"file": "", "index": title,
                          "label": _clean_wiki(m.group(2) or title), "section": sec})
    return items


def _norm(s: str) -> str:
    return re.sub(r"[\s\-_/.,()]+", " ", (s or "").lower()).strip()


# Cyrillic (Ukrainian + Russian) → a coarse Latin form so a place typed in
# Russian or English matches the Ukrainian section name (и/і/ї/й/ы/y → i,
# е/є/э → e, etc.). The site is Ukrainian — nobody types Ukrainian.
_TRANSLIT = {
    "а": "a", "б": "b", "в": "v", "г": "g", "ґ": "g", "д": "d", "е": "e", "є": "e",
    "ё": "e", "ж": "zh", "з": "z", "и": "i", "і": "i", "ї": "i", "й": "i", "к": "k",
    "л": "l", "м": "m", "н": "n", "о": "o", "п": "p", "р": "r", "с": "s", "т": "t",
    "у": "u", "ф": "f", "х": "h", "ц": "c", "ч": "ch", "ш": "sh", "щ": "sch",
    "ъ": "", "ы": "i", "ь": "", "э": "e", "ю": "iu", "я": "ia", "'": "", "’": "",
}
_LOC_STOP = {"povit", "uezd", "guberniia", "guberniya", "gorod", "misto", "selo",
             "raion", "okrug", "voevodstvo", "voievodstvo", "the", "of"}


def _translit(s: str) -> str:
    s = (s or "").lower()
    t = "".join(_TRANSLIT.get(ch, ch) for ch in s)
    # fold English spellings onto the Ukrainian transliteration
    t = t.replace("shch", "sch").replace("kh", "h").replace("ts", "c")
    t = t.replace("y", "i").replace("j", "i")          # English y/j → i
    t = re.sub(r"[^a-z0-9]+", " ", t)
    return re.sub(r"\s+", " ", t).strip()


def _loc_match(item: dict, locality: str) -> bool:
    if not locality:
        return True
    sec = _translit(item["section"] + " " + item["label"])
    toks = [t for t in _translit(locality).split()
            if len(t) >= 3 and t not in _LOC_STOP]
    if not toks:
        return True
    # every meaningful query token's stem must appear in the section (handles
    # Ukrainian declensions: «Сураж» → «Суражського», «Сосниця» → «Сосницького»)
    for tok in toks:
        if tok[:6] not in sec:
            return False
    return True


def _type_match(item: dict, types: list, all_docs: bool, jewish_only: bool) -> bool:
    lab = (item["label"] + " " + item["file"]).lower()
    if all_docs:
        return (not jewish_only) or any(k in lab for k in _JEWISH)
    if not types:                                   # nothing ticked → any Jewish doc
        return any(k in lab for k in _JEWISH)
    for t in types:
        for kw in DOC_TYPES.get(t, []):
            if kw in lab:
                return True
    return False


# ── Download (streaming — PDFs can be > 1 GB; Commons rate-limits hard) ──────
# Wikimedia answers rapid bulk full-file pulls with HTTP 429. We stay polite:
# a minimum gap between downloads that grows whenever a 429 is seen, and a
# back-off (honouring «Retry-After») that retries instead of giving up.
_dl_gap = [3.0]        # current min seconds between download starts (adaptive)
_dl_last = [0.0]


def _polite_wait():
    dt = time.time() - _dl_last[0]
    if dt < _dl_gap[0]:
        time.sleep(_dl_gap[0] - dt)


def _stream(url: str, dest: Path):
    """Stream url → dest. Returns (bytes, status, retry_after). status:
    'ok' | '429' | '404' | 'httpNNN' | '<ExceptionName>'."""
    _polite_wait()
    req = urllib.request.Request(url, headers={
        "User-Agent": UA, "Referer": "https://uk.wikisource.org/"})
    try:
        with urllib.request.urlopen(req, timeout=180) as r:
            tmp = dest.with_suffix(dest.suffix + ".part")
            done = 0
            with open(tmp, "wb") as f:
                while True:
                    chunk = r.read(1 << 20)          # 1 MB
                    if not chunk:
                        break
                    f.write(chunk)
                    done += len(chunk)
            tmp.replace(dest)
            _dl_last[0] = time.time()
            return done, "ok", None
    except urllib.error.HTTPError as e:
        _dl_last[0] = time.time()
        if e.code == 429:
            return 0, "429", e.headers.get("Retry-After")
        return 0, ("404" if e.code == 404 else f"http{e.code}"), None
    except Exception as e:
        _dl_last[0] = time.time()
        try:
            dest.with_suffix(dest.suffix + ".part").unlink()
        except Exception:
            pass
        return 0, type(e).__name__, None


def _download_file(filename: str, dest: Path, preview: bool, log=print,
                   max_tries: int = 6) -> int:
    for attempt in range(max_tries):
        if preview:
            url = _imageinfo_thumb(filename, 1024, log) or _file_preview_url(filename, 960)
        else:
            url = _file_dl_url(filename)
        nbytes, status, ra = _stream(url, dest)
        if nbytes:
            return nbytes
        if status == "429":
            wait = int(ra) if (ra and str(ra).isdigit()) else min(90, 6 * (2 ** attempt))
            _dl_gap[0] = min(15.0, _dl_gap[0] + 1.5)      # grow the gap, stay polite
            log(f"      ⏳ 429 — waiting {wait}s, retry {attempt + 1}/{max_tries}…")
            time.sleep(wait)
            continue
        if status == "404" and not preview:               # odd shard or moved file
            alt = _imageinfo_url(filename, log)
            if alt:
                nb, _, _ = _stream(alt, dest)
                if nb:
                    return nb
            new = _moved_target(filename, log)            # moved without redirect
            if new and new != filename:
                log(f"      ↪ file moved → {new[:60]}")
                nb, _, _ = _stream(_file_dl_url(new), dest)
                if nb:
                    return nb
            log(f"      !! file not found: {filename}")
            return 0
        if attempt < max_tries - 1:                       # transient → brief retry
            time.sleep(4)
            continue
        log(f"      !! couldn't fetch ({status}): {filename}")
        return 0
    return 0


# ── Word manifest ───────────────────────────────────────────────────────────
def _add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    h = OxmlElement("w:hyperlink"); h.set(qn("r:id"), r_id)
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "1155CC"); rPr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    new_run = OxmlElement("w:r"); new_run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; new_run.append(t)
    h.append(new_run); paragraph._p.append(h)


def write_docx(path, recs, qlines, append=False):
    if not _DOCX_OK:
        return
    if append and Path(path).exists():
        doc = Document(str(path)); doc.add_page_break()
        ap = doc.add_paragraph(); ar = ap.add_run(f"➕ Добавлено {len(recs)}")
        ar.bold = True; ar.font.size = Pt(13)
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width = Mm(210); s.page_height = Mm(297)
        s.left_margin = s.right_margin = Mm(16)
        ht = doc.add_paragraph(); htr = ht.add_run("Викиджерела — скачанные документы")
        htr.bold = True; htr.font.size = Pt(14)
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ln in qlines:
            doc.add_paragraph(ln)
        doc.add_paragraph(f"Всего: {len(recs)}")
        doc.add_paragraph("")
    tbl = doc.add_table(rows=1, cols=4); tbl.style = "Light Grid Accent 1"
    tbl.autofit = False; tbl.allow_autofit = False
    # «№» — узкая; «Источник» — чуть шире (чтобы слово влезало); «Название
    # документа» и «Файл» — широкие (A4 ≈ 7").
    widths = (Inches(0.4), Inches(3.3), Inches(2.35), Inches(0.95))
    for c, txt, wd in zip(tbl.rows[0].cells, ("№", "Название документа", "Файл", "Источник"), widths):
        c.text = txt; c.width = wd
        for run in c.paragraphs[0].runs:
            run.bold = True
    for i, r in enumerate(recs, 1):
        cells = tbl.add_row().cells
        cells[0].text = str(i)
        cells[1].text = r.get("label", "")
        cells[2].text = (str(Path(r["path"]).resolve())   # точный путь куда сгружен
                         if r.get("path") else "— (не скачано)")
        _add_hyperlink(cells[3].paragraphs[0], "Commons", r.get("page", ""))
        for cell, wd in zip(cells, widths):
            cell.width = wd
    doc.save(str(path))


# ── Main entry point ────────────────────────────────────────────────────────
async def run_scraper(*,
    mode="locality", query="", gubernia="", locality="",
    types=None, all_docs=False, jewish_only=True,
    preview_only=False, output_folder=".",
    log=print, progress=None, cancel_event=None, ask_file_conflict=None,
    max_docs=400, **_,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _stop():
        return bool(cancel_event and cancel_event.is_set())

    types = types or []
    out_dir = Path(output_folder); out_dir.mkdir(parents=True, exist_ok=True)
    summary = {"ok": False, "n_records": 0}

    # 1) collect candidate documents ────────────────────────────────────────
    _prog(2, "Starting…")
    items = []
    if mode == "code":
        code = query.strip()
        if not code:
            _prog(100, "Empty code."); return summary
        qkey = code
        _prog(6, f"Search by code: {code}")
        titles = [f"Архів:{code}"] + _search(code, 200, log)
        seen_t, code_n = set(), _norm(code)
        for ti, t in enumerate(titles):
            if _stop():
                break
            if t in seen_t:
                continue
            seen_t.add(t)
            _prog(6 + int(20 * ti / max(1, len(titles))), f"Reading: {t}")
            for it in _parse_tree(_wikitext(t, log)):
                if (code_n in _norm(it.get("file") or "")
                        or code_n in _norm(it.get("index") or "")
                        or code_n in _norm(it["label"])):
                    items.append(it)
    else:
        loc = locality.strip()
        qkey = " · ".join(x for x in (_region_ru(gubernia) if gubernia else "", loc) if x) or "Еврейские местечки"
        if not gubernia and not loc:
            _prog(100, "Specify a region or place."); return summary
        if gubernia and gubernia in REGIONS:
            gubs, root_wt = [gubernia], ""
        else:
            gubs = _live_regions(log)               # ALL regions, incl. root-only ones
            root_wt = _wikitext(ROOT, log)
        for gi, g in enumerate(gubs):
            if _stop():
                break
            _prog(6 + int(24 * gi / len(gubs)), f"Region: {_region_en(g)}")
            tree = _parse_tree(_region_wikitext(g, root_wt, log))
            for it in tree:
                if _loc_match(it, loc) and _type_match(it, types, all_docs, jewish_only):
                    it["gubernia"] = g
                    items.append(it)

    # de-dupe (by direct file or by index-page title), cap, then resolve the
    # «Архів:CODE» index pages to their Commons file.
    uniq, seen = [], set()
    for it in items:
        key = it.get("file") or it.get("index")
        if key and key not in seen:
            seen.add(key); uniq.append(it)
    items = uniq[:max_docs]
    for it in items:
        if not it.get("file") and it.get("index"):
            it["file"] = _index_to_file(it["index"], log)
    items = [it for it in items if it.get("file")]
    log(f"  Found {len(items)} document(s)")
    if not items:
        _prog(100, "Nothing found for these criteria.")
        summary["ok"] = True
        return summary

    # 2) download ────────────────────────────────────────────────────────────
    sub = out_dir / safe_fn(qkey)
    sub.mkdir(parents=True, exist_ok=True)
    recs, n = [], len(items)
    for i, it in enumerate(items, 1):
        if _stop():
            break
        label = it["label"] or it["file"]
        _prog(32 + int(64 * i / n), f"[{i}/{n}] {label[:60]}")
        ext = ".jpg" if preview_only else (Path(it["file"]).suffix or ".pdf")
        dest = sub / (safe_fn(label) + ext)
        page = _file_page_url(it["file"])
        if dest.exists() and dest.stat().st_size > 10000:   # resume: keep what's done
            log(f"      • already downloaded: {dest.name}")
            recs.append({"label": label, "file": it["file"], "page": page, "path": str(dest)})
            await asyncio.sleep(0)
            continue
        got = _download_file(it["file"], dest, preview_only, log)
        rec = {"label": label, "file": it["file"], "page": page,
               "path": str(dest) if got else ""}
        if got:
            log(f"      ✓ {dest.name} ({got / 1048576:.1f} MB)")
        else:
            log(f"      ✗ couldn't download: {it['file']}")
        recs.append(rec)
        await asyncio.sleep(0)

    # 3) Word manifest ───────────────────────────────────────────────────────
    _prog(98, "Writing manifest…")
    qlines = [f"Запрос: {qkey}",
              f"Режим: {'архивный шифр' if mode == 'code' else 'местность'}"]
    docx_p = out_dir / (safe_fn("wikisource_" + qkey) + ".docx")
    decision = "overwrite"
    if docx_p.exists() and ask_file_conflict:
        decision = ask_file_conflict([docx_p.name]) or "overwrite"
    try:
        if decision != "skip":
            write_docx(docx_p, recs, qlines, append=(decision == "append"))
    except PermissionError:
        docx_p = out_dir / (safe_fn("wikisource_" + qkey) + f"_{int(time.time())}.docx")
        write_docx(docx_p, recs, qlines, append=False)
    except Exception as e:
        log(f"   !! Word: {type(e).__name__}: {e}")

    n_ok = sum(1 for r in recs if r["path"])
    _prog(100, f"Done — downloaded {n_ok} of {len(recs)}.")
    summary.update({"ok": True, "n_records": n_ok, "n_total": len(recs),
                    "output_folder": str(sub)})
    return summary
