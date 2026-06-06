#!/usr/bin/env python3
"""
JewishGen mass-search scraper.

The script asks you, via the terminal, for the values you want to put into the
search form, then drives the JewishGen site through Playwright. It never uses
CSS attribute selectors (no [name=...], no [id=...], no [class=...]). It
clicks elements by their visible names/labels exactly the way a human would.

Flow
----
1. Terminal asks for your search rows (up to four) and the keyword used to
   filter result rows.
2. A Chromium window opens at https://jewishgen.org.
3. The script fills the form by choosing dropdown options by their visible
   labels (Surname / Given Name / Town / Any Field, Phonetically Like / …,
   ALL COUNTRIES / Ukraine / …) and clicks the button labelled
   "Search our Collection!".
4. If JewishGen asks for a login, log in manually in the browser window. The
   script never asks for or stores your password. Press Enter in the terminal
   when you're back on the results page.
5. The script then clicks every "List N records" button on the results page
   in order (by its visible label). For each one it scrapes the rows of the
   result table, filters by your keyword (case-insensitive), follows the
   "Next" link at the bottom if there is one, and appends matching rows to
   jewishgen_results.docx under the database's own title.

Requirements
------------
    pip install playwright python-docx
    playwright install chromium
"""

import asyncio
import json
import os
import random
import re
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Mm
from docx.enum.section import WD_ORIENT

# Fix Playwright browser path inside PyInstaller bundle
if getattr(sys, "frozen", False):
    exe_dir = Path(sys.executable).resolve().parent
    browser_dir = exe_dir / "ms-playwright"

    if browser_dir.exists():
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(browser_dir)

try:
    from playwright.async_api import async_playwright, TimeoutError as PWTimeout
except ImportError:
    sys.stderr.write(
        "Playwright is not installed. Run:\n"
        "    pip install playwright python-docx openpyxl\n"
        "    playwright install chromium\n"
    )
    sys.exit(1)

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Excel output is optional at runtime — if openpyxl is missing the user can
# still produce .docx files. We only complain when they ask for xlsx.
try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    _OPENPYXL_AVAILABLE = True
except ImportError:
    _OPENPYXL_AVAILABLE = False


# When the script runs from inside a PyInstaller bundle, Playwright must look
# for its bundled browsers next to the executable, not in the user's cache. We
# set PLAYWRIGHT_BROWSERS_PATH before importing playwright would have read it,
# but here it's still fine because async_playwright reads the env on launch.
def _wire_playwright_bundled_browsers():
    if not getattr(sys, "frozen", False):
        return
    bundle_dir = getattr(sys, "_MEIPASS", None) or os.path.dirname(sys.executable)
    candidate = os.path.join(bundle_dir, "ms-playwright")
    if os.path.isdir(candidate):
        os.environ.setdefault("PLAYWRIGHT_BROWSERS_PATH", candidate)


_wire_playwright_bundled_browsers()


# ---------- server-error helpers -------------------------------------------- #

class TemporaryServerError(RuntimeError):
    """Raised when JewishGen returns a 503 / 524 / 'Service Unavailable' page."""


async def ensure_not_503(page):
    """Raise TemporaryServerError if the current page looks like a 5xx error.
    Covers both 503 (overload) and Cloudflare 524 (origin timeout)."""
    try:
        title = (await page.title()).lower()
    except Exception:
        title = ""
    try:
        html = (await page.content())[:8000].lower()
    except Exception:
        html = ""
    if (
        "503" in title
        or "524" in title
        or "service unavailable" in html
        or "temporarily unavailable" in html
        or "bad gateway" in html
        or "error code 524" in html
        or "connection timed out" in html
        or ("cloudflare" in html and ("524" in html or "timeout" in html))
    ):
        raise TemporaryServerError(
            f"JewishGen returned a server/Cloudflare error page (title={title!r})"
        )


HYPERLINK_REL = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


def _add_hyperlink(paragraph, text, url, bold=False):
    """Insert a clickable hyperlink into a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, HYPERLINK_REL, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text or url
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


HOME_URL = "https://jewishgen.org/"
HERE = Path(__file__).resolve().parent
OUTPUT_DOCX = HERE / "jewishgen_results.docx"
# Persistent Chromium profile so JewishGen cookies live between runs — you
# log in once, the next runs are already authenticated.
PROFILE_DIR = HERE / ".profile"

DATA_TYPES = ["Surname", "Given Name", "Town", "Any Field"]
# JewishGen's dropdown uses "GivenName" without a space — translate when
# we select it.
DATA_TYPE_TO_JG_LABEL = {
    "Surname": "Surname",
    "Given Name": "GivenName",
    "Town": "Town",
    "Any Field": "Any Field",
}
SEARCH_TYPES = [
    "Phonetically Like", "Sounds Like", "Starts with", "is Exactly",
    "Fuzzy Match", "Fuzzier Match", "Fuzziest Match",
]
COUNTRIES = [
    "ALL COUNTRIES", "Algeria", "Argentina", "Austria / Czechia", "Belarus",
    "Bulgaria", "Canada", "CryptoJews", "Egypt", "France", "Germany",
    "Hungary / Slovakia", "India", "Iraq", "Ireland", "Israel", "Italy",
    "Latvia", "Latin America", "Lebanon", "Libya", "Lithuania", "Morocco",
    "Netherlands", "Poland", "Portugal", "Romania / Moldova", "Sephardic",
    "Scandinavia", "Spain", "South Africa", "Syria", "Tunisia", "Turkey",
    "Ukraine", "United Kingdom", "United States", "Venezuela", "Yugoslavia",
    "Holocaust records", "Cemetery records", "Search by Cemetery",
]


# ---------- user input ------------------------------------------------------ #
def _menu_pick(prompt, options, allow_blank=False):
    print(prompt)
    for i, opt in enumerate(options, 1):
        print(f"  {i:>2}. {opt}")
    while True:
        ans = input("  number (or text)" + (" — blank to skip" if allow_blank else "") + ": ").strip()
        if not ans:
            if allow_blank:
                return None
            continue
        if ans.isdigit() and 1 <= int(ans) <= len(options):
            return options[int(ans) - 1]
        for opt in options:
            if opt.lower() == ans.lower():
                return opt
        print("  unrecognised, try again.")


def _required_text(prompt, missing_msg):
    """Ask for a non-empty text — keeps re-prompting on blank instead of
    crashing the program."""
    while True:
        s = input(prompt).strip()
        if s:
            return s
        print(f"  ({missing_msg})")


def ask_inputs():
    print("=== JewishGen mass-search ===\n")
    rows = []
    for i in range(1, 5):
        is_required = (i == 1)
        print(f"\n--- Search row {i} "
              f"{'(required)' if is_required else '(optional — leave blank to skip)'} ---")
        # Data Type: required for row 1 (loops until valid), optional for the
        # rest (blank skips the whole row, no further prompts for this row).
        data_type = _menu_pick("Data Type:", DATA_TYPES, allow_blank=not is_required)
        if data_type is None:
            # Optional row left blank — stop asking for further rows too.
            break
        search_type = _menu_pick("Search Type:", SEARCH_TYPES, allow_blank=False)
        term = _required_text(
            "  text to search: ",
            "text is required for this row",
        )
        rows.append((data_type, search_type, term))

    print()
    country = _menu_pick("Region / Country:", COUNTRIES, allow_blank=False)
    keyword = _required_text(
        "\nKeyword to filter result rows: ",
        "a keyword is required",
    )
    email = input("\nYour JewishGen e-mail (optional, used to pre-fill the login\n"
                  "form — leave blank to type it yourself): ").strip()
    return rows, country, keyword, email


# ---------- output format prompt -------------------------------------------- #
OUTPUT_FORMATS = {
    "1": "docx",
    "2": "xlsx",
    "3": "both",
}


def ask_output_format(default="3"):
    """Ask the user whether to write .docx (one per DB), .xlsx (one workbook
    with a sheet per DB), or both. Defaults to BOTH."""
    print("\nOutput format:")
    print("  1. Word (.docx) only — one file per database in results/")
    print("  2. Excel (.xlsx) only — one workbook, sheet per database")
    print("  3. Both (Recommended)")
    while True:
        ans = input(f"  number [{default}]: ").strip() or default
        if ans in OUTPUT_FORMATS:
            return OUTPUT_FORMATS[ans]
        # Accept the word too, for convenience.
        if ans.lower() in ("docx", "word"):
            return "docx"
        if ans.lower() in ("xlsx", "excel"):
            return "xlsx"
        if ans.lower() in ("both", "all"):
            return "both"
        print("  unrecognised, try again (1 / 2 / 3).")


# ---------- docx output ----------------------------------------------------- #
def safe_filename(name, max_len=80):
    """Convert an arbitrary database name into a safe filename."""
    s = re.sub(r"[^\w\s-]", "", name, flags=re.UNICODE).strip()
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return (s[:max_len] or "unnamed")


def _set_cell_lines(cell, lines, bold=False):
    """Render a cell's content into a docx table cell.

    `lines` is a list of lines. Each line is either:
      * a plain string                            — plain text
      * a list of {"t": text, "h": href|null}      — mixed text + hyperlinks
    Each line becomes its own paragraph inside the cell. Segments with `h`
    set are rendered as clickable blue underlined hyperlinks."""
    cell.text = ""
    if not lines:
        return
    # Normalise to a list-of-lines-of-segments form.
    norm = []
    for line in lines:
        if isinstance(line, str):
            t = line.strip()
            if t:
                norm.append([{"t": t, "h": None}])
        elif isinstance(line, list):
            segs = [s for s in line if (s.get("t") or s.get("h"))]
            if segs:
                norm.append(segs)
    if not norm:
        return

    def render(paragraph, segs):
        for seg in segs:
            text = seg.get("t") or ""
            href = seg.get("h")
            if href:
                _add_hyperlink(paragraph, text or href, href, bold=bold)
            elif text:
                r = paragraph.add_run(text)
                if bold:
                    r.bold = True

    render(cell.paragraphs[0], norm[0])
    for segs in norm[1:]:
        render(cell.add_paragraph(), segs)


def write_database_docx(out_path, db_name, header_lines, columns, rows, query_lines):
    """Write ONE .docx file for a single database. Each matched row from the
    JewishGen page becomes its own row in the .docx table, cell-for-cell —
    multi-line cells stay multi-line, single-line cells stay single-line."""
    doc = Document()

    # A4 paper, landscape, minimal margins (5 mm) so wide JG tables fit.
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.top_margin = Mm(5)
    section.bottom_margin = Mm(5)
    section.left_margin = Mm(5)
    section.right_margin = Mm(5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(db_name)
    tr.bold = True
    tr.font.size = Pt(16)

    for line in header_lines:
        if line.strip().lower() == db_name.strip().lower():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.bold = True
        r.font.size = Pt(11)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(" | ".join(query_lines))
    sr.italic = True

    doc.add_paragraph()

    if not rows:
        doc.save(out_path)
        return

    def row_text_len(row):
        total = 0
        for c in row:
            if isinstance(c, list):
                for line in c:
                    if isinstance(line, str):
                        total += len(line.strip())
                    elif isinstance(line, list):
                        for s in line:
                            if isinstance(s, dict):
                                total += len((s.get("t") or "").strip())
            elif isinstance(c, str):
                total += len(c.strip())
        return total

    # Drop any trailing empty rows so the table doesn't end with blanks.
    while rows and row_text_len(rows[-1]) == 0:
        rows.pop()
    if not rows:
        doc.save(out_path)
        return

    def cell_text(c):
        """Flatten a cell into a single string (for synthetic-column check)."""
        if isinstance(c, list):
            out = []
            for line in c:
                if isinstance(line, str):
                    out.append(line)
                elif isinstance(line, list):
                    out.append("".join(s.get("t", "") for s in line if isinstance(s, dict)))
            return " ".join(out).strip()
        return str(c).strip()

    real_cols = bool(columns) and not all(
        re.match(r"^Col\s+\d+$", cell_text(c)) for c in columns
    )

    def cell_lines(c):
        """Normalise a cell into a list of lines (each line: str or [segs])."""
        if isinstance(c, list):
            # Already a list of lines (each a str or list-of-segments).
            return c
        if isinstance(c, str):
            return [c]
        return [str(c)]

    ncols = max(
        len(columns or []),
        max(len(r) for r in rows),
    )

    n_data = len(rows)
    n_total = n_data + (1 if real_cols else 0)
    table = doc.add_table(rows=n_total, cols=ncols)
    table.style = "Light Grid Accent 1"

    start_row = 0
    if real_cols:
        cols_padded = list(columns) + [""] * (ncols - len(columns))
        for j, c in enumerate(cols_padded):
            _set_cell_lines(table.rows[0].cells[j], cell_lines(c), bold=True)
        start_row = 1

    for i, row in enumerate(rows):
        for j in range(ncols):
            cell_val = row[j] if j < len(row) else ""
            _set_cell_lines(table.rows[start_row + i].cells[j], cell_lines(cell_val))

    doc.save(out_path)


# ---------- xlsx output ----------------------------------------------------- #
# Excel limits a sheet name to 31 chars and forbids: \ / ? * [ ] :
_BAD_SHEET_CHARS = re.compile(r"[\\/\?\*\[\]:]")


def safe_sheet_name(name, used, max_len=31):
    """Return a unique, Excel-safe sheet name. `used` is the set of names
    already taken in this workbook — collisions get '_2', '_3' suffixes."""
    s = _BAD_SHEET_CHARS.sub(" ", name or "").strip()
    s = re.sub(r"\s+", " ", s).strip()
    if not s:
        s = "Sheet"
    base = s[:max_len]
    candidate = base
    n = 2
    while candidate in used:
        suffix = f"_{n}"
        candidate = (base[: max_len - len(suffix)]) + suffix
        n += 1
    used.add(candidate)
    return candidate


def _cell_to_text_and_link(cell_val):
    """Reduce a docx-style cell value into (plain text with '\\n' between
    lines, first_url_or_None). Multiple hyperlinks in one cell are joined
    into the text as 'Anchor (https://...)' — Excel only supports one
    hyperlink per cell, so we keep the first one as the click target."""
    lines = []
    first_url = None

    def visit_line(line):
        nonlocal first_url
        if isinstance(line, str):
            t = line.strip()
            if t:
                lines.append(t)
            return
        if isinstance(line, list):
            parts = []
            for seg in line:
                if not isinstance(seg, dict):
                    continue
                txt = (seg.get("t") or "").strip()
                href = seg.get("h")
                if href:
                    if first_url is None:
                        first_url = href
                        parts.append(txt or href)
                    else:
                        # Inline secondary links as "text (url)" so the
                        # information isn't lost.
                        parts.append(f"{txt} ({href})" if txt else href)
                elif txt:
                    parts.append(txt)
            joined = " ".join(p for p in parts if p).strip()
            if joined:
                lines.append(joined)

    if isinstance(cell_val, list):
        for line in cell_val:
            visit_line(line)
    elif isinstance(cell_val, str):
        t = cell_val.strip()
        if t:
            lines.append(t)
    elif cell_val is not None:
        lines.append(str(cell_val))

    return ("\n".join(lines), first_url)


def _cell_row_text_len(row):
    total = 0
    for c in row:
        text, _ = _cell_to_text_and_link(c)
        total += len(text)
    return total


def write_xlsx_all(out_path, databases, query_lines):
    """Write ONE workbook with one sheet per database. Each entry in
    `databases` is a dict: {name, header_lines, columns, rows}.

    - Title row (bold, 14pt) holds the database name.
    - Header lines (italic) below it.
    - The query summary (italic, joined by ' | ') sits one row below.
    - One blank row, then the column header row (bold, light-blue fill).
    - Data rows follow. Multi-line cells use wrap_text. The first hyperlink
      in each cell becomes the cell's clickable link.
    - Column widths are auto-sized (capped) from the longest line in each
      column.
    """
    if not _OPENPYXL_AVAILABLE:
        raise RuntimeError(
            "openpyxl is not installed. Run:\n"
            "    pip install openpyxl"
        )

    wb = Workbook()
    # The default empty sheet — remove it after we add the first real one.
    default_sheet = wb.active
    used_names = set()

    title_font = Font(bold=True, size=14)
    header_font = Font(italic=True, size=10, color="555555")
    columns_font = Font(bold=True, color="FFFFFF")
    columns_fill = PatternFill("solid", fgColor="4472C4")
    link_font = Font(color="0563C1", underline="single")
    wrap_align = Alignment(wrap_text=True, vertical="top")
    centre_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border_side = Side(style="thin", color="BFBFBF")
    grid_border = Border(left=border_side, right=border_side,
                         top=border_side, bottom=border_side)

    saved_any = False
    skipped = []   # ((db_name, reason), …) — surfaced via print() at the end.
    for db in databases:
        name = db.get("name") or "Sheet"
        header_lines = db.get("header_lines") or []
        columns = db.get("columns") or []
        rows = list(db.get("rows") or [])

        # Trim trailing fully-empty rows so the table doesn't end with blanks.
        while rows and _cell_row_text_len(rows[-1]) == 0:
            rows.pop()
        if not rows:
            skipped.append((name, "all rows empty after trim"))
            continue

        # Reduce columns to plain strings — Excel header rows don't need
        # multi-line formatting, just bold text.
        col_texts = []
        for c in columns:
            t, _ = _cell_to_text_and_link(c)
            col_texts.append(t)
        # Detect real vs synthetic ("Col 1", "Col 2" …) headers.
        has_real_cols = bool(col_texts) and not all(
            re.match(r"^Col\s+\d+$", t.strip()) for t in col_texts
        )

        ncols = max(
            len(col_texts),
            max((len(r) for r in rows), default=1),
        )
        if ncols < 1:
            ncols = 1

        sheet_title = safe_sheet_name(name, used_names)
        try:
            ws = wb.create_sheet(title=sheet_title)
        except Exception as exc:
            skipped.append((name, f"create_sheet failed: {exc}"))
            continue

        # Whole-sheet write inside try/except — a single bad cell value
        # (e.g. an unrepresentable Unicode character) shouldn't kill the
        # whole workbook. We drop just that one sheet.
        try:
            current_row = 1
            # 1. Title.
            ws.cell(row=current_row, column=1, value=name).font = title_font
            ws.cell(row=current_row, column=1).alignment = centre_align
            if ncols > 1:
                ws.merge_cells(start_row=current_row, start_column=1,
                               end_row=current_row, end_column=ncols)
            ws.row_dimensions[current_row].height = 22
            current_row += 1

            # 2. Header lines (skip ones identical to the title).
            for line in header_lines:
                if (line or "").strip().lower() == name.strip().lower():
                    continue
                ws.cell(row=current_row, column=1, value=line).font = header_font
                ws.cell(row=current_row, column=1).alignment = centre_align
                if ncols > 1:
                    ws.merge_cells(start_row=current_row, start_column=1,
                                   end_row=current_row, end_column=ncols)
                current_row += 1

            # 3. Query summary.
            if query_lines:
                qline = " | ".join(query_lines)
                ws.cell(row=current_row, column=1, value=qline).font = header_font
                ws.cell(row=current_row, column=1).alignment = centre_align
                if ncols > 1:
                    ws.merge_cells(start_row=current_row, start_column=1,
                                   end_row=current_row, end_column=ncols)
                current_row += 1

            # Blank spacer row.
            current_row += 1

            # 4. Column header row.
            header_row_idx = current_row
            if has_real_cols:
                for j in range(ncols):
                    val = col_texts[j] if j < len(col_texts) else ""
                    cell = ws.cell(row=header_row_idx, column=j + 1, value=val)
                    cell.font = columns_font
                    cell.fill = columns_fill
                    cell.alignment = centre_align
                    cell.border = grid_border
                current_row += 1

            # 5. Data rows.
            data_start = current_row
            max_col_widths = [0] * ncols
            for row in rows:
                for j in range(ncols):
                    raw = row[j] if j < len(row) else ""
                    text, url = _cell_to_text_and_link(raw)
                    cell = ws.cell(row=current_row, column=j + 1, value=text)
                    cell.alignment = wrap_align
                    cell.border = grid_border
                    if url:
                        try:
                            cell.hyperlink = url
                            cell.font = link_font
                        except Exception:
                            pass
                    # Track the widest single line in this column so we can
                    # size columns later (cap at a reasonable max).
                    widest = max((len(s) for s in text.split("\n")), default=0)
                    if widest > max_col_widths[j]:
                        max_col_widths[j] = widest
                current_row += 1

            # 6. Auto-ish column widths. Excel widths are measured in
            # approximate-character units; cap to 60 so a single ugly cell
            # doesn't blow the layout out.
            for j in range(ncols):
                target = max(
                    max_col_widths[j],
                    len(col_texts[j]) if j < len(col_texts) else 0,
                )
                width = min(max(target + 2, 10), 60)
                ws.column_dimensions[get_column_letter(j + 1)].width = width

            # 7. Freeze the top + header rows so they stay visible while
            # scrolling large databases.
            ws.freeze_panes = ws.cell(row=data_start, column=1)
            saved_any = True
        except Exception as exc:
            # One sheet broke — drop it and keep going.
            try:
                wb.remove(ws)
            except Exception:
                pass
            skipped.append((name, f"sheet build failed: {exc}"))
            continue

    if skipped:
        print("  xlsx: skipped sheets:")
        for db_name, reason in skipped:
            print(f"      - {db_name}: {reason}")

    if not saved_any:
        # Nothing to write — bail without producing an empty workbook.
        print("  xlsx: nothing to write — no sheet produced any rows.")
        return False

    # Remove the placeholder empty sheet (it's still around because
    # create_sheet appended new ones at the end).
    try:
        wb.remove(default_sheet)
    except Exception:
        pass

    # Save, with one fallback name in case the .xlsx is open in Excel
    # (PermissionError on Windows; macOS / OneDrive can do this too).
    out_path = Path(out_path)
    save_target = out_path
    try:
        wb.save(str(save_target))
    except PermissionError:
        alt = out_path.with_name(out_path.stem + "_new" + out_path.suffix)
        print(
            f"  xlsx: {out_path.name} is locked (open in Excel?). "
            f"Saving as {alt.name} instead."
        )
        wb.save(str(alt))
        save_target = alt
    except Exception as exc:
        raise RuntimeError(
            f"openpyxl could not save {out_path}: {type(exc).__name__}: {exc}"
        ) from exc

    return str(save_target)


# ---------- form filling (by visible labels) -------------------------------- #
class PaidFeatureError(RuntimeError):
    """Raised when the user tries to use search rows that aren't available
    on their JewishGen plan. Free accounts only see 2 search rows on the
    home form; rows 3 and 4 are a paid feature.

    The GUI catches this and shows a friendly popup with the form intact.
    """


async def fill_search_form(page, rows, country):
    """Each row in the JewishGen home form is: Data Type combo, Search Type
    combo, free-text input. The country combo is the last combo on the page.

    The number of available rows depends on the user's JewishGen subscription:
        - Free: 2 rows  → 4 dropdowns + 1 country = 5 combos, 2 textboxes
        - Paid: 4 rows  → 8 dropdowns + 1 country = 9 combos, 4 textboxes

    Before we touch any combo we measure how many search rows the page
    actually has (= (combo_count - 1) // 2). If the caller asked for more
    than that, we raise PaidFeatureError so the GUI can show a clear popup
    instead of a 30-second TimeoutError.
    """
    combos = page.get_by_role("combobox")
    textboxes = page.get_by_role("textbox")

    total_combos = await combos.count()
    total_textboxes = await textboxes.count()
    # `total_combos` includes the country dropdown at the end, so divide by 2
    # after subtracting it. Clamp at 0 in case the page layout changes.
    available_rows = max(0, (total_combos - 1) // 2)
    # Also bound by the number of text inputs in case the layout changes.
    available_rows = min(available_rows, total_textboxes)

    if len(rows) > available_rows:
        raise PaidFeatureError(
            f"JewishGen only shows {available_rows} search row(s) on this page, "
            f"but you filled in {len(rows)}. Rows 3 and 4 are a paid feature — "
            f"either log in with a paid account, or clear rows beyond row "
            f"{available_rows} and try again."
        )

    async def _safe_select(combo, label_text):
        """select_option by label — with apostrophe-safe fallback.

        Playwright's label= selector uses XPath contains() which can choke on
        labels that include a single-quote (apostrophe).  We catch any error
        and retry by iterating all <option> texts ourselves via JavaScript."""
        try:
            await combo.select_option(label=label_text)
        except Exception:
            # Fallback: find matching option index via JS and select by value.
            try:
                value = await combo.evaluate(
                    """(el, txt) => {
                        for (const o of el.options) {
                            if (o.text.trim() === txt.trim()) return o.value;
                        }
                        // case-insensitive second pass
                        const low = txt.trim().toLowerCase();
                        for (const o of el.options) {
                            if (o.text.trim().toLowerCase() === low) return o.value;
                        }
                        return null;
                    }""",
                    label_text,
                )
                if value is not None:
                    await combo.select_option(value=value)
                else:
                    raise ValueError(
                        f"Option {label_text!r} not found in dropdown"
                    )
            except Exception as exc2:
                raise RuntimeError(
                    f"Could not select option {label_text!r}: {exc2}"
                ) from exc2

    for idx, (data_type, search_type, term) in enumerate(rows):
        await _safe_select(
            combos.nth(idx * 2),
            DATA_TYPE_TO_JG_LABEL.get(data_type, data_type),
        )
        await _safe_select(combos.nth(idx * 2 + 1), search_type)
        # .fill() handles apostrophes correctly — no patch needed there
        clean_term = term.strip()
        await textboxes.nth(idx).fill(clean_term)

    if country and country != "ALL COUNTRIES":
        await _safe_select(combos.last, country)

    await page.get_by_role("button", name=re.compile(r"Search our Collection", re.I)).first.click()
    await page.wait_for_load_state("domcontentloaded")


async def detect_login_form(page):
    """Return the locator for an e-mail-like input if the current page
    looks like JewishGen's login form, else None."""
    try:
        # 1) explicit type=email
        email_inputs = page.locator("input[type='email']")
        if await email_inputs.count() > 0:
            return email_inputs.first
        # 2) input whose accessible name/placeholder/label mentions e-mail
        labelled = page.get_by_label(re.compile(r"e[\-\s]?mail", re.I))
        if await labelled.count() > 0:
            return labelled.first
        ph = page.get_by_placeholder(re.compile(r"e[\-\s]?mail", re.I))
        if await ph.count() > 0:
            return ph.first
        # 3) presence of a password field alongside any textbox => take the
        #    textbox right before it.
        if await page.locator("input[type='password']").count() > 0:
            tb = page.get_by_role("textbox")
            if await tb.count() > 0:
                return tb.first
    except Exception:
        pass
    return None


RESULTS_URL = "https://jewishgen.org/databases/jgform.php"


_AUTH0_HOSTS = ("login.jewishgen.org", "auth.jewishgen.org")


async def _looks_like_login(page):
    """True if the current page seems to be a JewishGen / Auth0 login page."""
    try:
        host = (page.url or "").lower()
        if any(h in host for h in _AUTH0_HOSTS):
            return True
        if await page.locator("input[type='password']").count() > 0:
            return True
    except Exception:
        pass
    return False


async def try_autofill_login(page, email, password, timeout_s=15):
    """Best-effort: if we're on JewishGen / Auth0 login and BOTH credentials
    are provided, try to fill the form and submit. Returns True if we
    successfully left the login page, False otherwise (caller falls back to
    passive wait). Always exception-safe — never raises."""
    if not (email and password):
        return False
    try:
        # Auth0 universal login: email step first, sometimes password on
        # the same form, sometimes on a separate step.
        email_input = None
        for sel in (
            "input[type='email']",
            "input[name='email']",
            "input[name='username']",
            "input[id='username']",
            "input[autocomplete='username']",
        ):
            loc = page.locator(sel).first
            try:
                if await loc.count() > 0:
                    try:
                        if await loc.is_visible(timeout=1500):
                            email_input = loc
                            break
                    except Exception:
                        email_input = loc
                        break
            except Exception:
                continue
        if email_input is None:
            return False
        await email_input.fill(email)

        # Auth0 buttons commonly used: Continue / Next / Log in / Sign in.
        async def _click_label(label_res):
            for pat in label_res:
                btn = page.get_by_role("button", name=pat)
                try:
                    if await btn.count() > 0:
                        await btn.first.click(timeout=5_000)
                        return True
                except Exception:
                    continue
            # Fall back to plain submit button.
            try:
                submit = page.locator("button[type='submit'], input[type='submit']").first
                if await submit.count() > 0:
                    await submit.click(timeout=5_000)
                    return True
            except Exception:
                pass
            return False

        await _click_label([
            re.compile(r"^\s*Continue\s*$", re.I),
            re.compile(r"^\s*Next\s*$", re.I),
            re.compile(r"^\s*Log in\s*$", re.I),
            re.compile(r"^\s*Sign in\s*$", re.I),
        ])

        # Wait for the password field (might already have been on screen).
        pw_input = None
        try:
            pw_input = page.locator("input[type='password']").first
            await pw_input.wait_for(state="visible", timeout=timeout_s * 1000)
        except Exception:
            return False
        await pw_input.fill(password)
        await _click_label([
            re.compile(r"^\s*Log in\s*$", re.I),
            re.compile(r"^\s*Sign in\s*$", re.I),
            re.compile(r"^\s*Continue\s*$", re.I),
            re.compile(r"^\s*Submit\s*$", re.I),
        ])

        try:
            await page.wait_for_function(
                "() => !/login\\.jewishgen\\.org|auth\\.jewishgen\\.org/i.test(location.host)",
                timeout=timeout_s * 1000,
            )
        except Exception:
            # Still on login — maybe MFA / OTP / captcha. Hand back to user.
            return False
        try:
            await page.wait_for_load_state("domcontentloaded", timeout=10_000)
        except Exception:
            pass
        return True
    except Exception:
        return False


async def wait_for_results_page(context, page, email, password=None,
                                timeout_seconds=900):
    """Wait for the browser to land on the JewishGen results URL. If
    email+password are provided AND we hit the login page, try to autofill;
    otherwise pause passively while the user logs in by hand."""
    if "jgform" in (page.url or "").lower():
        return page

    # Give Chromium a beat to land somewhere stable before we sniff.
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=10_000)
    except Exception:
        pass

    if await _looks_like_login(page):
        ok = await try_autofill_login(page, email, password)
        if ok:
            print(">>> Auto-login succeeded.")
        else:
            print(">>> Please log in to JewishGen in the browser window.")
            print(">>> The script will continue when the URL becomes the results page.")
    else:
        # Not on login (yet) — quiet by default; let the URL watcher do its job.
        pass

    try:
        await page.wait_for_url(
            re.compile(r"jgform", re.I),
            timeout=timeout_seconds * 1000,
        )
    except PWTimeout:
        if "jgform" not in (page.url or "").lower():
            print("Navigating to the results page…")
            await page.goto(RESULTS_URL, wait_until="domcontentloaded")
    return page


# ---------- iterate list buttons by their visible label --------------------- #
LIST_BUTTON_RE = re.compile(r"^List\s+[\d,]+\s+record", re.I)


async def collect_list_button_labels(page):
    """Return a list of {label, desc} dicts — one per 'List N records' button.
    `desc` is the human-readable database name taken from the surrounding row
    (e.g. 'Family Finder', 'Belarus Births Database'). It is used both as
    file name and as the section title in the .docx output."""
    return await page.evaluate(
        """() => {
            const items = [];
            for (const el of document.querySelectorAll('input[type=submit], button')) {
                const label = (el.value || el.textContent || '').trim();
                if (!/^List\\s+[\\d,]+\\s+record/i.test(label)) continue;
                let desc = '';
                // First try the previous sibling row(s) for an obvious header
                // (JewishGen often puts a logo + DB name above a group of
                // related 'List N records' rows). Otherwise fall back to the
                // first cell of the row that contains the button.
                let row = el.closest('tr');
                if (row && row.cells && row.cells[0]) {
                    desc = row.cells[0].textContent.replace(/\\s+/g, ' ').trim();
                }
                if (!desc || desc.length < 3) {
                    // Look back for a row that contains a heading-like label.
                    let prev = row ? row.previousElementSibling : null;
                    while (prev) {
                        const t = prev.textContent.replace(/\\s+/g, ' ').trim();
                        if (t && t.length < 80 && /jewish|belarus|austria|poland|ukraine|romania|lithuania|latvia|israel|cemetery|gen|family/i.test(t)) {
                            desc = t; break;
                        }
                        prev = prev.previousElementSibling;
                    }
                }
                items.push({ label, desc: desc.slice(0, 200) });
            }
            return items;
        }"""
    )


DONATE_BUTTON_RE = re.compile(r"Maybe Later|Continue to Site", re.I)
DISCUSS_BUTTON_RE = re.compile(r"^\s*No Thanks\s*$", re.I)


def _is_popup_url(url):
    u = (url or "").lower()
    return ("/cure/" in u) or ("menu.asp" in u) or ("jgdg" in u)


async def dismiss_popups(page):
    """Click 'Maybe Later, Continue to Site' / 'No Thanks' — but ONLY on
    JewishGen's donation / discussion-group popup URLs. On any other page
    this is a no-op, so we never accidentally click a button somewhere else
    on the site that happens to share that text."""
    for _ in range(3):
        if not _is_popup_url(page.url):
            return
        clicked = False
        for pattern in (DONATE_BUTTON_RE, DISCUSS_BUTTON_RE):
            btn = page.get_by_role("button", name=pattern)
            if await btn.count() == 0:
                btn = page.get_by_role("link", name=pattern)
            if await btn.count() > 0:
                try:
                    await btn.first.click()
                    await page.wait_for_load_state("domcontentloaded", timeout=10_000)
                    clicked = True
                    break
                except Exception:
                    pass
        if not clicked:
            break


# run_scraper() sets these so the per-popup login waiter can try the same
# autofill that wait_for_results_page does. Restored to None in the
# `finally` block of run_scraper. Treated as best-effort hints — every read
# tolerates them being None.
_CURRENT_CREDS = {"email": None, "password": None}


async def wait_for_login_to_finish(page, timeout_seconds=600):
    """If the page is on login.jewishgen.org (Auth0 universal login), try
    an optional autofill (when credentials were provided to run_scraper);
    otherwise wait passively until the URL leaves that host."""
    if "login.jewishgen.org" not in (page.url or "").lower():
        return

    email = _CURRENT_CREDS.get("email")
    password = _CURRENT_CREDS.get("password")
    if email and password:
        ok = await try_autofill_login(page, email, password)
        if ok:
            print("    >> Auto-login succeeded on popup login page.")
            return

    print("    !! Login required. Please log in in the newly-opened tab.")
    print("    !! The script is waiting (it will NOT close the tab).")
    try:
        await page.wait_for_function(
            "() => !location.host.includes('login.jewishgen.org')",
            timeout=timeout_seconds * 1000,
        )
    except PWTimeout:
        print(f"    !! Still on login after {timeout_seconds}s — continuing anyway.")
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=15_000)
    except PWTimeout:
        pass


async def click_list_button_open_new_page(context, results_page, label, attempts=4):
    """Click a 'List N records' button by its exact visible label. If the
    opened tab is the Auth0 login page, pause until the user has logged in
    (the tab is NOT closed). Dismiss any donation / discussion popups that
    appear after that. Retries on 503 / server error pages."""
    for attempt in range(attempts):
        button = results_page.get_by_role("button", name=label, exact=True).first
        try:
            async with context.expect_page(timeout=10_000) as new_page_event:
                await button.click()
            new_page = await new_page_event.value
            await new_page.wait_for_load_state("domcontentloaded", timeout=15_000)
        except PWTimeout:
            # Same-tab navigation — back-button later.
            await results_page.wait_for_load_state("domcontentloaded", timeout=15_000)
            new_page = results_page

        await wait_for_login_to_finish(new_page)
        await dismiss_popups(new_page)

        # Check for 503 / server error page.
        title = await new_page.title()
        if re.search(r"503|service\s+unavailable|bad\s+gateway|error\s+page", title, re.I):
            wait_s = 10 * (2 ** attempt)
            print(f"    503/error page on click (attempt {attempt+1}/{attempts}), "
                  f"waiting {wait_s}s before retry…")
            if new_page is not results_page:
                try:
                    await new_page.close()
                except Exception:
                    pass
            await asyncio.sleep(wait_s)
            # Reload the results page in case it also got stale.
            try:
                await results_page.reload(wait_until="domcontentloaded", timeout=30_000)
            except Exception:
                pass
            continue

        return new_page

    raise RuntimeError(
        f"Could not open '{label}' after {attempts} attempts — server kept returning errors."
    )


# ---------- scraping per database page -------------------------------------- #

# JS helpers injected separately to keep each piece small and testable.
# Playwright serialises them as strings — keeping them short reduces the
# risk of execution-context loss after long sessions.

_JS_UTILS = r"""
window.__jg = window.__jg || {};

__jg.clean = function(s) {
    return (s || '')
        .replace(/[\u00A0\u2000-\u200B\u2028\u2029]/g, ' ')
        .replace(/[ \t]+/g, ' ');
};

__jg.linesByDom = function(el) {
    var LINE = '\u0001';
    var clone = el.cloneNode(true);
    clone.querySelectorAll('br').forEach(function(b) {
        b.replaceWith(document.createTextNode(LINE));
    });
    clone.querySelectorAll(
        'div,p,li,tr,table,thead,tbody,tfoot,' +
        'h1,h2,h3,h4,h5,h6,header,footer,section,' +
        'article,blockquote,pre,dl,dt,dd,ol,ul,' +
        'fieldset,legend,form,address'
    ).forEach(function(b) {
        b.insertAdjacentText('beforebegin', LINE);
        b.insertAdjacentText('afterend', LINE);
    });
    var text = clone.textContent.replace(/[\u00A0\u2000-\u200B\u2028\u2029]/g, ' ');
    return text.split(LINE)
        .map(function(s) { return s.replace(/[ \t\r\n]+/g, ' ').trim(); })
        .filter(Boolean);
};

__jg.getCellSegments = function(el) {
    var raw = (el.innerText !== undefined ? el.innerText : el.textContent) || '';
    var lines = raw.split(/\r?\n/).map(function(l) {
        return __jg.clean(l).trim();
    }).filter(Boolean);
    var domLines = __jg.linesByDom(el);
    if (domLines.length > lines.length) lines = domLines;

    var links = [];
    el.querySelectorAll('a[href]').forEach(function(a) {
        var href = a.href || a.getAttribute('href');
        if (!href || /^javascript:|^#/i.test(href)) return;
        var t = __jg.clean(
            (a.innerText !== undefined ? a.innerText : a.textContent)
        ).trim();
        if (t) links.push({t: t, h: href});
    });

    var usedLinkIdx = new Set();
    return lines.map(function(line) {
        var segs = [];
        var remaining = line;
        while (remaining.length > 0) {
            var bestIdx = remaining.length;
            var bestLink = null;
            var bestLinkPos = -1;
            for (var i = 0; i < links.length; i++) {
                if (usedLinkIdx.has(i)) continue;
                var lt = links[i].t;
                if (!lt) continue;
                var idx = remaining.indexOf(lt);
                if (idx >= 0 && idx < bestIdx) {
                    bestIdx = idx; bestLink = links[i]; bestLinkPos = i;
                }
            }
            if (bestLink) {
                if (bestIdx > 0) segs.push({t: remaining.slice(0, bestIdx), h: null});
                segs.push({t: bestLink.t, h: bestLink.h});
                usedLinkIdx.add(bestLinkPos);
                remaining = remaining.slice(bestIdx + bestLink.t.length);
            } else {
                segs.push({t: remaining, h: null});
                remaining = '';
            }
        }
        return segs.filter(function(s) { return s.t || s.h; });
    }).filter(function(l) { return l.length > 0; });
};

__jg.getLines = function(el) {
    return __jg.getCellSegments(el).map(function(line) {
        return line.map(function(s) { return s.t; }).join('');
    });
};

__jg.isLeafRow = function(tr) {
    for (var c of tr.children) {
        if (c.querySelector && c.querySelector('table')) return false;
    }
    return true;
};

__jg.gridifyByTable = function(trs) {
    var tables = new Map();
    trs.forEach(function(tr) {
        var t = tr.closest('table');
        if (!tables.has(t)) tables.set(t, []);
        tables.get(t).push(tr);
    });
    var out = [];
    tables.forEach(function(ts) {
        var covered = [];
        ts.forEach(function(tr) {
            var kids = Array.from(tr.children);
            var row = [];
            var i = 0, col = 0;
            while (i < kids.length || (col < covered.length && covered[col] > 0)) {
                if (col < covered.length && covered[col] > 0) {
                    row.push(null); covered[col]--; col++;
                } else if (i < kids.length) {
                    var c = kids[i];
                    var rs = parseInt(c.getAttribute('rowspan') || '1', 10) || 1;
                    var cs = parseInt(c.getAttribute('colspan') || '1', 10) || 1;
                    while (covered.length < col + cs) covered.push(0);
                    for (var k = 0; k < cs; k++) {
                        row.push(c);
                        if (rs > 1) covered[col + k] = rs - 1;
                    }
                    col += cs; i++;
                } else { break; }
            }
            out.push({tr: tr, cells: row});
        });
    });
    return out;
};
"""

_JS_EXTRACT = r"""
(args) => {
    var keywords = args.keywords;
    var mode     = args.mode || 'OR';   // 'OR' or 'AND'

    var escape = function(s) { return s.replace(/[.*+?^${}()|[\]\\]/g, '\\$&'); };
    var regexes = keywords.map(function(k) { return new RegExp(escape(k), 'i'); });

    var jg = window.__jg;

    // ── header lines ──────────────────────────────────────────────────── //
    var headerLines = [];
    var h = document.querySelector('h2, h1');
    if (h) {
        var t = jg.getLines(h).join(' ').trim();
        if (t) headerLines.push(t);
    }
    document.querySelectorAll('center, p, h3, h4').forEach(function(c) {
        var lines = jg.getLines(c);
        var joined = lines.join(' ');
        if (!joined || joined.length > 600) return;
        if (!/Searching for|Number of (hits|matches)|matching records|total matches|records found/i.test(joined)) return;
        lines.forEach(function(t) {
            if (t && t.length < 240 && !headerLines.includes(t)) headerLines.push(t);
        });
    });

    // ── leaf rows ─────────────────────────────────────────────────────── //
    var leafTrs = Array.from(document.querySelectorAll('tr')).filter(jg.isLeafRow);
    var grid = jg.gridifyByTable(leafTrs);

    var trToCells = new Map();
    grid.forEach(function(g) { trToCells.set(g.tr, g.cells); });

    // ── column header row ─────────────────────────────────────────────── //
    var columns = [];
    var thRow = null, thCount = 0;
    grid.forEach(function(g) {
        var n = 0;
        g.cells.forEach(function(c) { if (c && c.tagName === 'TH') n++; });
        if (n > thCount) { thCount = n; thRow = g.tr; }
    });
    var debug = {columnsHTML: [], firstRowHTML: []};
    if (thRow) {
        var cells = trToCells.get(thRow) || [];
        columns = cells.map(function(c) { return c ? jg.getCellSegments(c) : []; });
        debug.columnsHTML = cells.slice(0, 12).map(function(c) {
            return c ? c.outerHTML.slice(0, 500) : '<covered/>';
        });
    }

    // ── noise filter ──────────────────────────────────────────────────── //
    var NOISE_RE = /Edmond\s+J\.?\s+Safra|36\s+Battery\s+Place|info@jewishgen\.org|JewishGen\s+is\s+actively\s+raising|actively\s+raising\s+funds|Mission\s+Our\s+Team\s+Get\s+Involved|Learn\s+more\s+about\s+our\s+org|©\s*JewishGen|All\s+Rights\s+Reserved.*JewishGen|Sign\s+In\s*\/\s*Register/i;

    function gridRowText(cells) {
        var t = '';
        cells.forEach(function(c) { if (c) t += ' ' + (c.textContent || ''); });
        return t;
    }
    function isRepeatedCells(cells) {
        var texts = cells.filter(function(c) { return c !== null; })
            .map(function(c) { return (c.textContent || '').replace(/\s+/g, ' ').trim(); });
        if (texts.length < 2) return false;
        var first = texts[0];
        return first.length > 10 && texts.every(function(t) { return t === first; });
    }
    function rowHasContent(cells) {
        return cells.some(function(c) {
            return c && (c.textContent || '').replace(/\s+/g, '').length > 0;
        });
    }
    function gridCellText(c) {
        if (!c) return '';
        return (c.textContent || '').replace(/\s+/g, ' ').trim();
    }

    var dataGrid = grid
        .filter(function(g) { return g.tr !== thRow && g.cells.length > 0; })
        .filter(function(g) { return !NOISE_RE.test(gridRowText(g.cells)); })
        .filter(function(g) { return !isRepeatedCells(g.cells); });

    // ── STRICT per-row matching ───────────────────────────────────────── //
    // EVERY output row must itself satisfy the filter — a row that does not
    // match is never included (no family/group can drag a non-matching row in;
    // that was the leak). Each row also INHERITS the values of cells that a
    // rowspan above spans into it, so a person listed under a rowspanned
    // Town/Year column is still judged WITH that town/year. Net effect: a word
    // filter is strict — e.g. people filed under Odessa are dropped by a
    // «Bendery» filter, while everyone genuinely under a rowspanned «Bendery»
    // cell is kept. With no keywords nothing is kept (the GUI requires one).
    var seenKeys = new Set();
    var matched = [];
    var inherited = [];   // per-column text currently spanning down from above

    dataGrid.forEach(function(g) {
        if (!rowHasContent(g.cells)) return;

        // Full logical text of THIS row = own cells + inherited (covered) cells.
        var fullParts = [];
        for (var ci = 0; ci < g.cells.length; ci++) {
            var c = g.cells[ci];
            if (c) {
                var ctext = (c.textContent || '');
                fullParts.push(ctext);
                inherited[ci] = ctext;            // this cell spans down from here
            } else if (inherited[ci]) {
                fullParts.push(inherited[ci]);    // covered by a rowspan above
            }
        }
        var rowText = fullParts.join(' ');

        if (!regexes.length) return;
        var hit = mode === 'AND'
            ? regexes.every(function(re) { return re.test(rowText); })
            : regexes.some(function(re) { return re.test(rowText); });
        if (!hit) return;

        var cellsSegs = g.cells.map(function(c) {
            return c ? jg.getCellSegments(c) : [];
        });
        var total = 0;
        cellsSegs.forEach(function(cell) {
            cell.forEach(function(line) {
                line.forEach(function(s) { total += (s.t || '').length; });
            });
        });
        // Drop single-cell noise rows
        if (cellsSegs.length === 1 && total > 200) return;
        var rowFlat = '';
        cellsSegs.forEach(function(cell) {
            cell.forEach(function(line) {
                line.forEach(function(s) { rowFlat += ' ' + (s.t || ''); });
            });
        });
        if (NOISE_RE.test(rowFlat)) return;

        // JS-side deduplication by row content key
        var rowKey = rowFlat.replace(/\s+/g, ' ').trim().slice(0, 500);
        if (seenKeys.has(rowKey)) return;
        seenKeys.add(rowKey);

        if (!matched.length) {
            debug.firstRowHTML = g.cells.slice(0, 12).map(function(c) {
                return c ? c.outerHTML.slice(0, 500) : '<covered/>';
            });
        }
        matched.push(cellsSegs);
    });

    if (!columns.length && matched.length) {
        var counts = {};
        matched.forEach(function(r) { counts[r.length] = (counts[r.length] || 0) + 1; });
        var bestN = 0, bestC = 0;
        Object.entries(counts).forEach(function(e) {
            var nn = parseInt(e[0], 10);
            if (e[1] > bestC) { bestC = e[1]; bestN = nn; }
        });
        columns = Array.from({length: bestN}, function(_, i) {
            return [[{t: 'Col ' + (i + 1), h: null}]];
        });
    }

    return {
        headerLines: headerLines,
        columns: columns,
        rows: matched,
        debug: debug,
        totalLeafRows: leafTrs.length,
        totalDataRows: dataGrid.length,
    };
}
"""


async def _inject_utils(page):
    """Inject __jg utilities into the page. Safe to call multiple times —
    the guard `window.__jg = window.__jg || {}` is a no-op on re-injection."""
    try:
        await page.evaluate(_JS_UTILS)
    except Exception as exc:
        # Execution context can die mid-navigation; log but don't crash.
        print(f"    [warn] could not inject JS utils: {exc}")


async def extract_matches(page, keywords, keyword_mode="OR"):
    """Walk every <tr>, return rows matching the keyword filter.

    keyword_mode: 'OR'  — row kept when ANY keyword matches (default)
                  'AND' — row kept only when ALL keywords match
    """
    if isinstance(keywords, str):
        keywords = [keywords]
    keywords = [k for k in (keywords or []) if (k or "").strip()]
    if not keywords:
        keywords = ["\x00__no_match__\x00"]

    await _inject_utils(page)
    return await page.evaluate(
        _JS_EXTRACT,
        {"keywords": keywords, "mode": keyword_mode},
    )


async def follow_next_pages(page, keywords, acc, keyword_mode="OR",
                           pause=2.5, max_pages=100):
    """Walk pagination collecting matching rows into *acc*.

    Fixes applied:
    - seen_page_sigs  : stops if the same page signature appears twice
    - Python-side dedup: rows already in acc are not added again
    - Page\s+\d+ removed from next-button detection (caused cyclic clicks)
    - ensure_not_503 after every click + retry loop
    - acc["rows"] is cleared of the heavy JS objects after each page to
      limit memory; only the Python-level fingerprint set grows.
    """
    SIG_JS = (
        "(() => {"
        "  var best = null, bestRows = 0;"
        "  document.querySelectorAll('table').forEach(function(t) {"
        "    var rc = t.querySelectorAll('tr').length;"
        "    if (rc > bestRows) { bestRows = rc; best = t; }"
        "  });"
        "  if (!best) return '';"
        "  var tds = Array.from(best.querySelectorAll('td')).slice(0, 12);"
        "  return tds.map(function(t) {"
        "    return (t.textContent || '').replace(/\\s+/g, ' ').trim().slice(0, 60);"
        "  }).join('|');"
        "})()"
    )

    def _row_fingerprint(row):
        parts = []
        for cell in row:
            for line in cell:
                for seg in line:
                    t = (seg.get("t") or "").strip()
                    if t:
                        parts.append(t)
        return "|".join(parts)

    seen_page_sigs   = set()
    seen_fingerprints = set()
    page_idx = 0

    while page_idx < max_pages:
        page_idx += 1
        # Human-like pause with jitter — reduces 524 rate-limit triggers.
        jitter = random.uniform(0.5, 2.0)
        await asyncio.sleep(pause + jitter)

        # ── 503 guard ───────────────────────────────────────────────── #
        try:
            await ensure_not_503(page)
        except TemporaryServerError as exc:
            print(f"    [page {page_idx}] {exc} — stopping pagination.")
            break

        # ── page signature ──────────────────────────────────────────── #
        current_sig = await page.evaluate(SIG_JS)
        if current_sig and current_sig in seen_page_sigs:
            print(f"    [page {page_idx}] duplicate page signature — loop detected, stopping.")
            break
        if current_sig:
            seen_page_sigs.add(current_sig)

        # ── extract & inject utils ──────────────────────────────────── #
        await _inject_utils(page)
        result = await extract_matches(page, keywords, keyword_mode)
        print(f"    [page {page_idx}] scanned: "
              f"{result.get('totalLeafRows','?')} leaf rows, "
              f"{result.get('totalDataRows','?')} data rows, "
              f"{len(result['rows'])} matches")

        if result["rows"]:
            if not acc["headerLines"]:
                acc["headerLines"] = result["headerLines"]
                acc["columns"]     = result["columns"]
                dbg = result.get("debug") or {}
                if dbg.get("columnsHTML"):
                    print("    --- DEBUG: first <th> ---")
                    for ht in dbg["columnsHTML"][:2]:
                        print("      " + ht.replace("\n", " "))

            new_rows = 0
            for row in result["rows"]:
                fp = _row_fingerprint(row)
                if fp and fp in seen_fingerprints:
                    continue
                seen_fingerprints.add(fp)
                acc["rows"].append(row)
                new_rows += 1
            if new_rows < len(result["rows"]):
                print(f"    [page {page_idx}] deduped: "
                      f"kept {new_rows}/{len(result['rows'])}")

        # ── find next button — NO /^Page\s+\d+/ ─────────────────────── #
        next_info = await page.evaluate(
            """() => {
                for (var el of document.querySelectorAll('[data-pw-next]'))
                    el.removeAttribute('data-pw-next');

                var candidates = Array.from(document.querySelectorAll(
                    'a, button, input[type=submit], input[type=button]'
                ));
                for (var el of candidates) {
                    var text = (el.value || el.innerText || el.textContent || '')
                        .replace(/\\s+/g, ' ').trim();
                    if (!text || text.length > 40) continue;
                    if (/previous|prev\\b|^back\\b|^first\\b/i.test(text)) continue;
                    if (
                        /^next\\b/i.test(text) ||
                        /^(>>+|»+|→+)$/.test(text)
                    ) {
                        el.setAttribute('data-pw-next', '1');
                        return {text: text, tag: el.tagName};
                    }
                }
                return null;
            }"""
        )
        if not next_info:
            break

        print(f"    -> clicking next: {next_info['text']!r}")
        next_locator = page.locator('[data-pw-next="1"]').first
        try:
            await next_locator.scroll_into_view_if_needed(timeout=5_000)
        except Exception:
            pass

        # ── click + 503 retry ───────────────────────────────────────── #
        clicked = False
        for retry in range(5):
            try:
                await next_locator.click(timeout=10_000)
            except Exception as exc:
                print(f"    -> click failed ({exc}); stopping.")
                return

            for _ in range(3):
                try:
                    await page.wait_for_load_state("domcontentloaded", timeout=15_000)
                    break
                except Exception:
                    pass

            try:
                await ensure_not_503(page)
                clicked = True
                break
            except TemporaryServerError:
                wait_s = 5 * (retry + 1)
                print(f"    -> 503 after next click (retry {retry+1}/5), "
                      f"waiting {wait_s}s…")
                await asyncio.sleep(wait_s)
                try:
                    await page.reload(wait_until="domcontentloaded", timeout=30_000)
                except Exception:
                    pass

        if not clicked:
            print("    -> server kept returning 503 — stopping pagination.")
            break

        # ── wait for content change ──────────────────────────────────── #
        wait_js = (
            "(before) => {"
            "  var best = null, bestRows = 0;"
            "  document.querySelectorAll('table').forEach(function(t) {"
            "    var rc = t.querySelectorAll('tr').length;"
            "    if (rc > bestRows) { bestRows = rc; best = t; }"
            "  });"
            "  if (!best) return false;"
            "  var tds = Array.from(best.querySelectorAll('td')).slice(0, 12);"
            "  var sig = tds.map(function(t) {"
            "    return (t.textContent || '').replace(/\\s+/g,' ').trim().slice(0,60);"
            "  }).join('|');"
            "  return sig && sig !== before;"
            "}"
        )
        try:
            await page.wait_for_function(wait_js, arg=current_sig, timeout=20_000)
        except PWTimeout:
            print("    -> page didn't change — stopping pagination.")
            break
        except Exception as exc:
            print(f"    -> wait error ({exc}); stopping.")
            break


# ─────────────────────────────────────────────────────────────────────────────
# IMAGE DOWNLOADS — triggered by links in JewishGen result rows
# ─────────────────────────────────────────────────────────────────────────────

def _get_fs_credentials():
    """Read FamilySearch login from the FamilySearch GUI autosave file."""
    try:
        p = HERE / "gui" / ".fs_autosave.json"
        if p.exists():
            d = json.loads(p.read_text("utf-8"))
            return d.get("username", ""), d.get("password", "")
    except Exception:
        pass
    return "", ""


def _safe_part(s, maxlen=40):
    """Clean a string for use in a filename / dir name."""
    s = re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip())
    s = re.sub(r"\s+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s[:maxlen]


def _query_prefix(rows):
    """Build a short prefix like 'Sanders_Alexander' from search rows."""
    parts = [_safe_part(term) for _, _, term in (rows or []) if term.strip()]
    return "_".join(p for p in parts[:3] if p)[:50] or "search"


def _row_label(row, max_len=80):
    """
    Build a human-readable filename from the first few cells of a row.
    Grabs the first meaningful text from each of the first 5 columns.
    """
    parts = []
    for cell in (row or [])[:5]:
        for line in (cell or []):
            for seg in (line or []):
                if isinstance(seg, dict):
                    t = (seg.get("t") or "").strip()
                    # skip URLs and very long fragments
                    if t and not t.startswith("http") and len(t) < 60:
                        parts.append(_safe_part(t, 30))
                        break
            else:
                continue
            break
    label = "_".join(p for p in parts if p)
    label = re.sub(r"_+", "_", label).strip("_")
    return label[:max_len] or "row"


def _cell_links(cell):
    """Yield (text, href) for every link segment in a cell."""
    for line in (cell or []):
        for seg in (line or []):
            if isinstance(seg, dict):
                t = (seg.get("t") or "").strip()
                h = seg.get("h")
                if h:
                    yield t, h


def _row_image_tasks(row):
    """
    Analyse a result row and return image-download tasks.

    Each task is a dict:
      type      "fhl"     FamilySearch microfilm — open viewer, go to img_num, download
                "direct"  click link → image (bracketed / plain number, or "Image" word)
                "ukraine" slash number — DON'T click, link already in docx/xlsx
      url       the href
      img_num   frame number for FHL (int or None)
    """
    if not row:
        return []

    tasks = []

    # ── JOWBR cemetery record → the gravestone PHOTO lives on its detail page
    # (…/cemetery/jowbr.php?rec=…), which the row's name links to. That link is
    # neither a number nor an «Image» word, so it was ignored and the photo was
    # never saved. Route it through the «direct» handler, which opens the page
    # and grabs the largest image (the tombstone). ──
    # Also: ANY link whose href OR text ends in an image extension (e.g. a grave
    # photo shown as «IMG_5963.JPG»). The «direct» handler opens it and saves it.
    _img_ext = re.compile(r"\.(jpe?g|png|gif|tiff?)(\?|$)", re.I)
    _seen_jowbr = set()
    for cell in (row or []):
        for line in (cell or []):
            for seg in (line or []):
                if not isinstance(seg, dict):
                    continue
                h = (seg.get("h") or "")
                if not h:
                    continue
                hl = h.lower()
                t = (seg.get("t") or "")
                if h in _seen_jowbr:
                    continue
                if ("jowbr.php" in hl or "/cemetery/" in hl
                        or _img_ext.search(hl) or _img_ext.search(t)):
                    _seen_jowbr.add(h)
                    tasks.append({"type": "direct", "url": h, "img_num": None})

    # ── second-to-last column: look for clickable word "Image" ────── #
    if len(row) >= 2:
        for text, href in _cell_links(row[-2]):
            if text.lower() == "image":
                tasks.append({"type": "direct", "url": href, "img_num": None})

    # ── last column ───────────────────────────────────────────────── #
    last_cell = row[-1] if row else []
    all_lines = list(last_cell or [])

    for li, line in enumerate(all_lines):
        for seg in (line or []):
            if not isinstance(seg, dict):
                continue
            text = (seg.get("t") or "").strip()
            href = seg.get("h")
            if not href:
                continue

            # FHL microfilm link (text starts with "FHL")
            if text.upper().startswith("FHL"):
                img_num = None
                for lj in range(li + 1, len(all_lines)):
                    for s2 in (all_lines[lj] or []):
                        if isinstance(s2, dict):
                            n = (s2.get("t") or "").strip().strip("[]").strip()
                            if n.isdigit():
                                img_num = int(n)
                                break
                    if img_num is not None:
                        break
                tasks.append({"type": "fhl", "url": href, "img_num": img_num})
                break

            # Ukrainian slash reference — save link only, do NOT open
            if re.search(r"\d+\s*/\s*\d+", text):
                tasks.append({"type": "ukraine", "url": href, "img_num": None})
                break

            # Plain or bracketed number
            clean_num = text.strip("[] ").replace(" ", "")
            if re.match(r"^\d+$", clean_num):
                num = int(clean_num)
                # Large numbers (6+ digits) are archive/microfilm refs —
                # save link only, do NOT open.
                # This also catches the multi-line pattern (2423959 / 3
                # on separate lines) because the main number is large.
                if num >= 100_000:
                    tasks.append({"type": "ukraine", "url": href, "img_num": None})
                else:
                    tasks.append({"type": "direct", "url": href, "img_num": None})
                break

    return tasks


# JS: src of the largest content <img> on the page (skip chrome/logos/thumbs).
_BEST_IMG_JS = r"""() => {
    const SKIP = ['icon','logo','sprite','avatar','pixel','placeholder','.svg',
                  'fscdn.org','thumb'];
    let best = '', area = 0;
    for (const im of document.querySelectorAll('img[src]')) {
        const src = (im.src || '').trim();
        if (!/^https?:/i.test(src)) continue;
        const low = src.toLowerCase();
        if (SKIP.some(b => low.includes(b))) continue;
        const a = (im.naturalWidth || 0) * (im.naturalHeight || 0);
        if (a > area) { area = a; best = src; }
    }
    return best;
}"""

# JS: every *.jpg/png/gif/tif the page LINKS to — via href OR via the link's text
# (JOWBR grave photos appear as a link whose text is «IMG_5963.JPG»). De-duped.
_IMG_LINKS_JS = r"""() => {
    const ext = /\.(jpe?g|png|gif|tiff?)(\?|$)/i;
    const bad = /\.svg|sprite|logo|icon|thumb/i;
    const out = [], seen = new Set();
    for (const a of document.querySelectorAll('a[href]')) {
        const h = a.href || '';
        const t = (a.textContent || '').trim();
        let u = '';
        if (ext.test(h)) u = h;
        else if (ext.test(t)) u = h;     // text «IMG_5963.JPG» → file is in href
        if (u && /^https?:/i.test(u) && !bad.test(u) && !seen.has(u)) {
            seen.add(u); out.push(u);
        }
    }
    return out;
}"""


async def _save_image_url(context, url, dest, log) -> bool:
    """Open an image URL in a throwaway tab; if it IS an image save its body, else
    save the largest image on the page it opened. Returns True on success."""
    p = await context.new_page()
    try:
        r = await p.goto(url, timeout=20000, wait_until="domcontentloaded")
        ct = (r.headers.get("content-type", "") if r else "")
        if r and r.ok and "image" in ct:
            body = await r.body()
            if len(body) > 3000:
                dest.write_bytes(body)
                log(f"    Сохранено: {dest.name} ({len(body)//1024}KB)")
                return True
        await asyncio.sleep(1)
        best = await p.evaluate(_BEST_IMG_JS)
        if best:
            r2 = await context.request.get(best, timeout=15000)
            if r2.ok:
                body = await r2.body()
                if len(body) > 3000:
                    dest.write_bytes(body)
                    log(f"    Сохранено: {dest.name} ({len(body)//1024}KB)")
                    return True
    except Exception:
        pass
    finally:
        try:
            await p.close()
        except Exception:
            pass
    return False


async def _do_direct_download(context, url, dest, log):
    """Open url and save the image(s) it leads to:
      1) the URL itself if it is an image;
      2) otherwise EVERY *.jpg/png/gif the page LINKS to (e.g. JOWBR grave photos
         «IMG_5963.JPG») — one file each (dest, dest_2, …);
      3) failing that, the largest image on the page."""
    dest = Path(dest)
    dest.parent.mkdir(parents=True, exist_ok=True)
    page = await context.new_page()
    saved_paths = []

    def _nth_dest():
        n = len(saved_paths)
        return dest if n == 0 else dest.with_name(f"{dest.stem}_{n+1}{dest.suffix}")

    try:
        resp = await page.goto(url, wait_until="domcontentloaded", timeout=30000)
        await asyncio.sleep(2)

        # 1) the URL itself is an image
        ct = (resp.headers.get("content-type", "") if resp else "")
        if resp and "image" in ct:
            body = await resp.body()
            if len(body) > 1000:
                dest.write_bytes(body)
                log(f"    Сохранено: {dest.name} ({len(body)//1024}KB)")
                return str(dest)

        # 2) every *.jpg/png the page links to (grave photos etc.)
        try:
            links = await page.evaluate(_IMG_LINKS_JS)
        except Exception:
            links = []
        for u in links[:12]:
            d = _nth_dest()
            if await _save_image_url(context, u, d, log):
                saved_paths.append(str(d))
        if saved_paths:
            return saved_paths[0]

        # 3) fallback: the largest image on the page itself
        best = await page.evaluate(_BEST_IMG_JS)
        if best and await _save_image_url(context, best, dest, log):
            return str(dest)

        log(f"    !! direct: изображение не найдено ({url[:60]})")
        return None
    except Exception as e:
        log(f"    !! direct download error: {e}")
        return None
    finally:
        try:
            await page.close()
        except Exception:
            pass


async def _do_fhl_download(fs_url, img_num, dest, fs_email, fs_password, log):
    """
    Open FamilySearch microfilm viewer in a DEDICATED browser.

    KEY INSIGHT: the URL parameter ?i=N selects image N+1 (0-indexed).
    So for Снимок 61 we navigate to ?i=60 — the viewer opens with that
    thumbnail already highlighted in blue. No need to type in the input.

    After the page loads (with login if needed), the thumbnail at position
    img_num-1 is pre-selected. We click it → full-size viewer opens →
    download JPG.
    """
    from playwright.async_api import async_playwright as _apw

    dest = Path(dest)
    dest.parent.mkdir(parents=True, exist_ok=True)

    # Build target URL with correct frame index (?i= is 0-based)
    base_url = fs_url.split("?")[0]
    if img_num and img_num >= 1:
        target_url = f"{base_url}?i={img_num - 1}"
    else:
        target_url = fs_url
    log(f"    FHL: целевой URL {target_url}")

    async with _apw() as pw:
        # Exact same browser setup as familysearch_scraper.py
        browser = await pw.chromium.launch(
            headless=False,
            args=["--start-maximized",
                  "--disable-blink-features=AutomationControlled"],
        )
        ctx = await browser.new_context(no_viewport=True, accept_downloads=True)
        page = await ctx.new_page()
        downloaded = None

        try:
            await page.goto(target_url, wait_until="domcontentloaded", timeout=45000)

            # ── Wait up to 30s for login form OR film grid ────────── #
            log("    FHL: жду загрузки (до 30с)...")
            for _t in range(30):
                await asyncio.sleep(1)
                if await page.locator("#userName").count():
                    log(f"    FHL: форма логина ({_t+1}с)")
                    break
                # Film grid is loaded when thumbnails are visible
                if await page.locator("img[src]").count() > 5:
                    log(f"    FHL: сетка загружена ({_t+1}с)")
                    break
            else:
                log("    FHL: страница не загрузилась за 30с")

            # ── Login if needed ───────────────────────────────────── #
            pt = (await page.title()).lower()
            is_login = ("login" in page.url or "sign in" in pt
                        or await page.locator("#userName").count() > 0)
            if is_login:
                if not fs_email or not fs_password:
                    log("    !! FHL: нет FamilySearch credentials — пропуск")
                    return None
                log("    FHL: логин в FamilySearch...")
                try:
                    await page.locator("#userName").wait_for(state="visible", timeout=20000)
                    await asyncio.sleep(2)
                    for _ in range(3):
                        await page.fill("#userName", fs_email)
                        await asyncio.sleep(0.5)
                        if (await page.locator("#userName").input_value()).strip():
                            break
                    for _ in range(3):
                        await page.fill("#password", fs_password)
                        await asyncio.sleep(0.5)
                        if (await page.locator("#password").input_value()).strip():
                            break
                    u_ok = (await page.locator("#userName").input_value()).strip()
                    p_ok = (await page.locator("#password").input_value()).strip()
                    log(f"    Проверка: user={'OK' if u_ok else '!ПУСТО'}, "
                        f"pass={'OK' if p_ok else '!ПУСТО'}")
                    await page.click("#login")
                    # After login FamilySearch redirects back to target_url via state= param
                    await page.wait_for_url(
                        lambda u: "familysearch.org" in u and "login" not in u,
                        timeout=30000)
                    log(f"    FHL: логин OK ✓  URL: {page.url[:60]}")
                    # If the redirect didn't land on target_url, navigate there manually
                    if f"?i={img_num - 1}" not in page.url:
                        log("    FHL: навигируюсь на целевой снимок...")
                        await page.goto(target_url, wait_until="domcontentloaded",
                                        timeout=45000)
                    # Wait for film grid to appear
                    log("    FHL: жду сетку после логина (до 30с)...")
                    for _t in range(30):
                        await asyncio.sleep(1)
                        if await page.locator("img[src]").count() > 5:
                            log(f"    FHL: сетка готова ({_t+1}с)")
                            break
                    else:
                        log("    FHL: сетка не появилась — продолжаю")
                except Exception as e:
                    log(f"    !! FHL логин провалился: {e}")
                    return None

            # ── Find and click the selected (highlighted) thumbnail ── #
            # When ?i=N is in the URL, thumbnail N is pre-selected with a
            # blue border. We find it by class/aria and click it.
            log("    FHL: ищу подсвеченную миниатюру...")
            clicked = False
            for sel in [
                '[class*="selected"]',
                '[class*="highlight"]',
                '[aria-selected="true"]',
                '[aria-current="true"]',
                '[class*="active"]',
            ]:
                try:
                    els = page.locator(sel)
                    n = await els.count()
                    for idx in range(min(n, 20)):
                        el = els.nth(idx)
                        try:
                            tag = await el.evaluate("e => e.tagName.toLowerCase()")
                        except Exception:
                            continue
                        if tag in ("input", "button", "select", "textarea",
                                   "nav", "header", "footer", "a"):
                            continue
                        try:
                            if await el.is_visible():
                                await el.scroll_into_view_if_needed(timeout=3000)
                                await asyncio.sleep(0.5)
                                await el.click(timeout=5000)
                                await asyncio.sleep(5)  # Wait for full-size view
                                clicked = True
                                log(f"    Миниатюра кликнута (sel={sel})")
                                break
                        except Exception:
                            continue
                    if clicked:
                        break
                except Exception:
                    continue

            if not clicked:
                log("    !! FHL: подсвеченная миниатюра не найдена")
                return None

            # ── Download from the full-size viewer ───────────────── #
            log("    FHL: жду кнопку Download (до 15с)...")
            try:
                await page.wait_for_selector(
                    'button[aria-label*="Download" i]', timeout=15000)
            except Exception:
                await asyncio.sleep(5)

            dl_dir = Path.home() / "Downloads"
            before = set(dl_dir.glob("*.jpg")) | set(dl_dir.glob("*.jpeg"))

            try:
                async with page.expect_download(timeout=45000) as dl_info:
                    dl_ok = False
                    for sel in ['button[aria-label*="Download" i]',
                                'button[title*="Download" i]']:
                        try:
                            el = page.locator(sel).first
                            if await el.count() and await el.is_visible():
                                await el.click(timeout=4000)
                                await asyncio.sleep(1.5)
                                dl_ok = True
                                log("    Download нажат")
                                break
                        except Exception:
                            continue
                    if not dl_ok:
                        raise RuntimeError("кнопка Download не найдена")

                    for lbl in ("JPG Only", "JPG only"):
                        try:
                            el = page.get_by_text(lbl, exact=True).first
                            if await el.count():
                                await el.click(timeout=3000)
                                await asyncio.sleep(0.5)
                                log("    JPG Only выбран")
                                break
                        except Exception:
                            continue

                    for sel in ['[data-testid="full-text-confirm-download"]',
                                'button:has-text("Download")']:
                        try:
                            btn = page.locator(sel).first
                            if await btn.count() and await btn.is_visible():
                                await btn.click(timeout=5000)
                                log("    Confirm нажат")
                                break
                        except Exception:
                            continue

                dl = await dl_info.value
                await dl.save_as(str(dest))
                downloaded = str(dest)
                log(f"    Сохранено: {dest.name} ({dest.stat().st_size//1024}KB)")

            except Exception as exc:
                log(f"    expect_download: {exc} → жду в Downloads...")
                for _ in range(15):
                    await asyncio.sleep(1)
                    after = set(dl_dir.glob("*.jpg")) | set(dl_dir.glob("*.jpeg"))
                    nw = after - before
                    if nw:
                        src_f = max(nw, key=lambda p: p.stat().st_mtime)
                        shutil.move(str(src_f), str(dest))
                        downloaded = str(dest)
                        log(f"    Перемещено: {dest.name}")
                        break

        except Exception as e:
            log(f"    !! FHL error: {e}")
        finally:
            try:
                await ctx.close()
            except Exception:
                pass
            try:
                await browser.close()
            except Exception:
                pass

    return downloaded


async def _process_images_for_db(context, rows, images_dir, fs_email, fs_password, log):
    """
    Walk matched rows, detect image tasks, execute downloads.
    Images saved to images_dir/{row_label}[_tN].jpg
    """
    if not rows:
        return
    for ri, row in enumerate(rows):
        tasks = _row_image_tasks(row)
        if not tasks:
            continue
        label = _row_label(row) or f"row_{ri+1:03d}"
        for ti, task in enumerate(tasks):
            sfx = f"_t{ti+1}" if len(tasks) > 1 else ""
            dest = images_dir / f"{label}{sfx}.jpg"
            if task["type"] == "fhl":
                log(f"    [строка {ri+1}] FHL снимок {task['img_num']} → {dest.name}")
                await _do_fhl_download(
                    task["url"], task["img_num"],
                    dest, fs_email, fs_password, log)
            elif task["type"] == "direct":
                log(f"    [строка {ri+1}] Direct image → {dest.name}")
                await _do_direct_download(context, task["url"], dest, log)
            elif task["type"] == "ukraine":
                log(f"    [строка {ri+1}] Ukrainian ref (ссылка сохранена, не скачиваем): "
                    f"{task['url'][:60]}")


# ---------- main ------------------------------------------------------------ #
async def goto_with_retries(page, url, attempts=6, backoff=3.0):
    """Navigate, retrying on 5xx errors (502 Bad Gateway, 503 Service
    Unavailable). 503 from JewishGen usually means the server is temporarily
    overloaded — we wait progressively longer between retries."""
    last_err = None
    for i in range(attempts):
        try:
            response = await page.goto(url, wait_until="domcontentloaded", timeout=45_000)
            if response is not None and 500 <= response.status < 600:
                wait = backoff * (2 ** i)  # 3s, 6s, 12s, 24s, 48s, 96s
                print(f"  HTTP {response.status} from server — retrying in {wait:.0f}s… "
                      f"(attempt {i+1}/{attempts})")
                await asyncio.sleep(wait)
                continue
            return response
        except Exception as exc:
            last_err = exc
            wait = backoff * (2 ** i)
            print(f"  navigation error ({exc}); retrying in {wait:.0f}s… "
                  f"(attempt {i+1}/{attempts})")
            await asyncio.sleep(wait)
    if last_err:
        raise last_err


async def run_scraper(
    rows,
    country,
    keywords,
    output_format,
    output_folder,
    *,
    keyword_mode="OR",
    email=None,
    password=None,
    log=print,
    cancel_event=None,
    progress=None,
):
    """Run the JewishGen search and write the matched rows into the chosen
    output folder. All parameters are explicit so this function can be
    called from a GUI as well as from the CLI wrapper below.

    Parameters
    ----------
    rows : list of (data_type, search_type, term) tuples (1..4 entries)
    country : str — one of COUNTRIES (or 'ALL COUNTRIES')
    keywords : list[str] | str — one or more filter keywords. A row is kept
        when its combined cell text contains AT LEAST ONE of them
        (case-insensitive OR-logic). A bare string is accepted too and is
        treated as a single-keyword list.
    output_format : 'docx' | 'xlsx' | 'both'
    output_folder : str or Path — directory to write results into
    email, password : optional, used for best-effort Auth0 autofill
    log : callable(str) — anywhere a CLI would print(...) the function uses
        log(...). Defaults to print.
    cancel_event : optional threading.Event-like with .is_set() — checked
        between databases so a GUI Stop button can interrupt cleanly.
    progress : optional callable(value:int, label:str) — called at every
        major phase. `value` is 0..100. `label` is a short user-facing
        description (e.g. "[3/12] Belarus Births"). Safe to omit.
    """
    def _progress(value, label):
        if progress is None:
            return
        try:
            progress(int(value), str(label))
        except Exception:
            pass
    output_folder = Path(output_folder).expanduser().resolve()
    output_folder.mkdir(parents=True, exist_ok=True)
    # Log the resolved path so it's obvious where files will go (Alla was
    # confused once because of OneDrive / iCloud Documents redirection).
    log(f"Output folder (resolved): {output_folder}")
    if not os.access(str(output_folder), os.W_OK):
        raise PermissionError(
            f"Cannot write into {output_folder}. Pick a different folder "
            f"in the GUI or check the permissions."
        )

    # Default summary returned at the end so the GUI can show a nice popup.
    summary_result = {
        "ok": False,
        "docx_count": 0,
        "xlsx_path": None,
        "n_databases": 0,
        "output_folder": str(output_folder),
    }

    want_docx = output_format in ("docx", "both")
    want_xlsx = output_format in ("xlsx", "both")

    # Prefix for filenames and directories: search terms come first
    q_prefix = _query_prefix(rows)

    if want_xlsx and not _OPENPYXL_AVAILABLE:
        log("\n  !! openpyxl is not installed — Excel output disabled.")
        log("     Install it with: pip install openpyxl")
        want_xlsx = False
        if not want_docx:
            log("     Nothing left to write. Exiting.")
            return summary_result

    def _cancelled():
        return cancel_event is not None and cancel_event.is_set()

    # Best-effort autofill hints — picked up by wait_for_login_to_finish()
    # when a per-popup login appears mid-scraping.
    _CURRENT_CREDS["email"] = email or None
    _CURRENT_CREDS["password"] = password or None

    _progress(2, "Opening browser…")

    async with async_playwright() as p:
        # Persistent context: cookies + storage live in PROFILE_DIR, so once
        # you've logged in to JewishGen the next runs are already authenticated.
        PROFILE_DIR.mkdir(parents=True, exist_ok=True)
        context = await p.chromium.launch_persistent_context(
            str(PROFILE_DIR),
            headless=False,
            accept_downloads=True,
        )
        try:
            page = context.pages[0] if context.pages else await context.new_page()

            _progress(6, "Loading JewishGen home page…")
            await goto_with_retries(page, HOME_URL)
            _progress(10, "Filling the search form…")
            try:
                await fill_search_form(page, rows, country)
            except PaidFeatureError as exc:
                log(f"\n  !! {exc}")
                _progress(100, "Stopped — paid feature required.")
                summary_result["error"] = "paid_feature"
                summary_result["message"] = str(exc)
                return summary_result
            _progress(15, "Waiting for results page (log in if asked)…")
            page = await wait_for_results_page(
                context, page, email, password=password
            )

            if _cancelled():
                log("Cancelled by user before results page.")
                return summary_result

            _progress(18, "Scanning result buttons…")
            items = await collect_list_button_labels(page)
            if not items:
                log("No 'List N records' buttons on the results page.")
                _progress(100, "No databases on the results page.")
                return summary_result

            # Normalise keywords to a list of non-empty strings for display.
            if isinstance(keywords, str):
                kw_list = [keywords] if keywords.strip() else []
            else:
                kw_list = [k for k in (keywords or []) if (k or "").strip()]
            kw_label = " OR ".join(kw_list) if kw_list else "(none)"
            query_lines = [f"Region: {country}", f"Keywords (OR): {kw_label}"]
            for dt, st, term in rows:
                query_lines.append(f"{dt} ({st.lower()}): {term}")

            formats_label = "/".join(
                f for f in (("docx" if want_docx else None),
                            ("xlsx" if want_xlsx else None)) if f
            )
            log(f"\nFound {len(items)} 'List records' buttons. "
                f"Output format: {formats_label}. Writing into:")
            log(f"   {output_folder}\n")
            _progress(20, f"Found {len(items)} databases — scanning…")

            saved_docx_count = 0
            used_filenames = set()
            databases_for_xlsx = []  # one entry per DB that produced matches
            n_items = len(items)
            consecutive_unavailable = 0   # run of DBs that all 503'd → give up
            for i, item in enumerate(items, 1):
                if _cancelled():
                    log("Cancelled by user.")
                    break
                label = item["label"]
                desc = (item.get("desc") or "").strip() or label
                # Progress climbs from 20% to 90% across the databases.
                pct = 20 + int(70 * (i - 1) / max(n_items, 1))
                _progress(pct, f"[{i}/{n_items}] {desc}")
                log(f"  [{i}/{n_items}] {desc} — {label}")

                # Human-like pause between databases (3–8 s) to avoid 524.
                if i > 1:
                    inter_pause = random.uniform(3.0, 8.0)
                    log(f"    (pausing {inter_pause:.1f}s before next database…)")
                    await asyncio.sleep(inter_pause)

                # Open database page with 524 retry.
                result_page = None
                open_error = None
                browser_dead = False
                for attempt in range(3):
                    try:
                        result_page = await click_list_button_open_new_page(
                            context, page, label
                        )
                        await ensure_not_503(result_page)
                        break
                    except TemporaryServerError:
                        open_error = "503"
                        wait_s = 20 * (attempt + 1)  # 20, 40, 60 s
                        log(f"    524/503 opening database (attempt {attempt+1}/3) "
                            f"— waiting {wait_s}s…")
                        if result_page and result_page is not page:
                            try:
                                await result_page.close()
                            except Exception:
                                pass
                        result_page = None
                        await asyncio.sleep(wait_s)
                    except Exception as exc:
                        msg = str(exc)
                        open_error = msg
                        # Browser / context gone → nothing more will open; abort
                        # the whole run instead of churning through every DB.
                        if ("has been closed" in msg.lower()
                                or "target page" in msg.lower()
                                or "target closed" in msg.lower()):
                            browser_dead = True
                        else:
                            log(f"    skipped (could not open): {exc}")
                        result_page = None
                        break

                if browser_dead:
                    log("    !! браузер/вкладка закрыты — прекращаю прогон "
                        "(дальше ничего не откроется); сохраняю собранное.")
                    break

                if result_page is None:
                    if open_error == "503":
                        consecutive_unavailable += 1
                        log(f"    skipped — сервер недоступен (503/524) "
                            f"[{consecutive_unavailable} подряд].")
                        if consecutive_unavailable >= 3:
                            log("    !! сервер стабильно отдаёт 503/524 — "
                                "прекращаю прогон, попробуйте позже; "
                                "сохраняю собранное.")
                            break
                    else:
                        log("    skipped — could not open after retries.")
                    continue

                consecutive_unavailable = 0   # opened OK → reset the 503 run

                log(f"     -> opened: {result_page.url}")
                acc = {"headerLines": [], "columns": [], "rows": []}
                try:
                    await follow_next_pages(
                        result_page, keywords, acc,
                        keyword_mode=keyword_mode,
                    )
                except Exception as exc:
                    log(f"    error during scan: {exc}")

                if result_page is not page:
                    try:
                        await result_page.close()
                    except Exception:
                        pass
                else:
                    try:
                        await page.go_back(wait_until="domcontentloaded")
                    except Exception:
                        pass

                if not acc["rows"]:
                    log("    no match — skipping")
                    continue

                # ── Download images linked from result rows ─────────── #
                # Dir: {search_terms}_{db_name}
                _db_part  = safe_filename(desc or label) or f"db_{i}"
                _img_dir  = output_folder / "images" / f"{q_prefix}_{_db_part}"
                _fs_u, _fs_p = _get_fs_credentials()
                try:
                    await _process_images_for_db(
                        context, acc["rows"], _img_dir, _fs_u, _fs_p, log)
                except Exception as _exc:
                    log(f"    !! image processing error: {_exc}")

                if want_xlsx:
                    databases_for_xlsx.append({
                        "name": desc,
                        "header_lines": acc["headerLines"],
                        "columns": acc["columns"],
                        "rows": list(acc["rows"]),
                    })

                if want_docx:
                    # Filename: {search_terms}_{db_name}.docx
                    base_name = f"{q_prefix}_{safe_filename(desc) or safe_filename(label)}"
                    file_name = f"{base_name}.docx"
                    n = 2
                    while file_name in used_filenames:
                        file_name = f"{base_name}_{n}.docx"
                        n += 1
                    used_filenames.add(file_name)

                    out_path = output_folder / file_name
                    write_database_docx(
                        out_path,
                        desc,
                        acc["headerLines"],
                        acc["columns"],
                        acc["rows"],
                        query_lines,
                    )
                    saved_docx_count += 1
                    log(f"    saved {len(acc['rows'])} row(s) → {file_name}")
                else:
                    log(f"    queued {len(acc['rows'])} row(s) for the workbook")

            xlsx_path = None
            if want_xlsx and databases_for_xlsx:
                _progress(95, "Writing Excel workbook…")
                target = output_folder / f"{q_prefix}_jewishgen_results.xlsx"
                try:
                    wrote = write_xlsx_all(
                        target, databases_for_xlsx, query_lines
                    )
                    # write_xlsx_all now returns either the path string it
                    # actually wrote to (so we can pick up the _new fallback
                    # name on PermissionError) or False if nothing was saved.
                    if wrote:
                        xlsx_path = Path(wrote)
                        log(f"\n  Excel workbook → {xlsx_path.name} "
                            f"({len(databases_for_xlsx)} sheet(s))")
                    else:
                        xlsx_path = None
                        log(
                            "\n  !! Excel workbook NOT written — no sheet "
                            "produced any rows."
                        )
                except Exception as exc:
                    log(f"\n  !! could not write xlsx: "
                        f"{type(exc).__name__}: {exc}")
                    xlsx_path = None
            elif want_xlsx and not databases_for_xlsx:
                log(
                    "\n  !! Excel was requested but no database produced "
                    "matches — nothing to put in the workbook."
                )

            log(f"\nDone. Output folder:")
            log(f"   {output_folder}")
            if want_docx:
                log(f"   .docx files: {saved_docx_count}")
            if want_xlsx:
                if xlsx_path:
                    log(f"   .xlsx: {xlsx_path.name}")
                else:
                    log("   .xlsx: none (no matches)")

            # Compose final, user-friendly summary for the GUI status label.
            parts = []
            if want_docx:
                parts.append(f"{saved_docx_count} Word file(s)")
            if want_xlsx and xlsx_path:
                parts.append(f"{len(databases_for_xlsx)}-sheet Excel workbook")
            summary = " + ".join(parts) if parts else "no matches"
            _progress(100, f"Done — {summary} in {output_folder}")

            summary_result.update({
                "ok": True,
                "docx_count": saved_docx_count,
                "xlsx_path": str(xlsx_path) if xlsx_path else None,
                "n_databases": len(databases_for_xlsx) if want_xlsx else saved_docx_count,
            })
        finally:
            try:
                await context.close()
            except Exception:
                pass
            # Wipe creds — they were only ever a hint for the per-popup
            # login waiter while this run was active.
            _CURRENT_CREDS["email"] = None
            _CURRENT_CREDS["password"] = None

    return summary_result


async def run():
    """Original CLI entry — asks everything in the terminal, then calls
    run_scraper(). The GUI uses run_scraper directly."""
    rows, country, keyword, email = ask_inputs()
    # `ask_inputs` only collects one keyword. Power users can chain a few
    # more here interactively without changing the prompt for everyone.
    extra = input(
        "Extra filter keyword (optional — leave blank to skip): "
    ).strip()
    extra2 = input(
        "One more filter keyword (optional — leave blank to skip): "
    ).strip() if extra else ""
    keywords = [k for k in (keyword, extra, extra2) if k]
    output_format = ask_output_format()
    output_folder = HERE / "results"
    await run_scraper(
        rows, country, keywords, output_format, output_folder,
        email=email or None,
        password=None,    # CLI never prompts for password
        log=print,
    )


if __name__ == "__main__":
    asyncio.run(run())