#!/usr/bin/env python3
"""
pamyat_scraper.py
=================
Scraper for «Память народа» (https://pamyat-naroda.ru) — WWII participant records.

Flow:
  1. Go straight to the advanced search page (?adv_search=y), fill the fields
     (all addressed by their language-independent id/name), submit «Уточнить и
     найти».
  2. Collect the result persons (a.heroes-list-item-name). When a name was given,
     keep only FUZZY-matching people — surname & first name may vary
     (Рогинский/Рагинский/Рочинский, Лев/Лейб/Л.) but the patronymic must match
     (Михайлович/Михелевич/М. — yes; Яковлевич/Григорьевич — no).
  3. Open each matched person; one person groups SEVERAL document cards — walk
     ALL of them, expand «Информация об архиве», copy every field to Word exactly
     as shown. Save the person photo and (best-effort) the document images.

GUI is always English with a ru/en language selector; the scraper opens the
chosen language version of the site. No login.

NEVER crash on a broken/empty page: every person and every card is wrapped in
try/except, urlparse().netloc (not .host), and bad links are skipped.
"""

import asyncio
import difflib
import io
import re
import sys
import time
from pathlib import Path
from docx_util import set_cell_lines
from urllib.parse import urlparse, urljoin

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        import os
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from PIL import Image
    _PIL_OK = True
except ImportError:
    _PIL_OK = False

# ── Constants ─────────────────────────────────────────────────────────────── #
HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")

# Branch-of-service (army_unit_rod) — English label → site option value.
BRANCH = {
    "": "", "Aviation": "1", "Armored troops": "2", "Motor transport": "3",
    "Artillery": "4", "Navy": "7", "Signal troops": "8", "Railway troops": "10",
    "Reserve units": "11", "Engineering troops": "12", "Cavalry": "13",
    "Medical": "14", "People's militia": "15", "NKVD": "16", "Air defence": "17",
    "Infantry": "20", "Tank troops": "21", "Technical troops": "22",
    "Training units": "23",
}
# Awards (award_ids) — a useful subset, English label → value.
AWARDS = {
    "": "", "Hero of the Soviet Union": "3", "Order of Lenin": "151",
    "Order of the Red Banner": "144", "Order of the Red Star": "146",
    "Order of the Patriotic War": "157",
    "Order of the Patriotic War 1st class": "158",
    "Order of the Patriotic War 2nd class": "159",
    "Order of Glory": "163", "Order of Glory 1st class": "165",
    "Order of Glory 2nd class": "166", "Order of Glory 3rd class": "167",
    "Order of Alexander Nevsky": "118", "Order of Kutuzov": "147",
    "Order of Suvorov": "168", "Order of Nakhimov": "153",
    "Order of Ushakov": "177", "Medal 'For Courage'": "78",
    "Medal 'For Battle Merit'": "48",
    "Medal 'For the Defence of Stalingrad'": "69",
    "Medal 'For the Defence of Leningrad'": "63",
    "Medal 'For the Defence of Moscow'": "64",
    "Medal 'For the Capture of Berlin'": "49",
    "Medal 'For Victory over Germany'": "84",
}


# ── Helpers ───────────────────────────────────────────────────────────────── #
def safe_fn(s: str, n: int = 80) -> str:
    s = re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip())
    return re.sub(r"\s+", "_", s)[:n].strip("_") or "record"


def _host(u: str) -> str:
    u = u or ""
    if not u or u.startswith(("chrome-error", "about:", "data:")):
        return ""
    try:
        return urlparse(u).netloc or ""
    except Exception:
        return ""


def _to_png(data: bytes):
    if not data:
        return None
    if (data[:3] == b"\xff\xd8\xff" or data[:8] == b"\x89PNG\r\n\x1a\n"
            or data[:4] == b"GIF8" or data[:2] == b"BM"):
        return data
    if _PIL_OK:
        try:
            im = Image.open(io.BytesIO(data)).convert("RGB")
            out = io.BytesIO(); im.save(out, format="PNG")
            return out.getvalue()
        except Exception:
            return None
    return None


_SUMMARY_LABELS = ["Дата рождения", "Место рождения", "Место призыва",
                   "Дата призыва", "Воинское звание", "Воинская часть", "Награды",
                   "Дата выбытия/смерти", "Место службы", "Дата смерти",
                   "Место выбытия", "Дата рождения/возраст"]


# Fixed labels used in «Память народа» document cards (from live cards). We split
# a card's text by these KNOWN labels — robust regardless of how label/value are
# separated, and it never absorbs the card name/type. Longest-first so «Дата и
# место призыва» wins over «Дата».
_CARD_LABELS = sorted([
    "Дата рождения", "Место рождения", "Дата и место призыва", "Место призыва",
    "Дата призыва", "Наименование военкомата", "Воинское звание", "Воинская часть",
    "Последнее место службы", "Военно-пересыльный пункт", "Прибыл в часть",
    "Выбытие из воинской части", "Куда выбыл", "Дата прибытия", "Откуда прибыл",
    "Причина выбытия", "Дата выбытия", "Госпиталь", "Судьба",
    "Номер свидетельства", "Дата свидетельства", "Наименование награды",
    "Номер документа", "Дата документа", "Автор документа",
    "Источник информации", "Фонд ист. информации", "Опись ист. информации",
    "Дело ист. информации", "Номер ящика", "Номер шкафа", "Архив", "Картотека",
    "Раздел", "Название фонда", "Расположение документа", "Лагерь военнопленных",
    "Место пленения", "Дата пленения", "Награды",
], key=len, reverse=True)

_CARD_LABEL_RE = re.compile(
    r"(" + "|".join(re.escape(_l) for _l in _CARD_LABELS) + r")\s*:?\s*")

# Archive-block labels (appear after «Информация об архиве»). The card TYPE
# («Картотека ранений», «Именной список части» …) sits before the first DATA
# label, so we split type↔fields at the first DATA label (a non-archive one).
_ARCHIVE_LABELS = {"Источник информации", "Фонд ист. информации",
                   "Опись ист. информации", "Дело ист. информации", "Номер ящика",
                   "Номер шкафа", "Архив", "Картотека", "Раздел", "Название фонда",
                   "Расположение документа"}
_DATA_LABEL_RE = re.compile(
    r"(" + "|".join(re.escape(_l) for _l in
                    sorted([x for x in _CARD_LABELS if x not in _ARCHIVE_LABELS],
                           key=len, reverse=True)) + r")")


def _parse_fields(txt: str):
    """Split a card's text by the KNOWN labels → ordered [(label, value)]."""
    txt = re.sub(r"[ \t ]+", " ", (txt or ""))
    txt = re.sub(r"\s*\n\s*", " ", txt).strip()
    txt = re.sub(r"Информаци[ия] об архиве\s*[+\-]?", " ", txt)
    marks = list(_CARD_LABEL_RE.finditer(txt))
    out, seen = [], set()
    for i, m in enumerate(marks):
        end = marks[i + 1].start() if i + 1 < len(marks) else len(txt)
        val = txt[m.end():end].strip(" ;,")
        k = m.group(1).strip()
        if val and len(val) < 400 and (k, val) not in seen:
            seen.add((k, val)); out.append((k, val))
    return out


_CHROME_MARKERS = ("О проекте", "Вопросы и ответы", "Как искать", "Обратная связь",
                   "Правовая информация", "Пользовательское соглашение",
                   "Министерство обороны", "© Министерство", "Урок Победы",
                   "Расскажи одноклассникам", "К странице")


def _clean_block(txt: str) -> str:
    """Strip the site header/footer/nav junk that leaks into extracted text."""
    if not txt:
        return ""
    cut = len(txt)
    for m in _CHROME_MARKERS:
        i = txt.find(m)
        if 0 <= i < cut:
            cut = i
    return txt[:cut].strip(" ;,—-·•\t\n")


def _add_hyperlink(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


# ── FIO (surname / first / patronymic) fuzzy matching ─────────────────────── #
def _sim(a: str, b: str) -> float:
    a, b = (a or "").lower().strip(), (b or "").lower().strip()
    if not a or not b:
        return 0.0
    return difflib.SequenceMatcher(None, a, b).ratio()


def _is_initial(token: str, full: str) -> bool:
    """token like «Л.» / «Л» matches full «Лев» (same first letter)."""
    t = (token or "").strip(". ").lower()
    f = (full or "").lower()
    return len(t) == 1 and bool(f) and f.startswith(t)


def _parse_fio(name: str):
    """«Рогинский Лев Михайлович» / «Рогинский Л.М.» / «Рагинский Л. М.» /
    «Рочинский Лев Мих.» → (last, first, middle)."""
    name = (name or "").strip()
    # split glued initials «Л.М.» → «Л. М.»
    name = re.sub(r"([А-ЯЁA-Za-zа-яё])\.([А-ЯЁA-Z])", r"\1. \2", name)
    parts = [p for p in re.split(r"\s+", name) if p]
    return (parts[0] if parts else "",
            parts[1] if len(parts) > 1 else "",
            parts[2] if len(parts) > 2 else "")


def _first_match(r: str, w: str) -> bool:
    """First name: «Лев» accepts «Лейб» / «Л.» (same initial / similar)."""
    r = (r or "").strip(". ").lower()
    w = (w or "").strip(". ").lower()
    if not r or not w:
        return True
    if _is_initial(r, w) or _is_initial(w, r):
        return True
    if r[:2] == w[:2]:                       # Лев / Лейб
        return True
    return _sim(r, w) >= 0.5


def _patr_match(r: str, w: str) -> bool:
    """Patronymic — the DISCRIMINATOR. «Михайлович» accepts «Михелевич», «Мих.»,
    «М.» (same stem) but REJECTS «Максимович»/«Захарович»/«Яковлевич» (same
    «-ович» ending, different stem)."""
    r = (r or "").strip(". ").lower()
    w = (w or "").strip(". ").lower()
    if not r or not w:
        return True
    short, long = (r, w) if len(r) <= len(w) else (w, r)
    if short and long.startswith(short):     # «Мих.»→«Михайлович», «М.»→…
        return True
    if r[:3] == w[:3]:                         # same stem start: мих == мих
        return True
    return _sim(r, w) >= 0.78


def _person_matches(result_name: str, want_last: str, want_first: str,
                    want_middle: str) -> bool:
    """Accept if surname differs by ~1-2 letters, first name matches/initials,
    and the patronymic STEM matches (the discriminator)."""
    if not want_last:
        return True
    r_last, r_first, r_middle = _parse_fio(result_name)
    if _sim(r_last, want_last) < 0.7:          # Рогинский/Рагинский/Рочинский/Роминский…
        return False
    if want_first and r_first and not _first_match(r_first, want_first):
        return False
    if want_middle and r_middle and not _patr_match(r_middle, want_middle):
        return False
    return True


# ── Search ────────────────────────────────────────────────────────────────── #
def _base(lang: str) -> str:
    return "https://pamyat-naroda.ru/en" if lang == "en" else "https://pamyat-naroda.ru"


async def _fill_by_id(page, fid: str, value: str, log):
    if not value:
        return
    try:
        el = page.locator(f"#{fid}").first
        if await el.count():
            await el.click(timeout=3000)
            await page.keyboard.press("Control+a")
            await page.keyboard.press("Delete")
            await page.keyboard.type(str(value), delay=20)
            log(f"  ✓ {fid} = {value!r}")
    except Exception as e:
        log(f"  !! поле {fid}: {e}")


async def _do_search(page, params, lang, log) -> bool:
    url = f"{_base(lang)}/heroes/?adv_search=y"
    log(f"  → Открываю расширенный поиск: {url}")
    await page.goto(url, wait_until="domcontentloaded", timeout=40000)
    await asyncio.sleep(2)

    # Ignore the (empty) main search string, keep documents grouped per person.
    for cid, want in (("use_main_string", True),
                      ("use_person", bool(params.get("group_person", True)))):
        try:
            cb = page.locator(f"#{cid}").first
            if await cb.count():
                checked = await cb.is_checked()
                if checked != want:
                    await cb.click(timeout=2500)
        except Exception:
            pass

    await _fill_by_id(page, "last_name",   params.get("last_name", ""), log)
    await _fill_by_id(page, "first_name",  params.get("first_name", ""), log)
    await _fill_by_id(page, "middle_name", params.get("middle_name", ""), log)
    await _fill_by_id(page, "date_birth_from", params.get("birth_from", ""), log)
    await _fill_by_id(page, "date_birth_to",   params.get("birth_to", ""), log)
    await _fill_by_id(page, "place_birth", params.get("place_birth", ""), log)
    await _fill_by_id(page, "rank",        params.get("rank", ""), log)
    await _fill_by_id(page, "poslednee_mesto_sluzhbi",
                      params.get("place_service", ""), log)
    await _fill_by_id(page, "data_i_mesto_priziva",
                      params.get("place_conscription", ""), log)
    # Dropdowns: set the underlying <select> by value (language-independent).
    for sid, val in (("army_unit_rod", params.get("branch", "")),
                     ("award_ids", params.get("award", ""))):
        if val:
            try:
                await page.select_option(f"#{sid}", value=str(val), timeout=3000)
                log(f"  ✓ {sid} = {val}")
            except Exception as e:
                log(f"  !! select {sid}: {e}")

    # Submit «Уточнить и найти» / «Refine and find».
    clicked = False
    for sel in ('input[type="submit"][value="Уточнить и найти"]',
                'input[type="submit"][value*="найти" i]',
                'input[type="submit"][value*="find" i]',
                'input.js-index-search-submit', 'input[type="submit"]'):
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                clicked = True
                log(f"  ✓ Отправил поиск ({sel})")
                break
        except Exception:
            continue
    if not clicked:
        await page.keyboard.press("Enter")
        log("  → Поиск: Enter")

    for _ in range(30):
        await asyncio.sleep(1)
        try:
            ok = await page.evaluate(
                "() => document.querySelectorAll('a.heroes-list-item-name').length > 0"
                " || /найдено|not found|ничего не/i.test(document.body.innerText)")
            if ok:
                return True
        except Exception:
            pass
    return True


async def _collect_results(page, params, log, max_pages, max_persons) -> list:
    """Collect person results across ALL pages; keep only FIO-matches (and, when
    a birth year is given, drop people whose card year clearly differs)."""
    want = (params.get("last_name", ""), params.get("first_name", ""),
            params.get("middle_name", ""))
    ym = re.search(r"(18|19|20)\d{2}", params.get("birth_from", "") or "")
    want_year = ym.group(0) if ym else ""
    out, seen = [], set()
    page_no = 1
    while page_no <= max_pages and len(out) < max_persons:
        rows = await page.evaluate(r"""() => {
            const norm = s => (s || '').replace(/\s+/g, ' ').trim();
            const out = [];
            for (const a of document.querySelectorAll('a.heroes-list-item-name')) {
                const name = norm(a.textContent);
                const href = a.href || '';
                // climb to the RESULT BOX (the ancestor whose text has the
                // description with a birth year) — not the link itself.
                let snippet = '', box = a;
                for (let i = 0; i < 6 && box.parentElement; i++) {
                    box = box.parentElement;
                    const t = norm(box.innerText);
                    if (/(18|19|20)\d{2}/.test(t) && t.length > name.length + 5) {
                        snippet = t.split(name).join(' ').replace(/\s+/g,' ').trim();
                        break;
                    }
                }
                let docs = '';
                const m = (snippet || '').match(/(\d+)\s*документ/i);
                if (m) docs = m[1];
                if (name && href) out.push({name, href, docs, snippet: snippet.slice(0,300)});
            }
            return out;
        }""")
        kept = 0
        for r in rows:
            if r["href"] in seen:
                continue
            if not _person_matches(r["name"], *want):
                continue
            if want_year:
                yrs = re.findall(r"(?:18|19|20)\d{2}", r.get("snippet", ""))
                # EXACT year: drop people whose card shows a year but not the
                # wanted one (1921 search → no 1922/1925). Keep cards with no year.
                if yrs and want_year not in yrs:
                    continue
            seen.add(r["href"]); out.append(dict(r)); kept += 1
            if len(out) >= max_persons:
                break
        log(f"  → Страница {page_no}: результатов {len(rows)}, подходящих {kept}")
        # First result href — to detect that the page actually advanced.
        first_href = rows[0]["href"] if rows else ""
        # Next page is «<a onclick="getPage(N)">→</a>», NOT a normal link.
        advanced = await page.evaluate(r"""(n) => {
            try { if (typeof getPage === 'function') { getPage(n); return true; } } catch(e){}
            const a = Array.from(document.querySelectorAll('a[onclick*="getPage"]'))
                .find(x => /→|»|next|след/i.test((x.textContent||'') + ' ' + x.className));
            if (a) { a.click(); return true; }
            return false;
        }""", page_no + 1)
        if not advanced:
            break
        # wait until the first result link changes (new page loaded)
        changed = False
        for _ in range(15):
            await asyncio.sleep(1)
            try:
                cur = await page.evaluate(
                    "() => { const a = document.querySelector('a.heroes-list-item-name');"
                    " return a ? a.href : ''; }")
                if cur and cur != first_href:
                    changed = True; break
            except Exception:
                pass
        if not changed:
            break
        page_no += 1
    return out


# ── Person page ───────────────────────────────────────────────────────────── #
# JS that parses a run of «Метка: значение Метка2: значение2…» text into pairs.
_PARSE_JS = r"""
    const norm = s => (s||'').replace(/ /g,' ').replace(/\s+/g,' ').trim();
    function parseFields(txt){
        txt = (txt||'').replace(/ /g,' ');
        const re = /([А-ЯЁ][А-Яа-яё.,№()\/\-«»"' ]{1,48}?):[ \t]*/g;
        const marks = []; let m;
        while((m=re.exec(txt))!==null)
            marks.push({l: m[1].replace(/\s+/g,' ').trim(), s: m.index, v: re.lastIndex});
        const out = [];
        for(let i=0;i<marks.length;i++){
            const end = (i+1<marks.length) ? marks[i+1].s : txt.length;
            let val = txt.slice(marks[i].v, end).replace(/\s+/g,' ').trim();
            if(val && val.length < 400) out.push([marks[i].l, val]);
        }
        return out;
    }
"""


async def _tab_block(page, words, wait_secs, log):
    """Click a tab (span.heroes_card__tabs-item-title) and return its visible
    content text. «Боевой путь» loads slowly → wait."""
    try:
        clicked = await page.evaluate(r"""(words) => {
            const norm = s => (s||'').replace(/\s+/g,' ').trim().toLowerCase();
            for (const sp of document.querySelectorAll(
                    '.heroes_card__tabs-item-title, [class*="tabs-item-title"]')) {
                if (words.some(w => norm(sp.textContent).includes(w))) {
                    (sp.closest('[class*="tabs-item"],li,a,button') || sp).click();
                    return true;
                }
            }
            return false;
        }""", words)
        if not clicked:
            return ""
        await asyncio.sleep(wait_secs)
        return await page.evaluate(_PARSE_JS + r"""
            () => {
                const sel = '[class*="tab-content"],[class*="tab-pane"],'
                    + '[class*="heroes_card__content"],[class*="card__content"],'
                    + '[class*="boevoy"],[class*="combat"]';
                let el = null, best = 0;
                document.querySelectorAll(sel).forEach(e => {
                    if (e.offsetParent !== null && (e.innerText||'').length > best) {
                        best = (e.innerText||'').length; el = e; }
                });
                return el ? norm(el.innerText).slice(0, 4000) : '';
            }""")
    except Exception:
        return ""


# JS: the scan is shown as an IMAGE-MAP image — «<img usemap=…>» with clickable
# <area> rows. Return its src (or, failing that, the largest loaded content img).
_SCAN_IMG_JS = r"""() => {
    // ONLY the image-map scan (what the «Документы» viewer actually shows). The
    // old "largest image" fallback grabbed the PORTRAIT when no real document
    // loaded → it ended up embedded huge as a fake "document". No fallback now:
    // no image-map → no scan (better nothing than the portrait).
    const ok = im => im && im.complete && im.naturalWidth > 150;
    for (const im of document.querySelectorAll('img[usemap], img[ismap]'))
        if (ok(im)) return im.src;
    return '';
}"""


async def _find_scan(page, secs: int) -> str:
    """Poll up to `secs` for the document scan image (image-map img) in ANY frame."""
    for _ in range(secs):
        await asyncio.sleep(1)
        for fr in page.frames:
            try:
                src = await fr.evaluate(_SCAN_IMG_JS)
            except Exception:
                src = ""
            if src:
                return src
    return ""


async def _grab_scan(page, images_dir, base, rec, log):
    """For each document card: select it, wait for its scan to load, then CLICK
    the «Сохранить изображение» button (#btnSaveLocalImage) and capture the
    download. Cards without a document are skipped fast. Returns
    (saved_any, saw_document) — saw_document=True means a real document was found,
    so the caller may reload once if the scan stalled; False means there's nothing
    to wait for and the caller must NOT reload (что и вешало на пустых страницах)."""
    cards = page.locator(".hero-card-docs-item")
    try:
        nc = await cards.count()
    except Exception:
        nc = 0
    saved, saw_doc = 0, False
    _vis_btn_js = ("() => { const b=document.querySelector('#btnSaveLocalImage');"
                   " return !!(b && b.offsetParent !== null); }")
    for k in range(min(nc, 12)):
        try:
            await cards.nth(k).scroll_into_view_if_needed(timeout=2000)
            await cards.nth(k).click(timeout=2500)
        except Exception:
            continue
        # Does this card actually carry a document? The save button must be
        # VISIBLE (offsetParent) — a hidden #btnSaveLocalImage sits in the DOM
        # even on document-less cards and made us wait/reload for nothing.
        # Check immediately, then poll ~2s max → empty card skipped fast.
        has_btn = False
        for _ in range(3):
            try:
                has_btn = await page.evaluate(_vis_btn_js)
            except Exception:
                has_btn = False
            if has_btn:
                break
            await asyncio.sleep(1)
        if not has_btn:
            continue                        # no document on this card → skip fast
        saw_doc = True
        # Visible button → the document is REAL and its scan WILL load («если
        # кнопка видима, скан надо ЖДАТЬ — и 6с может не хватить»). Wait up to
        # ~30s, continuing the moment it's decoded; only EMPTY cards skip fast.
        ready = False
        for _ in range(30):
            try:
                ready = await page.evaluate(
                    "() => { const b=document.querySelector('#btnSaveLocalImage');"
                    " const i=document.querySelector('img[usemap],img[ismap]');"
                    " return !!(b && b.offsetParent!==null && i && i.complete && i.naturalWidth>150); }")
            except Exception:
                ready = False
            if ready:
                break
            await asyncio.sleep(1)
        # Even if the image-map check didn't pass, the button IS visible →
        # attempt the save anyway (some documents may not be image-maps).
        # talking filename built from THIS card's own document name
        label = ""
        try:
            ne = cards.nth(k).locator(".hero-card-docs-item__name").first
            if await ne.count():
                label = safe_fn(re.sub(r"\s+", " ", await ne.inner_text()).strip())[:60]
        except Exception:
            label = ""
        label = label or f"документ {saved + 1}"

        def _dest(suffix):
            d = images_dir / f"{base}_{label}{suffix}"
            if d.exists():
                d = images_dir / f"{base}_{label}_{saved + 1}{suffix}"
            return d

        # click «Сохранить изображение» → capture the download
        try:
            btn = page.locator("#btnSaveLocalImage").first
            async with page.expect_download(timeout=15000) as dl:
                await btn.click(timeout=4000)
            d = await dl.value
            suf = Path(d.suggested_filename or "scan.jpg").suffix or ".jpg"
            dest = _dest(suf)
            await d.save_as(str(dest))
            rec["scans"].append(str(dest)); saved += 1
            log(f"      🖼 документ сохранён: {dest.name}")
        except Exception:
            # fallback: fetch the image-map src directly
            try:
                src = await _find_scan(page, 2)
                if src:
                    r = await page.request.get(src, timeout=30000)
                    if r.ok:
                        body = await r.body()
                        if len(body) > 15000:
                            dest = _dest(".jpg")
                            dest.write_bytes(body)
                            rec["scans"].append(str(dest)); saved += 1
                            log(f"      🖼 документ сохранён (img): {dest.name}")
            except Exception:
                pass
    return saved > 0, saw_doc


async def _save_img_url(page, url, dest) -> bool:
    """Download an image URL and write it to dest. A REFERER is essential — the
    pamyat CDN blocks hot-linking, so a plain fetch returns 403 (this is exactly
    why the photo silently never saved). Returns True on success."""
    if not url or not url.startswith("http"):
        return False
    try:
        r = await page.request.get(url, headers={"referer": page.url}, timeout=20000)
        if not r.ok:
            return False
        body = await r.body()
        if len(body) < 2000:
            return False
        dest.write_bytes(body)
        return True
    except Exception:
        return False


async def _img_bytes_via_goto(context, url, referer) -> bytes:
    """Fetch image bytes by NAVIGATING a throwaway tab to the URL. This uses the
    REAL browser (its cookies + a referer), which is what the pamyat CDN requires:
    a plain APIRequest gets 403, and a cross-origin fetch() is blocked by CORS, so
    both failed before. A browser navigation to the image works. Returns b'' on
    failure. Never raises."""
    p = await context.new_page()
    try:
        try:
            await p.set_extra_http_headers({"referer": referer})
        except Exception:
            pass
        resp = await p.goto(url, timeout=20000, wait_until="commit")
        if resp and resp.ok:
            body = await resp.body()
            if body and len(body) > 2000:
                return body
    except Exception:
        pass
    finally:
        try:
            await p.close()
        except Exception:
            pass
    return b""


async def _grab_photos(page, images_dir, base, rec, log):
    """Person photo(s): «<a class="hero-card-person-photo__item" href="…_photo.jpg"
    style="background-image:url(…)">». No save arrow — you click the photo and save
    it. The same portrait can repeat across cards → de-dupe by URL. Download via a
    real browser navigation (the CDN blocks plain fetches); if that fails, SCREENSHOT
    the photo element itself (its background-image renders the portrait). Saved to
    disk AND kept for Word. Never raises."""
    # We run BEFORE _save_media creates images_dir, so create it here too —
    # otherwise write_bytes() throws FileNotFoundError and the photo is lost
    # silently (this is exactly why the portrait never saved).
    try:
        images_dir.mkdir(parents=True, exist_ok=True)
    except Exception:
        pass
    try:
        n = await page.evaluate(r"""() => {
            const sel = '.hero-card-person-photo__item, .js-hero-person-card-photo,'
                + ' a[rel="hero-card-photo"], a[class*="person-photo" i]';
            const els = [...document.querySelectorAll(sel)];
            els.forEach((a, i) => a.setAttribute('data-pw-photo', i));
            return els.length;
        }""")
        if not n:
            log("      (фото на странице не найдено)")
            return
        seen = set()
        for i in range(min(n, 12)):
            a = page.locator(f'[data-pw-photo="{i}"]').first
            if not await a.count():
                continue
            try:
                url = await a.evaluate(r"""el => {
                    let u = el.getAttribute('href') || el.href || '';
                    if (!/\.(jpe?g|png|webp|gif)/i.test(u)) {
                        const bg = (el.style && el.style.backgroundImage) || '';
                        const m = bg.match(/url\(["']?(.*?)["']?\)/);
                        if (m) u = m[1];
                    }
                    return u || '';
                }""")
            except Exception:
                url = ""
            if url and url in seen:
                continue
            if url:
                seen.add(url)
            nidx = len(rec["photos"]) + 1
            tag = "_фото" if nidx == 1 else f"_фото{nidx}"
            saved = False

            # 1) real-browser navigation → original image bytes
            if url:
                body = await _img_bytes_via_goto(page.context, url, page.url)
                if body:
                    suf = Path(urlparse(url).path).suffix.lower()
                    if suf not in (".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"):
                        suf = ".jpg"
                    dest = images_dir / f"{base}{tag}{suf}"
                    try:
                        dest.write_bytes(body)
                        rec["photos"].append(str(dest)); saved = True
                        log(f"      📷 фото сохранено: {dest.name}")
                    except Exception as e:
                        log(f"      !! фото: не записалось на диск ({type(e).__name__})")

            # 2) screenshot the photo element itself (renders its background image)
            if not saved:
                try:
                    await a.scroll_into_view_if_needed(timeout=2000)
                    shot = await a.screenshot(timeout=6000)
                    if shot and len(shot) > 1000:
                        dest = images_dir / f"{base}{tag}.png"
                        dest.write_bytes(shot)
                        rec["photos"].append(str(dest)); saved = True
                        log(f"      📷 фото сохранено (снимок): {dest.name}")
                except Exception as e:
                    log(f"      !! фото-снимок не вышел ({type(e).__name__})")

            if not saved:
                log("      !! фото найдено, но не скачалось")
    except Exception as e:
        log(f"      !! фото: ошибка ({type(e).__name__})")


async def _save_media(page, rec, images_dir, base, log):
    """Person photo(s) → downloaded. Awards (order/medal images near «Орден/Медаль»
    text) → bytes for Word, NOT saved. Document SCAN (large, slow-loading archival
    page in «Документы») → waited for, refreshed if it stalls, then saved. Never
    touch «Скачать документы» (the album/booklet). Never raises."""
    images_dir.mkdir(parents=True, exist_ok=True)

    # ── Awards (orders/medals): images associated with award text or an award
    # container, ANYWHERE on the page. Exclude the portrait (src «_photo» / inside
    # a person-photo box) and banners/logos. Downloaded through the browser (the
    # CDN blocks plain requests, same as the photo). Kept as bytes for Word —
    # NOT written to the images folder. ──
    try:
        award_srcs = await page.evaluate(r"""() => {
            const out = [], seen = new Set();
            const bad = /\.svg|sprite|logo|icon|placeholder|data:|booklet|urok|_photo|person|avatar|banner/i;
            for (const im of document.querySelectorAll('img[src]')) {
                const s = im.src || '';
                if (!s || seen.has(s) || bad.test(s)) continue;
                if (im.closest('[class*="person-photo" i],[class*="hero-card-person" i]')) continue;
                // associated with an award? walk up a few ancestors (class or text)
                let el = im, ok = false;
                for (let up = 0; up < 5 && el; up++, el = el.parentElement) {
                    const cls = (el.className && el.className.toString)
                        ? el.className.toString() : '';
                    if (/award|nagrad|orden|medal/i.test(cls)) { ok = true; break; }
                    const t = el.textContent || '';
                    if (t.length < 220 && /орден|медал|наград/i.test(t)) { ok = true; break; }
                }
                if (ok) { seen.add(s); out.push(s); }
            }
            return out.slice(0, 8);
        }""")
        for src in award_srcs:
            body = await _img_bytes_via_goto(page.context, src, page.url)
            if 1200 < len(body) < 1500000:
                rec["awards"].append(body)
        if rec["awards"]:
            log(f"      🎖 орденов в Word (мелко): {len(rec['awards'])}")
        else:
            log("      (орденов-изображений на странице не найдено)")
    except Exception:
        pass

    # ── Document scan (image-map img). Go to «Документы», grab it; if it stalls
    # → ONE refresh + retry. Fast: we don't click through cards (that wasted a
    # minute on cards without documents). ──
    try:
        await _tab_block(page, ["документ", "documents"], 1.5, log)
        await asyncio.sleep(1.5)
        # ONE pass, NO reload-retry: страницы без документов уходят за секунды,
        # незагрузившийся скан просто пропускается (перезагрузка удваивала
        # ожидание на каждой пустой записи — пользователь требовала убрать).
        ok, saw_doc = await _grab_scan(page, images_dir, base, rec, log)
        if not ok:
            log("      (документов у этой записи нет — пропускаю)" if not saw_doc
                else "      (скан документа не загрузился — пропускаю)")
    except Exception:
        pass


async def _extract_person(page, url, images_dir, log, result_name="") -> dict:
    """Copy EVERYTHING from a person page: tabs (Сводная / Боевой путь / Доп.) in
    the header + a structured table for EACH document card + the document scans
    + the person's photo(s)."""
    rec = {"name": "", "url": url, "summary": [], "combat": "", "additional": "",
           "docs": [], "scans": [], "awards": [], "photos": []}
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=30000)
        await asyncio.sleep(2.5)
        rec["name"] = await page.evaluate(
            "() => { const h=document.querySelector('h1');"
            " return h ? h.textContent.replace(/\\s+/g,' ').trim() : ''; }")
        if not rec["name"]:
            rec["name"] = result_name        # h1 missing → use the result-list name

        # ── Summary (default «Сводная информация» tab). Its labels may have NO
        # colon (a two-column layout), so split the text by the KNOWN labels.
        try:
            rec["summary"] = await page.evaluate(r"""(labels) => {
                const norm = s => (s||'').replace(/ /g,' ').replace(/\s+/g,' ').trim();
                const txt = norm(document.body.innerText);
                const esc = s => s.replace(/[.*+?^${}()|[\]\\]/g, '\\$&');
                const re = new RegExp('(' + labels.map(esc).join('|') + ')\\s*:?\\s*', 'g');
                const marks = []; let m;
                while ((m = re.exec(txt)) !== null)
                    marks.push({l: m[1], s: m.index, v: re.lastIndex});
                const out = [], seen = new Set();
                for (let i = 0; i < marks.length; i++) {
                    const end = (i+1 < marks.length) ? marks[i+1].s
                                                     : Math.min(marks[i].v + 250, txt.length);
                    const val = txt.slice(marks[i].v, end).replace(/\s+/g,' ').trim();
                    if (val && val.length < 250 && !seen.has(marks[i].l)) {
                        seen.add(marks[i].l); out.push([marks[i].l, val]);
                    }
                }
                return out;
            }""", _SUMMARY_LABELS)
        except Exception:
            pass

        # ── Person photo(s): grab NOW, while we are on «Сводная информация» and
        # before tab switches / card clicks can remove the portrait from view.
        base = safe_fn(rec.get("name") or result_name or "person")
        await _grab_photos(page, images_dir, base, rec, log)

        # ── Tabs for the header. «Боевой путь» loads slowly → wait longer.
        rec["combat"]     = await _tab_block(page, ["боевой путь", "combat"], 6, log)
        rec["additional"] = await _tab_block(page, ["дополнительн", "additional"], 2, log)
        await _tab_block(page, ["сводная", "summary"], 1.2, log)

        # ── Document cards. The EXACT container is .hero-card-docs-item (the
        # [class*=…] wildcard also grabbed the name div + archive button → 85
        # fake cards). Each card loads its full fields only when CLICKED, so:
        # click → expand «Информация об архиве» → read.
        cards = page.locator(".hero-card-docs-item")
        try:
            ncards = await cards.count()
        except Exception:
            ncards = 0
        docs = []
        for k in range(min(ncards, 40)):
            c = cards.nth(k)
            try:
                nm = ""
                ne = c.locator(".hero-card-docs-item__name").first
                if await ne.count():
                    nm = re.sub(r"\s+", " ", await ne.inner_text()).strip()
                await c.scroll_into_view_if_needed(timeout=2500)
                await c.click(timeout=2500)
                await asyncio.sleep(0.35)
                ab = c.locator(".hero-card-docs-item-archive-button").first
                if await ab.count():
                    t = (await ab.inner_text()) or ""
                    if not t.strip().endswith("-"):
                        try:
                            await ab.click(timeout=1500)
                            await asyncio.sleep(0.3)
                        except Exception:
                            pass
                txt = await c.inner_text()
                # split: name | TYPE (up to the first DATA label) | fields
                body = re.sub(r"\s+", " ", txt)
                if nm:
                    body = body.replace(nm, " ", 1)
                dm = _DATA_LABEL_RE.search(body)
                if dm:
                    typ = body[:dm.start()].strip(" -+:")
                    fields = _parse_fields(body[dm.start():])
                else:
                    typ = ""
                    fields = _parse_fields(body)
                if fields or len(txt) > 10:
                    docs.append({"name": nm, "type": typ, "fields": fields,
                                 "text": re.sub(r"\s+", " ", txt).strip()[:1800]})
            except Exception:
                continue
        seen, uniq = set(), []
        for d in docs:
            key = (d.get("name", ""), tuple(d.get("fields", [])[:2]))
            if key not in seen:
                seen.add(key); uniq.append(d)
        rec["docs"] = uniq
        log(f"      сводка: {len(rec['summary'])} полей, карточек: {len(rec['docs'])}, "
            f"боевой путь: {'да' if rec['combat'] else 'нет'}")

        await _save_media(page, rec, images_dir, base, log)   # base set above
    except Exception as e:
        log(f"      !! страница человека ({type(e).__name__})")
    # strip leaked site chrome (footer/nav) from every text block
    rec["combat"] = _clean_block(rec.get("combat", ""))
    rec["additional"] = _clean_block(rec.get("additional", ""))
    rec["summary"] = [(k, _clean_block(v)) for k, v in rec.get("summary", [])
                      if _clean_block(v)]
    for d in rec.get("docs", []):
        d["fields"] = [(k, _clean_block(v)) for k, v in d.get("fields", [])
                       if _clean_block(v)]
    return rec


# ── Word output ───────────────────────────────────────────────────────────── #
def _kv_table(doc, pairs):
    if not pairs:
        return
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    for k, v in pairs:
        r = tbl.add_row().cells
        r[0].text = str(k); set_cell_lines(r[1], v)
        for run in r[0].paragraphs[0].runs:
            run.bold = True


def _docx_add_person(doc, i, rec):
    name = rec.get("name") or f"Запись {i}"
    hp = doc.add_paragraph(); hr = hp.add_run(f"{i}. {name}")
    hr.bold = True; hr.font.size = Pt(13)

    # Person photo(s): medium size, right under the name.
    for ph in rec.get("photos", []):
        try:
            png = _to_png(Path(ph).read_bytes())
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(2.2))
        except Exception:
            pass

    # ── Header: Сводная / Боевой путь / Дополнительная ──
    summary = rec.get("summary") or []
    if summary:
        doc.add_paragraph().add_run("Сводная информация").bold = True
        _kv_table(doc, summary)
    if rec.get("combat"):
        doc.add_paragraph().add_run("Боевой путь").bold = True
        doc.add_paragraph(rec["combat"])
    if rec.get("additional"):
        doc.add_paragraph().add_run("Дополнительная информация").bold = True
        doc.add_paragraph(rec["additional"])

    # Awards (orders): embed SMALL in Word, do NOT save as files.
    for aw in rec.get("awards", []):
        try:
            png = _to_png(aw)
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(1.3))
        except Exception:
            pass

    # ── A separate table for EVERY document card (with spacing between) ──
    docs = rec.get("docs") or []
    if docs:
        doc.add_paragraph().add_run(f"Документы ({len(docs)})").bold = True
    for j, d in enumerate(docs, 1):
        title = d.get("name") or "Документ"
        if d.get("type"):
            title += f" — {d['type']}"
        doc.add_paragraph().add_run(f"{j}. {title}").bold = True
        if d.get("fields"):
            _kv_table(doc, d["fields"])
        elif d.get("text"):
            doc.add_paragraph(d["text"])
        doc.add_paragraph("")          # ← spacing between cards

    # Document SCANS (the slow-loading archival pages): saved to disk; embedded
    # at a moderate size (was too large before — «УМЕНЬШАЙ»).
    for img in rec.get("scans", []):
        try:
            png = _to_png(Path(img).read_bytes())
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(3.3))
        except Exception:
            pass

    if rec.get("url"):
        # the word «Источник» itself is the clickable link; the long URL (with
        # its huge backurl=… tail) is hidden in the href, NOT shown as text.
        clean = re.sub(r"\?.*$", "", rec["url"]) or rec["url"]
        p = doc.add_paragraph()
        _add_hyperlink(p, "Источник", clean)
    doc.add_paragraph("")      # single blank line between people (never two)


def write_docx(path, persons, qlines, append=False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    existing = append and Path(path).exists()
    if existing:
        doc = Document(str(path)); doc.add_page_break()
        ap = doc.add_paragraph(); ar = ap.add_run(f"➕ Добавлено {len(persons)}")
        ar.bold = True; ar.font.size = Pt(13)
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width = Mm(210); s.page_height = Mm(297)
        s.left_margin = s.right_margin = Mm(18)
        s.top_margin = s.bottom_margin = Mm(18)
        ht = doc.add_paragraph(); htr = ht.add_run("Память народа — результаты поиска")
        htr.bold = True; htr.font.size = Pt(14)
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ln in qlines:
            doc.add_paragraph(ln)
        doc.add_paragraph(f"Найдено человек: {len(persons)}")
        doc.add_paragraph("")
    for i, rec in enumerate(persons, 1):
        _docx_add_person(doc, i, rec)
    doc.save(str(path))


# ── Main entry point ──────────────────────────────────────────────────────── #
async def run_scraper(*,
    lang="ru",
    last_name="", first_name="", middle_name="",
    birth_from="", birth_to="", place_birth="",
    rank="", branch="", place_service="", place_conscription="",
    award="",
    group_person=True,
    output_folder=Path("."),
    log=print,
    progress=None,
    cancel_event=None,
    ask_file_conflict=None,
    max_pages=10,
    max_persons=60,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    params = dict(last_name=last_name.strip(), first_name=first_name.strip(),
                  middle_name=middle_name.strip(), birth_from=birth_from.strip(),
                  birth_to=birth_to.strip(), place_birth=place_birth.strip(),
                  rank=rank.strip(), branch=branch, place_service=place_service.strip(),
                  place_conscription=place_conscription.strip(), award=award,
                  group_person=group_person)
    output_folder = Path(output_folder); output_folder.mkdir(parents=True, exist_ok=True)
    qkey = " ".join(p for p in (last_name, first_name, middle_name) if p) or "pamyat"
    images_dir = output_folder / "images" / (safe_fn(qkey) or "pamyat")
    summary = {"ok": False}

    if not any(v for k, v in params.items() if k != "group_person"):
        _prog(100, "Пустой запрос.")
        return summary

    qlines = [f"Запрос: {qkey}", f"Язык сайта: {lang}"]

    _prog(0, "Запускаю браузер…")
    async with async_playwright() as pw:
        browser = await pw.chromium.launch(
            headless=False,
            args=["--start-maximized", "--disable-blink-features=AutomationControlled"])
        ctx = await browser.new_context(no_viewport=True, accept_downloads=True)
        page = await ctx.new_page()
        try:
            _prog(5, "Поиск…")
            if not await _do_search(page, params, lang, log):
                summary["message"] = "Не удалось выполнить поиск."
                return summary
            if _done():
                return summary

            _prog(15, "Сбор результатов…")
            people = await _collect_results(page, params, log, max_pages, max_persons)
            log(f"  Подходящих человек: {len(people)}")
            if not people:
                _prog(100, "Ничего подходящего не найдено.")
                summary.update({"ok": True, "n_records": 0})
                return summary

            persons = []
            n = len(people)
            for i, pr in enumerate(people, 1):
                if _done():
                    break
                _prog(20 + int(70 * i / n), f"[{i}/{n}] {pr['name'][:50]}…")
                log(f"  [{i}/{n}] {pr['name']}  "
                    f"({pr.get('docs','?')} документ.)")
                # one bad person must NOT abort the run
                try:
                    dp = await ctx.new_page()
                    try:
                        rec = await _extract_person(dp, pr["href"], images_dir, log,
                                                    result_name=pr["name"])
                    finally:
                        await dp.close()
                    if not rec.get("name"):
                        rec["name"] = pr["name"]
                    persons.append(rec)
                except Exception as _e:
                    log(f"      !! пропускаю человека ({type(_e).__name__})")
                    persons.append({"name": pr["name"], "url": pr["href"],
                                    "summary": [], "docs": [], "scans": [],
                                    "awards": [], "photos": []})
                await asyncio.sleep(0.3)

            _prog(92, "Сохранение…")
            base = safe_fn(f"pamyat_{qkey}") or "pamyat_results"
            docx_p = output_folder / f"{base}.docx"
            decision = "overwrite"
            if docx_p.exists() and ask_file_conflict:
                try:
                    decision = (ask_file_conflict([docx_p.name]) or "overwrite").lower()
                except Exception:
                    decision = "overwrite"
                log(f"  → Файл существует → {decision}")
            if decision != "skip" and persons:
                try:
                    write_docx(docx_p, persons, qlines, append=(decision == "append"))
                    log(f"  → Word: {docx_p.name}")
                except PermissionError:
                    # the .docx is open in Word → save to a fresh name, don't lose data
                    alt = output_folder / f"{base}_{time.strftime('%H%M%S')}.docx"
                    write_docx(alt, persons, qlines, append=False)
                    docx_p = alt
                    log(f"  !! файл занят (открыт в Word) → сохранил как {alt.name}")

            _prog(100, f"Готово — {len(persons)} человек.")
            summary.update({"ok": True, "n_records": len(persons),
                            "output_folder": str(output_folder)})
        except Exception as exc:
            summary["message"] = f"{type(exc).__name__}: {exc}"
            log(f"  !! {exc}")
        finally:
            try:
                await ctx.close()
            except Exception:
                pass
            try:
                await browser.close()
            except Exception:
                pass
    return summary
