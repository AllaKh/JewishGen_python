"""hryc_scraper.py — search hryc.by (archives, newspapers and books).

Runs a VISIBLE (headed) Chromium via Playwright so the user can watch it work
(login, search, paging) — like the other scrapers. A persistent profile
(.hryc_profile) keeps the login between runs.

Login is required for results; the site gates the session behind an ASP.NET
cookie-consent banner («Принять») which we click first. The search is the
server-rendered GET form; we navigate the full results URL (R.Q + all 164
R.S[i].Chk/.Id source toggles + R.Page + the option flags) and walk every page.

For now results are plain lists → Word + Excel. Opening/saving the actual
documents needs a paid account and will be added later.
"""
import sys, os, re, json, time, asyncio, base64, hashlib
import html as _html
from pathlib import Path
from urllib import parse as _up

try:
    from playwright.async_api import async_playwright
    _PW_OK = True
except ImportError:
    _PW_OK = False

try:
    from docx import Document
    from docx.shared import Mm, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

BASE_URL    = "https://hryc.by"
LOGIN_URL   = BASE_URL + "/Identity/Account/Login"
SEARCH_URL  = BASE_URL + "/search"
SITE_NAME   = "hryc.by"
SRC_FILE    = Path(__file__).resolve().parent / "config" / "hryc_sources.json"
from paths_util import user_data_dir
PROFILE_DIR = user_data_dir() / ".hryc_profile"   # writable even in a packaged install

HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")


# OCR snippets carry stray control characters that python-docx / openpyxl reject
# ("All strings must be XML compatible … no … control characters"). Strip everything
# illegal for XML — _CTRL_KEEP preserves our \x02/\x03 bold sentinels during parsing;
# _CTRL_ALL removes everything (used at write time, after sentinels are consumed).
_CTRL_KEEP = re.compile(r'[\x00\x01\x04-\x08\x0b\x0c\x0e-\x1f]')
_CTRL_ALL  = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f]')


def _safe(s) -> str:
    return _CTRL_ALL.sub('', s or '')


# ── helpers ─────────────────────────────────────────────────────────────────── #
def safe_fn(s: str, n: int = 100) -> str:
    return re.sub(r"\s+", " ",
                  re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip()))[:n].strip() or "hryc"


def _strip_tags(html: str) -> str:
    txt = re.sub(r"<[^>]+>", " ", html or "")
    return re.sub(r"\s+", " ", _html.unescape(txt)).strip()


def load_sources() -> list:
    """Flat list of every searchable source: {idx, id, label, en}, idx-sorted."""
    tree = json.loads(SRC_FILE.read_text(encoding="utf-8"))
    flat = []

    def walk(ns):
        for n in ns:
            if "idx" in n:
                flat.append(n)
            else:
                walk(n.get("children", []))

    walk(tree)
    return sorted(flat, key=lambda x: x["idx"])


def build_search_url(query: str, selected_ids, page: int = 1, *,
                     no_stemming: bool = False, no_fuzziness: bool = False,
                     show_experts: bool = False, fund: str = "", inventory: str = "",
                     record: str = "", doc_dates: str = "", added_since: str = "") -> str:
    """Build the results URL EXACTLY like the site's CURRENT search (verified against the
    user's live working URL). The site switched the source filter from the old per-index
    array (R.S[i].Chk/.Id for all 164) to a SINGLE «R.Sources=<comma-joined ids>» param — the
    old array is ignored now and returns 0. Other params: R.Q, R.Page, R.SearchInRecords, the
    option flags, R.MustBeInPlace/R.MustBeInSurname, and the archive-reference / date fields."""
    sel = [s for s in (selected_ids or []) if s]
    params = [
        ("R.Q", query),
        ("R.Page", str(page)),
        ("R.Sources", ",".join(sel)),                  # comma-joined ticked source ids
        ("R.ShowExperts", "true" if show_experts else "false"),
        ("R.SearchInRecords", "True"),
        ("R.NoStemming",  "true" if no_stemming  else "false"),
        ("R.NoFuzziness", "true" if no_fuzziness else "false"),
        ("R.MustBeInPlace", "false"),
        ("R.MustBeInSurname", "false"),
        # always present (empty = no filter); R.UpdateStartDate empty = «0001-01-01» sentinel.
        ("R.Fund", fund),
        ("R.Inventory", inventory),
        ("R.Record", record),
        ("R.DocDateRange", doc_dates),
        ("R.UpdateStartDate", added_since or "0001-01-01"),
    ]
    # keep «*» literal in R.Q (the site uses raw * for wildcards, not %2A)
    return SEARCH_URL + "?" + _up.urlencode(params, safe="*,")


def parse_results(html: str, log) -> dict:
    """Each result on a hryc results page is a text-snippet block
    `<div style="margin: 1em"> …matched OCR text with <b>highlights</b>… <span>Для
    открытия документа …</span></div>` (≈20 per page; «Total: N» up the top, capped at
    1000; the document itself opens only on the paid «Эксперт» tariff). We keep the
    snippet text as the list row. Returns {rows, total}."""
    rows, seen = [], set()

    m = re.search(r'Total:.{0,60}?(\d[\d\s ,]*)', html, re.S)
    total = int(re.sub(r'\D', '', m.group(1))) if m else 0

    # Each result is  <div style="margin:1em"><div>SNIPPET</div><a href="/document?id=…">
    # TITLE</a></div>.  Split on the block marker (a non-greedy «…</div>» would stop at the
    # INNER </div> and miss the link → that was the «0 ссылок» bug). Take the inner <div> as
    # the OCR snippet and the first /document?id= link as the document.
    for chunk in re.split(r'<div style="margin: ?1em">', html)[1:]:
        # document link (paid account): <a href="/document?id=…">Title</a>
        doc_url, doc_title = "", ""
        dm = re.search(r'<a[^>]+href="(/document\?id=[^"]+)"[^>]*>(.*?)</a>', chunk, re.S)
        if dm:
            doc_url = BASE_URL + _html.unescape(dm.group(1))
            doc_title = _strip_tags(dm.group(2))
        # snippet = the inner OCR <div> (fall back to everything before the link)
        sm = re.search(r'<div\b[^>]*>(.*?)</div>', chunk, re.S)
        inner = sm.group(1) if sm else (chunk[:dm.start()] if dm else chunk)
        # drop the «pay to open» note (free tier) from the snippet text
        inner = re.sub(r'<span[^>]*>\s*(Для открытия|To open|Каб адкрыць).*$', '',
                       inner, flags=re.S)
        inner = re.sub(r'<br\s*/?>', ' ', inner)
        # keep the site's <b>…</b> match highlights as sentinels (\x02…\x03) so Word
        # can render the searched word in bold; everything else → plain text
        inner = re.sub(r'<\s*b\s*>', '\x02', inner, flags=re.I)
        inner = re.sub(r'<\s*/\s*b\s*>', '\x03', inner, flags=re.I)
        inner = re.sub(r'<[^>]+>', ' ', inner)
        inner = re.sub(r'[ \t\r\n]+', ' ', _html.unescape(inner)).strip()
        inner = _CTRL_KEEP.sub('', inner)                           # drop XML-illegal ctrl chars
        text = inner.replace('\x02', '').replace('\x03', '')        # plain (Excel/dedup)
        if (not text or len(text) < 4) and not doc_url:
            continue
        key = text + "|" + doc_url
        if key in seen:
            continue
        seen.add(key)
        rows.append({"source": doc_title, "text": text, "rich": inner, "url": doc_url})

    return {"rows": rows, "total": total}


# ── Word ────────────────────────────────────────────────────────────────────── #
def _add_link(para, text, url):
    part = para.part
    r_id = part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    c = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1"); rpr.append(c)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rpr.append(u)
    run.append(rpr)
    t = OxmlElement("w:t"); t.text = text or url; run.append(t)
    hl.append(run); para._p.append(hl)


# landscape A4 with 0.7" L/R margins → ~10.3" usable; keep the table inside it so the
# right margin survives (sum ≈ 9.8"). #, Source, File, Link narrow — Record wide.
# python-docx ignores table-level widths → set on EVERY cell per row.
_COLW = [Inches(0.4), Inches(1.4), Inches(5.6), Inches(1.7), Inches(0.7)]


def _set_widths(cells):
    for c, w in zip(cells, _COLW):
        c.width = w


def _add_rich(cell, rich):
    """Render the snippet into the cell, bolding the searched word(s) — the site marked
    them with <b>…</b>, kept here as \\x02…\\x03 sentinels."""
    para = cell.paragraphs[0]
    plain = lambda s: _safe(s.replace("\x02", "").replace("\x03", ""))
    i = 0
    for m in re.finditer(r'\x02(.*?)\x03', rich):
        if m.start() > i:
            para.add_run(plain(rich[i:m.start()])).font.size = Pt(8)
        para.add_run(_safe(m.group(1))).bold = True
        para.runs[-1].font.size = Pt(8)
        i = m.end()
    if i < len(rich):
        para.add_run(plain(rich[i:])).font.size = Pt(8)


def _docx_table(doc, rows):
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Light Grid Accent 1"
    tbl.autofit = False
    tbl.allow_autofit = False
    hdr = tbl.rows[0].cells
    for c, h in zip(hdr, ("#", "Источник", "Запись", "Файл", "Ссылка")):
        c.text = ""
        c.paragraphs[0].add_run(h).bold = True
    _set_widths(hdr)
    for i, r in enumerate(rows, 1):
        cells = tbl.add_row().cells
        cells[0].text = str(i)
        cells[1].text = r.get("source", "") or ""
        _add_rich(cells[2], r.get("rich") or r.get("text", "") or "")
        # saved document scan(s) — full path(s), one per line
        files = r.get("files") or []
        cells[3].paragraphs[0].add_run("\n".join(files)).font.size = Pt(8)
        if r.get("url"):
            _add_link(cells[4].paragraphs[0], "Открыть", r["url"])
        _set_widths(cells)


def _add_page_numbers(doc):
    """Centered «Стр. PAGE из NUMPAGES» footer on every page. Idempotent — skips if a
    page field already exists (so appending doesn't duplicate it)."""
    for section in doc.sections:
        footer = section.footer
        if footer._element.findall('.//' + qn('w:fldChar')):
            continue                                   # already numbered
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _field(instr):
            b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
            t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = instr
            e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
            run = p.add_run(); run._r.append(b); run._r.append(t); run._r.append(e)

        p.add_run("Стр. "); _field("PAGE"); p.add_run(" из "); _field("NUMPAGES")


def _doc_preview(fp):
    """A small JPEG (BytesIO) of a saved scan for embedding — the FULL file stays on disk;
    only a downscaled preview goes into the Word so it doesn't bloat to hundreds of MB."""
    try:
        from PIL import Image
        import io
        im = Image.open(fp)
        im.thumbnail((1100, 1500))
        if im.mode not in ("RGB", "L"):
            im = im.convert("RGB")
        buf = io.BytesIO(); im.save(buf, "JPEG", quality=80); buf.seek(0)
        return buf
    except Exception:
        return None


def write_docx(path: Path, rows: list, qlines: list, append: bool = False):
    if not _DOCX_OK:
        return
    if append and path.exists():
        doc = Document(str(path))
        doc.add_page_break()
        doc.add_paragraph().add_run(f"➕ Добавлено ещё {len(rows)} записей").bold = True
    else:
        doc = Document()
        for s in doc.sections:
            s.orientation = 1
            s.page_width, s.page_height = Mm(297), Mm(210)
            s.left_margin = s.right_margin = Inches(0.7)   # keep real left/right margins
        h = doc.add_heading("hryc.by — результаты поиска", level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for q in qlines:
            doc.add_paragraph(q)
        doc.add_paragraph(f"Найдено записей: {len(rows)}")
    _docx_table(doc, rows)
    # Embed each saved scan with its FULL path beside it, so «где находится файл» is
    # unmissable (the per-row «Файл» column stays too, for quick reference).
    saved = [(r, fp) for r in rows for fp in (r.get("files") or [])]
    if saved:
        doc.add_paragraph()
        doc.add_paragraph().add_run(f"Документы (сканов: {len(saved)}):").bold = True
        for r, fp in saved:
            cap = r.get("source") or ""
            if cap:
                doc.add_paragraph().add_run(cap).bold = True
            prev = _doc_preview(fp)
            if prev is not None:
                try:
                    doc.add_picture(prev, width=Inches(4.2))
                except Exception:
                    pass
            doc.add_paragraph().add_run("Файл: " + str(fp)).font.size = Pt(8)
            doc.add_paragraph()
    _add_page_numbers(doc)
    doc.save(str(path))


# ── Excel ─────────────────────────────────────────────────────────────────── #
def write_xlsx(path: Path, rows: list, qlines: list, append: bool = False):
    if not _OPENPYXL_OK:
        return
    cols = ["#", "База", "Источник", "Запись", "Файл", "URL"]
    if append and path.exists():
        wb = load_workbook(str(path)); ws = wb.active
        n0 = ws.max_row - 1
    else:
        wb = Workbook(); ws = wb.active; ws.title = "hryc.by"
        ws.append(cols)
        for c in range(1, len(cols) + 1):
            ws.cell(row=1, column=c).font = Font(bold=True)
            ws.cell(row=1, column=c).fill = PatternFill("solid", fgColor="DCE6F1")
        n0 = 0
    for i, r in enumerate(rows, 1):
        ws.append([n0 + i, SITE_NAME, _safe(r.get("source", "")),
                   _safe(r.get("text", "")), _safe("\n".join(r.get("files") or [])),
                   _safe(r.get("url", ""))])
        if r.get("url"):
            cell = ws.cell(row=ws.max_row, column=6)
            cell.hyperlink = r["url"]; cell.value = "Открыть"
            cell.font = Font(color="0563C1", underline="single")
    # #, Database, Source, File, URL narrow — Record wide
    for i, wd in enumerate([5, 10, 10, 90, 40, 10], 1):
        ws.column_dimensions[get_column_letter(i)].width = wd
    # wrap the long snippet + file columns
    for row in ws.iter_rows(min_row=2, min_col=4, max_col=5):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    ws.freeze_panes = "A2"
    wb.save(str(path))


# ── browser: cookies / login ────────────────────────────────────────────────── #
async def _wait_if_captcha(page, log, timeout=300) -> None:
    """If an «I'm not a robot» / reCAPTCHA challenge appears, PAUSE and let the USER solve it
    in the visible window (we never click/solve a CAPTCHA ourselves). Polls until it's gone."""
    async def _present():
        try:
            return await page.evaluate(r"""() => {
                const vis = el => el && el.offsetParent !== null
                                  && el.getBoundingClientRect().width > 20
                                  && el.getBoundingClientRect().height > 20;
                for (const f of document.querySelectorAll('iframe')) {
                    const s = (f.src || '') + ' ' + (f.title || '');
                    if (/recaptcha|hcaptcha|captcha|challenge|turnstile/i.test(s) && vis(f))
                        return true;
                }
                if (document.querySelector('#cf-challenge-running, .g-recaptcha:not(:empty)'))
                    return true;
                const t = (document.body && document.body.innerText) || '';
                return /я\s*не\s*робот|i'?m not a robot|подтвердите.{0,4}что вы не робот|verify you are human|are you a robot/i.test(t);
            }""")
        except Exception:
            return False
    if not await _present():
        return
    log("  ⛔ КАПЧА «Я не робот» — РЕШИ ЕЁ В ОКНЕ БРАУЗЕРА. Скрипт ждёт, не торопит.")
    try:
        await page.bring_to_front()
    except Exception:
        pass
    waited = 0
    while waited < timeout:
        await asyncio.sleep(2); waited += 2
        if not await _present():
            log("  ✓ Капча решена — продолжаю.")
            await asyncio.sleep(1)
            return
        if waited % 20 == 0:
            log(f"  … жду капчу ({waited}/{timeout}s) — нажми «Я не робот» в окне.")
    log("  !! капча не решена за отведённое время — продолжаю как есть.")


async def _accept_cookies(page, log):
    """Click the «Принять» cookie-consent banner (#accept-cookies) so the session
    sticks — «сначала прими куки, потом войти»."""
    for sel in ('#accept-cookies', 'button:has-text("Принять")',
                'button:has-text("Accept")'):
        try:
            btn = page.locator(sel).first
            if await btn.count() and await btn.is_visible():
                await btn.click(timeout=3000)
                log("  ✓ Куки приняты")
                await asyncio.sleep(0.5)
                return
        except Exception:
            pass


async def _is_logged_in(page) -> bool:
    try:
        return await page.evaluate(
            """() => !!document.querySelector('a[href*="Account/Logout"], a[href*="Account/Manage"]')
                     || !document.querySelector('a[href="/Identity/Account/Login"]')""")
    except Exception:
        return False


async def _login(page, email, password, log) -> bool:
    await page.goto(LOGIN_URL, wait_until="domcontentloaded", timeout=30000)
    await _accept_cookies(page, log)
    try:
        await page.fill("#Input_Email", email, timeout=8000)
        await page.fill("#Input_Password", password, timeout=8000)
        try:
            await page.check("#Input_RememberMe", timeout=2000)
        except Exception:
            pass
        await page.click('button[type="submit"]:has-text("Войти"), #login-submit, '
                         'form#account button[type="submit"]', timeout=8000)
    except Exception as e:
        log(f"  !! не удалось заполнить форму логина: {e}")
        return False
    try:
        await page.wait_for_load_state("domcontentloaded", timeout=10000)
    except Exception:
        pass
    ok = await _is_logged_in(page)
    log("  ✓ Вход выполнен" if ok else
        "  !! вход не подтверждён — проверь email/пароль (или войди вручную в окне)")
    return ok


def _clear_singleton_locks():
    for n in ("SingletonLock", "SingletonCookie", "SingletonSocket"):
        try:
            (PROFILE_DIR / n).unlink()
        except Exception:
            pass


def _looks_image(b: bytes) -> bool:
    return (b[:3] == b"\xff\xd8\xff" or b[:8] == b"\x89PNG\r\n\x1a\n"
            or b[:6] in (b"GIF87a", b"GIF89a")
            or (b[:4] == b"RIFF" and b[8:12] == b"WEBP"))


async def _save_doc_scan(page, images_dir, base, log, seen_hashes=None):
    """Save the document's scanned page image. On hryc the scan is
    `<img id="mainCanvas" src="https://downloader.disk.yandex.ru/…">` inside #mainCanvasDiv
    (the src is a no-referrer, token-signed Yandex Disk URL). Strategy:
      1. wait for #mainCanvas to actually finish loading;
      2. fetch the bytes IN-PAGE with referrerPolicy:'no-referrer' (full resolution — same
         as the user's «save the page, grab the scan locally» idea, done in one step);
      3. fall back to an element screenshot of the rendered scan (always works).
    """
    saved = []
    # wait for the scan <img> to load (Yandex Disk image, loads after domcontentloaded)
    try:
        await page.wait_for_function(
            "() => { const c = document.querySelector('#mainCanvas');"
            "        return c && c.complete && (c.naturalWidth || 0) > 50; }",
            timeout=15000)
    except Exception:
        pass
    src = await page.evaluate(
        "() => { const c = document.querySelector('#mainCanvas');"
        "        return c ? (c.currentSrc || c.src || '') : ''; }")
    img_bytes = None
    if src:
        try:
            data = await page.evaluate(
                """async (u) => {
                    try {
                        const r = await fetch(u, {referrerPolicy: 'no-referrer'});
                        if (!r.ok) return '';
                        const a = new Uint8Array(await r.arrayBuffer());
                        let s = ''; const CH = 0x8000;
                        for (let i = 0; i < a.length; i += CH)
                            s += String.fromCharCode.apply(null, a.subarray(i, i + CH));
                        return btoa(s);
                    } catch (e) { return ''; }
                }""", src)
            if data:
                img_bytes = base64.b64decode(data)
        except Exception:
            img_bytes = None
    if not (img_bytes and len(img_bytes) > 3000 and _looks_image(img_bytes)):
        # fallback: screenshot the rendered scan element
        try:
            el = page.locator("#mainCanvas").first
            if await el.count():
                img_bytes = await el.screenshot(timeout=12000)
        except Exception:
            img_bytes = None
    if img_bytes and len(img_bytes) > 3000 and _looks_image(img_bytes):
        # content dedup: never save the SAME scan twice in one run (a broad query returns
        # the same page in many result rows → otherwise we'd write identical files)
        if seen_hashes is not None:
            h = hashlib.md5(img_bytes).hexdigest()
            if h in seen_hashes:
                log(f"    ⏩ дубликат скана — уже сохранён как {Path(seen_hashes[h]).name}")
                return [seen_hashes[h]]
        images_dir.mkdir(parents=True, exist_ok=True)
        ext = ".png" if img_bytes[:4] == b"\x89PNG" else ".jpg"
        fp = images_dir / f"{base}{ext}"
        fp.write_bytes(img_bytes)
        full = str(fp.resolve())
        saved.append(full)
        if seen_hashes is not None:
            seen_hashes[h] = full
        log(f"    ✓ скан сохранён: {fp.name} ({len(img_bytes)//1024} КБ)")
    else:
        log("    !! скан не сохранён (нет #mainCanvas или не картинка)")
    return saved


async def _open_document(page, url, images_dir, base, dump, log, seen_hashes=None):
    """Open a document page (/document?id=…) and save its scanned page image.
    Returns (saved_file_paths, page_text)."""
    saved, text = [], ""
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=40000)
    except Exception as e:
        log(f"    !! документ не открылся: {e}")
        return saved, text
    if dump is not False and os.environ.get("JGS_DEBUG"):   # diagnostic dump, opt-in
        try:
            (Path(dump) / "hryc_document_sample.html").write_text(
                await page.content(), encoding="utf-8")
        except Exception:
            pass
    saved = await _save_doc_scan(page, images_dir, base, log, seen_hashes)
    return saved, text


# ── orchestrator ────────────────────────────────────────────────────────────── #
async def run_scraper(
    *,
    email:         str       = "",
    password:      str       = "",
    query:         str       = "",
    sources:       list|None = None,
    no_stemming:   bool      = False,
    no_fuzziness:  bool      = False,
    show_experts:  bool      = False,
    fund:          str       = "",
    inventory:     str       = "",
    record:        str       = "",
    doc_dates:     str       = "",
    added_since:   str       = "",
    open_documents: bool     = True,    # open & save the linked documents (paid account)
    max_docs:      int       = 60,
    max_pages:     int       = 50,
    output_format: str       = "both",
    output_folder            = Path("."),
    log                      = print,
    progress                 = None,
    cancel_event             = None,
    ask_file_conflict        = None,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _cancelled():
        return bool(cancel_event and cancel_event.is_set())

    want_docx = output_format in ("docx", "both")
    want_xlsx = output_format in ("xlsx", "both")
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    qlines = [f"Запрос: {query}", f"Выбрано источников: {len(sources or [])}"]
    summary = {"ok": False}

    if not _PW_OK:
        summary.update({"error": "playwright",
                        "message": "Playwright is not installed."})
        return summary

    PROFILE_DIR.mkdir(parents=True, exist_ok=True)
    _clear_singleton_locks()

    _prog(2, "Запускаю браузер…")
    async with async_playwright() as pw:
        # The profile can be locked (another window of this app, or a previous/crashed Chrome,
        # still holds .hryc_profile → TargetClosedError). Clear stale locks and retry a few
        # times; then give a clear message instead of a raw crash.
        ctx = None
        for _attempt in range(1, 4):
            _clear_singleton_locks()
            # Use REAL Google Chrome (channel="chrome"); fall back to the bundled Chromium
            # only if Chrome isn't installed (so a distributed copy still works).
            for _channel in ("chrome", None):
                try:
                    kw = dict(headless=False, accept_downloads=True, no_viewport=True,
                              args=["--start-maximized",
                                    "--disable-blink-features=AutomationControlled"])
                    if _channel:
                        kw["channel"] = _channel
                    ctx = await pw.chromium.launch_persistent_context(str(PROFILE_DIR), **kw)
                    log(f"  → браузер: {'Google Chrome' if _channel else 'Chromium (Chrome не найден)'}")
                    break
                except Exception as e:
                    log(f"  !! {_channel or 'chromium'}: {type(e).__name__}")
            if ctx:
                break
            await asyncio.sleep(3)
        if ctx is None:
            summary.update({"error": "browser", "message":
                            "Не удалось открыть браузер. Закрой другие окна этого приложения "
                            "и любые окна Chrome, использующие профиль hryc, затем повтори."})
            return summary
        page = ctx.pages[0] if ctx.pages else await ctx.new_page()
        for extra in list(ctx.pages):
            if extra is not page:
                try: await extra.close()
                except Exception: pass

        try:
            # ── 1. cookies + login ───────────────────────────────────── #
            await page.goto(BASE_URL + "/", wait_until="domcontentloaded", timeout=30000)
            await _accept_cookies(page, log)
            if email and password and not await _is_logged_in(page):
                _prog(8, "Вход на hryc.by…")
                if not await _login(page, email, password, log):
                    # leave the window open a bit so the user can log in manually
                    for _ in range(60):
                        if _cancelled() or await _is_logged_in(page):
                            break
                        await asyncio.sleep(1)
            logged = await _is_logged_in(page)
            log("  → Статус: залогинен" if logged else
                "  → Статус: гость (большинство источников будут недоступны)")
            await _wait_if_captcha(page, log)        # pause for «я не робот» if shown

            # ── 2. SEARCH (walk every page) ──────────────────────────── #
            # Results are server-rendered in the page (≈20 snippet blocks per page),
            # so we only need domcontentloaded — NO networkidle/sleep (that 10s wait
            # per page was the hang the user complained about).
            _prog(20, "Поиск…")
            rows, seen, total, pages = [], set(), 0, max_pages
            for pno in range(1, max_pages + 1):
                if _cancelled():
                    break
                url = build_search_url(query, sources or [], pno,
                                       no_stemming=no_stemming, no_fuzziness=no_fuzziness,
                                       show_experts=show_experts, fund=fund,
                                       inventory=inventory, record=record,
                                       doc_dates=doc_dates, added_since=added_since)
                await page.goto(url, wait_until="domcontentloaded", timeout=30000)
                if pno == 1:
                    await _wait_if_captcha(page, log)    # search may trigger «я не робот»
                html = await page.content()
                if pno == 1 and os.environ.get("JGS_DEBUG"):   # diagnostic dump, opt-in
                    try:
                        (output_folder / "hryc_last_results.html").write_text(
                            html, encoding="utf-8")
                    except Exception:
                        pass
                res = parse_results(html, log)
                if pno == 1:
                    total = res.get("total", 0)
                    if total:
                        pages = min(max_pages, (total + 19) // 20)  # ≈20 per page
                new = 0
                for r in res["rows"]:
                    if r["text"] in seen:
                        continue
                    seen.add(r["text"]); rows.append(r); new += 1
                _prog(20 + min(60, pno * 60 // max(pages, 1)),
                      f"  стр.{pno}/{pages}: +{new} (всего {len(rows)})")
                if new == 0 or pno >= pages:
                    break

            _prog(85, f"Найдено (Total на сайте: {total}): собрано {len(rows)}")

            if not rows:
                summary.update({"ok": True, "n_records": 0, "total": total,
                                "message": "Ничего не найдено (или нужен платный доступ)."})
                _prog(100, "Ничего не найдено.")
                return summary

            # ── 2b. Open & save the linked documents (paid account) ─────── #
            doc_rows = [r for r in rows if r.get("url")]
            n_docs = 0
            if open_documents and doc_rows:
                # put the search YEAR (Doc dates) in the folder name when given
                year_tag = "_" + safe_fn(doc_dates) if doc_dates else ""
                images_dir = output_folder / "images" / "hryc.by" / (safe_fn(query) + year_tag)
                # A broad query (*а* + year) returns the SAME /document?id= in many snippet
                # rows. Open each UNIQUE url ONCE, and dedup by image content too — so the
                # folder never gets duplicate scans. Duplicate rows reuse the same saved path.
                rep, uniq_urls = {}, []
                for r in doc_rows:
                    u = r["url"]
                    if u not in rep:
                        rep[u] = r; uniq_urls.append(u)
                uniq_urls = uniq_urls[:max_docs]
                seen_hashes, url_files = {}, {}
                _prog(85, f"Открываю документы: {len(uniq_urls)} уникальных "
                          f"(из {len(doc_rows)} ссылок)…")
                for i, u in enumerate(uniq_urls, 1):
                    if _cancelled():
                        break
                    r0 = rep[u]
                    nm = safe_fn(r0.get("source") or query) or "doc"
                    m = re.search(r"id=(HAID[\w]+)", u)        # stable, unique per document
                    base = f"{nm}_{m.group(1)[-10:]}" if m else f"{nm}_{i}"
                    files, dtext = await _open_document(
                        page, u, images_dir, base,
                        dump=(output_folder if i == 1 else False), log=log,
                        seen_hashes=seen_hashes)
                    url_files[u] = files
                    if files:
                        n_docs += 1
                    if dtext and len(dtext) > len(r0.get("text", "")):
                        r0["text"] = dtext       # richer text from the opened document
                    _prog(85 + min(8, i * 8 // max(len(uniq_urls), 1)),
                          f"  документ {i}/{len(uniq_urls)}: {len(files)} скан(ов)")
                # attach the saved path(s) to EVERY row sharing that url (no re-download)
                for r in doc_rows:
                    f = url_files.get(r["url"])
                    if f:
                        r["files"] = f
                log(f"  → Уникальных документов со сканами: {n_docs} "
                    f"(сканов на диске: {len(seen_hashes)})")

            # ── 3. SAVE ──────────────────────────────────────────────── #
            _prog(88, "Сохранение файлов…")
            # year (Doc dates) goes into the file names too, like the scans folder
            year_tag = "_" + safe_fn(doc_dates) if doc_dates else ""
            base   = (safe_fn(f"hryc_{query}") or "hryc_results") + year_tag
            docx_p = output_folder / f"{base}.docx"
            xlsx_p = output_folder / f"{base}.xlsx"
            existing = [p.name for p, want in ((docx_p, want_docx), (xlsx_p, want_xlsx))
                        if want and p.exists()]
            decision = "overwrite"
            if existing and ask_file_conflict:
                try:
                    decision = (ask_file_conflict(existing) or "overwrite").lower()
                except Exception:
                    decision = "overwrite"
                log(f"  → Файл(ы) существуют {existing} → {decision}")
            append = (decision == "append")

            sd = sx = False
            if decision == "skip":
                log("  → Сохранение пропущено.")
            else:
                if want_docx:
                    write_docx(docx_p, rows, qlines, append=append); sd = True
                    log(f"  Word: {docx_p}")
                if want_xlsx:
                    write_xlsx(xlsx_p, rows, qlines, append=append); sx = True
                    log(f"  Excel: {xlsx_p}")

            _prog(100, f"Готово — {len(rows)} записей.")
            summary.update({"ok": True, "n_records": len(rows), "total": total,
                            "documents": n_docs,
                            "docx_count": 1 if sd else 0,
                            "xlsx_path": str(xlsx_p) if sx else None,
                            "output_folder": str(output_folder)})
        except Exception as exc:
            summary.update({"error": "exception",
                            "message": f"{type(exc).__name__}: {exc}"})
            log(f"  !! {exc}")
        finally:
            try:
                await ctx.close()
            except Exception:
                pass
    return summary


if __name__ == "__main__":
    asyncio.run(run_scraper(query=sys.argv[1] if len(sys.argv) > 1 else "Иванов",
                            output_folder=Path("results")))
