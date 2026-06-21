#!/usr/bin/env python3
"""
pogroms_scraper.py
===================
jewishpogroms.info — Memorial site about Jewish pogroms in the Russian Empire
(WWI / Civil War). Fully open WordPress site, no login.

The search is a plain WP-AJAX POST (action=pogroms_search) that returns the
results table as ready HTML, and every person card is a plain GET page — so this
scraper needs NO browser (urllib only, like wikisource_scraper).

Flow:
1. POST admin-ajax.php with family_name (latin), optional soundslike=yes and the
   dropdown filters (region / city / rec_type / incident / notes / found /
   register / case_number / archive / is_alive).
2. Parse the returned table: 15 columns, every cell links to the person's card
   (/pogroms_people/pogroms-post/?id=NNNNN).
3. The table goes to Word AND Excel.
4. Each card (h1 name + Short info labels + Description + Notes) goes ONLY to
   Word. Labels with no value are still written with an empty right cell.
"""

import json, re, sys, time
from pathlib import Path
from urllib import request as _rq, parse as _up
from docx_util import set_cell_lines, add_page_numbers

try:
    from docx import Document
    from docx.shared import Mm, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

BASE_URL  = "https://jewishpogroms.info"
AJAX_URL  = BASE_URL + "/wp-admin/admin-ajax.php"
SITE_NAME = "Jewish Pogroms"

# the form's dropdown field names, in the site's own order
FILTER_KEYS = ["region", "city", "rec_type", "incident", "notes",
               "found", "register", "case_number", "archive", "is_alive"]

# results-table columns: (cell css-class suffix, header)
TABLE_COLS = [
    ("region",      "Region"),
    ("page",        "Page"),
    ("city",        "City"),
    ("rec_type",    "Record type"),
    ("family_name", "Family name"),
    ("name",        "Name"),
    ("patronymic",  "Patronymic"),
    ("age",         "Age"),
    ("incident",    "Incident"),
    ("is_alive",    "Is Alive"),
    ("notes",       "Notes"),
    ("found",       "Found"),
    ("register",    "Register"),
    ("case_number", "Case"),        # the cell class is pogroms-case_number
    ("archive",     "Archive"),
]

# card "Short info" labels — ALWAYS written to Word, even when empty
CARD_LABELS = ["Registration", "Age", "Incident", "Is alive", "Source",
               "Description", "Notes"]

HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")

_HDRS = {
    "User-Agent": ("Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                   "AppleWebKit/537.36 (KHTML, like Gecko) "
                   "Chrome/141.0.0.0 Safari/537.36"),
    "Accept": "*/*",
    "Accept-Language": "en-US,en;q=0.9,ru;q=0.8",
    "Referer": BASE_URL + "/",
}


def safe_fn(s: str, n: int = 100) -> str:
    return re.sub(r"\s+", " ",
                  re.sub(r'[\\/*?:"<>|]', "_", s.strip()))[:n].strip() or "document"


def _http(url: str, data: dict | None = None, timeout: int = 40) -> str:
    """GET (data=None) or POST (urlencoded) with browser-like headers — the site
    returns 406 to bare urllib requests."""
    hdrs = dict(_HDRS)
    body = None
    if data is not None:
        body = _up.urlencode(data).encode()
        hdrs["X-Requested-With"] = "XMLHttpRequest"
    req = _rq.Request(url, data=body, headers=hdrs)
    return _rq.urlopen(req, timeout=timeout).read().decode("utf-8", "replace")


def _http_bytes(url: str, timeout: int = 40) -> bytes:
    req = _rq.Request(url, headers=dict(_HDRS))
    return _rq.urlopen(req, timeout=timeout).read()


def _strip_tags(html: str) -> str:
    txt = re.sub(r"<[^>]+>", " ", html)
    txt = (txt.replace("&nbsp;", " ").replace("&amp;", "&")
              .replace("&lt;", "<").replace("&gt;", ">")
              .replace("&quot;", '"').replace("&#039;", "'"))
    return re.sub(r"\s+", " ", txt).strip()


def _looks_image(b: bytes) -> bool:
    return (b[:3] == b"\xff\xd8\xff" or b[:8] == b"\x89PNG\r\n\x1a\n"
            or b[:6] in (b"GIF87a", b"GIF89a")
            or (b[:4] == b"RIFF" and b[8:12] == b"WEBP"))


# ── Search ────────────────────────────────────────────────────────────────── #

def search(family_name: str, soundslike: bool, filters: dict, log) -> tuple:
    """POST pogroms_search → (persons_found, rows). Each row is a dict
    {col_key: value} + 'url' of the person card."""
    data = {"action": "pogroms_search", "family_name": family_name.strip()}
    if soundslike:
        # an unchecked checkbox is simply absent from the form data
        data["soundslike"] = "yes"
    for k in FILTER_KEYS:
        data[k] = (filters or {}).get(k, "") or ""
    log(f"  → POST pogroms_search: family_name={family_name!r}, "
        f"soundslike={'yes' if soundslike else 'no (exact)'}")
    used = {k: v for k, v in (filters or {}).items() if v}
    if used:
        log(f"    фильтры: {used}")
    html = _http(AJAX_URL, data)

    m = re.search(r"(\d+)\s*persons?\s*found", html)
    found = int(m.group(1)) if m else 0

    rows = []
    for trm in re.finditer(r"<tr>(.*?)</tr>", html, re.S):
        tr = trm.group(1)
        if "<td" not in tr:
            continue                      # the header row
        row, url = {}, ""
        for tdm in re.finditer(
                r"<td class='pogroms-optional pogroms-([\w]+)'>(.*?)</td>",
                tr, re.S):
            key, cell = tdm.group(1), tdm.group(2)
            am = re.search(r"<a href='([^']+)'[^>]*>(.*?)</a>", cell, re.S)
            if am:
                url = url or am.group(1)
                row[key] = _strip_tags(am.group(2))
            else:
                sm = re.search(r"<span[^>]*>(.*?)</span>", cell, re.S)
                row[key] = _strip_tags(sm.group(1) if sm else cell)
        if row:
            row["url"] = url
            rows.append(row)
    return found, rows


# ── Person card ───────────────────────────────────────────────────────────── #

def parse_card(html: str) -> dict:
    """Card page → {'name', 'fields' {label: value, ALL CARD_LABELS present},
    'photo_urls' [...]}. Empty values stay as '' (the user wants the label row
    written with a blank right cell)."""
    card = {"name": "", "fields": {k: "" for k in CARD_LABELS}, "photo_urls": []}

    m = re.search(r'<h1 class="pogroms-person-title">(.*?)</h1>', html, re.S)
    if m:
        card["name"] = _strip_tags(m.group(1))

    # Short info: <li><span class="pogroms-person-info-type">Label: </span> value</li>
    for lm in re.finditer(
            r'<li>\s*<span class="pogroms-person-info-type">([^<]+?):?\s*</span>(.*?)</li>',
            html, re.S):
        label = _strip_tags(lm.group(1)).rstrip(":")
        value = _strip_tags(lm.group(2))
        for known in CARD_LABELS:
            if known.lower() == label.lower():
                card["fields"][known] = value
                break
        else:
            if label:
                card["fields"][label] = value

    for cls in ("description", "notes"):
        dm = re.search(
            rf'<div class="pogroms-person-{cls}">(.*?)</div>', html, re.S)
        if dm:
            card["fields"]["Description" if cls == "description" else "Notes"] = \
                _strip_tags(dm.group(1))

    pm = re.search(r'<ul class="pogroms-person-photos">(.*?)</ul>', html, re.S)
    if pm:
        for im in re.finditer(r'<img src="([^"]+)"', pm.group(1)):
            card["photo_urls"].append(im.group(1))
    return card


def fetch_card(url: str, log) -> dict:
    try:
        return parse_card(_http(url))
    except Exception as e:
        log(f"    !! карточка не открылась ({type(e).__name__}) — "
            f"в Word пойдёт строка таблицы")
        return {"name": "", "fields": {k: "" for k in CARD_LABELS},
                "photo_urls": []}


# ── Word ──────────────────────────────────────────────────────────────────── #

def _add_link(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


def _docx_results_table(doc, rows):
    tbl = doc.add_table(rows=1, cols=len(TABLE_COLS))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for ci, (_k, name) in enumerate(TABLE_COLS):
        hdr[ci].text = name
        for run in hdr[ci].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = tbl.add_row().cells
        for ci, (k, _name) in enumerate(TABLE_COLS):
            set_cell_lines(cells[ci], row.get(k, ""))


def _docx_add_person(doc, i, rec):
    """One person: heading, card table (ALL labels, blanks kept), photo, link."""
    title = rec.get("card_name") or " ".join(
        p for p in (rec.get("family_name"), rec.get("name"),
                    rec.get("patronymic")) if p) or "—"
    doc.add_heading(f"{i}. {title}", level=2)

    if rec.get("url"):
        pp = doc.add_paragraph()
        pp.add_run(f"Источник ({SITE_NAME}): ").bold = True
        _add_link(pp, "Открыть запись", rec["url"])

    fields = rec.get("card_fields") or {k: "" for k in CARD_LABELS}
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    for label in fields:
        r = tbl.add_row().cells
        r[0].text = label
        for run in r[0].paragraphs[0].runs:
            run.bold = True
        set_cell_lines(r[1], fields.get(label, ""))

    if rec.get("photo_path") and Path(rec["photo_path"]).exists():
        doc.add_paragraph("")
        try:
            doc.add_picture(rec["photo_path"], width=Inches(2.5))
        except Exception:
            pass
        p = doc.add_paragraph(); p.add_run("Файл: ").bold = True
        p.add_run(str(Path(rec["photo_path"]).resolve()))   # exact path where it was saved
    doc.add_paragraph("")


def write_docx(path: Path, rows: list, qlines: list, append: bool = False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    existing = append and Path(path).exists()
    if existing:
        doc = Document(str(path))
        doc.add_page_break()
        sep = doc.add_heading(
            f"➕ Добавлено ещё {len(rows)} — {time.strftime('%Y-%m-%d %H:%M')}",
            level=1)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Запрос:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph("")
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width  = Mm(297); s.page_height = Mm(210)
        s.left_margin = s.right_margin = Mm(12)
        s.top_margin  = s.bottom_margin = Mm(14)
        h = doc.add_heading(f"{SITE_NAME} — результаты поиска", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Запрос:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph(f"Найдено: {len(rows)}")
        doc.add_paragraph("")

    # the popup table, as on the site
    _docx_results_table(doc, rows)
    doc.add_paragraph("")

    # one section per person — the card data goes ONLY to Word
    for i, rec in enumerate(rows, 1):
        _docx_add_person(doc, i, rec)
    add_page_numbers(doc)
    doc.save(path)


# ── Excel ─────────────────────────────────────────────────────────────────── #

def write_xlsx(path: Path, rows: list, qlines: list, append: bool = False):
    if not _OPENPYXL_OK:
        raise RuntimeError("openpyxl не установлен")
    HF = PatternFill("solid", fgColor="8B6914")
    HN = Font(bold=True, color="FFFFFF", size=11)
    LINKF = Font(color="0563C1", underline="single")
    TS = Side(style="thin", color="C8B878")
    T  = Border(left=TS, right=TS, top=TS, bottom=TS)

    base_cols = ["#", "База"] + [name for _k, name in TABLE_COLS] + ["URL"]

    existing = append and Path(path).exists()
    if existing:
        wb = load_workbook(str(path)); ws = wb.active
        header = [ws.cell(row=1, column=c).value
                  for c in range(1, ws.max_column + 1)]
        for name in base_cols:
            if name not in header:
                header.append(name)
                c = ws.cell(row=1, column=len(header), value=name)
                c.font = HN; c.fill = HF; c.border = T
                c.alignment = Alignment(horizontal="center",
                                        vertical="center", wrap_text=True)
        start_row = ws.max_row + 1
        start_num = ws.max_row - 1
    else:
        wb = Workbook(); ws = wb.active; ws.title = SITE_NAME
        header = base_cols
        for ci, cn in enumerate(header, 1):
            c = ws.cell(row=1, column=ci, value=cn)
            c.font = HN; c.fill = HF; c.border = T
            c.alignment = Alignment(horizontal="center",
                                    vertical="center", wrap_text=True)
        start_row = 2
        start_num = 0
    col_idx = {name: i + 1 for i, name in enumerate(header)}

    for n, rec in enumerate(rows):
        ri = start_row + n
        vals = {"#": start_num + n + 1, "База": SITE_NAME,
                "URL": rec.get("url", "")}
        for k, name in TABLE_COLS:
            vals[name] = rec.get(k, "")
        for name, val in vals.items():
            ci = col_idx.get(name)
            if not ci:
                continue
            c = ws.cell(row=ri, column=ci)
            if name == "URL" and val:
                c.value = "Открыть"; c.hyperlink = val; c.font = LINKF
            else:
                c.value = val
            c.border = T
            c.alignment = Alignment(wrap_text=True, vertical="top")

    for ci in range(1, len(header) + 1):
        ltr = get_column_letter(ci)
        mw  = max(len(str(header[ci-1] or "")),
                  *(len(str(ws.cell(row=r, column=ci).value or "").split("\n")[0])
                    for r in range(2, ws.max_row + 1)), 6)
        ws.column_dimensions[ltr].width = min(mw + 4, 50)
    wb.save(path)


# ── Main entry point ──────────────────────────────────────────────────────── #

def run_scraper(
    *,
    family_name:   str       = "",
    soundslike:    bool      = True,
    filters:       dict|None = None,
    output_format: str       = "both",
    output_folder            = Path("."),
    log                      = print,
    progress                 = None,
    cancel_event             = None,
    ask_file_conflict        = None,   # callable(list[str]) → overwrite/append/skip
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress: progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    want_docx = output_format in ("docx", "both")
    want_xlsx = output_format in ("xlsx", "both")
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    images_dir = output_folder / "images"

    qlines = [f"Family name: {family_name}",
              f"Sounds like: {'yes' if soundslike else 'no (exact)'}"]
    for k in FILTER_KEYS:
        v = (filters or {}).get(k, "")
        if v:
            qlines.append(f"{k}: {v}")
    summary = {"ok": False}

    try:
        # ── 1. SEARCH ─────────────────────────────────────────────────── #
        _prog(5, "Поиск…")
        found, rows = search(family_name, soundslike, filters or {}, log)
        _prog(15, f"Найдено: {found} (строк: {len(rows)})")
        if not rows:
            summary.update({"ok": True, "n_records": 0,
                            "message": "Ничего не найдено."})
            _prog(100, "Ничего не найдено.")
            return summary

        # ── 2. PERSON CARDS (each row links to its own card) ──────────── #
        for i, row in enumerate(rows, 1):
            if _done():
                break
            label = " ".join(p for p in (row.get("family_name"),
                                         row.get("name")) if p) or "—"
            _prog(15 + int(70 * i / len(rows)), f"[{i}/{len(rows)}] {label}…")
            row["card_fields"] = {k: "" for k in CARD_LABELS}
            row["card_name"]   = ""
            row["photo_path"]  = ""
            if not row.get("url"):
                continue
            try:
                card = fetch_card(row["url"], log)
                row["card_fields"] = card["fields"]
                row["card_name"]   = card["name"]
                # photo: best-effort — the photos block often holds an empty
                # uploads/ dir URL; only real image bytes are saved
                for pu in card["photo_urls"]:
                    try:
                        b = _http_bytes(pu)
                        if len(b) > 3000 and _looks_image(b):
                            images_dir.mkdir(parents=True, exist_ok=True)
                            ext = ".png" if b[:4] == b"\x89PNG" else ".jpg"
                            fp = images_dir / (safe_fn(card["name"] or label) + ext)
                            fp.write_bytes(b)
                            row["photo_path"] = str(fp)
                            log(f"    📷 фото сохранено: {fp.name}")
                            break
                    except Exception:
                        continue
            except Exception as e:
                log(f"    !! {type(e).__name__}: {e}")
            time.sleep(0.4)               # polite pacing for the small WP site

        # ── 3. SAVE ───────────────────────────────────────────────────── #
        _prog(90, "Сохранение файлов…")
        base   = safe_fn(f"pogroms_{family_name}") or "pogroms_results"
        docx_p = output_folder / f"{base}.docx"
        xlsx_p = output_folder / f"{base}.xlsx"

        existing_names = [p.name for p, want in
                          ((docx_p, want_docx), (xlsx_p, want_xlsx))
                          if want and rows and p.exists()]
        decision = "overwrite"
        if existing_names and ask_file_conflict:
            try:
                decision = (ask_file_conflict(existing_names)
                            or "overwrite").lower()
            except Exception as _e:
                log(f"  !! диалог конфликта файлов: {_e}")
                decision = "overwrite"
            log(f"  → Файл(ы) уже существуют {existing_names} → выбор: {decision}")
        append = (decision == "append")

        sd = sx = False
        if decision == "skip":
            log("  → Сохранение пропущено по выбору пользователя.")
        else:
            if want_docx and rows:
                write_docx(docx_p, rows, qlines, append=append)
                sd = True
                log(f"  Word: {docx_p}")
            if want_xlsx and rows:
                write_xlsx(xlsx_p, rows, qlines, append=append)
                sx = True
                log(f"  Excel: {xlsx_p}")

        _prog(100, f"Готово — {len(rows)} записей.")
        summary.update({
            "ok":            True,
            "docx_count":    1 if sd else 0,
            "xlsx_path":     str(xlsx_p) if sx else None,
            "output_folder": str(output_folder),
            "n_records":     len(rows),
        })
    except Exception as exc:
        summary.update({"error": "exception",
                        "message": f"{type(exc).__name__}: {exc}"})
        log(f"  !! {exc}")
    return summary


if __name__ == "__main__":
    run_scraper(family_name=sys.argv[1] if len(sys.argv) > 1 else "Lopatiner",
                output_folder=Path.home() / "Downloads" / "Pogroms_results")
