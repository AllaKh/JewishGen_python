#!/usr/bin/env python3
"""
yadvashem_scraper.py
====================
Scraper for the Yad Vashem «Central Database of Shoah Victims' Names»
(https://collections.yadvashem.org/ru/names/search-results-names).

Open database, no login. The search is performed by URL parameters
(gn_<field> for the value, gnt_<field> for the match type). The scraper builds
the search URL, collects the matching victim records (paginating), opens each
record and copies its fields to Word, saving any document image (Page of
Testimony scan) it finds.

NOTE: the exact URL parameter names and the result/detail DOM are being verified
against the live site (the tooling can't reach the domain). PARAM_MAP and the
result/detail selectors are best-effort and easy to correct.

NEVER crash on a broken/empty page: every record is wrapped in try/except.
"""

import asyncio
import io
import re
import sys
import time
from pathlib import Path
from docx_util import set_cell_lines, add_page_numbers
from urllib.parse import urlencode, urlparse, urljoin

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        import os
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    import browser_util
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from PIL import Image
    _PIL_OK = True
except ImportError:
    _PIL_OK = False

HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")

BASE = "https://collections.yadvashem.org"

# Search/UI languages — GUI label → code. The name params carry the code as a
# suffix (s_<field>_search_<lang>); the path uses /<lang>/.
LANGS = {
    "English": "en", "Hebrew": "he", "Russian": "ru",
    "Spanish": "es", "German": "de", "French": "fr",
}

# GUI «search type» label → URL value (t_*). Confirmed: yvSynonym.
SEARCH_TYPES = {
    "YV synonyms": "yvSynonym",
    "Literal":     "literal",
    "Phonetic":    "phonetic",
}
# GUI «birth/death year precision» → value. Confirmed: exactly.
YEAR_PREC = {
    "Exact": "exactly",
    "± 2":   "year2",
    "± 5":   "year5",
}

# GUI field key → the «<field>» part of the YV param. The real scheme is
# s_<field>_search_<lang> (value) + t_<field>_search_<lang> (type). CONFIRMED from
# a live search URL: last_name, first_name, father_first_name, place_birth,
# place_permanent (= «До войны»), year_birth (type «exactly»). The rest are
# best-effort by analogy.
PARAM_MAP = {
    "last_name":       "last_name",
    "first_name":      "first_name",
    "maiden_name":     "maiden_name",
    "birth_place":     "place_birth",
    "place_before":    "place_permanent",
    "place_during":    "place_during_war",
    "death_place":     "place_death",
    "father_name":     "father_first_name",
    "mother_name":     "mother_first_name",
    "mother_maiden":   "mother_maiden_name",
    "spouse_name":     "spouse_first_name",
    "spouse_maiden":   "spouse_maiden_name",
    "submitter_first": "submitter_first_name",
    "submitter_last":  "submitter_last_name",
}
_PLACE_KEYS = ("birth_place", "place_before", "place_during", "death_place")
# Year fields use no language suffix.
_YEAR_FIELDS = {"birth_year": "year_birth", "death_year": "year_death"}


# ── Helpers ───────────────────────────────────────────────────────────────── #
def safe_fn(s: str, n: int = 80) -> str:
    s = re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip())
    return re.sub(r"\s+", "_", s)[:n].strip("_") or "record"


def _host(u: str) -> str:
    u = u or ""
    if not u or u.startswith(("chrome-error", "about:", "data:")):
        return ""
    try:
        return urlparse(u).netloc or ""
    except Exception:
        return ""


def _to_png(data: bytes):
    if not data:
        return None
    if (data[:3] == b"\xff\xd8\xff" or data[:8] == b"\x89PNG\r\n\x1a\n"
            or data[:4] == b"GIF8" or data[:2] == b"BM"):
        return data
    if _PIL_OK:
        try:
            im = Image.open(io.BytesIO(data)).convert("RGB")
            out = io.BytesIO(); im.save(out, format="PNG")
            return out.getvalue()
        except Exception:
            return None
    return None


def _add_hyperlink(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


# ── Build the search URL from the GUI fields ──────────────────────────────── #
def _build_search_url(fields: dict, place_mode: str, global_text: str,
                      lang: str = "ru") -> str:
    q = [("page", "1")]
    for key, fld in PARAM_MAP.items():
        val = (fields.get(key) or "").strip()
        if not val:
            continue
        if place_mode == "anyplace" and key in _PLACE_KEYS:
            continue                              # collapsed into s_place_search
        q.append((f"s_{fld}_search_{lang}", val))
        t_label = fields.get(key + "_type")
        q.append((f"t_{fld}_search_{lang}",
                  SEARCH_TYPES.get(t_label, "yvSynonym")))
    for key, fld in _YEAR_FIELDS.items():
        val = (fields.get(key) or "").strip()
        if not val:
            continue
        q.append((f"s_{fld}_search", val))
        t_label = fields.get(key + "_type")
        q.append((f"t_{fld}_search", YEAR_PREC.get(t_label, "exactly")))
    if place_mode == "anyplace":
        anyplace = next((fields.get(k) for k in _PLACE_KEYS
                         if (fields.get(k) or "").strip()), "")
        if anyplace:
            q.append((f"s_place_search_{lang}", anyplace.strip()))
    if (global_text or "").strip():
        q.append((f"s_global_search_{lang}", global_text.strip()))
    return f"{BASE}/{lang}/names/search-results?{urlencode(q, doseq=True)}"


# ── Results ───────────────────────────────────────────────────────────────── #
# Generic: a record detail link points under /ru/names/<id> (not …search…). We
# collect those + the surrounding row text (name / year / place / fate / source).
_COLLECT_JS = r"""() => {
    const norm = s => (s || '').replace(/\s+/g, ' ').trim();
    const out = [], seen = new Set();
    const links = document.querySelectorAll(
        'a[href*="/names/"]:not([href*="search"]):not([href*="advanced"])');
    links.forEach(a => {
        const href = a.href || '';
        if (!/\/names\/\d/.test(href) || seen.has(href)) return;
        seen.add(href);
        let box = a;
        for (let up = 0; up < 6 && box.parentElement; up++) {
            box = box.parentElement;
            if (norm(box.innerText).length > 30) break;
        }
        out.push({href, name: norm(a.innerText),
                  snippet: norm(box ? box.innerText : '').slice(0, 400)});
    });
    return out;
}"""


async def _collect_results(page, log, max_pages, max_records) -> list:
    out, seen = [], set()
    base_url = page.url
    page_no = 1
    while page_no <= max_pages and len(out) < max_records:
        try:
            rows = await page.evaluate(_COLLECT_JS)
        except Exception:
            rows = []
        new = 0
        for r in rows:
            href = r.get("href", "")
            if not href or href in seen:
                continue
            seen.add(href)
            out.append({"name": r.get("name", ""), "href": href,
                        "snippet": r.get("snippet", "")})
            new += 1
            if len(out) >= max_records:
                break
        log(f"  → Страница {page_no}: ссылок {len(rows)}, новых {new}, всего {len(out)}")
        if len(out) >= max_records or new == 0:
            break
        # pagination by the page= query param
        try:
            nxt = re.sub(r"([?&]page=)\d+", lambda m: m.group(1) + str(page_no + 1),
                         base_url)
            if "page=" not in nxt:
                sep = "&" if "?" in nxt else "?"
                nxt = f"{nxt}{sep}page={page_no + 1}"
            first_before = rows[0]["href"] if rows else ""
            await page.goto(nxt, wait_until="domcontentloaded", timeout=30000)
            await asyncio.sleep(2)
            changed = False
            for _ in range(12):
                cur = await page.evaluate(
                    "() => { const a=document.querySelector("
                    "'a[href*=\"/names/\"]:not([href*=\"search\"])'); return a?a.href:''; }")
                if cur and cur != first_before:
                    changed = True; break
                await asyncio.sleep(1)
            if not changed:
                break
        except Exception:
            break
        page_no += 1
    return out


# ── Record detail ─────────────────────────────────────────────────────────── #
_BAD_IMG = (r"maps\.google|googleapis|gstatic|staticmap|/maps/|\.svg|sprite|"
            r"logo|icon|placeholder|avatar|flag|/thumb|_thumb|share|social")
_DETAIL_JS = r"""(BAD) => {
    const bad = new RegExp(BAD, 'i');
    const norm = s => (s || '').replace(/ /g, ' ')
        .replace(/[ \t]+/g, ' ').replace(/\s*\n\s*/g, '\n').trim();
    // person name = the page H1 («Любовь Шендерович»)
    const h1 = document.querySelector('h1');
    const name = h1 ? norm(h1.textContent).replace(/\n+/g, ' ').slice(0, 90) : '';
    // «Record Details» fields = innermost «label⏎value» blocks (label on one
    // line, value on the next). A 2-line block whose descendants are NOT 2-line
    // blocks is exactly one field.
    const twoLines = e => {
        const ls = norm(e.innerText).split('\n').map(x => x.trim()).filter(x => x);
        return ls.length === 2 ? ls : null;
    };
    const pairs = [], seen = new Set();
    document.querySelectorAll('div, li, section').forEach(e => {
        const ls = twoLines(e);
        if (!ls) return;
        const [label, value] = ls;
        if (!label || !value || label.length > 55 || value.length > 400) return;
        for (const c of e.querySelectorAll('div, li, section'))
            if (twoLines(c)) return;            // not innermost
        const key = label + '|' + value;
        if (seen.has(key)) return; seen.add(key);
        pairs.push([label, value]);
    });
    // document images (Pages of Testimony / scans). Be liberal: a lazy <img> has
    // naturalWidth 0 and its real URL may sit in data-src/srcset/currentSrc; the
    // scan can also live in a same-origin <iframe> or as a CSS background. Bytes
    // are validated on download, so a few extra candidates are harmless.
    const imgs = [];
    const okUrl = s => s && /^https?:/i.test(s) && !bad.test(s);
    const add = (s, area) => { if (okUrl(s)) imgs.push([area || 1, s]); };
    const scan = doc => {
        doc.querySelectorAll('img').forEach(im => {
            const w = Math.max(im.naturalWidth || 0, im.clientWidth || 0,
                               parseInt(im.getAttribute('width')) || 0);
            const h = Math.max(im.naturalHeight || 0, im.clientHeight || 0,
                               parseInt(im.getAttribute('height')) || 0);
            const big = (w >= 150 && h >= 150);
            const lazy = !im.naturalWidth;                 // not yet decoded
            const area = (w && h) ? w * h : 1e9;           // unknown size → large
            const cand = im.currentSrc || im.src || im.getAttribute('data-src')
                       || im.getAttribute('data-original') || im.getAttribute('data-lazy') || '';
            if (big || lazy) add(cand, area);
            const ss = im.getAttribute('srcset') || '';
            if (ss) add(ss.split(',').pop().trim().split(/\s+/)[0], area);
        });
        doc.querySelectorAll('*').forEach(el => {
            const m = (getComputedStyle(el).backgroundImage || '')
                .match(/url\(["']?(https?:[^"')]+)["']?\)/i);
            if (m) add(m[1], (el.clientWidth || 0) * (el.clientHeight || 0));
        });
    };
    scan(document);
    document.querySelectorAll('iframe').forEach(f => {
        try { if (f.contentDocument) scan(f.contentDocument); } catch (e) {}
    });
    imgs.sort((x, y) => y[0] - x[0]);
    const docLinks = [...document.querySelectorAll('a[href]')].map(a => a.href)
        .filter(h => /\.(jpe?g|png|tiff?|pdf)(\?|$)/i.test(h) && !bad.test(h));
    const main = document.querySelector('main, [class*="content" i], [role="main"]') || document.body;
    return {name, pairs: pairs.slice(0, 80),
            imgs: [...new Set(imgs.map(i => i[1]))].slice(0, 12),
            docLinks: [...new Set(docLinks)].slice(0, 12),
            bodyText: norm(main.innerText).replace(/\n+/g, ' ').slice(0, 3000)};
}"""


def _looks_image(b: bytes) -> bool:
    """True if bytes start with a known image / PDF magic — avoids saving a
    misnamed HTML/error page as «.jpg» (the «ДАТА_1.jpg» won't-open bug)."""
    if not b or len(b) < 200:
        return False
    return (b[:3] == b"\xff\xd8\xff" or b[:8] == b"\x89PNG\r\n\x1a\n"
            or b[:4] == b"GIF8" or b[:2] == b"BM" or b[:4] == b"%PDF"
            or b[:4] == b"RIFF" or b[:2] in (b"II", b"MM"))


# Site chrome that leaks into fields/text (nav, social, map widget, search bar,
# Material-icon names like «keyboard_arrow_down») — dropped.
_JUNK = ("соцсет", "яд вашем в соц", "начать заново", "print candle",
         "поделиться", "холокост", "коллекции и исследования", "увековечение",
         "образование", "музеи и выставки", "праведники", "посетителям",
         "map data", "keyboard shortcuts", "satellite", "©", "google",
         "deutsch", "français", "официальные названия мест",
         "база данных имен", "расширенный поиск", "очистить все",
         "очистить поля", "keyboard_arrow", "expand_more", "expand_less",
         "arrow_drop")

# A value that is a bare Material-icon token («keyboard_arrow_down»): all
# lowercase ASCII with an underscore, no spaces.
_ICON_RE = re.compile(r"^[a-z][a-z0-9]*_[a-z0-9_]+$")

# A bare document-tab label («Record 1», «Запись 2») — the tab strip leaks into
# the fields as a fake «label⏎value» pair; never a real genealogy field.
_REC_TAB_RE = re.compile(r"^(record|запис[ьи]?)\s*\d*$", re.I)


def _is_junk(s: str) -> bool:
    s = (s or "").lower()
    return any(j in s for j in _JUNK)


def _clean_fields(pairs):
    out, seen = [], set()
    for k, v in pairs:
        if not k or not v or _is_junk(k) or _is_junk(v):
            continue
        if _ICON_RE.match(v) or _ICON_RE.match(k):     # Material-icon leak
            continue
        if _REC_TAB_RE.match(k.strip()) or _REC_TAB_RE.match(v.strip()):
            continue                                    # «Record 1 / Record 2» tab strip
        if k.strip().lower() in ("record details", "record", "детали записи"):
            continue
        if len(k) > 55 or len(v) < 1 or (k, v) in seen:
            continue
        seen.add((k, v)); out.append((k, v))
    return out


def _clean_text(txt: str) -> str:
    """Cut the site chrome (social/nav) off a free-text fallback."""
    if not txt:
        return ""
    cut = len(txt)
    for m in ("ЯД ВАШЕМ В СОЦ", "Начать заново", "Print Candle", "Поделиться",
              "Официальные названия мест"):
        i = txt.find(m)
        if 0 <= i < cut:
            cut = i
    t = txt[:cut].strip()
    # drop leaked document-tab labels («Record 1 Record 2 Record Details»)
    t = re.sub(r"\b[Rr]ecord(\s+Details|\s*\d+)\b", " ", t)
    t = re.sub(r"\bЗапис[ьи]\s*\d+\b", " ", t)
    return re.sub(r"[ \t]{2,}", " ", t).strip()


async def _img_bytes_via_goto(context, url, referer) -> bytes:
    if not url or not url.startswith("http"):
        return b""
    p = await context.new_page()
    try:
        try:
            await p.set_extra_http_headers({"referer": referer})
        except Exception:
            pass
        resp = await p.goto(url, timeout=20000, wait_until="commit")
        if resp and resp.ok:
            body = await resp.body()
            if body and len(body) > 2000:
                return body
    except Exception:
        pass
    finally:
        try:
            await p.close()
        except Exception:
            pass
    return b""


_REC_TABS_JS = r"""() => {
    const t = [...document.querySelectorAll('.mdc-tab, [role="tab"]')];
    const rec = t.filter(e => /record|запис/i.test(e.textContent || ''));
    return (rec.length || t.length);
}"""
_REC_CLICK_JS = r"""(i) => {
    let t = [...document.querySelectorAll('.mdc-tab, [role="tab"]')]
        .filter(e => /record|запис/i.test(e.textContent || ''));
    if (!t.length) t = [...document.querySelectorAll('.mdc-tab, [role="tab"]')];
    if (t[i]) { t[i].click(); return true; }
    return false;
}"""


async def _extract_record(page, url, images_dir, log, result_name=""):
    rec = {"name": "", "url": url, "records": []}
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=30000)
        await asyncio.sleep(1.0)
        # Yad Vashem often renders garbage / empty on the first hit — ALWAYS reload
        # before reading, and reload once more if it still came back empty.
        for attempt in range(2):
            try:
                await page.reload(wait_until="domcontentloaded", timeout=30000)
            except Exception:
                pass
            await asyncio.sleep(2.5)
            name = await page.evaluate(
                "() => { const h=document.querySelector('h1');"
                " return h ? h.textContent.replace(/\\s+/g,' ').trim() : ''; }")
            rec["name"] = (name or result_name).strip()
            base = safe_fn(rec["name"] or result_name or "record")
            # «Record 1 / Record 2 / …» tabs — each is a DIFFERENT source document.
            try:
                ntabs = int(await page.evaluate(_REC_TABS_JS) or 1)
            except Exception:
                ntabs = 1
            ntabs = max(1, min(ntabs, 15))
            records = []
            for ti in range(ntabs):
                if ntabs > 1:
                    try:
                        await page.evaluate(_REC_CLICK_JS, ti)
                        await asyncio.sleep(1.2)
                    except Exception:
                        pass
                try:                           # walk the page so lazy scans decode
                    await page.evaluate("""async () => {
                        const H = document.body.scrollHeight;
                        const step = Math.max(300, window.innerHeight * 0.8);
                        for (let y = 0; y <= H; y += step) {
                            window.scrollTo(0, y);
                            await new Promise(r => setTimeout(r, 220));
                        }
                        window.scrollTo(0, 0);
                    }""")
                    await asyncio.sleep(1.2)
                except Exception:
                    pass
                info = await page.evaluate(_DETAIL_JS, _BAD_IMG)
                fields = _clean_fields(info.get("pairs", []))
                imgs, urls = [], list(info.get("imgs", [])) + list(info.get("docLinks", []))
                for u in urls[:8]:
                    body = await _img_bytes_via_goto(page.context, u, page.url)
                    if _looks_image(body):     # skip misnamed HTML/error pages
                        images_dir.mkdir(parents=True, exist_ok=True)
                        suf = Path(urlparse(u).path).suffix.lower()
                        if suf not in (".jpg", ".jpeg", ".png", ".gif", ".tif", ".tiff",
                                       ".pdf", ".webp"):
                            suf = ".jpg"
                        tag = f"_rec{ti + 1}_{len(imgs) + 1}" if ntabs > 1 else f"_{len(imgs) + 1}"
                        dest = images_dir / f"{base}{tag}{suf}"
                        dest.write_bytes(body)
                        imgs.append(str(dest))
                records.append({"label": f"Record {ti + 1}", "fields": fields,
                                "images": imgs,
                                "text": "" if fields else _clean_text(info.get("bodyText", ""))})
            nf = sum(len(r["fields"]) for r in records)
            nd = sum(len(r["images"]) for r in records)
            rec["records"] = records
            if nf or nd:
                log(f"      вкладок: {len(records)}, полей: {nf}, документов: {nd}")
                break
            if attempt == 0:
                log("      ↻ пусто — обновляю ещё раз…")
            else:
                log(f"      вкладок: {len(records)}, полей: 0, документов: 0")
    except Exception as e:
        log(f"      !! страница записи ({type(e).__name__})")
    return rec


# ── Word output ───────────────────────────────────────────────────────────── #
def _kv_table(doc, pairs):
    if not pairs:
        return
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    for k, v in pairs:
        r = tbl.add_row().cells
        r[0].text = str(k); set_cell_lines(r[1], v)
        for run in r[0].paragraphs[0].runs:
            run.bold = True


def _docx_add_record(doc, i, rec):
    name = rec.get("name") or f"Record {i}"
    hp = doc.add_paragraph(); hr = hp.add_run(f"{i}. {name}")
    hr.bold = True; hr.font.size = Pt(13)
    records = rec.get("records") or []
    # One block per document — its scans, then its OWN fields table.
    # No «Record N» captions between them, just a blank line.
    for di, r in enumerate(records):
        if di:
            doc.add_paragraph("")                 # one blank line between documents
        for img in r.get("images", []):
            if str(img).lower().endswith(".pdf"):
                p = doc.add_paragraph(); p.add_run("Файл: ").bold = True
                p.add_run(str(Path(img).resolve()))   # PDF can't embed — show path
                continue
            try:
                png = _to_png(Path(img).read_bytes())
                if png:
                    doc.add_picture(io.BytesIO(png), width=Inches(3.2))
                    p = doc.add_paragraph(); p.add_run("Файл: ").bold = True
                    p.add_run(str(Path(img).resolve()))   # exact path where it was saved
            except Exception:
                pass
        fields = r.get("fields") or []
        if fields:
            _kv_table(doc, fields)
        elif r.get("text"):
            doc.add_paragraph(r["text"])
    if rec.get("url"):
        p = doc.add_paragraph(); _add_hyperlink(p, "Source / Источник", rec["url"])
    doc.add_paragraph("")


def write_docx(path, records, qlines, append=False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    if append and Path(path).exists():
        doc = Document(str(path)); doc.add_page_break()
        ap = doc.add_paragraph(); ar = ap.add_run(f"➕ Added {len(records)}")
        ar.bold = True; ar.font.size = Pt(13)
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width = Mm(210); s.page_height = Mm(297)
        s.left_margin = s.right_margin = Mm(18)
        s.top_margin = s.bottom_margin = Mm(18)
        ht = doc.add_paragraph()
        htr = ht.add_run("Yad Vashem — Shoah Victims' Names — search results")
        htr.bold = True; htr.font.size = Pt(14)
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ln in qlines:
            doc.add_paragraph(ln)
        doc.add_paragraph(f"Найдено: {len(records)}")
        doc.add_paragraph("")
    for i, rec in enumerate(records, 1):
        _docx_add_record(doc, i, rec)
    add_page_numbers(doc)
    doc.save(str(path))


# ── Main entry point ──────────────────────────────────────────────────────── #
async def run_scraper(*,
    fields=None, place_mode="byfield", global_text="", lang="ru",
    output_folder=Path("."),
    log=print, progress=None, cancel_event=None, ask_file_conflict=None,
    max_pages=10, max_records=60,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    fields = fields or {}
    output_folder = Path(output_folder); output_folder.mkdir(parents=True, exist_ok=True)
    qkey = " ".join(v for k, v in fields.items()
                    if not k.endswith("_type") and v) or global_text or "yadvashem"
    images_dir = output_folder / "images" / (safe_fn(qkey) or "yadvashem")
    summary = {"ok": False}

    if not any(v for k, v in fields.items() if not k.endswith("_type") and v) \
            and not global_text:
        _prog(100, "Пустой запрос.")
        return summary

    url = _build_search_url(fields, place_mode, global_text, lang=lang)
    qlines = [f"Запрос: {qkey}"]

    _prog(0, "Запускаю браузер…")
    async with async_playwright() as pw:
        browser = await browser_util.launch(pw, 
            headless=False,
            args=["--start-maximized", "--disable-blink-features=AutomationControlled"])
        ctx = await browser.new_context(no_viewport=True, accept_downloads=True,
                                        ignore_https_errors=True)
        page = await ctx.new_page()
        try:
            _prog(5, "Поиск…")
            log(f"  → Открываю поиск: {url}")
            await page.goto(url, wait_until="domcontentloaded", timeout=45000)
            await asyncio.sleep(1.5)
            try:                              # YV glitches — refresh before reading
                await page.reload(wait_until="domcontentloaded", timeout=45000)
            except Exception:
                pass
            await asyncio.sleep(4)            # SPA renders results via XHR
            if _done():
                return summary

            _prog(15, "Сбор результатов…")
            recs_meta = await _collect_results(page, log, max_pages, max_records)
            log(f"  Записей: {len(recs_meta)}")
            if not recs_meta:
                _prog(100, "Ничего не найдено (или селекторы результата надо поправить).")
                summary.update({"ok": True, "n_records": 0})
                return summary

            records = []
            n = len(recs_meta)
            for i, rm in enumerate(recs_meta, 1):
                if _done():
                    break
                _prog(20 + int(70 * i / n), f"[{i}/{n}] {rm['name'][:50]}…")
                log(f"  [{i}/{n}] {rm['name']}")
                try:
                    dp = await ctx.new_page()
                    try:
                        rec = await _extract_record(dp, rm["href"], images_dir, log,
                                                    result_name=rm["name"])
                    finally:
                        await dp.close()
                    if not rec.get("name"):
                        rec["name"] = rm["name"]
                    records.append(rec)
                except Exception as _e:
                    log(f"      !! пропускаю запись ({type(_e).__name__})")
                    records.append({"name": rm["name"], "url": rm["href"],
                                    "records": []})
                await asyncio.sleep(0.3)

            _prog(92, "Сохранение…")
            base = safe_fn(f"yadvashem_{qkey}") or "yadvashem_results"
            docx_p = output_folder / f"{base}.docx"
            decision = "overwrite"
            if docx_p.exists() and ask_file_conflict:
                try:
                    decision = (ask_file_conflict([docx_p.name]) or "overwrite").lower()
                except Exception:
                    decision = "overwrite"
                log(f"  → Файл существует → {decision}")
            if decision != "skip" and records:
                try:
                    write_docx(docx_p, records, qlines, append=(decision == "append"))
                    log(f"  → Word: {docx_p.name}")
                except PermissionError:
                    alt = output_folder / f"{base}_{time.strftime('%H%M%S')}.docx"
                    write_docx(alt, records, qlines, append=False)
                    docx_p = alt
                    log(f"  !! файл занят → сохранил как {alt.name}")

            _prog(100, f"Готово — {len(records)} запис(ей).")
            summary.update({"ok": True, "n_records": len(records),
                            "output_folder": str(output_folder)})
        except Exception as exc:
            summary["message"] = f"{type(exc).__name__}: {exc}"
            log(f"  !! Ошибка: {exc}")
        finally:
            try:
                await browser.close()
            except Exception:
                pass
    return summary
