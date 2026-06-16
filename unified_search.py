"""
unified_search.py — «Unified Search» across every site that has a personal
(name + surname) search, EXCEPT Wikisource.
=============================================================================
The user enters a First name and a Last name once. For each site we:
  • translate (transliterate) the name to the script the site searches in
    (Latin or Cyrillic); sites that search in several languages at once
    (Yad Vashem, MyHeritage) get the name AS ENTERED — no translation;
  • run that site's own search in `list_only` mode — it collects ONLY the
    rows shown on the results page (no record is opened, nothing downloaded);
  • collect those rows.
Finally we write ONE Word + ONE Excel, GROUPED BY SITE, containing only the
results-page information, cleaned of junk.

No credentials are needed except MyHeritage (login fields in the GUI tab).

This orchestrator reuses each scraper's run_scraper(list_only=True). Adding a
site = implement list_only in its scraper (return summary["rows"]) and add one
SITES entry below.
"""
from __future__ import annotations
import asyncio
import re
from pathlib import Path


def _safe(s: str, n: int = 90) -> str:
    """Filesystem-safe stem, spaces → underscores."""
    s = re.sub(r"[^\w\- ]+", "", (s or ""), flags=re.UNICODE).strip()
    s = re.sub(r"\s+", "_", s)
    return s[:n] or "results"

# ── transliteration (Latin ↔ Cyrillic) — names aren't "translated", they're ──
# transliterated into the script the site searches in. Best-effort.
_CYR2LAT = {
    "а":"a","б":"b","в":"v","г":"g","д":"d","е":"e","ё":"e","ж":"zh","з":"z",
    "и":"i","й":"y","к":"k","л":"l","м":"m","н":"n","о":"o","п":"p","р":"r",
    "с":"s","т":"t","у":"u","ф":"f","х":"kh","ц":"ts","ч":"ch","ш":"sh",
    "щ":"shch","ъ":"","ы":"y","ь":"","э":"e","ю":"yu","я":"ya",
}
_LAT2CYR = {  # digraphs first (handled in code), then single letters
    "a":"а","b":"б","c":"к","d":"д","e":"е","f":"ф","g":"г","h":"х","i":"и",
    "j":"й","k":"к","l":"л","m":"м","n":"н","o":"о","p":"п","q":"к","r":"р",
    "s":"с","t":"т","u":"у","v":"в","w":"в","x":"кс","y":"ы","z":"з",
}
_LAT2CYR_DI = [("shch","щ"),("zh","ж"),("kh","х"),("ts","ц"),("ch","ч"),
               ("sh","ш"),("th","т"),("ph","ф"),("yu","ю"),("ya","я"),("yo","ё")]


def _is_cyrillic(s: str) -> bool:
    return any("Ѐ" <= c <= "ӿ" for c in (s or ""))


def _to_latin(s: str) -> str:
    if not _is_cyrillic(s):
        return s
    out = []
    for ch in s:
        low = ch.lower()
        rep = _CYR2LAT.get(low, ch)
        out.append(rep.capitalize() if ch.isupper() and rep else rep)
    return "".join(out)


def _to_cyrillic(s: str) -> str:
    if _is_cyrillic(s) or not s:
        return s
    t = s
    for lat, cyr in _LAT2CYR_DI:
        t = t.replace(lat, cyr).replace(lat.capitalize(), cyr.upper())
    out = []
    for ch in t:
        if "Ѐ" <= ch <= "ӿ":
            out.append(ch); continue
        low = ch.lower()
        rep = _LAT2CYR.get(low, ch)
        out.append(rep.capitalize() if ch.isupper() and rep else rep)
    return "".join(out)


def translate_name(name: str, lang: str) -> str:
    """lang: 'lat' → Latin script, 'cyr' → Cyrillic, 'multi' → unchanged."""
    if lang == "lat":
        return _to_latin(name)
    if lang == "cyr":
        return _to_cyrillic(name)
    return name


# ── site registry ───────────────────────────────────────────────────────────
# Each entry: a label, the target script, and an async adapter(first, last, mh)
# → list[dict] of results-page rows. `mh` carries MyHeritage credentials.
async def _run_list_only(module_name, log, **kw):
    """Import an ASYNC scraper and run its run_scraper(list_only=True, …)."""
    import importlib
    mod = importlib.import_module(module_name)
    res = await mod.run_scraper(list_only=True, output_format="docx",
                                log=log, **kw)
    return res.get("rows", []) if isinstance(res, dict) else []


async def _run_sync_list_only(module_name, log, **kw):
    """Import a SYNC scraper (urllib-based, e.g. Pogroms) and run it off the event
    loop so it doesn't block the others."""
    import importlib
    mod = importlib.import_module(module_name)
    res = await asyncio.to_thread(mod.run_scraper, list_only=True,
                                  output_format="docx", log=log, **kw)
    return res.get("rows", []) if isinstance(res, dict) else []


async def _ancestry(first, last, mh, strict, log):
    exact = {"name_forms": ["1"], "surname_forms": ["1"]} if strict else {}
    return await _run_list_only("ancestry_scraper", log,
                                first_names=first, last_names=last, exact=exact)


async def _familysearch(first, last, mh, strict, log):
    exact = {"name": True, "surname": True} if strict else {}
    return await _run_list_only("familysearch_scraper", log,
                                first_names=first, last_names=last, exact=exact)


async def _jewishgen(first, last, mh, strict, log):
    # JewishGen Unified Search by Surname / Given Name across ALL countries; phonetic
    # by default, exact when strict. run_scraper is positional, so call it directly.
    import importlib
    mod = importlib.import_module("jewishgen_scraper")
    st = "is Exactly" if strict else "Phonetically Like"
    rows = []
    if last:
        rows.append(("Surname", st, last))
    if first:
        rows.append(("Given Name", st, first))
    if not rows:
        return []
    res = await mod.run_scraper(rows, "ALL COUNTRIES", [], "docx", ".",
                                list_only=True, log=log)
    return res.get("rows", []) if isinstance(res, dict) else []


async def _pogroms(first, last, mh, strict, log):
    # Pogroms has only a surname field and the form accepts Latin only (lang="lat").
    # soundslike=fuzzy by default; strict → exact match.
    return await _run_sync_list_only("pogroms_scraper", log,
                                     family_name=(last or first),
                                     soundslike=not strict)


# SITES — order = output order. lang: lat / cyr / multi. Wikisource is excluded
# (no personal search). More sites are added here as their list_only lands.
SITES = [
    # Ancestry temporarily EXCLUDED (its browser run is slow and holds the profile);
    # re-add the entry once it's fast enough: {"key":"ancestry", … "run": _ancestry}.
    {"key": "jewishgen",    "label": "JewishGen",      "lang": "lat", "run": _jewishgen},
    {"key": "familysearch", "label": "FamilySearch",   "lang": "lat", "run": _familysearch},
    {"key": "pogroms",      "label": "Jewish Pogroms", "lang": "lat", "run": _pogroms},
    # TODO (need list_only in their scraper + the user's live test, the sites are
    # unreachable from here): MyHeritage (login), Skarb, Memorial/Memsearch,
    # Память народа, Памяти героев ВВ, Yad Vashem.
]


async def run_unified(*, first_name="", last_name="", strict=False,
                      mh_email=None, mh_password=None,
                      output_format="both", output_folder=Path("."),
                      log=print, progress=None, cancel_event=None,
                      ask_file_conflict=None):
    """Search every registered site and write ONE grouped Word + Excel."""
    def _prog(p, t):
        log(t)
        if progress:
            progress(int(p), str(t))

    def _stop():
        return bool(cancel_event and cancel_event.is_set())

    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    mh = {"email": mh_email, "password": mh_password}
    grouped = []                       # [(label, [rows])] in SITES order
    n = len(SITES) or 1
    for i, site in enumerate(SITES):
        if _stop():
            break
        fn = translate_name(first_name, site["lang"])
        ln = translate_name(last_name, site["lang"])
        _prog(5 + int(85 * i / n),
              f"[{site['label']}] поиск: {fn} {ln}".strip())
        try:
            rows = await site["run"](fn, ln, mh, strict, log)
        except Exception as exc:
            log(f"  !! {site['label']}: {type(exc).__name__}: {exc}")
            rows = []
        log(f"  {site['label']}: {len(rows)} строк(и)")
        grouped.append((site["label"], rows))

    _prog(92, "Сохранение результатов…")
    qlines = [f"First name: {first_name}", f"Last name: {last_name}"]
    # Filename carries WHAT was searched, e.g. unified_search_Rebecca_Sanders.docx
    stem = _safe(f"unified_search {first_name} {last_name}")
    docx_path = (output_folder / f"{stem}.docx") if output_format in ("docx", "both") else None
    xlsx_path = (output_folder / f"{stem}.xlsx") if output_format in ("xlsx", "both") else None
    total = sum(len(r) for _l, r in grouped)

    # Ask before clobbering existing files (overwrite / append / skip).
    append = False
    existing = [str(p) for p in (docx_path, xlsx_path) if p and p.exists()]
    if existing and ask_file_conflict:
        choice = (ask_file_conflict(existing) or "overwrite").lower()
        if choice == "skip":
            _prog(100, "Готово — файлы не сохранены (пропущено).")
            return {"ok": True, "n_records": total, "skipped": True,
                    "output_folder": str(output_folder),
                    "docx_path": None, "xlsx_path": None}
        append = (choice == "append")

    if docx_path:
        _write_docx(docx_path, grouped, qlines, append=append)
    if xlsx_path:
        _write_xlsx(xlsx_path, grouped, qlines, append=append)
    _prog(100, f"Готово — {total} строк(и) с {len(grouped)} сайт(ов).")
    return {"ok": True, "n_records": total,
            "output_folder": str(output_folder),
            "docx_path": str(docx_path) if docx_path else None,
            "xlsx_path": str(xlsx_path) if xlsx_path else None}


# ── grouped output ──────────────────────────────────────────────────────────
def _columns(rows):
    cols = []
    for r in rows:
        for k in r:
            if k not in cols:
                cols.append(k)
    return cols


def _set_cell_lines(cell, text, pt=8):
    """Write text into a table cell, turning each '\\n' into its OWN paragraph —
    otherwise python-docx glues the lines («Натурализация2 ноября 1987Detroit…»).
    Small font (pt) so wide multi-column tables stay readable."""
    from docx.shared import Pt
    lines = str(text or "").split("\n")
    cell.text = lines[0]
    for ln in lines[1:]:
        cell.add_paragraph(ln)
    for p in cell.paragraphs:
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            run.font.size = Pt(pt)


def _landscape(doc):
    """Landscape A4-ish page with minimal margins — so the wide results tables
    (Pogroms has 15 columns) fit without squeezing each column to one letter."""
    from docx.enum.section import WD_ORIENT
    from docx.shared import Inches
    sec = doc.sections[0]
    if sec.orientation != WD_ORIENT.LANDSCAPE:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.left_margin = sec.right_margin = Inches(0.3)
    sec.top_margin = sec.bottom_margin = Inches(0.3)


def _write_docx(path, grouped, qlines, append=False):
    from docx import Document
    from docx.shared import Pt
    if append and Path(path).exists():
        doc = Document(str(path))
        doc.add_page_break()
        ap = doc.add_paragraph(); ar = ap.add_run("➕ Добавлены новые результаты")
        ar.bold = True; ar.font.size = Pt(13)
    else:
        doc = Document()
        h = doc.add_paragraph(); r = h.add_run("Unified Search — результаты")
        r.bold = True; r.font.size = Pt(15)
    _landscape(doc)
    for ln in qlines:
        doc.add_paragraph(ln)
    doc.add_paragraph("")
    for label, rows in grouped:
        hp = doc.add_paragraph(); hr = hp.add_run(f"{label} — {len(rows)}")
        hr.bold = True; hr.font.size = Pt(13)
        if not rows:
            doc.add_paragraph("  (нет результатов)")
            continue
        cols = _columns(rows)
        tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = "Table Grid"
        tbl.autofit = True; tbl.allow_autofit = True
        for c, name in zip(tbl.rows[0].cells, cols):
            _set_cell_lines(c, name, pt=8)
            for run in c.paragraphs[0].runs:
                run.bold = True
        for row in rows:
            cells = tbl.add_row().cells
            for ci, name in enumerate(cols):
                _set_cell_lines(cells[ci], row.get(name, ""), pt=8)
        doc.add_paragraph("")
    doc.save(str(path))


def _write_xlsx(path, grouped, qlines, append=False):
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter
    HN = Font(bold=True, color="FFFFFF"); HF = PatternFill("solid", fgColor="2a4a2a")
    HA = Alignment(horizontal="center", wrap_text=True)
    wrap = Alignment(wrap_text=True, vertical="top")

    if append and Path(path).exists():               # append rows to per-site sheets
        wb = load_workbook(str(path))
        for label, rows in grouped:
            if not rows:
                continue
            title = (label or "Sheet")[:31]
            cols = _columns(rows)
            if title in wb.sheetnames:
                ws = wb[title]
                header = [ws.cell(row=1, column=c).value
                          for c in range(1, ws.max_column + 1)]
                header = [h for h in header if h]
                for name in cols:
                    if name not in header:
                        header.append(name)
                        c = ws.cell(row=1, column=len(header), value=name)
                        c.font = HN; c.fill = HF; c.alignment = HA
                start = ws.max_row + 1
            else:
                ws = wb.create_sheet(title=title); header = cols
                for ci, name in enumerate(header, 1):
                    c = ws.cell(row=1, column=ci, value=name)
                    c.font = HN; c.fill = HF; c.alignment = HA
                start = 2
            for ri, row in enumerate(rows, start):
                for ci, name in enumerate(header, 1):
                    c = ws.cell(row=ri, column=ci, value=str(row.get(name, "")))
                    c.alignment = wrap
        wb.save(str(path)); return

    wb = Workbook(); first = True
    for label, rows in grouped:
        title = (label or "Sheet")[:31]
        ws = wb.active if first else wb.create_sheet(title=title)
        if first:
            ws.title = title; first = False
        cols = _columns(rows) or ["—"]
        for ci, name in enumerate(cols, 1):
            c = ws.cell(row=1, column=ci, value=name)
            c.font = HN; c.fill = HF; c.alignment = HA
        for ri, row in enumerate(rows, 2):
            for ci, name in enumerate(cols, 1):
                c = ws.cell(row=ri, column=ci, value=str(row.get(name, "")))
                c.alignment = wrap
        for ci in range(1, len(cols) + 1):
            ws.column_dimensions[get_column_letter(ci)].width = 28
    wb.save(str(path))
