#!/usr/bin/env python3
"""
gwar_scraper.py
===============
Scraper for «Памяти героев Великой войны» (https://gwar.mil.ru/heroes/) — WWI
(1914–1918) participant records (same ministry/structure family as pamyat).

Flow:
  1. Open /heroes/, fill the search fields (all by language-independent id/name),
     set the «Разделы» source checkboxes, submit «Найти».
  2. Collect the result documents (links to /heroes/chelovek…). When a name was
     given, keep only FUZZY-matching people (surname ~1–2 letters, first name /
     initials, patronymic stem) and, if a birth year is given, drop other years.
  3. Open each matched record, copy the left-panel fields to Word, and save the
     scanned document image(s) — every page (Страница X из N) via #btnSaveImage,
     falling back to the displayed image.

The site is Russian-only; the GUI is English. No login.

NEVER crash on a broken/empty page: every record and every step is wrapped in
try/except, urlparse().netloc (not .host), bad links are skipped.
"""

import asyncio
import difflib
import io
import re
import sys
import time
from pathlib import Path
from docx_util import set_cell_lines, add_page_numbers
from urllib.parse import urlparse, urlencode, parse_qsl, urlsplit, urlunsplit

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        import os
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    import browser_util
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from PIL import Image
    _PIL_OK = True
except ImportError:
    _PIL_OK = False

HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")

# «Разделы» (sources) — English label → checkbox id on the page.
SECTIONS = {
    "Awards":        "award_tag",
    "Losses":        "dead_tag",
    "Personal data": "frc_tag",
    "Commanders":    "commander_tag",
    "Notable people": "prs_tag",
}

# «Событие» (event_id) — WWI battles/operations, English label → site value.
EVENTS = {
    "": "",
    "East Prussian Operation": "13",
    "Battle of Galicia": "14",
    "Rava-Russka Operation": "18",
    "Battle of Gumbinnen": "19",
    "Battle of Kraśnik": "9",
    "Lublin-Kholm Operation": "15",
    "Battle of Komarów": "10",
    "Battle of Tannenberg": "23",
    "Battle of the Golden Lipa": "11",
    "Battle of the Gnila Lipa": "12",
    "Capitulation of the 13th & 15th Corps": "157",
    "Austro-Hungarian capture of Komarów": "16",
    "Galich-Lviv Operation": "17",
    "Russian capture of Lviv": "20",
    "Battle of Gorodok": "22",
    "Russian capture of Halych": "21",
    "Battle of the Masurian Lakes": "25",
    "Siege of Przemyśl": "27",
    "Warsaw-Ivangorod Operation": "29",
    "Köprüköy Operation": "32",
    "Łódź Operation": "33",
    "Sarikamish Operation": "38",
    "Battle of Ardahan": "39",
    "Russian offensive in the Carpathians": "40",
    "Lasdehnen Operation": "41",
    "Battle of Kozevo": "47",
    "Augustów Operation": "43",
    "Capture of Mount Makówka": "48",
    "Gorlice-Tarnów breakthrough": "53",
    "1st Austro-German offensive": "50",
    "2nd Austro-German offensive": "57",
    "Alashkert Operation": "60",
    "Riga-Shavli Operation": "102",
    "Polish pocket (loss of Poland)": "62",
    "Vilnius Operation (loss of Lithuania)": "65",
    "Sventiany breakthrough": "67",
    "Erzurum campaign": "75",
    "Trebizond Operation": "78",
    "Lake Naroch Operation": "80",
    "Brusilov offensive": "93",
    "Battle of Lutsk": "89",
    "Battle of Dobronouts (Okna breakthrough)": "91",
    "Battle of Kolomyia": "96",
    "1st Battle of Kovel": "103",
    "2nd Battle of Kovel": "101",
    "August Operation (SW Front)": "107",
    "Kovel Battle (SW Front)": "110",
    "Mitau Operation": "114",
    "Capture of Khanaqin (Persia)": "121",
    "June (Kerensky) offensive": "127",
    "Battle of Zborov": "126",
    "Tarnopol disaster": "130",
    "2nd Battle of Mărășești (Romanian Front)": "131",
}


# ── Helpers ───────────────────────────────────────────────────────────────── #
def safe_fn(s: str, n: int = 80) -> str:
    s = re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip())
    return re.sub(r"\s+", "_", s)[:n].strip("_") or "record"


def _host(u: str) -> str:
    u = u or ""
    if not u or u.startswith(("chrome-error", "about:", "data:")):
        return ""
    try:
        return urlparse(u).netloc or ""
    except Exception:
        return ""


def _to_png(data: bytes):
    if not data:
        return None
    if (data[:3] == b"\xff\xd8\xff" or data[:8] == b"\x89PNG\r\n\x1a\n"
            or data[:4] == b"GIF8" or data[:2] == b"BM"):
        return data
    if _PIL_OK:
        try:
            im = Image.open(io.BytesIO(data)).convert("RGB")
            out = io.BytesIO(); im.save(out, format="PNG")
            return out.getvalue()
        except Exception:
            return None
    return None


def _add_hyperlink(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


# ── FIO fuzzy matching (same rules as pamyat) ─────────────────────────────── #
def _sim(a: str, b: str) -> float:
    a, b = (a or "").lower().strip(), (b or "").lower().strip()
    if not a or not b:
        return 0.0
    return difflib.SequenceMatcher(None, a, b).ratio()


def _is_initial(token: str, full: str) -> bool:
    t = (token or "").strip(". ").lower()
    f = (full or "").lower()
    return len(t) == 1 and bool(f) and f.startswith(t)


def _parse_fio(name: str):
    name = (name or "").strip()
    name = re.sub(r"([А-ЯЁA-Za-zа-яё])\.([А-ЯЁA-Z])", r"\1. \2", name)
    parts = [p for p in re.split(r"\s+", name) if p]
    return (parts[0] if parts else "",
            parts[1] if len(parts) > 1 else "",
            parts[2] if len(parts) > 2 else "")


def _first_match(r: str, w: str) -> bool:
    r = (r or "").strip(". ").lower(); w = (w or "").strip(". ").lower()
    if not r or not w:
        return True
    if _is_initial(r, w) or _is_initial(w, r):
        return True
    if r[:2] == w[:2]:                       # Герш / Гирш
        return True
    return _sim(r, w) >= 0.5


def _patr_match(r: str, w: str) -> bool:
    r = (r or "").strip(". ").lower(); w = (w or "").strip(". ").lower()
    if not r or not w:
        return True
    short, long = (r, w) if len(r) <= len(w) else (w, r)
    if short and long.startswith(short):
        return True
    if r[:3] == w[:3]:
        return True
    return _sim(r, w) >= 0.78


def _person_matches(result_name, want_last, want_first, want_middle,
                    exact=False) -> bool:
    if not want_last:
        return True
    r_last, r_first, r_middle = _parse_fio(result_name)
    if exact:
        # Strict: the given fields must match EXACTLY (Рубик ≠ Рубин, Герш ≠ Гирш).
        norm = lambda s: (s or "").strip().lower()
        if norm(r_last) != norm(want_last):
            return False
        if want_first and norm(r_first) != norm(want_first):
            return False
        if want_middle and norm(r_middle) != norm(want_middle):
            return False
        return True
    if _sim(r_last, want_last) < 0.7:
        return False
    if want_first and r_first and not _first_match(r_first, want_first):
        return False
    if want_middle and r_middle and not _patr_match(r_middle, want_middle):
        return False
    return True


# ── Detail-page field labels (gwar left info panel) ───────────────────────── #
_GWAR_LABELS = sorted([
    "Дата рождения", "Место рождения", "Должность/Звание", "Должность / Звание",
    "Должность/ Звание", "Должность",
    "Воинское звание", "Воинская часть", "Место службы", "Лагерь военнопленных",
    "Лагерь", "Дата водворения в лагерь", "Дата пленения", "Место пленения",
    "Дата начала события", "Дата окончания события", "Дата события", "Событие",
    "Место события", "Тип документа", "Картотека", "Архив", "Название фонда",
    "Фонд", "Опись", "Шкаф", "Дело", "Ящик", "Награда", "Дата документа",
    "Номер документа", "Источник информации", "Губерния", "Уезд", "Волость",
    "Населенный пункт",
], key=len, reverse=True)

_GWAR_LABEL_RE = re.compile(
    r"(" + "|".join(re.escape(_l) for _l in _GWAR_LABELS) + r")\s*:?\s*")

_CHROME_MARKERS = ("О проекте", "Вопросы и ответы", "Как искать", "Обратная связь",
                   "Правовая информация", "Пользовательское соглашение",
                   "Министерство обороны", "© Министерство", "Перейти к просмотру",
                   "Инструкция по поиску", "Видео-инструкция", "Главная страница",
                   "Вы используете слишком старую", "Уважаемые пользователи",
                   "Новое в 20", "Персоналий на образе", "Боевой путь строится",
                   "Нет данных о точках", "В случае обнаружения")


def _clean_block(txt: str) -> str:
    if not txt:
        return ""
    cut = len(txt)
    for m in _CHROME_MARKERS:
        i = txt.find(m)
        if 0 <= i < cut:
            cut = i
    return txt[:cut].strip(" ;,—-·•\t\n")


def _parse_fields(txt: str):
    """Split the info-panel text by the KNOWN labels → ordered [(label, value)]."""
    txt = re.sub(r"[ \t ]+", " ", (txt or ""))
    txt = re.sub(r"\s*\n\s*", " ", txt).strip()
    marks = list(_GWAR_LABEL_RE.finditer(txt))
    out, seen = [], set()
    for i, m in enumerate(marks):
        end = marks[i + 1].start() if i + 1 < len(marks) else len(txt)
        val = _clean_block(txt[m.end():end].strip(" ;,"))
        k = m.group(1).strip()
        if val and len(val) < 400 and (k, val) not in seen:
            seen.add((k, val)); out.append((k, val))
    return out


# ── Image download through the browser (CDN blocks plain requests) ────────── #
async def _img_bytes_via_goto(context, url, referer) -> bytes:
    if not url or not url.startswith("http"):
        return b""
    p = await context.new_page()
    try:
        try:
            await p.set_extra_http_headers({"referer": referer})
        except Exception:
            pass
        resp = await p.goto(url, timeout=20000, wait_until="commit")
        if resp and resp.ok:
            body = await resp.body()
            if body and len(body) > 2000:
                return body
    except Exception:
        pass
    finally:
        try:
            await p.close()
        except Exception:
            pass
    return b""


# ── Search ────────────────────────────────────────────────────────────────── #
BASE = "https://gwar.mil.ru"


# The form has DUPLICATE inputs (mobile + desktop) sharing the same id/name, and
# some live in collapsed sections — a normal click hits the hidden duplicate and
# times out. So set the value on EVERY matching element via the native setter and
# fire input/change/keyup so the site's jQuery picks it up. Whichever copy the
# search reads, it is filled; the visible one also shows it to the user.
_FILL_JS = r"""([fid, val]) => {
    const els = [...document.querySelectorAll('#' + CSS.escape(fid) + ', [name="' + fid + '"]')];
    if (!els.length) return 'missing';
    els.forEach(el => {
        const proto = (el.tagName === 'TEXTAREA') ? HTMLTextAreaElement : HTMLInputElement;
        try {
            Object.getOwnPropertyDescriptor(proto.prototype, 'value').set.call(el, val);
        } catch (e) { el.value = val; }
        ['input', 'change', 'keyup', 'blur'].forEach(t =>
            el.dispatchEvent(new Event(t, {bubbles: true})));
        if (el.classList) el.classList.remove('empty-field');
    });
    return 'ok:' + els.length;
}"""


async def _fill_field(page, fid: str, value: str, log):
    if not value:
        return
    try:
        r = await page.evaluate(_FILL_JS, [fid, str(value)])
        if isinstance(r, str) and r.startswith("ok"):
            log(f"  ✓ {fid} = {value!r}")
        else:
            log(f"  !! поле {fid}: не найдено ({r})")
    except Exception as e:
        log(f"  !! поле {fid}: {type(e).__name__}")


async def _do_search(page, params, log) -> bool:
    log(f"  → Открываю поиск: {BASE}/heroes/")
    await page.goto(f"{BASE}/heroes/", wait_until="domcontentloaded", timeout=40000)
    await asyncio.sleep(2)

    # Open BOTH collapsible sections so the user SEES the advanced fields used:
    # «Дополнительные параметры поиска» (toggle «Больше параметров поиска») and
    # «Место хранения документов» (toggle «Показать архивные реквизиты»). The
    # toggles are duplicated (mobile+desktop) like the inputs, so click the
    # VISIBLE one via JS. We match only «Больше …»/«Показать …», so an already
    # open section (which shows «Меньше …»/«Скрыть …») is never collapsed.
    total_opened = 0
    for _pass in range(3):                  # 2nd toggle may appear after the 1st
        try:
            n = await page.evaluate(r"""() => {
                // Collapsed toggles carry class «more-sign … _closed» (the archive
                // one is «more-sign-arc more-sign_closed visible», text «Поиск по
                // архивным реквизитам»). Also accept «Больше параметров поиска» by
                // text. Click each VISIBLE collapsed toggle once; once open the
                // class flips to «_opened» / the text changes, so it is never
                // re-collapsed.
                let clicked = 0;
                const hit = new Set();
                const want = ['Больше параметров', 'архивным реквизитам'];
                for (const e of document.querySelectorAll(
                        'span,a,div,button,p,li,[class*="more-sign"]')) {
                    if (e.offsetParent === null || hit.has(e)) continue;
                    const c = (e.className && e.className.toString)
                        ? e.className.toString() : '';
                    const t = (e.textContent || '').trim();
                    const byClass = /more-sign/.test(c) && /_closed/.test(c);
                    const byText = t.length < 60 && want.some(w => t.includes(w));
                    if (byClass || byText) { hit.add(e); e.click(); clicked++; }
                }
                return clicked;
            }""")
        except Exception:
            n = 0
        total_opened += n
        if n == 0:
            break
        await asyncio.sleep(0.8)
    log(f"  → раскрыл доп. секции: {total_opened}")

    # Text fields (id == name on the page).
    await _fill_field(page, "last_name",   params.get("last_name", ""), log)
    await _fill_field(page, "first_name",  params.get("first_name", ""), log)
    await _fill_field(page, "middle_name", params.get("middle_name", ""), log)
    # «Дата рождения» has a date+month mask — only fill a REAL date (with dots);
    # a bare year (1889) would be mangled by the mask, so we keep it for the
    # result-side year filter instead.
    _bd = params.get("birth_date", "")
    if "." in _bd:
        await _fill_field(page, "birth_date", _bd, log)
    elif _bd:
        log("  → год рождения учту при фильтре результатов (поле даты не трогаю)")
    await _fill_field(page, "birth_place_gubernia", params.get("gubernia", ""), log)
    await _fill_field(page, "birth_place_uezd",     params.get("uezd", ""), log)
    await _fill_field(page, "birth_place_volost",   params.get("volost", ""), log)
    await _fill_field(page, "birth_place",          params.get("settlement", ""), log)
    await _fill_field(page, "rank",                 params.get("rank", ""), log)
    await _fill_field(page, "military_unit_name",   params.get("unit", ""), log)
    await _fill_field(page, "data_vibitiya",        params.get("event_from", ""), log)
    await _fill_field(page, "data_vibitiya_end",    params.get("event_to", ""), log)
    await _fill_field(page, "event_place",          params.get("event_place", ""), log)
    # «Событие» is a custom selectric over <select id="event_id"> — set the value
    # on the underlying select (language-independent).
    _ev = params.get("event", "")
    if _ev:
        try:
            r = await page.evaluate(r"""([sid, val]) => {
                const els = [...document.querySelectorAll(
                    '#' + CSS.escape(sid) + ', [name="' + sid + '"]')]
                    .filter(e => e.tagName === 'SELECT');
                if (!els.length) return 'missing';
                els.forEach(s => { s.value = val;
                    s.dispatchEvent(new Event('change', {bubbles: true})); });
                return 'ok:' + els.length;
            }""", ["event_id", str(_ev)])
            log(f"  ✓ event_id = {_ev} ({r})")
        except Exception as e:
            log(f"  !! select event_id: {type(e).__name__}")
    await _fill_field(page, "fund",      params.get("fund", ""), log)
    await _fill_field(page, "inventory", params.get("inventory", ""), log)
    await _fill_field(page, "file",      params.get("file", ""), log)

    # «Разделы» checkboxes — set each to the requested state (default: all on).
    sections = params.get("sections") or {}
    for label, cid in SECTIONS.items():
        want = sections.get(label, True)
        try:
            cb = page.locator(f"#{cid}").first
            if await cb.count():
                is_on = await cb.is_checked()
                if is_on != want:
                    lbl = page.locator(f'label[for="{cid}"]').first
                    target = lbl if await lbl.count() else cb
                    await target.click(timeout=2000)
        except Exception:
            pass

    # Let the user see the filled, expanded form for a moment before submitting.
    await asyncio.sleep(2.0)

    # Submit «Найти».
    clicked = False
    for sel in ('input.button-search-big[value="Найти"]',
                'input.button-search-big', 'input[type="submit"][value="Найти"]',
                'input[type="submit"]'):
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                clicked = True
                log(f"  ✓ Отправил поиск ({sel})")
                break
        except Exception:
            continue
    if not clicked:
        await page.keyboard.press("Enter")
        log("  → Поиск: Enter")

    for _ in range(30):
        await asyncio.sleep(1)
        try:
            ok = await page.evaluate(
                "() => document.querySelectorAll('a[href*=\"/heroes/chelovek\"]')"
                ".length > 0 || /найдено|ничего не/i.test(document.body.innerText)")
            if ok:
                return True
        except Exception:
            pass
    return True


def _with_page(url: str, n: int) -> str:
    """Return url with the `page` query parameter set to n."""
    sp = urlsplit(url)
    q = [(k, v) for k, v in parse_qsl(sp.query, keep_blank_values=True)
         if k != "page"]
    q.append(("page", str(n)))
    return urlunsplit((sp.scheme, sp.netloc, sp.path, urlencode(q, doseq=True),
                       sp.fragment))


_COLLECT_JS = r"""() => {
    const norm = s => (s || '').replace(/\s+/g, ' ').trim();
    const out = [], seen = new Set();
    // Result links carry class «heroes-list-item-name» regardless of path — normal
    // records are /heroes/chelovek…, FAMOUS persons are /heroes/commander… (Военачальники)
    // or /heroes/person…. Match the CLASS first (covers every type), then href patterns.
    document.querySelectorAll(
            'a.heroes-list-item-name, a[href*="/heroes/chelovek"], '
            + 'a[href*="/heroes/person"], a[href*="/heroes/commander"]').forEach(a => {
        const href = a.href || '';
        const name = norm(a.textContent);
        if (!href || !name || seen.has(href)) return;
        seen.add(href);
        // snippet = the smallest ancestor that adds context beyond the name
        let box = a;
        for (let up = 0; up < 5 && box.parentElement; up++) {
            box = box.parentElement;
            if (norm(box.textContent).length > name.length + 15) break;
        }
        out.push({href, name, snippet: norm(box ? box.textContent : '').slice(0, 400)});
    });
    return out;
}"""


async def _collect_results(page, params, log, max_pages, max_records) -> list:
    want = (params.get("last_name", ""), params.get("first_name", ""),
            params.get("middle_name", ""))
    exact = bool(params.get("exact"))
    log(f"  → точное совпадение ФИО: {'ВКЛ' if exact else 'выкл'}")
    ym = re.search(r"(18|19|20)\d{2}", params.get("birth_date", "") or "")
    want_year = ym.group(0) if ym else ""
    out, seen = [], set()
    page_no = 1
    while page_no <= max_pages and len(out) < max_records:
        try:
            rows = await page.evaluate(_COLLECT_JS)
        except Exception:
            rows = []
        new = 0
        for r in rows:
            href = r.get("href", "")
            if not href or href in seen:
                continue
            seen.add(href)
            name = r.get("name", "")
            if not _person_matches(name, *want, exact=exact):
                continue
            if want_year:
                yrs = re.findall(r"(?:18|19|20)\d{2}", r.get("snippet", ""))
                if yrs and want_year not in yrs:
                    continue
            out.append({"name": name, "href": href, "snippet": r.get("snippet", "")})
            new += 1
            if len(out) >= max_records:
                break
        log(f"  → Страница {page_no}: ссылок {len(rows)}, подходящих +{new}, "
            f"всего {len(out)}")
        if len(out) >= max_records:
            break
        # Next page is an AJAX link «<a data-page="N">» (its href is empty). Click
        # the VISIBLE one (links are duplicated mobile+desktop). NO such link → it
        # was the last/only page → stop. Do NOT navigate by a ?page= URL: gwar
        # loads results via AJAX, so that just reopens the empty search form
        # (which is exactly the «возвращаешься к пустой форме» bug).
        first_before = rows[0]["href"] if rows else ""
        try:
            clicked = await page.evaluate(r"""(n) => {
                const els = [...document.querySelectorAll('a[data-page="' + n + '"]')];
                const el = els.find(e => e.offsetParent !== null) || els[0];
                if (el) { el.click(); return true; }
                return false;
            }""", str(page_no + 1))
        except Exception:
            clicked = False
        if not clicked:
            break                       # no «next page» link → last/only page
        changed = False
        for _ in range(15):
            await asyncio.sleep(1)
            try:
                cur = await page.evaluate(
                    "() => { const a = document.querySelector("
                    "'a[href*=\"/heroes/chelovek\"]'); return a ? a.href : ''; }")
            except Exception:
                cur = ""
            if cur and cur != first_before:
                changed = True; break
        if not changed:
            break
        page_no += 1
    return out


# ── Record (detail) page ──────────────────────────────────────────────────── #
_DETAIL_JS = r"""(labels) => {
    const norm = s => (s || '').replace(/ /g, ' ').replace(/\s+/g, ' ').trim();
    const h1 = document.querySelector('h1');
    const name = h1 ? norm(h1.textContent) : '';
    // card type = first non-empty short text right after the h1
    let typ = '';
    if (h1) {
        let n = h1.nextElementSibling;
        while (n && !norm(n.textContent)) n = n.nextElementSibling;
        if (n) typ = norm(n.textContent).slice(0, 90);
    }
    // FREE-TEXT blocks (Биография / Описание / История …): everything written on the
    // card, not just the labelled fields. Grab leaf prose blocks (≥ 80 chars, no block
    // children), drop site chrome; dedup; keep order.
    const seen = new Set(), descParts = [];
    // Site chrome / notices / system text that must NEVER land in the card — the user
    // pasted these verbatim from a bad gwar_Иванов.docx.
    const JUNK = new RegExp([
        'О проекте','Урок Победы','Министерств','обратной связи','©','cookie',
        'Памяти героев','слишком старую версию','обновите ваш браузер','устаревш',
        'Уважаемые пользователи','временно ограничен','доступ к личным архивам',
        'технологических работ','высокой нагрузк','Новое в 20\\d\\d',
        'Добавлены документы','защитников Отечества','обнаружения технических',
        'некорректных данных','Персоналий на образе','точках боевых действий',
        'Боевой путь строится','координат воинской','Приносим извинения'
    ].join('|'), 'i');
    for (const el of document.querySelectorAll('p, div, span, article, section')) {
        if (el.querySelector && el.querySelector('p, div, ul, ol, table, h1, h2, h3'))
            continue;                                    // containers → skip (leaf only)
        const t = norm(el.textContent);
        if (t.length < 80 || JUNK.test(t) || seen.has(t)) continue;
        // the page-number strip «1234567891011…755» = a leaf that is (almost) all digits
        const digits = (t.match(/\d/g) || []).length;
        if (digits > 40 && digits > t.length * 0.6) continue;
        seen.add(t); descParts.push(t);
    }
    return {name, typ, body: norm(document.body.innerText),
            desc: descParts.join('\n\n')};
}"""


async def _grab_doc_scans(page, images_dir, base, rec, log):
    """Save the scanned document — every page (Страница X из N). Per page: click
    #btnSaveImage (catch the download); if that fails, download the largest
    displayed image via the browser. Move pages with td.to-right-arrow."""
    images_dir.mkdir(parents=True, exist_ok=True)
    # how many pages does the document have?
    pages = 1
    try:
        t = await page.evaluate(
            "() => { const m=(document.body.innerText||'').match(/из\\s*(\\d+)/);"
            " return m ? m[1] : '1'; }")
        pages = max(1, min(int(t), 20))
    except Exception:
        pages = 1

    saved = 0
    for pi in range(pages):
        await asyncio.sleep(0.6)
        ok = False
        # 1) save button (icon-save-dropdown) → catch download. SHORT timeout: when
        # the dropdown doesn't produce a download we must NOT hang 15s — the image
        # fallback below is what actually works on these cards.
        try:
            btn = page.locator("#btnSaveImage").first
            if await btn.count():
                async with page.expect_download(timeout=5000) as dl:
                    await btn.click(timeout=4000)
                    await asyncio.sleep(0.4)
                    for sel in ('.icon-save-dropdown a', '.icon-save-dropdown li',
                                'a:has-text("Скачать")', 'a[href*="download" i]'):
                        it = page.locator(sel).first
                        try:
                            if await it.count() and await it.is_visible():
                                await it.click(timeout=2000); break
                        except Exception:
                            continue
                d = await dl.value
                suf = Path(d.suggested_filename or "scan.jpg").suffix or ".jpg"
                dest = images_dir / f"{base}_лист{pi + 1}{suf}"
                await d.save_as(str(dest))
                rec["scans"].append(str(dest)); saved += 1; ok = True
                log(f"      🖼 документ сохранён: {dest.name}")
        except Exception:
            ok = False
        # 2) fallback: largest displayed image → download via browser
        if not ok:
            try:
                src = await page.evaluate(r"""() => {
                    let best = '', area = 0;
                    for (const im of document.querySelectorAll('img[src]')) {
                        const s = (im.src || '').toLowerCase();
                        if (!s.startsWith('http')) continue;
                        if (/\.svg|sprite|logo|icon|placeholder|avatar/.test(s)) continue;
                        const a = (im.naturalWidth||0) * (im.naturalHeight||0);
                        if (a > area && a > 60000) { area = a; best = im.src; }
                    }
                    return best;
                }""")
                if src:
                    body = await _img_bytes_via_goto(page.context, src, page.url)
                    if body:
                        dest = images_dir / f"{base}_лист{pi + 1}.jpg"
                        dest.write_bytes(body)
                        rec["scans"].append(str(dest)); saved += 1; ok = True
                        log(f"      🖼 документ сохранён (img): {dest.name}")
            except Exception:
                pass
        if not ok:
            log(f"      (лист {pi + 1}: скан не сохранился)")
        # next page within the document
        if pi < pages - 1:
            try:
                nx = page.locator("td.to-right-arrow, .to-right-arrow").first
                if await nx.count():
                    await nx.click(timeout=3000)
                else:
                    break
            except Exception:
                break
    return saved > 0


async def _extract_record(page, url, images_dir, base_dir, log, result_name=""):
    rec = {"name": "", "type": "", "url": url, "fields": [], "scans": []}
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=30000)
        await asyncio.sleep(1.2)
        info = await page.evaluate(_DETAIL_JS, _GWAR_LABELS)
        rec["name"] = info.get("name") or result_name
        rec["type"] = info.get("typ") or ""
        rec["fields"] = _parse_fields(info.get("body") or "")
        rec["desc"] = _clean_block(info.get("desc") or "")     # full free text (Биография…)
        base = safe_fn(rec.get("name") or result_name or "record")
        log(f"      полей: {len(rec['fields'])}, тип: {rec['type'] or '—'}"
            + (f", описание: {len(rec['desc'])} симв." if rec.get("desc") else ""))
        await _grab_doc_scans(page, images_dir, base, rec, log)
    except Exception as e:
        log(f"      !! страница записи ({type(e).__name__})")
    rec["fields"] = [(k, _clean_block(v)) for k, v in rec.get("fields", [])
                     if _clean_block(v)]
    return rec


# ── Word output ───────────────────────────────────────────────────────────── #
def _kv_table(doc, pairs):
    if not pairs:
        return
    tbl = doc.add_table(rows=0, cols=2); tbl.style = "Table Grid"
    for k, v in pairs:
        r = tbl.add_row().cells
        r[0].text = str(k); set_cell_lines(r[1], v)
        for run in r[0].paragraphs[0].runs:
            run.bold = True


def _docx_add_record(doc, i, rec):
    name = rec.get("name") or f"Запись {i}"
    title = name + (f" — {rec['type']}" if rec.get("type") else "")
    hp = doc.add_paragraph(); hr = hp.add_run(f"{i}. {title}")
    hr.bold = True; hr.font.size = Pt(13)

    if rec.get("fields"):
        _kv_table(doc, rec["fields"])

    # full free text from the card (Биография / описание) — each paragraph kept
    if rec.get("desc"):
        for para in str(rec["desc"]).split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    for img in rec.get("scans", []):
        try:
            png = _to_png(Path(img).read_bytes())
            if png:
                doc.add_picture(io.BytesIO(png), width=Inches(4.5))
                p = doc.add_paragraph(); p.add_run("Файл: ").bold = True
                p.add_run(str(Path(img).resolve()))      # exact path where it was saved
        except Exception:
            pass

    if rec.get("url"):
        clean = re.sub(r"\?.*$", "", rec["url"]) or rec["url"]
        p = doc.add_paragraph()
        _add_hyperlink(p, "Источник", clean)
    doc.add_paragraph("")


def write_docx(path, records, qlines, append=False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    existing = append and Path(path).exists()
    if existing:
        doc = Document(str(path)); doc.add_page_break()
        ap = doc.add_paragraph(); ar = ap.add_run(f"➕ Добавлено {len(records)}")
        ar.bold = True; ar.font.size = Pt(13)
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width = Mm(210); s.page_height = Mm(297)
        s.left_margin = s.right_margin = Mm(18)
        s.top_margin = s.bottom_margin = Mm(18)
        ht = doc.add_paragraph()
        htr = ht.add_run("Памяти героев Великой войны — результаты поиска")
        htr.bold = True; htr.font.size = Pt(14)
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ln in qlines:
            doc.add_paragraph(ln)
        doc.add_paragraph(f"Найдено записей: {len(records)}")
        doc.add_paragraph("")
    for i, rec in enumerate(records, 1):
        _docx_add_record(doc, i, rec)
    add_page_numbers(doc)
    doc.save(str(path))


# ── Main entry point ──────────────────────────────────────────────────────── #
# gwar reflects sidebar facets in the RESULTS URL — the clickable facet rows do NOT
# toggle through Playwright .click() («Выбрано 0/10»). Map «Источники информации» →
# type code (+ group); «Известные личности» → filters=famous_type[…]. Verified against
# the user's working URL: /heroes/?groups=awd:ptr:frc:cmd:prs&types=awd_nagrady:…:
# prs_person&last_name=ivanov&filters=famous_type[Военачальники]  → Найдено: 1 (general).
SOURCE_TYPE_CODES = {
    "Документы о награждениях":          ("awd_nagrady", "awd"),
    "Наградная картотека":               ("awd_kart", "awd"),
    "Именные списки потерь":             ("potery_doneseniya_o_poteryah", "ptr"),
    "Картотека потерь":                  ("potery_gospitali", "ptr"),
    "Паспорта захоронений":              ("potery_spiski_zahoroneniy", "ptr"),
    "Картотека военнопленных":           ("potery_voennoplen", "ptr"),
    "Послужные списки":                  ("frc_list", "frc"),
    "В справочнике командного состава":  ("cmd_commander", "cmd"),
    "В известных личностях":             ("prs_person", "prs"),
}
DEFAULT_GROUPS = "awd:ptr:frc:cmd:prs"
DEFAULT_TYPES  = ":".join(c for c, _g in SOURCE_TYPE_CODES.values())

_TRANSLIT = {'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'e',
             'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'i', 'к': 'k', 'л': 'l', 'м': 'm',
             'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
             'ф': 'f', 'х': 'kh', 'ц': 'ts', 'ч': 'ch', 'ш': 'sh', 'щ': 'shch',
             'ъ': '', 'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya'}


def _translit(s: str) -> str:
    return ''.join(_TRANSLIT.get(c, c) for c in (s or '').lower())


async def _click_facet(page, value, log):
    """Toggle a gwar sidebar facet by its data-facet-value, with a REAL mouse click.
    gwar's custom checkbox does NOT toggle on a JS .click() (that was the «Выбрано
    0/10» bug) — a genuine Playwright mouse event does. Match is whitespace-normalized
    because the live site has double spaces (e.g. «Георгиевский крест  IV-й степени»)."""
    try:
        ok = await page.evaluate(r"""(want) => {
            const norm = s => (s||'').replace(/\s+/g,' ').trim();
            const w = norm(want);
            document.querySelectorAll('[data-pw-facet]').forEach(
                e => e.removeAttribute('data-pw-facet'));
            for (const el of document.querySelectorAll('[data-facet-value]')) {
                if (norm(el.getAttribute('data-facet-value')) === w) {
                    (el.closest('.field-check-box') || el).setAttribute('data-pw-facet','1');
                    return true;
                }
            }
            return false;
        }""", value)
    except Exception as e:
        log(f"    !! поиск фасета «{value}»: {type(e).__name__}")
        return False
    if not ok:
        log(f"    !! фасет не найден на странице: «{value}»")
        return False
    clicked = False
    try:
        await page.locator('[data-pw-facet="1"]').first.click(timeout=4000)   # REAL click
        clicked = True
        log(f"    ✓ фасет отмечен: «{value}»")
    except Exception as e:
        log(f"    !! клик по фасету «{value}»: {type(e).__name__}")
    try:
        await page.evaluate("() => document.querySelectorAll('[data-pw-facet]')"
                            ".forEach(e => e.removeAttribute('data-pw-facet'))")
    except Exception:
        pass
    return clicked


async def _apply_award_loss_clicks(page, awards, losses, log):
    """«Награды»/«Потери» have no URL key, so apply them EXACTLY like the user showed:
    real-click each data-facet-value, with a 5-sec pause BEFORE each («перед КАЖДЫМ
    фильтром 5 сек»), then ONE «Применить» so all picked filters apply in a single
    search (gwar's facets combine; cb79c8d = one search, not one-per-filter)."""
    to_click = list(awards) + list(losses)
    if not to_click:
        return
    log(f"  Награды/Потери — кликаю {len(to_click)} фасет(ов) по data-facet-value:")
    for val in to_click:
        await asyncio.sleep(5)                      # 5 sec BEFORE each filter
        log(f"    (жду 5с) кликаю: «{val}»")
        await _click_facet(page, val, log)
    # ONE «Применить» for everything together (famous_type from the URL stays checked
    # in the sidebar, so it is re-applied alongside the just-clicked awards/losses).
    applied = False
    for sel in ('input.button-search-big[value="Применить"]',
                'input[value="Применить"]', 'button:has-text("Применить")',
                'a:has-text("Применить")', '.heroes-filter-apply', '.filter-apply'):
        try:
            btn = page.locator(sel).first
            if await btn.count() and await btn.is_visible():
                await btn.click(timeout=4000)
                applied = True
                log(f"  «Применить» нажата ({sel}) — все фильтры одним поиском")
                break
        except Exception:
            continue
    if not applied:
        log("  «Применить» не найдена — фасеты применяются по клику (AJAX), это ок")
    await asyncio.sleep(5)                           # wait for the filtered re-search
    _LINK_SEL = ("() => document.querySelectorAll('a.heroes-list-item-name, "
                 "a[href*=\"/heroes/chelovek\"], a[href*=\"/heroes/person\"], "
                 "a[href*=\"/heroes/commander\"]').length")
    links = 0
    for _ in range(40):
        try:
            links = int(await page.evaluate(_LINK_SEL))
        except Exception:
            links = 0
        if links:
            break
        await asyncio.sleep(0.5)
    try:
        found = await page.evaluate(
            "() => (document.body.innerText.match(/Найдено документов:\\s*([\\d\\s]+)/)"
            " || [])[1] || '?'")
        log(f"  После наград/потерь: Найдено документов = {found}, ссылок = {links}")
    except Exception:
        log(f"  После наград/потерь: ссылок = {links}")


async def _apply_facets_url(page, params, sec, log):
    """Apply the chosen facets via the RESULTS URL (gwar's clickable facet rows don't
    toggle through Playwright). «Источники информации» → types=/groups=, «Известные
    личности» → filters=famous_type[…]. Rebuild the current results URL with those
    params and navigate, then wait for the AJAX result list.

    The URL carries ONLY last_name (+ facets) — first/patronymic are dropped here and
    applied fuzzily in _collect_results (this mirrors the user's working URL, which had
    just last_name + filters, and avoids the famous DB returning 0 on a strict ФИО)."""
    base = page.url or (BASE + "/heroes/")
    sp = urlsplit(base)
    q = dict(parse_qsl(sp.query, keep_blank_values=True))

    # «Источники информации» → types / groups (subset narrows; none → defaults)
    codes, groups = [], []
    for ru in (sec.get("info_sources") or []):
        cg = SOURCE_TYPE_CODES.get(ru)
        if cg and cg[0] not in codes:
            codes.append(cg[0])
            if cg[1] not in groups:
                groups.append(cg[1])
    q["types"]  = ":".join(codes)  if codes  else (q.get("types")  or DEFAULT_TYPES)
    q["groups"] = ":".join(groups) if groups else (q.get("groups") or DEFAULT_GROUPS)

    # name: KEEP last_name + first_name + middle_name (Cyrillic, NOT translit — gwar's
    # URL search is Cyrillic, the live URL was last_name=%D0%98%D0%B2%D0%B0%D0%BD…=Иванов).
    # The user SEARCHES by имя+отчество — dropping first/patronymic gave «100 325 Ивановых»
    # instead of «Иванов Григорий …». Use the form-search URL's value if present, else the
    # GUI param as-is; an empty field is removed so it doesn't blank-filter.
    for k in ("last_name", "first_name", "middle_name"):
        v = (params.get(k) or "").strip()
        if q.get(k):
            continue                       # keep what the form search already put in URL
        if v:
            q[k] = v                       # Cyrillic, exactly as typed
        else:
            q.pop(k, None)

    # «Известные личности» → filters=famous_type[…] in the URL (CONFIRMED working —
    # user's URL had filters=famous_type[Военачальники]). «Награды»/«Потери» have NO
    # URL key I know, so they are NOT put in the URL — they are applied right after by
    # REAL-clicking their data-facet-value spans (see below), exactly the elements the
    # user pointed at: <span class="field-check-box-name-value" data-facet-value="…">.
    filt = []
    notable = [v for v in (sec.get("notable") or []) if v]
    awards  = [v for v in (sec.get("awards")  or []) if v]
    losses  = [v for v in (sec.get("losses")  or []) if v]
    if notable:
        filt.append("famous_type[" + ",".join(notable) + "]")
    if filt:
        q["filters"] = ":".join(filt)
    q["page"] = "1"

    url = urlunsplit((sp.scheme or "https", sp.netloc or "gwar.mil.ru",
                      sp.path or "/heroes/", urlencode(q, safe="[],:"), ""))
    log(f"  Применяю фильтры через URL: {url}")
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=40000)
    except Exception as e:
        log(f"  !! goto фильтр-URL: {type(e).__name__}")
        return

    # gwar updates the «Найдено» COUNT first, then renders the result CARDS a moment
    # later. The «ссылок 0» bug was breaking the wait as soon as the count text showed
    # (Найдено=1) — before the card link existed. So: wait a fixed 5s for the AJAX to
    # UPDATE (user: «жди 5 секунд»), THEN poll for the actual result link, NOT the count.
    await asyncio.sleep(5)
    _LINK_SEL = ("() => document.querySelectorAll('a.heroes-list-item-name, "
                 "a[href*=\"/heroes/chelovek\"], a[href*=\"/heroes/person\"], "
                 "a[href*=\"/heroes/commander\"]').length")
    links = 0
    for _ in range(40):                        # up to ~20s more for the card to render
        try:
            links = int(await page.evaluate(_LINK_SEL))
        except Exception:
            links = 0
        if links:
            break
        await asyncio.sleep(0.5)
    try:
        found = await page.evaluate(
            "() => (document.body.innerText.match(/Найдено документов:\\s*([\\d\\s]+)/)"
            " || [])[1] || '?'")
        log(f"  После URL-фильтра: Найдено документов = {found}, ссылок = {links}")
    except Exception:
        pass
    if not links:
        # Diagnostic: dump the result-area links so we learn the famous-person URL
        # format (the famous filter returns a link my selector didn't match).
        try:
            dump = await page.evaluate(
                "() => Array.from(document.querySelectorAll('a[href*=\"/heroes/\"]'))"
                ".map(a => a.getAttribute('href') + ' :: ' + "
                "(a.textContent||'').replace(/\\s+/g,' ').trim().slice(0,45))"
                ".filter(s => !/\\/heroes\\/?($| ::)/.test(s)).slice(0, 20)")
            if dump:
                log("  !! ссылок 0 — ссылки /heroes/ на странице (диагностика):")
                for s in dump:
                    log(f"      {s}")
        except Exception:
            pass

    # Now apply «Награды»/«Потери» by real-clicking their data-facet-value (no URL
    # key known), 5 sec before each, then ONE «Применить» — the notable (famous_type)
    # filter from the URL stays applied alongside them.
    await _apply_award_loss_clicks(page, awards, losses, log)


async def run_scraper(*,
    last_name="", first_name="", middle_name="",
    birth_date="", gubernia="", uezd="", volost="", settlement="",
    rank="", unit="", event="", event_from="", event_to="", event_place="",
    fund="", inventory="", file="",
    sections=None, exact=False,
    info_sources=None, awards=None, losses=None, notable=None,  # sidebar facets
    output_folder=Path("."),
    log=print,
    progress=None,
    cancel_event=None,
    ask_file_conflict=None,
    max_pages=20,
    max_records=80,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    params = dict(
        last_name=last_name.strip(), first_name=first_name.strip(),
        middle_name=middle_name.strip(), birth_date=birth_date.strip(),
        gubernia=gubernia.strip(), uezd=uezd.strip(), volost=volost.strip(),
        settlement=settlement.strip(), rank=rank.strip(), unit=unit.strip(),
        event=event, event_from=event_from.strip(), event_to=event_to.strip(),
        event_place=event_place.strip(), fund=fund.strip(),
        inventory=inventory.strip(), file=file.strip(),
        sections=sections or {}, exact=bool(exact),
        info_sources=info_sources or [], awards=awards or [],
        losses=losses or [], notable=notable or [])

    output_folder = Path(output_folder); output_folder.mkdir(parents=True, exist_ok=True)
    qkey = " ".join(p for p in (last_name, first_name, middle_name) if p) or "gwar"
    images_dir = output_folder / "images" / "Памяти героев Великой войны" / (safe_fn(qkey) or "gwar")
    summary = {"ok": False}

    if not any(params.get(k) for k in
               ("last_name", "first_name", "middle_name", "rank", "unit",
                "settlement", "gubernia", "uezd")):
        _prog(100, "Пустой запрос.")
        return summary

    qlines = [f"Запрос: {qkey}", "Сайт: gwar.mil.ru (1914–1918)"]

    _prog(0, "Запускаю браузер…")
    async with async_playwright() as pw:
        browser = await browser_util.launch(pw, 
            headless=False,
            args=["--start-maximized", "--disable-blink-features=AutomationControlled"])
        # gwar.mil.ru uses a Russian government root CA that Chromium does not
        # trust → Page.goto throws ERR_CERT_AUTHORITY_INVALID. Ignore HTTPS errors
        # (this is why the browser shows «Не защищено» but still loads the site).
        ctx = await browser.new_context(no_viewport=True, accept_downloads=True,
                                        ignore_https_errors=True)
        page = await ctx.new_page()
        try:
            # ONE search → tick ALL chosen facets together (sources + awards + losses
            # + notable are multi-select checkboxes that combine) → Применить → collect
            # → ONE document. NOT one search per filter (that re-searched the same name
            # over and over and only ever applied 1/1).
            _sec = {k: [v for v in (params.get(k) or []) if v]
                    for k in ("info_sources", "awards", "losses", "notable")}
            all_facets = (_sec["info_sources"] + _sec["awards"]
                          + _sec["losses"] + _sec["notable"])
            # Show EXACTLY what each GUI section sent — so «выбрал всю первую секцию,
            # но не выбрал военачальников» is visible at a glance (and: gwar OR-s the
            # facets, so adding the 9 «Information sources» broadens to ~2000 and
            # drowns a narrow «Notable persons → Военачальники» pick).
            log(f"  Фильтры из GUI — источники: {len(_sec['info_sources'])}, "
                f"награды: {len(_sec['awards'])}, потери: {len(_sec['losses'])}, "
                f"известные личности: {len(_sec['notable'])}")
            if _sec["notable"]:
                log(f"    известные личности: {', '.join(_sec['notable'])}")

            _prog(5, "Поиск…")
            if not await _do_search(page, params, log):
                summary["message"] = "Не удалось выполнить поиск."
                return summary
            if _done():
                return summary
            if all_facets:
                log(f"  Применяю фильтров: {len(all_facets)}")
                await _apply_facets_url(page, params, _sec, log)
                if _done():
                    return summary

            _prog(15, "Сбор результатов…")
            recs_meta = await _collect_results(page, params, log, max_pages, max_records)
            log(f"  Подходящих записей: {len(recs_meta)}")

            records = []
            n = len(recs_meta)
            for i, rm in enumerate(recs_meta, 1):
                if _done():
                    break
                _prog(20 + int(70 * i / max(n, 1)), f"[{i}/{n}] {rm['name'][:50]}…")
                log(f"  [{i}/{n}] {rm['name']}")
                try:
                    dp = await ctx.new_page()
                    try:
                        rec = await _extract_record(dp, rm["href"], images_dir,
                                                    output_folder, log,
                                                    result_name=rm["name"])
                    finally:
                        await dp.close()
                    if not rec.get("name"):
                        rec["name"] = rm["name"]
                    records.append(rec)
                except Exception as _e:
                    log(f"      !! пропускаю запись ({type(_e).__name__})")
                    records.append({"name": rm["name"], "type": "", "url": rm["href"],
                                    "fields": [], "scans": []})
                await asyncio.sleep(0.3)

            if records:
                _prog(92, "Сохранение…")
                base = safe_fn(f"gwar_{qkey}") or "gwar_results"
                docx_p = output_folder / f"{base}.docx"
                decision = "overwrite"
                if docx_p.exists() and ask_file_conflict:
                    try:
                        decision = (ask_file_conflict([docx_p.name]) or "overwrite").lower()
                    except Exception:
                        decision = "overwrite"
                    log(f"  → Файл существует → {decision}")
                if decision != "skip":
                    try:
                        write_docx(docx_p, records, qlines, append=(decision == "append"))
                        log(f"  → Word: {docx_p.name}")
                    except PermissionError:
                        alt = output_folder / f"{base}_{time.strftime('%H%M%S')}.docx"
                        write_docx(alt, records, qlines, append=False)
                        log(f"  !! файл занят (открыт в Word) → сохранил как {alt.name}")

            if _done():
                log("  ⛔ Отменено пользователем — сохранил собранное.")
            _prog(100, f"Готово — {len(records)} запис(ей).")
            summary.update({"ok": True, "n_records": len(records),
                            "output_folder": str(output_folder)})
        except Exception as exc:
            summary["message"] = f"{type(exc).__name__}: {exc}"
            log(f"  !! Ошибка: {exc}")
        finally:
            try:
                await browser.close()
            except Exception:
                pass
    return summary
