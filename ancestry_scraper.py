#!/usr/bin/env python3
"""
ancestry_scraper.py
===================
Ancestry.com scraper — built on the SAME pattern as familysearch_scraper.py
(persistent login profile, sign in ONCE before opening records, per-field exact
match, one scrape pass, Word + Excel output with the site name, overwrite/append/
skip dialog, never crash on a broken page).

Ancestry is a paid, login-walled, anti-bot site (like MyHeritage / FamilySearch):
- PERSISTENT profile `.ancestry_profile` keeps the login between runs — sign in
  once, not every run. NO hardcoded user_agent (stale UA = bot tell).
- Sign in ONCE up front, before opening any record.

Flow:
  1. Home page form → #firstName, #lastName, #place, #birthYear → Search.
  2. The results URL (/search/?name=First+Middle_Last&birth=YEAR…) is then
     augmented for the requested EXACT flags / spouse-parent filters and reloaded
     — this narrows the result set BEFORE scraping (like FS advanced search).
  3. Collect result rows (record links), keep name-matches ≥ MIN_MATCH.
  4. Open each record on ONE page, copy every field, save the document image
     (best-effort via the image viewer).

NOTE (live tuning): the home-form fields, the login fields, the results/edit/
more-options buttons and the search-URL scheme are confirmed; the result-row,
record-detail and image-viewer selectors are defensive/generic and dump a
diagnostic on 0 results — refine them from the first live run's log.
"""

import asyncio, difflib, io, json, os, re, shutil, sys, time
from pathlib import Path
from urllib.parse import urlparse, parse_qsl, urlencode, quote
from docx_util import set_cell_lines

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

# ── Config ────────────────────────────────────────────────────────────────── #
_HERE = Path(__file__).resolve().parent
try:
    _CFG = json.loads((_HERE / "config" / "ancestry.json").read_text("utf-8"))
except Exception:
    _CFG = {}

HOME_URL   = _CFG.get("home_url", "https://www.ancestry.com/")
SIGNIN_URL = _CFG.get("signin_url", "https://www.ancestry.com/account/signin")
SEARCH_URL = "https://www.ancestry.com/search/"
ANC_BASE   = "https://www.ancestry.com"
SITE_NAME  = "Ancestry"
MIN_MATCH  = int(_CFG.get("min_match", 80))
MAX_PAGES  = int(_CFG.get("max_pages", 1))     # result pages (count=50 → page 1 is
                                               # plenty; >1 re-collected dups and
                                               # dropped click-applied filters)
MAX_SCRAPE = int(_CFG.get("max_scrape", 40))   # cap records actually opened
ANC_PROFILE_DIR = _HERE / ".ancestry_profile"  # persistent login/cookies
_dl           = _CFG.get("downloads_dir", "")
DOWNLOADS_DIR = Path(_dl) if _dl else Path.home() / "Downloads"
HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")
_IMG_SKIP = ("icon", "logo", "sprite", "avatar", "pixel", "placeholder", ".svg",
             "static.ancestry", "/css/", "gravatar")
# record links worth opening
_REC_PATTERNS = ("/discoveryui-content/view/", "/imageviewer/",
                 "/family-tree/person/", "recordpid", "/cgi-bin/sse.dll")


# ── Utils ─────────────────────────────────────────────────────────────────── #

def _sim(a: str, b: str) -> float:
    a = re.sub(r"\s+", " ", (a or "").strip().lower())
    b = re.sub(r"\s+", " ", (b or "").strip().lower())
    if not a or not b:
        return 0.0
    wa, wb = set(a.split()), set(b.split())
    return max(len(wa & wb) / max(len(wa), 1),
               difflib.SequenceMatcher(None, a, b).ratio()) * 100


def safe_fn(s: str, n: int = 100) -> str:
    return re.sub(r"\s+", " ",
                  re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip()))[:n].strip() \
        or "document"


def _abs(href: str) -> str:
    return ANC_BASE + href if href.startswith("/") else href


def _is_record(href: str) -> bool:
    h = (href or "").lower()
    return any(p in h for p in _REC_PATTERNS)


def _name_part(s: str) -> str:
    """«Alexander W» → «Alexander+W» (spaces → +, as in Ancestry's name= param)."""
    return "+".join((s or "").split())


# ── Name matching (surname-aware — score by the PERSON, not the collection) ──── #

def _ratio(a: str, b: str) -> float:
    return difflib.SequenceMatcher(None, a, b).ratio()


def _match_name(first_q: str, last_q: str, full: str) -> float:
    """Score a result's PERSON name against the query. Surname must match closely
    (≥0.90 — «Sanders» keeps «Alexander Wolf Sanders» but rejects «Snyder» /
    «Sanderson» / «Saunders»); the given name matches by initial / prefix / sim."""
    toks = re.findall(r"[A-Za-zА-Яа-яёЁ]+", (full or "").lower())
    if not toks:
        return 0.0
    lastq  = (last_q or "").strip().lower()
    firstq = (first_q or "").strip().lower().split()
    sur = max((_ratio(lastq, t) for t in toks), default=0.0) if lastq else 1.0
    # surname must match closely: «Sanders»=1.0 kept; «Saunders»≈0.93 /
    # «Sanderson»≈0.88 / «Snyder»≈0.6 rejected (user wants exactly this surname).
    if lastq and sur < 0.94:
        return round(sur * 60, 1)
    if not firstq:
        return round(85 + 14 * sur, 1)
    g = firstq[0]
    giv = any(
        t == g
        or (len(g) == 1 and t.startswith(g))
        or (len(g) > 1 and t.startswith(g[:2]))
        or _ratio(g, t) >= 0.7
        for t in toks)
    return 99.0 if giv else 72.0


# ── Search-URL building ────────────────────────────────────────────────────── #

def _build_search_url(fn, ln, place, year, year_range, exact, adv) -> str:
    """Ancestry /search/ URL — built by hand (spaces are a literal «+»).
    name=First+Middle_Last, birth=YEAR&birth_x=<N>-0-0 (N = ± year range, 0=exact),
    residence=…, father/mother/spouse/sibling/child=… (LISTS), count=50.
    `<field>_x=1` only for the fields whose EXACT box is ticked."""
    exact = exact or {}
    adv   = adv or {}
    parts = []

    if fn or ln:
        parts.append(("name", f"{_name_part(fn)}_{_name_part(ln)}".strip("_")))
        if exact.get("name") or exact.get("surname"):
            parts.append(("name_x", "1"))   # Ancestry name is one unit (given+surname)
    if year:
        y = re.sub(r"\D", "", str(year))[:4]
        if y:
            parts.append(("birth", y))
            n = 1 if year_range is None else int(year_range)
            parts.append(("birth_x", f"{n}-0-0"))     # ±N years (0 = this year)
    if place:
        parts.append(("residence", _name_part(place)))
        if exact.get("place"):
            parts.append(("residence_x", "1"))
    for rel in ("father", "mother", "spouse", "sibling", "child"):
        for person in adv.get(rel, []) or []:
            rf, rl = person.get("first", ""), person.get("last", "")
            if rf or rl:
                parts.append((rel, f"{_name_part(rf)}_{_name_part(rl)}".strip("_")))
                # Ancestry's <rel>_x is one unit → exact if either field is exact
                if (person.get("first_exact") or person.get("last_exact")
                        or person.get("exact")):
                    parts.append((f"{rel}_x", "1"))
    # life events (Marriage / Death / Lived-In) — year(+range) and/or place
    _EV = {"marriage": "marriage", "death": "death", "residence": "residence"}
    for e in adv.get("events", []) or []:
        key = _EV.get(e.get("type", ""))
        if not key:
            continue                       # «Any Event» param unknown → skip
        y = re.sub(r"\D", "", str(e.get("year", "")))[:4]
        if y:
            parts.append((key, y))
            parts.append((f"{key}_x", f"{int(e.get('range', 1))}-0-0"))
        if e.get("place"):
            parts.append((f"{key}place", _name_part(e["place"])))
    if adv.get("keyword"):
        parts.append(("keyword", _name_part(adv["keyword"])))
    g = (adv.get("gender") or "").strip().lower()
    if g.startswith("m"):
        parts.append(("gender", "m"))
    elif g.startswith("f"):
        parts.append(("gender", "f"))
    if adv.get("race"):
        parts.append(("race", _name_part(adv["race"])))
    parts.append(("count", "50"))                     # 50 results per page
    q = "&".join(f"{k}={quote(v, safe='+_-')}" for k, v in parts)
    return f"{SEARCH_URL}?{q}" if q else SEARCH_URL


# ── Field typing ──────────────────────────────────────────────────────────── #

async def _type_field(page, sel: str, val: str, label: str, log) -> bool:
    if not val:
        return True
    try:
        el = page.locator(sel).first
        if not await el.count():
            log(f"  !! поле {label} не найдено ({sel})")
            return False
        await el.scroll_into_view_if_needed(timeout=4000)
        await el.click(timeout=3000)
        await asyncio.sleep(0.1)
        await page.keyboard.press("Control+a")
        await page.keyboard.press("Delete")
        await page.keyboard.type(val, delay=35)
        await asyncio.sleep(0.2)
        log(f"  OK  {label} = {val!r}")
        return True
    except Exception as e:
        log(f"  !! {label}: {e}")
        return False


# ── 1. SEARCH (built URL — the home form selectors are unreliable) ──────────── #

async def _search(page, fn, ln, place, year, year_range, exact, adv, log) -> str:
    url = _build_search_url(fn, ln, place, year, year_range, exact, adv)
    log(f"  Поиск: {url}")
    await page.goto(url, wait_until="domcontentloaded", timeout=40000)
    await asyncio.sleep(5)
    log(f"  Результаты: {page.url}")
    return url


# ── 2. LOGIN ──────────────────────────────────────────────────────────────── #

async def _login(page, email: str, password: str, log) -> bool:
    """Fill #username + #password, submit. Call ONCE per session."""
    log(f"  Страница логина: {page.url[:80]}")
    try:
        await page.locator('#username, input[name="username"]').first.wait_for(
            state="visible", timeout=20000)
    except Exception:
        log("  !! поле username не появилось")
        return False
    await asyncio.sleep(2)

    async def _fill(sel, val, name):
        for attempt in range(3):
            try:
                el = page.locator(sel).first
                await el.fill(val)
                await asyncio.sleep(0.5)
                if (await el.input_value()).strip():
                    log(f"  {name} заполнен")
                    return True
            except Exception as e:
                log(f"  !! fill {name}: {e}")
            await asyncio.sleep(1)
        return False

    if not await _fill('#username, input[name="username"]', email, "username"):
        return False
    if not await _fill('#password, input[name="password"]', password, "password"):
        return False

    for sel in ['#signInBtn', 'button[type="submit"]',
                'button:has-text("Sign in")', 'button:has-text("Sign In")']:
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                log(f"  Sign in нажат ({sel})")
                break
        except Exception:
            continue

    try:
        await page.wait_for_url(
            lambda u: "ancestry.com" in u and "signin" not in u and "login" not in u,
            timeout=30000)
    except Exception:
        await asyncio.sleep(6)
    ok = "signin" not in page.url and "login" not in page.url
    log(f"  Логин {'OK ✓' if ok else 'ПРОВАЛИЛСЯ ✗'}  URL: {page.url[:80]}")
    return ok


async def _sign_in_if_needed(page, email, password, logged_in_ref, log) -> bool:
    """Sign in ONCE up front (or skip if the persistent profile already has the
    session). Sets logged_in_ref[0]=True only on a real login."""
    log("  Проверяю авторизацию...")
    # already logged in? a logged-in Ancestry nav has an account avatar, no LOG IN
    try:
        has_login = await page.evaluate(
            """() => {
                const t = document.body.innerText || '';
                const link = [...document.querySelectorAll('a,button')].some(
                    e => /log\\s*in|sign\\s*in/i.test((e.textContent||'').trim())
                         && (e.textContent||'').trim().length < 12);
                return link;
            }""")
    except Exception:
        has_login = True
    if not has_login:
        log("  Уже авторизованы — вход не нужен")
        return True
    if not email or not password:
        log("  !! Нет логина/пароля — пробую без входа (только дерево/превью)")
        return False

    log("  Не авторизованы — вхожу ОДИН раз до открытия записей...")
    results_url = page.url
    try:
        await page.goto(SIGNIN_URL, wait_until="domcontentloaded", timeout=30000)
        await asyncio.sleep(2)
        ok = await _login(page, email, password, log)
        if ok:
            logged_in_ref[0] = True
            if "/search" not in page.url and results_url:
                await page.goto(results_url, wait_until="domcontentloaded",
                                timeout=30000)
                await asyncio.sleep(3)
        return ok
    except Exception as e:
        log(f"  !! Sign in: {e}")
        return False


# ── 3. COLLECT RESULTS ────────────────────────────────────────────────────── #

_COLLECT_JS = r"""() => {
    const out = [], seenRow = new Set(), seenHref = new Set();
    const pats = ["/discoveryui-content/view/", "/imageviewer/",
                  "/family-tree/person/", "recordpid", "/cgi-bin/sse.dll"];
    const isRec = h => h && pats.some(p => h.toLowerCase().includes(p));
    const labRe = /\b(Name|Spouse|Birth|Residence|Death|Marriage|Relative)\b/;
    for (const a of document.querySelectorAll('a[href]')) {
        const href = a.href || '';
        if (!isRec(href) || seenHref.has(href)) continue;
        // climb to the RESULT ROW (an ancestor carrying labelled fields)
        let row = a, best = null, hops = 0;
        while (row && hops < 9) {
            const t = row.innerText || '';
            if (labRe.test(t)) best = row;
            if (best && (best.innerText || '').length > 80) break;
            row = row.parentElement; hops++;
        }
        const cont = best || a;
        if (seenRow.has(cont)) { seenHref.add(href); continue; }
        seenRow.add(cont); seenHref.add(href);
        const text = (cont.innerText || '').trim();
        // the PERSON name = the value on the line after a «Name» label
        const lines = text.split('\n').map(s => s.trim()).filter(Boolean);
        let name = '';
        for (let i = 0; i < lines.length - 1; i++)
            if (/^Name$/i.test(lines[i])) { name = lines[i + 1]; break; }
        if (!name) name = (a.innerText || '').trim();
        name = (name.split('\n')[0] || '')      // drop a glued «…\nGoodie Family Tree»
                   .replace(/^(Mr|Mrs|Dr|Miss|Mstr)\.?\s+/i, '')
                   .replace(/\[.*?\]/g, '').trim();
        // collection title (the record's label) = the first labelled line OR
        // the link text
        let coll = (a.innerText || '').trim();
        // pick the best record link inside this row
        const links = [...cont.querySelectorAll('a[href]')]
            .map(x => x.href).filter(isRec);
        const recUrl =
            links.find(h => /discoveryui-content\/view/.test(h)) ||
            links.find(h => /family-tree\/person/.test(h)) ||
            links.find(h => /imageviewer/.test(h)) || href;
        out.push({url: recUrl, name: name, coll: coll, text: text.slice(0, 2000)});
    }
    return out;
}"""


async def _collect(page, first_q, last_q, log) -> list:
    await asyncio.sleep(2)
    try:
        raw = await page.evaluate(_COLLECT_JS)
    except Exception as e:
        log(f"  !! сбор результатов: {type(e).__name__}")
        raw = []
    out = []
    for r in raw:
        if not _is_record(r["url"]) or not r.get("name"):
            continue
        r["score"] = _match_name(first_q, last_q, r["name"])
        out.append(r)
        log(f"    {r['score']:5.1f}%  {r['name'][:45]:45}  | {r.get('coll','')[:40]}")
    log(f"  Кандидатов на странице: {len(out)}")
    return out


async def _collect_all(page, first_q, last_q, base_url, log) -> list:
    """Walk up to MAX_PAGES result pages (Ancestry paginates with &page=N). The
    relevant records are top-ranked, so this rarely needs more than 1-2 pages."""
    all_rows, seen = [], set()
    for pg in range(1, MAX_PAGES + 1):
        if pg > 1:
            base = re.sub(r"([?&])page=\d+", r"\1", base_url).rstrip("?&")
            sep = "&" if "?" in base else "?"
            nu = f"{base}{sep}page={pg}"
            try:
                await page.goto(nu, wait_until="domcontentloaded", timeout=40000)
                await asyncio.sleep(4)
            except Exception:
                break
        rows = await _collect(page, first_q, last_q, log)
        new = [r for r in rows if r["url"] not in seen]
        for r in new:
            seen.add(r["url"])
        all_rows.extend(new)
        if not new:
            break
        if len([r for r in all_rows if r["score"] >= MIN_MATCH]) >= MAX_SCRAPE:
            break
    return all_rows


# ── Result filters (left panel) ────────────────────────────────────────────── #

_FILTER_FIND_JS = r"""(labels) => {
    const norm = s => (s||'').replace(/\s+/g,' ').trim();
    const cand = [...document.querySelectorAll('a,button,[role=button],label,span')];
    for (const e of cand) {
        // own text = text minus the count badge (e.g. «Military 193» → «Military»)
        let t = norm(e.innerText);
        const m = t.match(/^(.*?)(?:\s+[\d,]+\+?)?$/);
        const own = m ? norm(m[1]) : t;
        for (const label of labels) {
            if (own === label || t === label || own.toLowerCase() === label.toLowerCase()) {
                const r = e.getBoundingClientRect();
                if (r.width > 0 && r.height > 0) {
                    e.setAttribute('data-anc-filter','1'); return true;
                }
            }
        }
    }
    return false;
}"""


async def _expand_parent(page, parent, log):
    """Reveal a collapsed filter group (e.g. «Birth, marriage & death» before its
    «Marriage & divorce» sub-item, or a century before its decade). Click its
    expander chevron if present; otherwise click the parent row itself."""
    if not parent:
        return
    js = r"""(parent) => {
        const norm = s => (s||'').replace(/\s+/g,' ').trim();
        for (const e of document.querySelectorAll('a,button,[role=button],span,li,div')) {
            let t = norm(e.innerText);
            const m = t.match(/^(.*?)(?:\s+[\d,]+\+?)?$/);
            const own = m ? norm(m[1]) : t;
            if (own === parent) {
                if (e.getAttribute('aria-expanded') === 'false' ||
                    /(›|▸|>|chevron|expand|caret)/i.test(
                        (e.className||'') + ' ' + e.innerHTML)) {
                    e.setAttribute('data-anc-exp','1'); return true;
                }
            }
        }
        return false;
    }"""
    try:
        if await page.evaluate(js, parent):
            el = page.locator('[data-anc-exp="1"]').first
            await el.click(timeout=3000)
            await el.evaluate("e => e.removeAttribute('data-anc-exp')")
            await asyncio.sleep(1.5)
            log(f"    раскрыл «{parent}»")
    except Exception:
        pass


async def _apply_filter_click(page, label, log, parent=None) -> bool:
    """Click a result filter by its visible text in the left «Filters» panel.
    Tries the label and its ±«s» variant; if not found and a parent is given,
    expands the parent group first. Defensive — logs and continues, never crashes."""
    variants = [label]
    if label.endswith("s"):
        variants.append(label[:-1])          # «1890s» → «1890»
    else:
        variants.append(label + "s")
    try:
        ok = await page.evaluate(_FILTER_FIND_JS, variants)
        if not ok and parent:
            await _expand_parent(page, parent, log)
            ok = await page.evaluate(_FILTER_FIND_JS, variants)
        if not ok:
            log(f"    !! фильтр не найден: {label}")
            return False
        el = page.locator('[data-anc-filter="1"]').first
        await el.scroll_into_view_if_needed(timeout=2000)
        await el.click(timeout=4000)
        await el.evaluate("e => e.removeAttribute('data-anc-filter')")
        await asyncio.sleep(3)
        log(f"    фильтр применён: {label}")
        return True
    except Exception as e:
        log(f"    !! фильтр «{label}»: {type(e).__name__}")
        return False


async def _clear_all_filters(page, log):
    for sel in ['button.link:has-text("Clear all")',
                'button:has-text("Clear all")', 'a:has-text("Clear all")']:
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=4000)
                await asyncio.sleep(2)
                log("    фильтры сброшены (Clear all)")
                return
        except Exception:
            continue


def _filter_passes(filters: dict) -> list:
    """Build the filter passes as the CARTESIAN PRODUCT across the three areas:
    within an area the choices are OR alternatives, across areas they are AND.
    So each pass = one type-choice AND one location-choice AND one date-choice
    (any of them may be absent). Each pass is its own document. No filters → one
    empty pass (plain results)."""
    filters = filters or {}
    t = [x for x in filters.get("types", []) if x] or [None]
    l = [x for x in filters.get("locations", []) if x] or [None]
    d = [x for x in filters.get("dates", []) if x] or [None]
    passes = []
    for ti in t:
        for li in l:
            for di in d:
                p = [x for x in (ti, li, di) if x]
                if p not in passes:
                    passes.append(p)
    return passes or [[]]


async def _diag(page, log):
    """Dump page state when 0 results — so the real DOM can be inspected from the
    log (live site is anti-bot / login-walled for me)."""
    try:
        d = await page.evaluate(r"""() => ({
            url: location.href,
            title: document.title,
            anchors: document.querySelectorAll('a[href]').length,
            recAnchors: [...document.querySelectorAll('a[href]')]
                .filter(a => /discoveryui-content|imageviewer|family-tree\/person/i
                    .test(a.href)).length,
            hasLogin: /log\s*in|sign\s*in/i.test(document.body.innerText||''),
            sample: [...document.querySelectorAll('a[href]')]
                .map(a=>a.href).filter(h=>/ancestry\.com/.test(h)).slice(0,12),
        })""")
        log(f"  ДИАГНОСТИКА: {json.dumps(d, ensure_ascii=False)[:1200]}")
    except Exception as e:
        log(f"  ДИАГНОСТИКА не удалась: {e}")


# ── 4. ADVANCED (More options) — best-effort open, mostly URL-driven ───────── #

async def _open_more_options(page, log):
    for sel in ['[data-testid="moreOptionsBtnTestId"]',
                '[data-testid="btnEditForm"]']:
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=4000)
                await asyncio.sleep(1.5)
                log(f"  Расширенный поиск открыт ({sel})")
                return True
        except Exception:
            continue
    return False


# ── 5. IMAGES ─────────────────────────────────────────────────────────────── #

async def _best_img(page) -> str:
    """Largest document image; scroll first (lazy) and keep a viewer fallback."""
    try:
        await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        await asyncio.sleep(0.6)
        await page.evaluate("window.scrollTo(0, 0)")
        await asyncio.sleep(0.3)
    except Exception:
        pass
    best, area, fallback = "", 0, ""
    for el in await page.query_selector_all("img"):
        try:
            src = (await el.evaluate("e => e.currentSrc || e.src || ''") or "").strip()
            if not src.startswith("http"):
                continue
            if any(b.lower() in src.lower() for b in _IMG_SKIP):
                continue
            w = int(await el.evaluate("e => e.naturalWidth")  or 0)
            h = int(await el.evaluate("e => e.naturalHeight") or 0)
            if w * h > area:
                area, best = w * h, src
            if not fallback and any(k in src.lower() for k in
                                    ("imageviewer", "/media", "dms", "mediasvc",
                                     "interactive")):
                fallback = src
        except Exception:
            continue
    return best or fallback


async def _fetch_bytes(ctx, src: str) -> bytes | None:
    if not src or not src.startswith("http"):
        return None
    pg = await ctx.new_page()
    try:
        r = await pg.goto(src, timeout=20000)
        if r and r.ok:
            body = await r.body()
            if len(body) > 4000:
                return body
    except Exception:
        pass
    finally:
        try: await pg.close()
        except Exception: pass
    return None


async def _download_image(ctx, page, dest_dir: Path, title: str, log) -> str | None:
    """Ancestry record page: «Save» (top-right) → «Save to your computer», caught
    via expect_download (the exact flow the user gave). Falls back to the largest
    image on the page."""
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / (safe_fn(title) + ".jpg")
    try:
        async with page.expect_download(timeout=30000) as dl_info:
            # 1) open the «Save» dropdown (top-right of the record page)
            opened = False
            for sel in ['button.save-btn', 'button[data-tracklink="Save-Open"]',
                        'button.ancBtn:has-text("Save")', 'button:has-text("Save")']:
                try:
                    el = page.locator(sel).first
                    if await el.count() and await el.is_visible():
                        await el.click(timeout=4000)
                        await asyncio.sleep(1.2)
                        opened = True
                        break
                except Exception:
                    continue
            if not opened:
                raise RuntimeError("кнопка Save не найдена")
            # 2) «Save to your computer»
            done = False
            for sel in ['button.link.item:has-text("Save to your computer")',
                        'button:has-text("Save to your computer")',
                        'a:has-text("Save to your computer")',
                        'li:has-text("Save to your computer")']:
                try:
                    el = page.locator(sel).first
                    if await el.count() and await el.is_visible():
                        await el.click(timeout=4000)
                        done = True
                        break
                except Exception:
                    continue
            if not done:
                raise RuntimeError("«Save to your computer» не найдено")
        dl  = await dl_info.value
        suf = Path(dl.suggested_filename or "doc.jpg").suffix or ".jpg"
        dest = dest_dir / (safe_fn(title) + suf)
        await dl.save_as(str(dest))
        log(f"    🖼 документ сохранён: {dest.name} ({dest.stat().st_size//1024}KB)")
        return str(dest)
    except Exception as exc:
        log(f"    (Save→компьютер не вышло: {type(exc).__name__}) — беру превью")

    src = await _best_img(page)
    if src:
        body = await _fetch_bytes(ctx, src)
        if body:
            dest = dest_dir / (safe_fn(title) + ".jpg")
            dest.write_bytes(body)
            log(f"    🖼 превью сохранено: {dest.name} ({len(body)//1024}KB)")
            return str(dest)
    log("    (изображение документа не найдено)")
    return None


# ── 6. SCRAPE ONE RECORD ──────────────────────────────────────────────────── #

_FIELDS_JS = r"""() => {
    const norm = s => (s||'').replace(/ /g,' ')
        .replace(/[ \t]+/g,' ').replace(/\n+/g,'\n').trim();
    const out = [], seen = new Set();
    const BAD = /^(sign in|search|menu|save|print|share|detail|source|view|home|trees|memories|dna|add or update|report a problem|listen|suggested|neighbo|ancestry|cart|help|browse|filter|hint)/i;
    const push = (k, v) => {
        k = norm(k).replace(/[: ]+$/,''); v = norm(v);
        if (!k || !v || k === v || k.length > 45 || v.length > 800) return;
        if (k.includes('\n') || !/[A-Za-z]/.test(k) || BAD.test(k)) return;
        const key = k.toLowerCase();
        if (seen.has(key)) return;
        seen.add(key); out.push([k, v]);
    };
    // Clean «label | value» rows only. A whole <table> (thead+tbody) or a wide
    // multi-column row must NOT be captured as one glued pair → require exactly 2
    // single-line children and skip anything wrapping table structure.
    for (const el of document.querySelectorAll('li, div, dl, tr')) {
        let kids;
        if (el.tagName === 'TR')
            kids = [...el.querySelectorAll(':scope > td, :scope > th')];
        else
            kids = [...el.children].filter(c => norm(c.innerText));
        if (kids.length !== 2) continue;
        if (kids.some(c => c.querySelector && c.querySelector('table,td,th,tr,li')))
            continue;
        const k = norm(kids[0].innerText), v = norm(kids[1].innerText);
        if (k.includes('\n') || v.includes('\n')) continue;   // not a clean field
        if (k.length < 45 && v.length < 200) push(k, v);
    }
    // Household members table — ONLY a real «Household Members» table, short cells
    // (a wide neighbours table would glue, so reject long cells).
    let household = [];
    for (const tb of document.querySelectorAll('table')) {
        if (!/Household Members/i.test(norm(tb.innerText))) continue;
        for (const tr of tb.querySelectorAll('tr')) {
            const c = [...tr.querySelectorAll(':scope > td, :scope > th')]
                .map(x => norm(x.innerText));
            if (c.length < 2 || c.length > 6) continue;
            if (c.some(x => x.length > 40)) continue;          // glued mega-cell
            if (/^name$/i.test(c[0]) && /relationship/i.test(c.join(' '))) continue;
            if (c[0]) household.push(c);
        }
    }
    // subtitle «in the 1930 United States Federal Census»
    let sub = "";
    const h1 = document.querySelector('h1');
    if (h1 && h1.nextElementSibling)
        sub = norm(h1.nextElementSibling.innerText).split('\n')[0];
    return {fields: out, household: household, subtitle: sub};
}"""


async def _scrape_page(ctx, page, url, name_hint, images_root,
                       logged_in_ref, email, password, log) -> dict:
    rec = {"url": url, "title": name_hint, "name": name_hint, "subtitle": "",
           "table_data": {}, "household": [], "images": [], "thumb_bytes": None,
           "collection": "", "score": 0}
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=40000)
        await asyncio.sleep(3)
    except Exception as e:
        log(f"    !! не открылась запись ({type(e).__name__})")
        return rec

    # login redirect fallback (early sign-in should already cover it)
    if ("signin" in page.url or "login" in page.url):
        if not logged_in_ref[0] and email and password:
            log("  → запись требует входа, вхожу...")
            if await _login(page, email, password, log):
                logged_in_ref[0] = True
                await page.goto(url, wait_until="domcontentloaded", timeout=40000)
                await asyncio.sleep(3)
        else:
            log("    !! запись за логином — пропускаю")
            return rec

    # title (h1 — drop a glued tree name on the 2nd line)
    for sel in ["h1", '[class*="title" i]', "h2"]:
        try:
            t = (await page.locator(sel).first.inner_text(timeout=2500) or "").strip()
            t = t.split("\n")[0].strip()
            if 2 < len(t) < 200:
                rec["title"] = t
                break
        except Exception:
            pass

    # fields + household + subtitle
    try:
        data = await page.evaluate(_FIELDS_JS)
    except Exception:
        data = {"fields": [], "household": [], "subtitle": ""}
    td = {}
    for k, v in data.get("fields", []):
        td.setdefault(k, v)
    rec["household"] = data.get("household", []) or []
    rec["subtitle"]  = data.get("subtitle", "") or ""

    # Print-page fallback: if the Detail panel gave too little, the print view
    # (#printPage, ?pf=true) is a clean record dump — parse it instead.
    if len(td) < 3:
        try:
            href = await page.evaluate(
                "() => { const a = document.querySelector('#printPage');"
                " return a ? a.href : ''; }")
            if href:
                pg = await ctx.new_page()
                try:
                    await pg.goto(href, wait_until="domcontentloaded", timeout=40000)
                    await asyncio.sleep(3)
                    d2 = await pg.evaluate(_FIELDS_JS)
                    for k, v in d2.get("fields", []):
                        td.setdefault(k, v)
                    if not rec["household"]:
                        rec["household"] = d2.get("household", []) or []
                    if not rec["subtitle"]:
                        rec["subtitle"] = d2.get("subtitle", "") or ""
                    log(f"    (поля со страницы печати: {len(td)})")
                finally:
                    try: await pg.close()
                    except Exception: pass
        except Exception as e:
            log(f"    !! print-страница: {type(e).__name__}")
    rec["table_data"] = td

    label = rec["title"] or name_hint
    img_dir = images_root / safe_fn(label)

    src = await _best_img(page)
    if src:
        rec["thumb_bytes"] = await _fetch_bytes(ctx, src)

    jp = await _download_image(ctx, page, img_dir, label, log)
    if jp:
        rec["images"] = [jp]
    elif rec["thumb_bytes"]:
        img_dir.mkdir(parents=True, exist_ok=True)
        fp = img_dir / (safe_fn(label) + "_preview.jpg")
        fp.write_bytes(rec["thumb_bytes"])
        rec["images"] = [str(fp)]
    log(f"    полей: {len(td)}, домочадцев: {len(rec['household'])}")
    return rec


# ── Word ──────────────────────────────────────────────────────────────────── #

def _add_link(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


def _docx_add_record(doc, i, rec):
    title = rec.get("title") or rec.get("name", "—")
    doc.add_heading(f"{i}. {title}", level=2)
    if rec.get("subtitle"):
        sp = doc.add_paragraph(); sr = sp.add_run(rec["subtitle"]); sr.italic = True
    if rec.get("url"):
        pp = doc.add_paragraph()
        pp.add_run(f"Источник ({SITE_NAME}): ").bold = True
        _add_link(pp, "Открыть запись", rec["url"])
    if rec.get("score"):
        p = doc.add_paragraph()
        p.add_run("Совпадение: ").bold = True
        p.add_run(f"{rec.get('score','?')}%")

    rows = []
    if rec.get("collection") and rec["collection"] != rec.get("subtitle"):
        rows.append(("Коллекция", rec["collection"]))
    for f, v in rec.get("table_data", {}).items():
        rows.append((str(f), str(v)))
    if rows:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Поле"; hdr[1].text = "Значение"
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for f, v in rows:
            r = tbl.add_row().cells
            r[0].text = f; set_cell_lines(r[1], v)
    doc.add_paragraph("")

    # Household members table
    hh = rec.get("household") or []
    if hh:
        doc.add_paragraph("Домочадцы:").runs[0].bold = True
        ncols = max(len(r) for r in hh)
        htbl = doc.add_table(rows=0, cols=ncols)
        htbl.style = "Table Grid"
        for r in hh:
            cells = htbl.add_row().cells
            for ci in range(ncols):
                set_cell_lines(cells[ci], r[ci] if ci < len(r) else "")
        doc.add_paragraph("")

    imgs = rec.get("images", []); tb = rec.get("thumb_bytes")
    if imgs and Path(imgs[0]).exists():
        doc.add_paragraph("Изображение документа:").runs[0].bold = True
        try:
            doc.add_picture(imgs[0], width=Inches(4))
        except Exception:
            doc.add_paragraph(f"  [{Path(imgs[0]).name}]")
    elif tb:
        doc.add_paragraph("Превью документа:").runs[0].bold = True
        try:
            doc.add_picture(io.BytesIO(tb), width=Inches(4))
        except Exception:
            doc.add_paragraph("  [не удалось вставить]")
    doc.add_paragraph("")


def write_docx(path: Path, records: list, qlines: list, append: bool = False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    existing = append and Path(path).exists()
    if existing:
        doc = Document(str(path))
        doc.add_page_break()
        sep = doc.add_heading(
            f"➕ Добавлено ещё {len(records)} (совпадение ≥{MIN_MATCH}%) — "
            f"{time.strftime('%Y-%m-%d %H:%M')}", level=1)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Параметры:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph("")
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width  = Mm(297); s.page_height = Mm(210)
        s.left_margin = s.right_margin = Mm(15)
        s.top_margin  = s.bottom_margin = Mm(15)
        h = doc.add_heading("Ancestry — Результаты", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Параметры:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph(f"Найдено: {len(records)} (совпадение ≥{MIN_MATCH}%)")
        doc.add_paragraph("")
    for i, rec in enumerate(records, 1):
        _docx_add_record(doc, i, rec)
    doc.save(path)


# ── Excel ─────────────────────────────────────────────────────────────────── #

def write_xlsx(path: Path, records: list, qlines: list, append: bool = False):
    if not _OPENPYXL_OK:
        raise RuntimeError("openpyxl не установлен")
    HF = PatternFill("solid", fgColor="6B8E23")
    HN = Font(bold=True, color="FFFFFF", size=11)
    LINKF = Font(color="0563C1", underline="single")
    TS = Side(style="thin", color="C0D0A0")
    T  = Border(left=TS, right=TS, top=TS, bottom=TS)

    aff = []
    for rec in records:
        for k in rec.get("table_data", {}):
            if k not in aff:
                aff.append(k)
    base_cols = ["#", "База", "Имя", "Запись", "Совп. %", "Домочадцы", "URL"]

    existing = append and Path(path).exists()
    if existing:
        wb = load_workbook(str(path)); ws = wb.active
        header = [ws.cell(row=1, column=c).value
                  for c in range(1, ws.max_column + 1)]
        for name in base_cols + aff:
            if name not in header:
                header.append(name)
                c = ws.cell(row=1, column=len(header), value=name)
                c.font = HN; c.fill = HF; c.border = T
                c.alignment = Alignment(horizontal="center",
                                        vertical="center", wrap_text=True)
        start_row = ws.max_row + 1
        start_num = ws.max_row - 1
    else:
        wb = Workbook(); ws = wb.active; ws.title = "Ancestry"
        header = base_cols + aff
        for ci, cn in enumerate(header, 1):
            c = ws.cell(row=1, column=ci, value=cn)
            c.font = HN; c.fill = HF; c.border = T
            c.alignment = Alignment(horizontal="center",
                                    vertical="center", wrap_text=True)
        start_row = 2
        start_num = 0
    col_idx = {name: i + 1 for i, name in enumerate(header)}

    for n, rec in enumerate(records):
        ri = start_row + n
        td = rec.get("table_data", {})
        hh = "; ".join(" ".join(c for c in r if c) for r in rec.get("household", []))
        row = {"#": start_num + n + 1, "База": SITE_NAME,
               "Имя": rec.get("title", rec.get("name", "")),
               "Запись": rec.get("subtitle", "") or rec.get("collection", ""),
               "Совп. %": rec.get("score", ""),
               "Домочадцы": hh,
               "URL": rec.get("url", "")}
        for f in aff:
            row[f] = td.get(f, "")
        for name, val in row.items():
            ci = col_idx.get(name)
            if not ci:
                continue
            c = ws.cell(row=ri, column=ci)
            if name == "URL" and val:
                c.value = "Открыть"; c.hyperlink = val; c.font = LINKF
            else:
                c.value = val
            c.border = T
            c.alignment = Alignment(wrap_text=True, vertical="top")

    for ci in range(1, len(header) + 1):
        ltr = get_column_letter(ci)
        mw  = max(len(str(header[ci-1] or "")),
                  *(len(str(ws.cell(row=r, column=ci).value or "").split("\n")[0])
                    for r in range(2, ws.max_row + 1)), 8)
        ws.column_dimensions[ltr].width = min(mw + 4, 60)
    wb.save(path)


# ── Main entry point ──────────────────────────────────────────────────────── #

async def run_scraper(
    *,
    first_names:   str       = "",
    last_names:    str       = "",
    place_lived:   str       = "",
    birth_year:    str       = "",
    year_range:    int       = 1,    # ± years (0 = exact); birth_x=<N>-0-0
    advanced:      dict|None = None,
    exact:         dict|None = None,
    filters:       dict|None = None,  # {"types":[…], "locations":[…], "dates":[…]}
    output_format: str       = "both",
    output_folder            = Path("."),
    email:         str|None  = None,
    password:      str|None  = None,
    log                      = print,
    progress                 = None,
    cancel_event             = None,
    ask_file_conflict        = None,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress: progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    adv     = advanced or {}
    exact   = exact or {}
    filters = filters or {}
    want_docx = output_format in ("docx", "both")
    want_xlsx = output_format in ("xlsx", "both")
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    images_root = output_folder / "images"

    qname  = " ".join(p for p in (first_names, last_names) if p)
    qlines = [ln for ln in [
        f"First Names: {first_names}", f"Last Names: {last_names}",
        f"Place: {place_lived}", f"Birth Year: {birth_year}",
    ] if not ln.endswith(": ")]
    if birth_year:
        qlines.append(f"Year range: ±{year_range}" if year_range else "Year: exact")
    for rel in ("father", "mother", "spouse", "sibling", "child"):
        for person in adv.get(rel, []) or []:
            nm = " ".join(p for p in (person.get("first", ""),
                                      person.get("last", "")) if p)
            if nm:
                qlines.append(f"{rel.capitalize()}: {nm}")
    summary   = {"ok": False}
    file_base = safe_fn(f"ancestry_{qname}") if qname else "ancestry_results"
    logged_in_ref = [False]

    _prog(0, "Запускаю браузер...")
    async with async_playwright() as pw:
        # PERSISTENT profile → the Ancestry login is kept between runs.
        for _lk in ("SingletonLock", "SingletonCookie", "SingletonSocket"):
            try: (ANC_PROFILE_DIR / _lk).unlink()
            except Exception: pass
        ANC_PROFILE_DIR.mkdir(parents=True, exist_ok=True)
        ctx = await pw.chromium.launch_persistent_context(
            str(ANC_PROFILE_DIR),
            headless=False,
            no_viewport=True,
            accept_downloads=True,
            args=["--start-maximized",
                  "--disable-blink-features=AutomationControlled"],
        )
        browser = ctx
        page = ctx.pages[0] if ctx.pages else await ctx.new_page()
        for _p in list(ctx.pages)[1:]:
            try: await _p.close()
            except Exception: pass

        try:
            # 1. SIGN IN ONCE up front (persistent profile usually already has it)
            _prog(5, "Открываю Ancestry...")
            await page.goto(HOME_URL, wait_until="domcontentloaded", timeout=40000)
            await asyncio.sleep(3)
            _prog(12, "Sign in...")
            await _sign_in_if_needed(page, email or "", password or "",
                                     logged_in_ref, log)
            if _done(): return summary

            # 2. SEARCH (built URL, count=50)
            _prog(20, "Поиск...")
            search_url = await _search(page, first_names, last_names, place_lived,
                                       birth_year, year_range, exact, adv, log)

            # 3-5. Per FILTER PASS: collect + scrape + SAVE its OWN document and
            # its OWN images folder (named by the filter), so every filter is
            # separable. Passes = cartesian product across areas (see
            # _filter_passes): the all-checkbox AND a specific child run as
            # separate passes. No filters → one plain pass.
            passes   = _filter_passes(filters)
            fparents = filters.get("parents", {})
            if len(passes) > 1 or passes[0]:
                log(f"  Фильтры: {len(passes)} проход(ов) → отдельный документ на каждый")
            total, saved_docx, last_xlsx = 0, 0, None
            for pi, pass_labels in enumerate(passes, 1):
                if _done(): break
                # fresh search URL each pass = filters reset (implicit Clear all)
                try:
                    await page.goto(search_url, wait_until="domcontentloaded",
                                    timeout=40000)
                    await asyncio.sleep(4)
                except Exception:
                    pass
                pass_name = " · ".join(pass_labels)
                if pass_labels:
                    _prog(25, f"Проход {pi}/{len(passes)} — фильтры: {pass_labels}")
                    ok_all = True
                    for lbl in pass_labels:
                        if not await _apply_filter_click(page, lbl, log,
                                                         parent=fparents.get(lbl)):
                            ok_all = False
                    if not ok_all:
                        # a filter that didn't apply would dump UNFILTERED results
                        # under this filter's name — skip the whole pass instead.
                        log(f"  → фильтр(ы) не применились — пропускаю проход "
                            f"«{pass_name}» (без сохранения)")
                        continue
                    await asyncio.sleep(2)

                raw, seen, qual = (await _collect_all(
                    page, first_names, last_names, page.url, log)), set(), []
                for r in raw:
                    if r["score"] >= MIN_MATCH and r["url"] not in seen:
                        seen.add(r["url"]); qual.append(r)
                qual = qual[:MAX_SCRAPE]
                log(f"  Подходящих (≥{MIN_MATCH}%): {len(qual)}")

                suffix    = safe_fn(pass_name) if pass_name else ""
                pass_imgs = (images_root / suffix) if suffix else images_root
                pass_recs = []
                for i, r in enumerate(qual, 1):
                    if _done(): break
                    _prog(25 + int(65 * pi / len(passes)),
                          f"[проход {pi}] [{i}/{len(qual)}] {r['name'][:50]}...")
                    det = await _scrape_page(ctx, page, r["url"], r["name"],
                                             pass_imgs, logged_in_ref,
                                             email or "", password or "", log)
                    det["score"] = r["score"]
                    if r.get("coll") and not det.get("collection"):
                        det["collection"] = r["coll"]
                    pass_recs.append(det)
                    log(f"  ✓  {det.get('title','')[:70]}  ({r['score']}%)")

                if not pass_recs:
                    if pass_labels:
                        log(f"  (по фильтру «{pass_name}» подходящих записей нет)")
                    else:
                        await _diag(page, log)
                    continue

                # save THIS pass to its own files (the filter is in the name)
                base   = file_base + (f"__{suffix}" if suffix else "")
                docx_p = output_folder / f"{base}.docx"
                xlsx_p = output_folder / f"{base}.xlsx"
                plines = qlines + ([f"Фильтр: {pass_name}"] if pass_name else [])
                existing = [p.name for p, want in
                            ((docx_p, want_docx), (xlsx_p, want_xlsx))
                            if want and p.exists()]
                decision = "overwrite"
                if existing and ask_file_conflict:
                    try:
                        decision = (ask_file_conflict(existing) or "overwrite").lower()
                    except Exception:
                        decision = "overwrite"
                    log(f"  → {existing} существуют → {decision}")
                if decision == "skip":
                    log(f"  → пропуск сохранения «{base}»")
                    continue
                append = (decision == "append")
                if want_docx:
                    write_docx(docx_p, pass_recs, plines, append=append)
                    saved_docx += 1
                    log(f"  Word: {docx_p.name}")
                if want_xlsx:
                    write_xlsx(xlsx_p, pass_recs, plines, append=append)
                    last_xlsx = str(xlsx_p)
                    log(f"  Excel: {xlsx_p.name}")
                total += len(pass_recs)

            _prog(100, f"Готово — {total} записей в {saved_docx} документ(ах).")
            summary.update({
                "ok":            True,
                "docx_count":    saved_docx,
                "xlsx_path":     last_xlsx,
                "output_folder": str(output_folder),
                "n_records":     total,
            })
        except Exception as exc:
            summary.update({"error": "exception",
                            "message": f"{type(exc).__name__}: {exc}"})
            log(f"  !! {exc}")
        finally:
            try: await ctx.close()
            except Exception: pass
            try: await browser.close()
            except Exception: pass
    return summary


if __name__ == "__main__":
    asyncio.run(run_scraper(
        first_names="Alexander W", last_names="Sanders", birth_year="1897",
        output_folder=Path.home() / "Downloads" / "Ancestry_results"))
