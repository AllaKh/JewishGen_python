# -*- coding: utf-8 -*-
"""Shared docx helpers used by every scraper's Word writer."""
import re

try:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _DOCX_OK = True
except Exception:
    _DOCX_OK = False


def add_page_numbers(doc, label="Стр."):
    """Add a centered «<label> PAGE из NUMPAGES» footer to every section, so Word numbers
    the pages automatically. Idempotent — if a page field already exists (e.g. when appending
    to an existing file) it does nothing, so the footer is never duplicated. Safe no-op if
    python-docx isn't available."""
    if not _DOCX_OK or doc is None:
        return
    try:
        for section in doc.sections:
            footer = section.footer
            if footer._element.findall('.//' + qn('w:fldChar')):
                continue                                   # already numbered
            footer.is_linked_to_previous = False
            p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            def _field(instr):
                b = OxmlElement('w:fldChar'); b.set(qn('w:fldCharType'), 'begin')
                t = OxmlElement('w:instrText'); t.set(qn('xml:space'), 'preserve'); t.text = instr
                e = OxmlElement('w:fldChar'); e.set(qn('w:fldCharType'), 'end')
                run = p.add_run(); run._r.append(b); run._r.append(t); run._r.append(e)

            if label:
                p.add_run(label + " ")
            _field("PAGE"); p.add_run(" из "); _field("NUMPAGES")
    except Exception:
        pass

# A 4-digit year glued DIRECTLY (no space) to a following capital letter is the
# common "glued-together rows" symptom in family / event lists, e.g.
#   "Моисей … 1858 - 1931Сора-Лея … 1860 - 1936"
# A space after the year ("1950 Detroit") is left alone — that's a real date+place.
_YEAR_GLUE = re.compile(r"(1[5-9]\d\d|20\d\d)(?=[A-ZА-ЯЁ])")


def split_glued(value: str) -> list:
    """Return the value as a list of lines: split on existing newlines AND right
    after a year that is stuck to a capital letter (one person / event per line)."""
    text = _YEAR_GLUE.sub(lambda m: m.group(1) + "\n", str(value or ""))
    lines = [ln.strip() for ln in text.split("\n")]
    lines = [ln for ln in lines if ln]
    return lines or [""]


def set_cell_lines(cell, value):
    """Write `value` into a python-docx table cell, one item per LINE (not all
    glued into one run). First line replaces the cell text, the rest become extra
    paragraphs in the same cell."""
    lines = split_glued(value)
    cell.text = lines[0]
    for extra in lines[1:]:
        cell.add_paragraph(extra)
