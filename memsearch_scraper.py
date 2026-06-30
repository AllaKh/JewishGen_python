#!/usr/bin/env python3
"""
memsearch_scraper.py
====================
Scraper for Memsearch (https://memsearch.org/ru) — a meta-search over Soviet
repression databases (built by GULAG.CZ + Memorial).

Flow (everything via real clicks — the SPA's URL mechanics are not relied on):
  1. Open https://memsearch.org/ru, type the query into the search box, click
     "ИСКАТЬ".
  2. Select the entity tab (Все типы / Люди / Места / Предметы / Документы).
  3. If advanced fields were given for that tab, fill them and click "Показать".
  4. Collect result cards (name = <h4 class="Title">) across all pages.
     Each card links to an EXTERNAL source (ru.openlist.wiki, base.memo.ru,
     jdoc.org.il, …) that holds the FULL record.
  5. For every card, open its external source page and scrape the full info
     (generic "Label: value" extraction + photos). Broken links are tolerated —
     we fall back to the card's own summary and never crash on one bad record.
  6. Save everything to Word (.docx). No Excel — the records are individual.

No login required.
"""

import asyncio
import io
import os
import re
import sys
from pathlib import Path
from docx_util import set_cell_lines, add_page_numbers
from urllib.parse import urljoin, urlparse

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        import os
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    import browser_util
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from PIL import Image
    _PIL_OK = True
except ImportError:
    _PIL_OK = False

# ── Constants ─────────────────────────────────────────────────────────────── #
HOME_URL = "https://memsearch.org/ru"
HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")

# Entity tabs — CANONICAL ENGLISH values (the GUI is always English). The
# scraper translates them to the SITE language (ru/en) chosen in the GUI.
TABS = ["All types", "People", "Places", "Objects", "Documents"]

# Dropdown option sets (English labels ↔ the site's Russian values, from its bundle).
REGION_TYPE  = ["", "Any", "Place of birth", "Place of living or arrest"]
PLACE_AMONG  = ["", "Any", "Burial places", "Prisons and camps",
                "Monuments and memorials"]
OBJECT_AMONG = ["", "Any", "Photos", "Museum items", "Descriptive texts",
                "Monuments and memorials"]

# ── Multi-language UI strings (the memsearch site exists in ru AND en) ──────── #
# concept (canonical EN) → {ru, en}. We match page text by trying BOTH variants:
# only the loaded site's language exists on the page, so this works on either.
_I18N = {
    "All types":  {"ru": "Все типы",   "en": "All types"},
    "People":     {"ru": "Люди",       "en": "People"},
    "Places":     {"ru": "Места",      "en": "Places"},
    "Objects":    {"ru": "Предметы",   "en": "Objects"},
    "Documents":  {"ru": "Документы",  "en": "Documents"},
    "Search":     {"ru": "ИСКАТЬ",     "en": "Search"},      # main search button
    "Show":       {"ru": "Показать",   "en": "Show"},        # apply-advanced button
    "Surname":    {"ru": "Фамилия",    "en": "Surname"},
    "Name":       {"ru": "Имя",        "en": "Name"},
    "Patronymic": {"ru": "Отчество",   "en": "Patronymic"},
    "Year of birth": {"ru": "Год рождения", "en": "Year of birth"},
    "Region":     {"ru": "Регион",     "en": "Region"},
    "Region type": {"ru": "Тип региона", "en": "Region type"},
    "Place name": {"ru": "Название места",   "en": "Place name"},
    "Object name": {"ru": "Название предмета", "en": "Object name"},
    "Document name": {"ru": "Название документа", "en": "Document name"},
    "Search among": {"ru": "Искать среди", "en": "Search among"},
    "Type of region": {"ru": "Тип региона", "en": "Type of region"},
    "Date of repression": {"ru": "Дата репрессии", "en": "Date of repression"},
    "First name": {"ru": "Имя", "en": "First name"},
    "Title of document": {"ru": "Название документа", "en": "Title of document"},
    "Date of issue": {"ru": "Дата создания", "en": "Date of issue"},
    "From (year)": {"ru": "От (год)", "en": "From (year)"},
    "Till (year)": {"ru": "До (год)", "en": "Till (year)"},
    # dropdown option values (English label ↔ the site's Russian value, from its bundle)
    "Any":        {"ru": "Любой",      "en": "Any"},          # region type
    "Any place":  {"ru": "Везде",      "en": "Any"},          # place/object "search among"
    "Place of birth": {"ru": "Место рождения", "en": "Place of birth"},
    "Place of living or arrest":
        {"ru": "Место жительства/репрессий", "en": "Place of living or arrest"},
    "Burial places": {"ru": "Места захоронений", "en": "Burial places"},
    "Prisons and camps": {"ru": "Места заключений", "en": "Prisons and camps"},
    "Photos": {"ru": "Фотографий", "en": "Photos"},
    "Museum items": {"ru": "Музейных предметов", "en": "Museum items"},
    "Descriptive texts": {"ru": "Сопроводительных текстов", "en": "Descriptive texts"},
    "Monuments and memorials": {"ru": "Памятников", "en": "Monuments and memorials"},
}

def _t(concept, lang):
    """Site-language string for a canonical-English concept."""
    d = _I18N.get(concept, {})
    return d.get(lang) or d.get("en") or concept

def _t_all(concept):
    """Both language variants of a concept (for robust text matching)."""
    d = _I18N.get(concept)
    if not d:
        return [concept]
    out = []
    for v in (d.get("ru"), d.get("en")):
        if v and v not in out:
            out.append(v)
    return out

# Images that are never content (site chrome / logos / icons).
_IMG_SKIP = ("logo", "icon", "sprite", "favicon", "avatar", "placeholder",
             "/assets/", "base64", "gulag.cz", "memsearch")


# ── Helpers ───────────────────────────────────────────────────────────────── #
def safe_fn(s: str, n: int = 80) -> str:
    s = re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip())
    return re.sub(r"\s+", "_", s)[:n].strip("_") or "record"


def _host(u: str) -> str:
    """Host of a URL — empty for blank / chrome-error / about: pages. Uses
    urlparse().netloc (Python has NO .host attribute — that mistake crashed a
    whole run on a broken link)."""
    u = u or ""
    if not u or u.startswith(("chrome-error", "about:", "data:")):
        return ""
    try:
        return urlparse(u).netloc or ""
    except Exception:
        return ""


# memsearch chrome (header menu / type tabs / filter labels / footer) that must be
# stripped from a card's brief when a result can't be opened — the user pasted these
# exact strings («приведи его в нормальный вид: отформатируй и убери …»).
_MS_NAV = [
    "О проекте", "Обратная связь", "Источники", "Помощь", "Memsearch",
    "ВСЕ ТИПЫ", "ЛЮДИ", "МЕСТА", "ПРЕДМЕТЫ", "ДОКУМЕНТЫ", "Тип региона",
    "Дата репрессии", "ПОКАЗАТЬ",
    "About the project", "About", "Feedback", "Sources", "Help",
    "ALL TYPES", "PEOPLE", "PLACES", "OBJECTS", "DOCUMENTS",
    "Type of region", "Date of repression", "SHOW",
]

def _clean_brief(text: str) -> str:
    """Strip the memsearch page chrome from a card brief and tidy it, so a result that
    couldn't be opened still reads cleanly (instead of «Memsearch О проекте … ВСЕ ТИПЫ
    ЛЮДИ … Техническая реализация:»)."""
    t = re.sub(r"\s+", " ", text or "")
    # drop the footer and everything after it
    t = re.split(r"Техническ\w*\s+реализаци|Technical implementation", t)[0]
    # drop the "(Люди/Места/…) найдено N" / "… found N" hit counter
    t = re.sub(r"(Люд\w+|Мест|Предмет\w+|Документ\w+|Запис\w+|People|Places|Objects|"
               r"Documents|Records?)?\s*(найдено|found)\s*\d+", " ", t, flags=re.I)
    t = re.sub(r"\b(EN\s+RU|RU\s+EN)\b", " ", t)        # language toggle
    for j in _MS_NAV:
        t = re.sub(r"(?<![\wА-Яа-яЁё])" + re.escape(j) + r"(?![\wА-Яа-яЁё])", " ",
                   t, flags=re.I)
    return re.sub(r"\s+", " ", t).strip(" .,;:—-")


_BRIEF_LABELS = [
    "Дата осуждения", "Дата реабилитации", "Дата ареста", "Дата рождения",
    "Дата смерти", "Дата репрессии", "Место репрессии/жительства", "Место рождения",
    "Место смерти", "Место репрессии", "Место жительства", "Место ареста",
    "Национальность", "Образование", "Профессия", "Приговор", "Статья",
    "Возраст", "Пол",
]

def _brief_to_fields(text: str) -> dict:
    """Split a cleaned card brief («Дата ареста: 23.8.1937 Место рождения: г. Могилев …»)
    into Label→value pairs so it renders as a tidy table, not one run-on line. Best-effort,
    Russian labels (the user's site language)."""
    if not text:
        return {}
    pat = "|".join(re.escape(l) for l in sorted(_BRIEF_LABELS, key=len, reverse=True))
    ms = list(re.finditer(r"(" + pat + r")\s*:?\s*", text))
    out = {}
    for i, mt in enumerate(ms):
        end = ms[i + 1].start() if i + 1 < len(ms) else len(text)
        val = text[mt.end():end].strip(" .,;:")
        val = re.sub(r"\s*[a-z0-9.-]+\.(?:ru|org|com|by|cz|il|net)\s*$", "", val).strip(" .,;:")
        if val:
            out[mt.group(1)] = val
    return out


def _to_png(data: bytes):
    """Convert any image (incl. WEBP) to PNG for python-docx; pass JPEG/PNG/GIF/BMP."""
    if not data:
        return None
    if (data[:3] == b"\xff\xd8\xff" or data[:8] == b"\x89PNG\r\n\x1a\n"
            or data[:4] == b"GIF8" or data[:2] == b"BM"):
        return data
    if _PIL_OK:
        try:
            im = Image.open(io.BytesIO(data)).convert("RGB")
            out = io.BytesIO(); im.save(out, format="PNG")
            return out.getvalue()
        except Exception:
            return None
    return None


def _add_hyperlink(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


# ── Search + navigation ───────────────────────────────────────────────────── #
async def _apply_sources(page, sources, log) -> None:
    """Limit the search to the chosen databases. memsearch builds its results URL as
    /search?query=…&page=1&entityTypes=…&searchSource=<comma-joined source keys> (verified
    from the site's JS bundle). So after the search lands on /search we just rewrite the
    searchSource param to the selected keys and reload — server-side filtering, no clicks."""
    if not sources:
        return                                   # all databases → no filter
    try:
        from urllib.parse import urlparse, parse_qs, urlencode, urlunparse
        u = page.url
        if "/search" not in u:
            log("  !! Фильтр источников: страница результатов не /search — пропуск")
            return
        pr = urlparse(u)
        qs = parse_qs(pr.query)
        qs["searchSource"] = [",".join(sources)]
        new = urlunparse(pr._replace(query=urlencode(qs, doseq=True)))
        if new != u:
            await page.goto(new, wait_until="domcontentloaded", timeout=40000)
            # wait for the filtered result cards to render + hydrate, otherwise the card
            # click won't open the source tab (the «filtered result not opened» bug).
            try:
                await page.wait_for_selector("h4.Title", timeout=15000)
            except Exception:
                pass
            await asyncio.sleep(2.5)
        log(f"  ✓ Ограничил поиск базами: {len(sources)}")
    except Exception as e:
        log(f"  !! Фильтр источников не применён: {e}")


async def _wait_if_captcha(page, log, timeout=300) -> None:
    """If an «I'm not a robot» / reCAPTCHA challenge appears, PAUSE and let the USER solve it
    in the visible browser window (we never click or solve a CAPTCHA ourselves — that's
    bot-detection bypass). Polls until the challenge is gone, so the run waits for the user
    instead of rushing past it («мельтешишь, я не успеваю»)."""
    async def _present():
        try:
            return await page.evaluate(r"""() => {
                const vis = el => el && el.offsetParent !== null
                                  && el.getBoundingClientRect().width > 20
                                  && el.getBoundingClientRect().height > 20;
                for (const f of document.querySelectorAll('iframe')) {
                    const s = (f.src || '') + ' ' + (f.title || '');
                    if (/recaptcha|hcaptcha|captcha|challenge|turnstile/i.test(s) && vis(f))
                        return true;
                }
                if (document.querySelector('#cf-challenge-running, .g-recaptcha:not(:empty)'))
                    return true;
                const t = (document.body && document.body.innerText) || '';
                return /я\s*не\s*робот|i'?m not a robot|подтвердите.{0,4}что вы не робот|verify you are human|are you a robot/i.test(t);
            }""")
        except Exception:
            return False
    if not await _present():
        return
    log("  ⛔ КАПЧА «Я не робот» — РЕШИ ЕЁ В ОКНЕ БРАУЗЕРА. Скрипт ждёт, не торопит.")
    try:
        await page.bring_to_front()
    except Exception:
        pass
    waited = 0
    while waited < timeout:
        await asyncio.sleep(2); waited += 2
        if not await _present():
            log("  ✓ Капча решена — продолжаю.")
            await asyncio.sleep(1)
            return
        if waited % 20 == 0:
            log(f"  … жду капчу ({waited}/{timeout}s) — нажми «Я не робот» в окне.")
    log("  !! капча не решена за отведённое время — продолжаю как есть.")


async def _do_search(page, query: str, lang: str, log) -> bool:
    """Open the LANGUAGE-specific home page, type the query, click the search
    button. True if results show."""
    url = f"https://memsearch.org/{lang}"
    log(f"  → Открываю {url}")
    await page.goto(url, wait_until="domcontentloaded", timeout=40000)
    await asyncio.sleep(2)

    # Type into the search box (placeholder is «Поиск» / «Search»).
    typed = False
    for sel in ('input.SearchInput', 'input[placeholder="Поиск"]',
                'input[placeholder="Search"]', 'input[type="text"]'):
        try:
            inp = page.locator(sel).first
            if await inp.count():
                await inp.click(timeout=4000)
                await page.keyboard.press("Control+a")
                await page.keyboard.press("Delete")
                await page.keyboard.type(query, delay=30)
                typed = True
                log(f"  ✓ Ввёл запрос «{query}» ({sel})")
                break
        except Exception:
            continue
    if not typed:
        log("  !! Поле поиска не найдено")
        return False

    # Click the search button (ИСКАТЬ / SEARCH), else press Enter.
    clicked = False
    words = "|".join(re.escape(w) for w in _t_all("Search"))
    for sel in ('button:has(span.RedButtonText)', 'span.RedButtonText',
                'button'):
        try:
            el = page.locator(sel, has_text=re.compile(words, re.I)).first
            if await el.count():
                await el.click(timeout=4000)
                clicked = True
                log("  ✓ Нажал «Искать»")
                break
        except Exception:
            continue
    if not clicked:
        await page.keyboard.press("Enter")
        log("  → Поиск: Enter")

    # Wait for results (cards or the «найдено / found» count header).
    for _ in range(30):
        await asyncio.sleep(1)
        try:
            ok = await page.evaluate(
                "() => document.querySelectorAll('h4.Title').length > 0 ||"
                " /найдено|found/i.test(document.body.innerText)")
            if ok:
                return True
        except Exception:
            pass
    return True


async def _select_tab(page, tab: str, log):
    """Click the entity-type tab button (button.EntityTypeMenuItem → h6 text).
    `tab` is the canonical-English concept; match BOTH language variants. Skip
    pagination buttons, whose h6 is a NUMBER (same class)."""
    if not tab or tab == "All types":
        return                      # default — nothing to click
    names = _t_all(tab)
    ok = await page.evaluate(r"""(names) => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim().toLowerCase();
        const want = names.map(norm);
        for (const b of document.querySelectorAll('button.EntityTypeMenuItem')) {
            const h = b.querySelector('h6');
            const t = norm(h ? h.textContent : b.textContent);
            if (/^\d+$/.test(t)) continue;          // a page number, not a tab
            if (want.includes(t)) { b.setAttribute('data-pw-tab', '1'); return true; }
        }
        return false;
    }""", names)
    if ok:
        try:
            await page.locator('[data-pw-tab="1"]').first.click(timeout=4000)
            await asyncio.sleep(1.5)
            log(f"  ✓ Таб: {tab}")
        except Exception as e:
            log(f"  !! таб «{tab}»: {e}")
    else:
        log(f"  !! таб «{tab}» не найден (оставляю как есть)")


async def _open_dropdown(page, trigger_concept: str, option_concept: str, log):
    """Custom dropdown: click the div trigger, then the option button. Both are
    given as canonical-English concepts and matched in BOTH site languages."""
    if not option_concept:
        return
    triggers = _t_all(trigger_concept)
    options  = _t_all(option_concept)
    try:
        opened = await page.evaluate(r"""(triggers) => {
            const norm = s => (s || '').replace(/\s+/g, ' ').trim().toLowerCase();
            const want = triggers.map(norm);
            for (const d of document.querySelectorAll('div')) {
                if (d.children.length > 3) continue;
                const t = norm(d.textContent);
                if (want.some(w => t.startsWith(w))) {
                    d.setAttribute('data-pw-dd', '1'); return true;
                }
            }
            return false;
        }""", triggers)
        if not opened:
            return
        await page.locator('[data-pw-dd="1"]').first.click(timeout=3000)
        await asyncio.sleep(0.6)
        picked = await page.evaluate(r"""(options) => {
            const norm = s => (s || '').replace(/\s+/g, ' ').trim().toLowerCase();
            const want = options.map(norm);
            for (const b of document.querySelectorAll('button, [role=option], li')) {
                if (want.includes(norm(b.textContent))) {
                    b.setAttribute('data-pw-opt', '1'); return true;
                }
            }
            return false;
        }""", options)
        if picked:
            await page.locator('[data-pw-opt="1"]').first.click(timeout=3000)
            await asyncio.sleep(0.4)
            log(f"  ✓ {trigger_concept}: {option_concept}")
    except Exception as e:
        log(f"  !! дропдаун «{trigger_concept}»: {e}")


async def _fill_field(page, concept: str, value: str, log):
    """Fill an advanced input. The site labels the input by its placeholder/id in
    the site language — try BOTH languages (id and placeholder)."""
    if not value:
        return
    sels = []
    for lab in _t_all(concept):
        sels += [f'input[id="{lab}"]', f'input[placeholder="{lab}"]']
    for sel in sels:
        try:
            el = page.locator(sel).first
            if await el.count():
                await el.click(timeout=3000)
                await page.keyboard.press("Control+a")
                await page.keyboard.press("Delete")
                await page.keyboard.type(str(value), delay=25)
                log(f"  ✓ {concept} = {value!r}")
                return
        except Exception:
            continue
    log(f"  !! поле «{concept}» не найдено")


async def _fill_advanced(page, tab: str, params: dict, log):
    """Fill the advanced fields for the selected tab, then click the apply button.
    `tab` and all option values are canonical-English concepts."""
    used = False
    if tab == "People":
        for concept, key in (("Surname", "last_name"), ("Name", "first_name"),
                             ("Patronymic", "patronymic"),
                             ("Year of birth", "birth_year"), ("Region", "region")):
            if params.get(key):
                await _fill_field(page, concept, params[key], log); used = True
        if params.get("region_type"):
            await _open_dropdown(page, "Region type", params["region_type"], log)
            used = True
    elif tab == "Places":
        if params.get("place_name"):
            await _fill_field(page, "Place name", params["place_name"], log)
            used = True
        if params.get("place_among"):
            await _open_dropdown(page, "Search among", params["place_among"], log)
            used = True
    elif tab == "Objects":
        if params.get("object_name"):
            await _fill_field(page, "Object name", params["object_name"], log)
            used = True
        if params.get("object_among"):
            await _open_dropdown(page, "Search among", params["object_among"], log)
            used = True
    elif tab == "Documents":
        if params.get("doc_name"):
            await _fill_field(page, "Document name", params["doc_name"], log)
            used = True

    if not used:
        return
    # Click the apply button (Показать / Show).
    words = "|".join(re.escape(w) for w in _t_all("Show"))
    for sel in ('button.RedButton:has(span)', 'button:has(span.RedButtonText)',
                'button'):
        try:
            el = page.locator(sel, has_text=re.compile(words, re.I)).first
            if await el.count():
                await el.click(timeout=4000)
                await asyncio.sleep(1.8)
                log("  ✓ Нажал «Показать»")
                return
        except Exception:
            continue


# ── Result cards ──────────────────────────────────────────────────────────── #
async def _tag_cards(page) -> list:
    """TAG each result card box with data-pw-card=<idx> so we can click it, and
    return [{index, name, summary}]. The card has NO usable <a href> (it opens
    the source via a JS click → a new tab), so we don't extract a URL here — the
    URL comes from the captured tab in _open_card."""
    return await page.evaluate(r"""() => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
        document.querySelectorAll('[data-pw-card]')
                .forEach(e => e.removeAttribute('data-pw-card'));
        const out = [];
        let idx = 0;
        // The real result card is the <button class="…Card…"> that holds the h4.Title —
        // clicking IT opens the source. (Do NOT walk up to a wrapper: that wrapper also
        // contains the advanced-search form, so clicking it hit «Показать», not the card —
        // which is exactly why a single result never opened while two did.)
        let cards = [...document.querySelectorAll('button.Card')]
                        .filter(c => c.querySelector('h4.Title'));
        if (!cards.length)
            cards = [...document.querySelectorAll('.Card, [class*="Card"]')]
                        .filter(c => c.querySelector('h4.Title') &&
                                     !c.querySelector('input, .SearchInput'));
        for (const card of cards) {
            const h = card.querySelector('h4.Title');
            card.setAttribute('data-pw-card', String(idx));
            const name = norm(h.textContent);
            const link = card.querySelector('.Link');           // source domain (red text)
            const source = link ? norm(link.textContent) : '';
            const desc = card.querySelector('.Description, p');  // the person/record text
            let summary = norm(desc ? (desc.innerText || desc.textContent) : (card.innerText || ''));
            if (name) summary = summary.split(name).join(' ');
            if (source) summary = summary.split(source).join(' ');
            summary = norm(summary);
            out.push({index: idx, name, source, summary: summary.slice(0, 700)});
            idx++;
        }
        return out;
    }""")


async def _goto_next_page(page, log) -> bool:
    """Pagination: numeric buttons share class EntityTypeMenuItem (square,
    h6=number). Click the one after the active page."""
    nxt = await page.evaluate(r"""() => {
        const norm = s => (s || '').replace(/\s+/g, ' ').trim();
        const nums = [];
        for (const b of document.querySelectorAll('button.EntityTypeMenuItem')) {
            const t = norm(b.textContent);
            if (/^\d+$/.test(t)) nums.push([parseInt(t, 10), b]);
        }
        if (!nums.length) return 0;
        // active page = the highlighted (invert) numeric button, else the max
        // currently realised; we want active+1 if it exists.
        let active = 0;
        for (const [n, b] of nums) if (b.className.includes('invert')) active = n;
        const want = active + 1;
        for (const [n, b] of nums) {
            if (n === want) { b.setAttribute('data-pw-next', '1'); return n; }
        }
        return 0;
    }""")
    if not nxt:
        return False
    try:
        await page.locator('[data-pw-next="1"]').first.click(timeout=4000)
        await asyncio.sleep(1.6)
        log(f"  → Страница {nxt}")
        return True
    except Exception as e:
        log(f"  → пагинация остановлена: {e}")
        return False


async def _diag_cards(page, log):
    """Dump the first card's structure when nothing is collected, so the real
    DOM can be confirmed (the live site is not directly inspectable)."""
    try:
        info = await page.evaluate(r"""() => {
            const h = document.querySelector('h4.Title');
            let card = h;
            if (h) for (let i=0;i<6&&card.parentElement;i++){card=card.parentElement;
                if (card.querySelector('a[href^="http"]')) break;}
            return {
                url: location.href,
                titles: document.querySelectorAll('h4.Title').length,
                anyCard: document.querySelectorAll('[class*="card" i]').length,
                header: (document.querySelector('h1,h2')||{}).textContent || '',
                cardHTML: card ? card.outerHTML.slice(0, 1200) : '(no h4.Title)'
            };
        }""")
        log("  🔎 ДИАГНОСТИКА memsearch:")
        log(f"     url: {info.get('url','')[:140]}")
        log(f"     h4.Title={info.get('titles')} anyCard={info.get('anyCard')} "
            f"header={info.get('header','')[:60]}")
        log(f"     cardHTML: {info.get('cardHTML','')[:700]}")
    except Exception as e:
        log(f"  🔎 диагностика не вышла: {e}")


# ── External source page (the FULL record) ────────────────────────────────── #
async def _extract_page(dp, images_dir: Path, log) -> dict:
    """Scrape the FULL record from an ALREADY-OPEN source page (the new tab the
    card opened): heading, every "Label: value" pair, and the content photo. The
    caller owns the page's lifecycle (it is NOT closed here)."""
    rec = {"fields": {}, "heading": "", "photo": None, "thumb_bytes": None}
    try:
        await asyncio.sleep(1.2)
        data = await dp.evaluate(r"""() => {
            const norm = s => (s || '').replace(/\s+/g, ' ').trim();
            const out = {heading: '', pairs: []};
            const seen = new Set();
            const push = (k, v) => {
                k = norm(k).replace(/:$/, '').trim();
                v = norm(v);
                if (!k || !v || k.length > 80 || v.length > 600) return;
                const key = k.toLowerCase();
                if (seen.has(key)) return; seen.add(key);
                out.pairs.push([k, v]);
            };
            const h = document.querySelector('h1, h2, h3');
            if (h) out.heading = norm(h.textContent);
            // dl / table pairs
            document.querySelectorAll('dl').forEach(dl => {
                const dt = dl.querySelectorAll('dt'), dd = dl.querySelectorAll('dd');
                for (let i = 0; i < Math.min(dt.length, dd.length); i++)
                    push(dt[i].textContent, dd[i].textContent);
            });
            document.querySelectorAll('table tr').forEach(tr => {
                const c = tr.querySelectorAll('td, th');
                if (c.length === 2) push(c[0].textContent, c[1].textContent);
            });
            // "<b>Label:</b> value"
            document.querySelectorAll('b, strong').forEach(b => {
                const lab = norm(b.textContent);
                if (!lab.includes(':')) return;
                let val = '', n = b.nextSibling;
                while (n) {
                    if (n.nodeType === 3) val += n.textContent;
                    else if (n.nodeType === 1) {
                        if (/^(B|STRONG|BR|DIV|P)$/.test(n.tagName)) break;
                        val += n.textContent;
                    }
                    n = n.nextSibling;
                }
                push(lab, val);
            });
            // plain "Label: value" lines
            document.querySelectorAll('p, li, div, td').forEach(el => {
                if (el.children.length > 2) return;
                const t = norm(el.textContent);
                const m = t.match(/^([^:]{2,60}):\s*(.+)$/);
                if (m) push(m[1], m[2]);
            });
            return out;
        }""")

        rec["heading"] = data.get("heading", "")
        NOISE = ("вход", "регистрац", "поиск", "меню", "категори", "навигац",
                 "помоч проект", "контакт", "©", "cookie", "поделит")
        for k, v in data.get("pairs", []):
            kl = k.lower()
            if any(nz in kl for nz in NOISE):
                continue
            rec["fields"][k] = v
        log(f"      Полей со страницы источника: {len(rec['fields'])}")

        # ── Photo: largest content image on the page ─────────────────── #
        best = await dp.evaluate(r"""() => {
            let best = '', area = 0;
            for (const im of document.querySelectorAll('img[src]')) {
                const s = (im.src || '').toLowerCase();
                if (!s || /\.svg|data:image\/svg/.test(s)) continue;
                const skip = %s;
                if (skip.some(x => s.includes(x))) continue;
                const a = (im.naturalWidth||im.width||0) * (im.naturalHeight||im.height||0);
                if (a > area) { area = a; best = im.src; }
            }
            return best;
        }""" % repr(list(_IMG_SKIP)))
        if best:
            try:
                r = await dp.request.get(best, timeout=20000)
                if r.ok:
                    body = await r.body()
                    if len(body) > 1500:
                        rec["thumb_bytes"] = body
                        images_dir.mkdir(parents=True, exist_ok=True)
                        ext = ".jpg"
                        fn = safe_fn(rec.get("heading") or "photo") + ext
                        (images_dir / fn).write_bytes(body)
                        rec["photo"] = str(images_dir / fn)
                        log(f"      📷 фото {len(body)//1024}KB")
            except Exception:
                pass
    except Exception as e:
        log(f"      !! не смог прочитать страницу источника ({type(e).__name__})")
    return rec


async def _open_card(ctx, page, idx: int, images_dir: Path, log, dump_dir=None):
    """Open result card #idx's external source and scrape it. memsearch opens the source by a
    REACT click handler (no href, no onclick attribute) — so clicking a wrapper <div> won't
    fire it; we must click the element whose computed cursor is 'pointer' with a REAL MOUSE
    click at its centre. Handle a new tab OR same-tab navigation, and never abort on a broken
    link (then fall back to the cleaned brief). The first card's HTML is dumped to
    results/memsearch_card_sample.html so the open mechanism can be verified offline."""
    ext = {"fields": {}, "heading": "", "thumb_bytes": None}
    src_url = ""
    card = page.locator(f'[data-pw-card="{idx}"]').first
    before_url = page.url
    before = set(ctx.pages)
    np = None
    same_tab = False

    if dump_dir is not None and os.environ.get("JGS_DEBUG"):   # diagnostic dump, opt-in
        try:
            html = await page.evaluate(
                "(i) => { const c = document.querySelector('[data-pw-card=\"'+i+'\"]');"
                "         return c ? c.outerHTML : ''; }", str(idx))
            (Path(dump_dir) / "memsearch_card_sample.html").write_text(html or "", encoding="utf-8")
        except Exception:
            pass

    async def _try_click():
        # REAL mouse click on the card's clickable (cursor:pointer) element — clicking a
        # wrapper div with page.click won't trigger a child's React onClick.
        try:
            box = await page.evaluate(
                """(i) => {
                    const c = document.querySelector('[data-pw-card="'+i+'"]');
                    if (!c) return null;
                    c.scrollIntoView({block: 'center'});
                    let t = null;
                    for (const el of [c, ...c.querySelectorAll('*')]) {
                        const cs = getComputedStyle(el);
                        if (cs.cursor === 'pointer' && el.offsetParent !== null) { t = el; break; }
                    }
                    const r = (t || c).getBoundingClientRect();
                    if (r.width < 2 || r.height < 2) return null;
                    return {x: r.x + r.width / 2, y: r.y + Math.min(r.height / 2, 36)};
                }""", str(idx))
            if box:
                await page.mouse.click(box["x"], box["y"])
                return True
        except Exception:
            pass
        # fallbacks: locator clicks (card, then its title)
        for loc in (card, card.locator('h4.Title').first):
            try:
                if await loc.count():
                    await loc.click(timeout=3000)
                    return True
            except Exception:
                continue
        return False

    try:
        try:
            async with ctx.expect_page(timeout=9000) as info:
                await _try_click()
            np = await info.value
        except Exception:
            np = None
        if np is None:                              # a tab that expect_page didn't catch
            extra = [p for p in ctx.pages if p not in before]
            if extra:
                np = extra[-1]
        if np is None:                              # maybe it navigated the SAME tab
            try:
                await page.wait_for_load_state("domcontentloaded", timeout=3000)
            except Exception:
                pass
            if page.url != before_url and "/search" not in page.url:
                np, same_tab = page, True
        if np is None:
            # last resort: a source URL hidden in the card's attributes (data-*, etc.).
            # EXCLUDE the known traps: SVG xmlns (w3.org), memsearch itself, asset CDNs.
            try:
                cand = await page.evaluate(
                    r"""(i) => {
                        const c = document.querySelector('[data-pw-card="'+i+'"]');
                        if (!c) return '';
                        const bad = /w3\.org|memsearch\.org|\.svg|googleapis|gstatic|fonts|schema\.org/i;
                        for (const el of [c, ...c.querySelectorAll('*')]) {
                            for (const a of (el.attributes || [])) {
                                const m = (a.value || '').match(/https?:\/\/[^\s"')]+/g);
                                if (m) for (const u of m) if (!bad.test(u)) return u;
                            }
                        }
                        return '';
                    }""", str(idx))
            except Exception:
                cand = ""
            if cand:
                log(f"      ↪ ссылка из карточки: {cand[:70]}")
                try:
                    np = await ctx.new_page()
                    await np.goto(cand, wait_until="domcontentloaded", timeout=30000)
                except Exception:
                    np = None
        if np is None:
            log("      !! карточка не открыла источник — беру краткое из карточки "
                "(дамп: results/memsearch_card_sample.html)")
            return "", ext

        try:
            await np.wait_for_load_state("domcontentloaded", timeout=20000)
        except Exception:
            pass
        await asyncio.sleep(0.5)
        src_url = np.url or ""
        if src_url.startswith("chrome-error"):      # often a transient hiccup → one reload
            await asyncio.sleep(1.5)
            try:
                await np.reload(wait_until="domcontentloaded", timeout=15000)
                await asyncio.sleep(1)
            except Exception:
                pass
            src_url = np.url or ""
        log(f"      ↪ источник: {src_url[:75]}")
        if src_url.startswith(("chrome-error", "about:")) or not _host(src_url):
            log("      !! ссылка-источник битая — беру краткое из карточки")
        else:
            ext = await _extract_page(np, images_dir, log)
    except Exception as e:
        log(f"      !! карточка не открылась ({type(e).__name__}) — беру краткое")
    finally:
        if np is not None and not same_tab:
            try: await np.close()
            except Exception: pass
        elif same_tab:                              # restore the results page for next cards
            try: await page.go_back(wait_until="domcontentloaded", timeout=15000)
            except Exception: pass
    return src_url, ext


# ── Word output ───────────────────────────────────────────────────────────── #
def _docx_add_record(doc, i, rec):
    name = rec.get("name") or rec.get("heading") or f"Запись {i}"
    # Modest record heading (a bold 12pt line — NOT a giant Heading-2).
    hp = doc.add_paragraph()
    hr = hp.add_run(f"{i}. {name}")
    hr.bold = True; hr.font.size = Pt(12)
    fields = rec.get("fields") or {}
    if fields:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hd = tbl.rows[0].cells
        hd[0].text = "Поле"; hd[1].text = "Значение"
        for cell in hd:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for k, v in fields.items():
            r = tbl.add_row().cells
            r[0].text = str(k); set_cell_lines(r[1], v)
    elif rec.get("summary"):
        doc.add_paragraph(rec["summary"])
    tb = rec.get("thumb_bytes")
    if tb:
        png = _to_png(tb)
        if png:
            try:
                doc.add_picture(io.BytesIO(png), width=Inches(2.6))
                ph = rec.get("photo")
                if ph and Path(ph).exists():
                    p = doc.add_paragraph(); p.add_run("Файл: ").bold = True
                    p.add_run(str(Path(ph).resolve()))   # exact path where it was saved
            except Exception:
                pass
    if rec.get("url"):
        p = doc.add_paragraph()
        p.add_run("Источник: ").bold = True
        _add_hyperlink(p, rec.get("source") or "Открыть запись", rec["url"])
    doc.add_paragraph("")


def write_docx(path, records, qlines, append=False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    existing = append and Path(path).exists()
    if existing:
        doc = Document(str(path))
        doc.add_page_break()
        ap = doc.add_paragraph(); ar = ap.add_run(f"➕ Добавлено {len(records)} записей")
        ar.bold = True; ar.font.size = Pt(13)
        ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width = Mm(210); s.page_height = Mm(297)
        s.left_margin = s.right_margin = Mm(18)
        s.top_margin = s.bottom_margin = Mm(18)
        # Compact title — a bold 14pt line, not the huge Heading-0/Title style.
        ht = doc.add_paragraph(); htr = ht.add_run("Memsearch — результаты поиска")
        htr.bold = True; htr.font.size = Pt(14)
        ht.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for ln in qlines:
            doc.add_paragraph(ln)
        doc.add_paragraph(f"Найдено записей: {len(records)}")
        doc.add_paragraph("")
    for i, rec in enumerate(records, 1):
        _docx_add_record(doc, i, rec)
    add_page_numbers(doc)
    doc.save(str(path))


# ── Main entry point ──────────────────────────────────────────────────────── #
async def run_scraper(*,
    query="",
    lang="ru",                       # site language: "ru" or "en"
    entity_tab="All types",          # canonical English tab
    sources=None,                    # source keys to limit to (None/[] = all databases)
    # People
    last_name="", first_name="", patronymic="", birth_year="",
    region="", region_type="", repress_from="", repress_to="",
    # Places / Objects
    place_name="", place_among="",
    object_name="", object_among="",
    # Documents
    doc_name="", doc_from="", doc_to="",
    output_folder=Path("."),
    log=print,
    progress=None,
    cancel_event=None,
    ask_file_conflict=None,
    max_pages=20,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress:
            progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    params = dict(last_name=last_name, first_name=first_name,
                  patronymic=patronymic, birth_year=birth_year,
                  region=region, region_type=region_type,
                  repress_from=repress_from, repress_to=repress_to,
                  place_name=place_name, place_among=place_among,
                  object_name=object_name, object_among=object_among,
                  doc_name=doc_name, doc_from=doc_from, doc_to=doc_to)
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    images_dir = output_folder / "images" / "Мемориал" / (safe_fn(query) or "memsearch")
    summary = {"ok": False}

    if not query.strip() and not any(params.values()):
        _prog(100, "Пустой запрос.")
        return summary

    qlines = [f"Запрос: {query}", f"Тип: {entity_tab}"]
    for k, v in params.items():
        if v:
            qlines.append(f"{k}: {v}")

    _prog(0, "Запускаю браузер…")
    async with async_playwright() as pw:
        browser = await browser_util.launch(pw, 
            headless=False,
            args=["--start-maximized",
                  "--disable-blink-features=AutomationControlled"])
        ctx = await browser.new_context(no_viewport=True, accept_downloads=True)
        page = await ctx.new_page()
        try:
            _prog(5, f"Поиск: {query} ({lang})")
            if not await _do_search(page, query, lang, log):
                summary["message"] = "Не удалось выполнить поиск."
                return summary
            await _wait_if_captcha(page, log)          # search may trigger «я не робот»
            if _done():
                return summary

            await _select_tab(page, entity_tab, log)
            await _fill_advanced(page, entity_tab, params, log)
            await _apply_sources(page, sources, log)   # limit to chosen databases
            await _wait_if_captcha(page, log)          # filtering navigates → may trigger it
            await asyncio.sleep(1.0)

            # ── Collect + OPEN each card (click → new tab) page by page ── #
            _prog(15, "Сбор карточек…")
            records, seen = [], set()
            total, page_no = 0, 1
            while page_no <= max_pages and not _done():
                await _wait_if_captcha(page, log)       # pause for a challenge on any page
                cards = await _tag_cards(page)
                log(f"  → Страница {page_no}: карточек {len(cards)}")
                if page_no == 1 and not cards:
                    await _diag_cards(page, log)
                for c in cards:
                    if _done():
                        break
                    sig = (c["name"], c["summary"][:60])
                    if not c["name"] or sig in seen:
                        continue
                    seen.add(sig)
                    total += 1
                    _prog(min(90, 20 + total * 3), f"[{total}] {c['name'][:50]}…")
                    log(f"  [{total}] {c['name']}")
                    # NEVER let one bad/broken source page abort the whole run —
                    # on ANY error keep the card's own summary and move on.
                    src_url, ext = "", {}
                    try:
                        src_url, ext = await _open_card(ctx, page, c["index"],
                                                        images_dir, log,
                                                        dump_dir=(output_folder if total == 1 else None))
                    except Exception as _e:
                        log(f"      !! карточка пропущена ({type(_e).__name__}) "
                            f"— беру краткое из карточки")
                    ext = ext or {}
                    fields = ext.get("fields") or {}
                    cleaned = _clean_brief(c.get("summary", ""))
                    if fields:
                        log(f"      Полей: {len(fields)}")
                    else:
                        # source page empty/blocked → parse the cleaned card brief into
                        # fields so the record still reads as a tidy table, not nav junk.
                        fields = _brief_to_fields(cleaned)
                        log(f"      (источник пуст/битый → краткое из карточки, полей: {len(fields)})")
                    rec = {"name": c["name"] or ext.get("heading", ""),
                           "url": ("" if (not src_url
                                          or src_url.startswith("chrome-error")
                                          or src_url.startswith("about:")) else src_url),
                           "source": _host(src_url) or c.get("source", ""),
                           "summary": cleaned,
                           "fields": fields,
                           "thumb_bytes": ext.get("thumb_bytes"),
                           "photo": ext.get("photo")}
                    records.append(rec)
                    await asyncio.sleep(0.3)
                if not await _goto_next_page(page, log):
                    break
                page_no += 1
            log(f"  Всего записей: {len(records)}")
            if not records:
                _prog(100, "Ничего не найдено.")
                summary.update({"ok": True, "n_records": 0})
                return summary

            # ── Save Word (ask on conflict) ─────────────────────────── #
            _prog(92, "Сохранение…")
            base = safe_fn(f"memsearch_{query}") or "memsearch_results"
            docx_p = output_folder / f"{base}.docx"
            decision = "overwrite"
            if docx_p.exists() and ask_file_conflict:
                try:
                    decision = (ask_file_conflict([docx_p.name]) or "overwrite").lower()
                except Exception:
                    decision = "overwrite"
                log(f"  → Файл существует → {decision}")
            if decision != "skip" and records:
                write_docx(docx_p, records, qlines, append=(decision == "append"))
                log(f"  → Word: {docx_p.name}")

            _prog(100, f"Готово — {len(records)} записей.")
            summary.update({"ok": True, "n_records": len(records),
                            "output_folder": str(output_folder)})
        except Exception as exc:
            summary["message"] = f"{type(exc).__name__}: {exc}"
            log(f"  !! {exc}")
        finally:
            try:
                await ctx.close()
            except Exception:
                pass
            try:
                await browser.close()
            except Exception:
                pass
    return summary
