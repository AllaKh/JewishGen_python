#!/usr/bin/env python3
"""
skarb_scraper.py
================
Scraper for АИС Скарб (https://archiveskarb.by/db/)

Search by one or more surnames (OR logic — joined with " или ").
Filter results by keyword(s).
For each matching record: scrape all fields, download photos if available.
Save to Word (.docx) and/or Excel (.xlsx).

No login required.
"""

import asyncio
import io
import re
import sys
from pathlib import Path
from urllib.parse import urljoin, urlencode, quote

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        import os
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

# ── Constants ─────────────────────────────────────────────────────────────── #
BASE_URL    = "https://archiveskarb.by/db/"
LOGO_SRCS   = ("/bitrix/templates/", "/local/templates/")  # skip these images
HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")


# ── Helpers ───────────────────────────────────────────────────────────────── #

def safe_fn(s: str, n: int = 80) -> str:
    s = re.sub(r'[\\/*?:"<>|]', "_", (s or "").strip())
    return re.sub(r"\s+", "_", s)[:n].strip("_") or "document"


def _matches(text: str, keywords: list, mode: str) -> bool:
    """Return True if text satisfies keyword filter."""
    if not keywords:
        return True
    text_l = text.lower()
    hits = [k.lower() in text_l for k in keywords]
    return all(hits) if mode == "AND" else any(hits)


def _add_hyperlink(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


# ── Page scraping ─────────────────────────────────────────────────────────── #

async def _collect_rows(page) -> list:
    """
    Extract all record links from the current results page.
    Uses JavaScript to reliably find all a[href*=ELEMENT_ID] anchors.
    Returns list of dicts: {name, birth, category, url}
    """
    raw = await page.evaluate("""() => {
        const out = [];
        for (const a of document.querySelectorAll('a[href*="ELEMENT_ID"]')) {
            const name = (a.textContent || '').trim();
            const href = a.getAttribute('href') || '';
            if (!name || !href) continue;
            let birth = '', category = '';
            const row = a.closest('tr');
            if (row) {
                const cells = row.querySelectorAll('td');
                if (cells.length > 1) birth    = (cells[1].textContent || '').trim();
                if (cells.length > 2) category = (cells[2].textContent || '').trim();
            }
            out.push({name, href, birth, category});
        }
        return out;
    }""")
    rows = []
    for r in raw:
        rows.append({
            "name":     r["name"],
            "birth":    r["birth"],
            "category": r["category"],
            "url":      urljoin(BASE_URL, r["href"]),
        })
    return rows


async def _get_next_url(page) -> str | None:
    """
    Find the 'След.' pagination link.
    АИС Скарб uses: <a class="modern-page-next" href="...">След.</a>
    """
    href = await page.evaluate("""() => {
        // Primary: class="modern-page-next"
        let a = document.querySelector('a.modern-page-next');
        if (a) return a.getAttribute('href');
        // Fallback: any link with text "След."
        for (const el of document.querySelectorAll('a')) {
            if ((el.textContent || '').trim() === 'След.') {
                return el.getAttribute('href');
            }
        }
        return null;
    }""")
    if href:
        return urljoin(BASE_URL, href)
    return None


async def _scrape_record(page, url: str, images_dir: Path, log) -> dict:
    """
    Open a record page and scrape all fields.
    If "Наличие фото: да" — download the photo.
    """
    rec = {"url": url, "name": "", "fields": {}, "image_path": None}
    try:
        await page.goto(url, wait_until="domcontentloaded", timeout=30000)
        await asyncio.sleep(1.5)

        # ── Record name = the page heading (e.g. "Шур Шая Залмановна") ─ #
        rec["name"] = await page.evaluate(r"""() => {
            for (const sel of ['h1', 'h2', '.detail-title', '.page-title',
                               '[class*="title"]']) {
                const el = document.querySelector(sel);
                if (el) {
                    const t = (el.textContent || '').replace(/\s+/g, ' ').trim();
                    if (t && t.length < 120) return t;
                }
            }
            return '';
        }""")

        # ── Extract ALL "Label: value" pairs from the record page ──── #
        # JS-based — robust across the various Bitrix markup variants.
        # Returns an ordered list of [label, value] pairs.
        pairs = await page.evaluate(r"""() => {
            const out = [];
            const seen = new Set();
            const push = (k, v) => {
                k = (k || '').replace(/\s+/g, ' ').replace(/:$/, '').trim();
                v = (v || '').replace(/\s+/g, ' ').trim();
                if (!k || !v || k.length > 80) return;
                const key = k.toLowerCase();
                if (seen.has(key)) return;
                seen.add(key);
                out.push([k, v]);
            };

            // 1) Property tables / definition lists
            document.querySelectorAll('dl').forEach(dl => {
                const dts = dl.querySelectorAll('dt');
                const dds = dl.querySelectorAll('dd');
                for (let i = 0; i < Math.min(dts.length, dds.length); i++)
                    push(dts[i].textContent, dds[i].textContent);
            });
            document.querySelectorAll('table tr').forEach(tr => {
                const c = tr.querySelectorAll('td, th');
                if (c.length === 2) push(c[0].textContent, c[1].textContent);
            });

            // 2) "<b>/<strong> Label:</b> value" inside any block
            document.querySelectorAll('b, strong').forEach(b => {
                const lab = (b.textContent || '').trim();
                if (!lab.includes(':')) return;
                // value = text right after the bold element within its parent
                let val = '';
                let n = b.nextSibling;
                while (n) {
                    if (n.nodeType === 3) val += n.textContent;      // text node
                    else if (n.nodeType === 1) {
                        if (/^(B|STRONG|BR)$/.test(n.tagName)) break;
                        val += n.textContent;
                    }
                    n = n.nextSibling;
                }
                push(lab, val);
            });

            // 3) Plain "Label: value" lines in content blocks
            document.querySelectorAll('p, li, div').forEach(el => {
                // only leaf-ish elements (avoid huge containers)
                if (el.children.length > 2) return;
                const t = (el.textContent || '').replace(/\s+/g, ' ').trim();
                const m = t.match(/^([^:]{2,60}):\s*(.+)$/);
                if (m) push(m[1], m[2]);
            });

            return out;
        }""")

        # Preserve order, drop site-noise labels
        NOISE = ("каталог", "поиск", "меню", "вход", "контакт",
                 "анкета", "©", "all rights", "cookie")
        for k, v in pairs:
            kl = k.lower()
            if any(nz in kl for nz in NOISE):
                continue
            rec["fields"][k] = v

        log(f"      Полей: {len(rec['fields'])}")

        # ── Photo ────────────────────────────────────────────────── #
        has_photo = False
        for k, v in rec["fields"].items():
            if "фото" in k.lower() and "да" in v.lower():
                has_photo = True
                break

        if has_photo:
            log("      Наличие фото: да — ищу изображение...")
            # Find the photo img tag (skip logos/templates)
            for img in await page.query_selector_all("img[src]"):
                src = (await img.get_attribute("src") or "").strip()
                if not src or not src.startswith("/"):
                    if not src.startswith("http"):
                        continue
                if any(skip in src for skip in LOGO_SRCS):
                    continue
                # This looks like a content image
                abs_src = urljoin(BASE_URL, src)
                # Download it
                img_path = await _download_image(page, abs_src, images_dir,
                                                 url, log)
                if img_path:
                    rec["image_path"] = img_path
                    break
        else:
            # Even without "Наличие фото: да", check for content images
            for img in await page.query_selector_all("img[src]"):
                src = (await img.get_attribute("src") or "").strip()
                if not src:
                    continue
                if any(skip in src for skip in LOGO_SRCS):
                    continue
                abs_src = urljoin(BASE_URL, src)
                img_path = await _download_image(page, abs_src, images_dir,
                                                 url, log)
                if img_path:
                    rec["image_path"] = img_path
                    break

    except Exception as e:
        log(f"      !! record error: {e}")
    return rec


async def _download_image(page, src: str, images_dir: Path,
                          record_url: str, log) -> str | None:
    """Download an image by opening its URL in a new page."""
    try:
        images_dir.mkdir(parents=True, exist_ok=True)
        # Get filename from URL
        fname = src.rstrip("/").split("/")[-1]
        if not fname or "." not in fname:
            fname = "photo.jpg"
        dest = images_dir / safe_fn(fname, 60)

        # Open new page for image
        img_page = await page.context.new_page()
        try:
            resp = await img_page.goto(src, timeout=20000)
            if resp and resp.ok:
                body = await resp.body()
                if len(body) > 1000:
                    dest.write_bytes(body)
                    log(f"      Фото: {dest.name} ({len(body)//1024}KB)")
                    return str(dest)
        finally:
            await img_page.close()
    except Exception as e:
        log(f"      !! image download: {e}")
    return None


async def _get_thumbnail(page, record_url: str, log) -> bytes | None:
    """Get a small preview of the photo for Word document."""
    try:
        rec_page = await page.context.new_page()
        try:
            await rec_page.goto(record_url, wait_until="domcontentloaded", timeout=15000)
            for img in await rec_page.query_selector_all("img[src]"):
                src = (await img.get_attribute("src") or "").strip()
                if not src or any(skip in src for skip in LOGO_SRCS):
                    continue
                abs_src = urljoin(BASE_URL, src)
                resp_p = await rec_page.context.new_page()
                try:
                    resp = await resp_p.goto(abs_src, timeout=10000)
                    if resp and resp.ok:
                        body = await resp.body()
                        if len(body) > 500:
                            return body
                finally:
                    await resp_p.close()
        finally:
            await rec_page.close()
    except Exception:
        pass
    return None


# ── Word output ───────────────────────────────────────────────────────────── #

def write_docx(path: Path, records: list, query_info: dict):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    doc = Document()
    s = doc.sections[0]
    s.page_width  = Mm(210); s.page_height = Mm(297)
    s.left_margin = s.right_margin = Mm(20)
    s.top_margin  = s.bottom_margin = Mm(20)

    # Title
    h = doc.add_heading("АИС Скарб — Результаты поиска", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Query info
    for k, v in query_info.items():
        p = doc.add_paragraph()
        p.add_run(f"{k}: ").bold = True
        p.add_run(str(v))
    doc.add_paragraph(f"Найдено записей: {len(records)}")
    doc.add_paragraph("")

    for i, rec in enumerate(records, 1):
        name = rec.get("name") or f"Запись {i}"
        doc.add_heading(f"{i}. {name}", level=2)

        # Full record info — all fields from the opened page (link goes LAST)
        fields = rec.get("fields", {})
        if fields:
            tbl = doc.add_table(rows=1, cols=2)
            tbl.style = "Table Grid"
            hdr = tbl.rows[0].cells
            hdr[0].text = "Поле"; hdr[1].text = "Значение"
            for cell in hdr:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for k, v in fields.items():
                r = tbl.add_row().cells
                r[0].text = str(k); r[1].text = str(v)

        # Image
        img_path = rec.get("image_path")
        thumb    = rec.get("thumb_bytes")
        if img_path and Path(img_path).exists():
            doc.add_paragraph("Изображение:").runs[0].bold = True
            try:
                doc.add_picture(img_path, width=Inches(3.5))
            except Exception:
                doc.add_paragraph(f"  [{Path(img_path).name}]")
        elif thumb:
            doc.add_paragraph("Превью:").runs[0].bold = True
            try:
                doc.add_picture(io.BytesIO(thumb), width=Inches(3.5))
            except Exception:
                pass

        # Source link — at the very end
        if rec.get("url"):
            pp = doc.add_paragraph()
            pp.add_run("Источник: ").bold = True
            _add_hyperlink(pp, rec["url"], rec["url"])

        doc.add_paragraph("")

    doc.save(str(path))


# ── Excel output ──────────────────────────────────────────────────────────── #

def write_xlsx(path: Path, records: list, query_info: dict):
    if not _OPENPYXL_OK:
        raise RuntimeError("openpyxl не установлен")
    wb = Workbook(); ws = wb.active; ws.title = "Скарб"
    HF = PatternFill("solid", fgColor="2A4A7F")
    HN = Font(bold=True, color="FFFFFF", size=11)
    TS = Side(style="thin", color="B0B8C8")
    T  = Border(left=TS, right=TS, top=TS, bottom=TS)

    # Collect all field names
    all_fields: list = []
    for rec in records:
        for k in rec.get("fields", {}):
            if k not in all_fields:
                all_fields.append(k)

    # Columns: #, Имя, <all record fields…>, Фото, URL (link LAST)
    cols = ["#", "Имя"] + all_fields + ["Фото", "URL"]
    for ci, cn in enumerate(cols, 1):
        c = ws.cell(row=1, column=ci, value=cn)
        c.font = HN; c.fill = HF; c.border = T
        c.alignment = Alignment(horizontal="center", vertical="center",
                                 wrap_text=True)

    for ri, rec in enumerate(records, 2):
        fields = rec.get("fields", {})
        name = rec.get("name", "")
        has_img = "да" if rec.get("image_path") else "нет"
        vals = ([ri-1, name]
                + [fields.get(f, "") for f in all_fields]
                + [has_img, rec.get("url", "")])
        for ci, val in enumerate(vals, 1):
            c = ws.cell(row=ri, column=ci, value=val)
            c.border = T
            c.alignment = Alignment(wrap_text=True, vertical="top")

    for ci in range(1, len(cols)+1):
        ltr = get_column_letter(ci)
        mw = max(len(str(cols[ci-1])),
                 *(len(str(ws.cell(row=r, column=ci).value or ""))
                   for r in range(2, ws.max_row+1)), 8)
        ws.column_dimensions[ltr].width = min(mw + 4, 60)
    wb.save(str(path))


# ── Main scraper ──────────────────────────────────────────────────────────── #

async def run_scraper(
    *,
    surnames:      list        = None,   # list of surnames to search
    keywords:      list        = None,   # filter keywords
    keyword_mode:  str         = "OR",
    output_format: str         = "both",
    output_folder              = Path("."),
    log                        = print,
    progress                   = None,
    cancel_event               = None,
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress: progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    surnames      = [s.strip() for s in (surnames or []) if s.strip()]
    keywords      = [k.strip() for k in (keywords or []) if k.strip()]
    want_docx     = output_format in ("docx", "both")
    want_xlsx     = output_format in ("xlsx", "both")
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    images_dir    = output_folder / "images"
    summary       = {"ok": False}

    if not surnames:
        _prog(100, "Не указаны фамилии для поиска.")
        return summary

    # Build search query: "Фамилия1 или Фамилия2 или ..."
    # АИС Скарб search form: <input type="text" name="q" ...>
    query = " или ".join(surnames)
    q_prefix = safe_fn("_".join(surnames[:3]), 50)

    query_info = {
        "Фамилии":  query,
        "Ключевые слова": (" " + keyword_mode + " ").join(keywords) if keywords else "(нет)",
    }

    _prog(0, "Запускаю браузер...")

    async with async_playwright() as pw:
        browser = await pw.chromium.launch(
            headless=False,
            args=["--start-maximized",
                  "--disable-blink-features=AutomationControlled"],
        )
        ctx  = await browser.new_context(no_viewport=True, accept_downloads=True)
        page = await ctx.new_page()

        try:
            # ── 1. Search ──────────────────────────────────────────── #
            # АИС Скарб: GET param name="q" (from HTML: <input type="text" name="q">)
            search_url = BASE_URL + "?" + urlencode({"q": query})
            _prog(5, f"Поиск: {query}")
            await page.goto(search_url, wait_until="domcontentloaded", timeout=30000)
            await asyncio.sleep(2)

            # ── 2. Collect all pages ───────────────────────────────── #
            _prog(10, "Сбор результатов по всем страницам...")
            all_rows: list = []
            page_num = 1

            while not _done():
                rows = await _collect_rows(page)
                log(f"  Страница {page_num}: {len(rows)} строк")

                for r in rows:
                    row_text = f"{r['name']} {r['birth']} {r['category']}"
                    if _matches(row_text, keywords, keyword_mode):
                        all_rows.append(r)

                next_url = await _get_next_url(page)
                if not next_url:
                    break
                page_num += 1
                await page.goto(next_url, wait_until="domcontentloaded", timeout=30000)
                await asyncio.sleep(1.5)

            log(f"  Итого совпадений: {len(all_rows)}")
            if not all_rows:
                _prog(100, "Ничего не найдено по заданным критериям.")
                summary.update({"ok": True, "n_records": 0})
                return summary

            # ── 3. Scrape individual records ───────────────────────── #
            records: list = []
            n = len(all_rows)

            for i, row in enumerate(all_rows, 1):
                if _done(): break
                _prog(20 + int(70 * i / n),
                      f"[{i}/{n}] {row['name'][:60]}...")
                log(f"  [{i}/{n}] {row['name']}")

                rec = await _scrape_record(page, row["url"], images_dir, log)
                # Name: page heading, else the search-result link text
                if not rec.get("name"):
                    rec["name"] = row["name"]
                # Merge search-table data into fields if not already there
                if "Дата рождения" not in rec["fields"] and row["birth"]:
                    rec["fields"]["Дата рождения"] = row["birth"]
                if "Категория" not in rec["fields"] and row["category"]:
                    rec["fields"]["Категория"] = row["category"]
                # Also try to get thumbnail for Word
                if rec.get("image_path") and Path(rec["image_path"]).exists():
                    try:
                        rec["thumb_bytes"] = Path(rec["image_path"]).read_bytes()
                    except Exception:
                        pass
                records.append(rec)

            # ── 4. Save ────────────────────────────────────────────── #
            _prog(92, "Сохранение файлов...")
            if want_docx and records:
                p = output_folder / f"{q_prefix}_skarb.docx"
                write_docx(p, records, query_info)
                log(f"  Word: {p}")
            if want_xlsx and records:
                p = output_folder / f"{q_prefix}_skarb.xlsx"
                write_xlsx(p, records, query_info)
                log(f"  Excel: {p}")

            _prog(100, f"Готово — {len(records)} записей.")
            summary.update({"ok": True, "n_records": len(records),
                            "output_folder": str(output_folder)})

        except Exception as exc:
            summary["message"] = f"{type(exc).__name__}: {exc}"
            log(f"  !! {exc}")
        finally:
            try: await ctx.close()
            except Exception: pass
            try: await browser.close()
            except Exception: pass

    return summary
