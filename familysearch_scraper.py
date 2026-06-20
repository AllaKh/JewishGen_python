#!/usr/bin/env python3
"""
familysearch_scraper.py
========================
АЛГОРИТМ (точно по требованию):

1.  Домашняя страница → First Names, Last Names, Place, Birth Year → SEARCH.
2.  Ждать 5 секунд результаты.
3.  Кликнуть таб Historical Records [data-testid="hr-tab"].
    URL меняется на ?tab=records — ждать ИМЕННО ЭТО, не networkidle.
4.  ВОЙТИ ОДИН РАЗ ДО открытия записей (_sign_in_if_needed: nav-кнопка
    [data-testid="no-loggedin-sign-in-button"] → #userName/#password/#login).
    С персистентным профилем на 2-м прогоне уже залогинены → пропуск.
5.  Если ЕСТЬ advanced — ДО скрапинга: reload → HR tab → Advanced Search попап
    [data-testid="advanced-search-form-button"] → заполнить поля/супругу →
    Search → заново собрать (advanced сужает выдачу ДО скрапинга).
6.  Собрать результаты ≥80%.
7.  Все записи открываются на ОДНОЙ главной странице (не в новых вкладках) —
    сессия одна, логин не повторяется. После каждой → назад на results_url.
8.  На каждой странице записи: текст + превьюшку → Word.
9.  Кликнуть картинку → вьюер → download → JPG Only →
    [data-testid="full-text-confirm-download"] → сохранить JPG.
10. Имена файлов: {FirstNames} {LastNames}.docx / .xlsx
"""

import asyncio, difflib, io, json, os, re, shutil, sys, time
from pathlib import Path
from docx_util import set_cell_lines

if getattr(sys, "frozen", False):
    bd = Path(sys.executable).resolve().parent / "ms-playwright"
    if bd.exists():
        os.environ["PLAYWRIGHT_BROWSERS_PATH"] = str(bd)

try:
    from playwright.async_api import async_playwright
except ImportError:
    sys.stderr.write("pip install playwright && playwright install chromium\n")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Mm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    _DOCX_OK = True
except ImportError:
    _DOCX_OK = False

try:
    from openpyxl import Workbook, load_workbook
    from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
    from openpyxl.utils import get_column_letter
    _OPENPYXL_OK = True
except ImportError:
    _OPENPYXL_OK = False

# ── Конфиг ────────────────────────────────────────────────────────────────── #
_HERE     = Path(__file__).resolve().parent
try:
    _CFG = json.loads((_HERE / "config" / "familysearch.json").read_text("utf-8"))
except Exception:
    _CFG = {}

HOME_URL      = _CFG.get("home_url", "https://www.familysearch.org/en/global")
FS_BASE       = "https://www.familysearch.org"
MIN_MATCH     = int(_CFG.get("min_match", 80))
BAD_PATHS     = ["/records/images", "/search/linker", "/linker",
                 "/en/tree/", "/tree/person/", "/tree/",
                 "/catalog", "/wiki", "/books", "/films"]
_dl           = _CFG.get("downloads_dir", "")
DOWNLOADS_DIR = Path(_dl) if _dl else Path.home() / "Downloads"
FS_PROFILE_DIR = _HERE / ".fs_profile"     # persistent login/cookies between runs
HYPERLINK_REL = ("http://schemas.openxmlformats.org/"
                 "officeDocument/2006/relationships/hyperlink")
_IMG_SKIP     = ("icon", "logo", "sprite", "avatar", "pixel", "placeholder",
                 ".svg", "fscdn.org")


# ── Утилиты ───────────────────────────────────────────────────────────────── #

def _sim(a: str, b: str) -> float:
    a = re.sub(r"\s+", " ", a.strip().lower())
    b = re.sub(r"\s+", " ", b.strip().lower())
    if not a or not b:
        return 0.0
    wa, wb = set(a.split()), set(b.split())
    return max(len(wa & wb) / max(len(wa), 1),
               difflib.SequenceMatcher(None, a, b).ratio()) * 100


def safe_fn(s: str, n: int = 100) -> str:
    return re.sub(r"\s+", " ",
                  re.sub(r'[\\/*?:"<>|]', "_", s.strip()))[:n].strip() or "document"


def _is_record(href: str) -> bool:
    if not href:
        return False
    for bad in BAD_PATHS:
        if bad in href:
            return False
    return "/ark:" in href


def _abs(href: str) -> str:
    return FS_BASE + href if href.startswith("/") else href


def _year_range(year, span: int = 2):
    """(from, to) for the FS birth-year filter — ±span around the entered year (FS
    treats a birth year as a small range), so «1897» also returns the 1898 census
    records. With span=0 (Exact ticked) it's from=to=year. (None, None) if no year."""
    try:
        y = int(re.sub(r"\D", "", str(year))[:4])
    except Exception:
        y = 0
    return (y - span, y + span) if y else (None, None)


def _fs_url(first_names, last_names, place_lived, birth_year, year_range,
            exact, adv) -> str:
    """Build the FamilySearch results URL with every q.*/f.* param straight from
    the form — the reliable, site-matching alternative to the advanced popup.
    Mirrors the live URL: q.givenName / q.surname (+.exact), q.birthLikeDate
    .from/.to/.exact, q.<event>Date/.Place, q.{spouse,father,mother,other}{GivenName
    ,Surname} (repeated for up to 3 each), q.recordCountry / q.recordSubcountry,
    q.batchNumber, q.filmNumber, q.isPrincipal, f.recordType=0..7."""
    from urllib.parse import quote
    exact = exact or {}
    adv = adv or {}
    parts = ["tab=records", "page=1", "results=60"]

    def add(k, v):
        parts.append(f"{k}={quote(str(v), safe='')}")

    if first_names:
        add("q.givenName", first_names)
        if exact.get("name"): parts.append("q.givenName.exact=on")
    if last_names:
        add("q.surname", last_names)
        if exact.get("surname"): parts.append("q.surname.exact=on")
    # alternate names of the principal → indexed «.1/.2/.3» on the own name params
    # (best-effort: this exact param wasn't in the user's sample URL).
    for i, a in enumerate(adv.get("alt_names") or [], start=1):
        if a.get("first"):
            add(f"q.givenName.{i}", a["first"])
            if a.get("first_exact"): parts.append(f"q.givenName.exact.{i}=on")
        if a.get("last"):
            add(f"q.surname.{i}", a["last"])
            if a.get("last_exact"): parts.append(f"q.surname.exact.{i}=on")
    sex = adv.get("sex")
    if sex and sex != "Unspecified": add("q.sex", sex)
    if adv.get("keywords"): add("q.text", adv["keywords"])

    # ── life events: Birth / Marriage / Residence / Death / Any ──────────────
    # each carries Place (+own exact) and a Year with a ± range (+own date exact).
    # «this year» / Exact (range 0) → only «.from» (no «.to»), matching the live
    # URL; ±N → from=year-N & to=year+N. The «Exact +/-» tick = «.exact=on», an
    # INDEPENDENT flag (birth «this year» has no .exact; ±2 death may have it).
    EVD = {"birth": "birthLikeDate", "marriage": "marriageLikeDate",
           "death": "deathLikeDate", "residence": "residenceDate", "any": "anyDate"}
    EVP = {"birth": "birthLikePlace", "marriage": "marriageLikePlace",
           "death": "deathLikePlace", "residence": "residencePlace", "any": "anyPlace"}
    events = dict(adv.get("events") or {})
    if birth_year and "birth" not in events:        # legacy basic Birth Year
        events["birth"] = {"year": birth_year, "range": year_range,
                           "place": place_lived, "place_exact": exact.get("place")}
    elif place_lived and "any" not in events:       # legacy basic Place → any place
        add("q.anyPlace", place_lived)
        if exact.get("place"): parts.append("q.anyPlace.exact=on")
    for key, ev in events.items():
        if key not in EVD:
            continue
        yr = str(ev.get("year") or "").strip()
        if yr.isdigit():
            y, n = int(yr), int(ev.get("range") or 0)
            if n <= 0:
                add(f"q.{EVD[key]}.from", y)
            else:
                add(f"q.{EVD[key]}.from", y - n); add(f"q.{EVD[key]}.to", y + n)
            if ev.get("date_exact"): parts.append(f"q.{EVD[key]}.exact=on")
        if ev.get("place"):
            add(f"q.{EVP[key]}", ev["place"])
            if ev.get("place_exact"): parts.append(f"q.{EVP[key]}.exact=on")

    # ── family members: up to 3 of each; 1st = base param, 2nd = «.1», 3rd = «.2».
    # exact flag is «q.<field>.exact», «q.<field>.exact.1», «q.<field>.exact.2».
    FAMQ = {"spouse": ("spouseGivenName", "spouseSurname"),
            "father": ("fatherGivenName", "fatherSurname"),
            "mother": ("motherGivenName", "motherSurname"),
            "other":  ("otherGivenName",  "otherSurname")}
    for key, people in (adv.get("family") or {}).items():
        gk, sk = FAMQ.get(key, (None, None))
        if not gk:
            continue
        for i, p in enumerate(people[:3]):
            sfx = "" if i == 0 else f".{i}"          # base / .1 / .2
            if p.get("first"):
                add(f"q.{gk}{sfx}", p["first"])
                if p.get("first_exact"): parts.append(f"q.{gk}.exact{sfx}=on")
            if p.get("last"):
                add(f"q.{sk}{sfx}", p["last"])
                if p.get("last_exact"): parts.append(f"q.{sk}.exact{sfx}=on")

    if adv.get("country"):
        add("q.recordCountry", adv["country"])
        if adv.get("state"):
            add("q.recordSubcountry", f"{adv['country']},{adv['state']}")
    if adv.get("batch"):     add("q.batchNumber", adv["batch"])
    if adv.get("film"):      add("q.filmNumber", adv["film"])
    if adv.get("principal"): parts.append("q.isPrincipal=true")
    for i in (adv.get("record_types") or []):
        parts.append(f"f.recordType={i}")
    return "https://www.familysearch.org/en/search/discovery/results?" + "&".join(parts)


def _apply_exact(url: str, exact: dict) -> str:
    """Add `.exact=on` to whichever q.* params FS already put in the results URL,
    for the fields the user ticked «Exact». Field-name-agnostic — we read what FS
    actually used (q.givenName / q.surname / q.anyPlace / q.birthLikePlace / …) and
    flag it, instead of guessing FS's param names."""
    if not exact:
        return url
    from urllib.parse import urlparse, parse_qsl, urlencode
    pr    = urlparse(url)
    pairs = parse_qsl(pr.query, keep_blank_values=True)
    keys  = {k for k, _ in pairs}
    want  = []
    if exact.get("name"):
        want += [k for k in keys if k.startswith("q.givenName")
                 and not k.endswith(".exact")]
    if exact.get("surname"):
        want += [k for k in keys if k.startswith("q.surname")
                 and not k.endswith(".exact")]
    if exact.get("place"):
        want += [k for k in keys
                 if (k.startswith("q.anyPlace") or "Place" in k)
                 and not k.endswith(".exact") and "birthLikeDate" not in k]
    for k in want:
        if f"{k}.exact" not in keys:
            pairs.append((f"{k}.exact", "on"))
            keys.add(f"{k}.exact")
    return pr._replace(query=urlencode(pairs)).geturl()


# Relationship roles glue to the name in the results cell («SpouseRebecca M
# Sanders»). innerText usually separates them, but if FS renders them inline this
# splits the role from the name and puts each relationship on its own line.
_REL_ROLES = ("Spouses", "Spouse", "Parents", "Parent", "Father", "Mother",
              "Children", "Child", "Wife", "Husband", "Son", "Daughter",
              "Siblings", "Sibling", "Brother", "Sister")
_REL_ALT   = "|".join(_REL_ROLES)
_REL_BREAK = re.compile(r"(?<=[A-Za-z.)])(?=(?:" + _REL_ALT + r")[A-Z])")
_REL_COLON = re.compile(r"^(" + _REL_ALT + r")(?=[A-Z])", re.M)


def _split_rels(text: str) -> str:
    if not text:
        return text
    t = _REL_BREAK.sub("\n", text)      # newline before each role
    t = _REL_COLON.sub(r"\1: ", t)      # «SpouseRebecca» → «Spouse: Rebecca»
    return t


# ── Ввод в поле поиска ────────────────────────────────────────────────────── #

async def _type_field(page, sel: str, val: str, label: str, log) -> bool:
    """Кликнуть, очистить Ctrl+A/Del, печатать посимвольно."""
    if not val:
        return True
    try:
        el = page.locator(sel).first
        if not await el.count():
            log(f"  !! поле {label} не найдено ({sel})")
            return False
        await el.scroll_into_view_if_needed(timeout=4000)
        await el.click(timeout=3000)
        await asyncio.sleep(0.1)
        await page.keyboard.press("Control+a")
        await page.keyboard.press("Delete")
        await page.keyboard.type(val, delay=40)
        await asyncio.sleep(0.2)
        got = (await el.input_value(timeout=2000)).strip()
        if got:
            log(f"  OK  {label} = {got!r}")
            return True
        log(f"  !! {label}: поле осталось пустым")
    except Exception as e:
        log(f"  !! {label}: {e}")
    return False


# ── 1. ПОИСК ──────────────────────────────────────────────────────────────── #

async def _search(page, fn, ln, place, year, log, exact=None):
    exact = exact or {}
    log(f"  Открываю {HOME_URL}")
    await page.goto(HOME_URL, wait_until="domcontentloaded", timeout=30000)
    await asyncio.sleep(3)

    # Ждём форму
    for sel in ['input[placeholder*="First and Middle" i]',
                'input[id*="givenName" i]']:
        try:
            await page.locator(sel).first.wait_for(state="visible", timeout=10000)
            break
        except Exception:
            continue

    await _type_field(page,
        'input[placeholder*="First and Middle" i], input[id*="givenName" i]',
        fn, "First Names", log)
    await _type_field(page,
        'input[placeholder*="Last or Maiden" i], input[id*="surname" i]',
        ln, "Last Names", log)
    if place:
        await _type_field(page,
            'input[placeholder*="City, County, State" i]',
            place, "Place Lived", log)
    if year:
        # FS renamed this field; the real placeholder is just "Year" (same as the
        # Advanced Search birth field). Try several, but the URL fallback below is
        # what actually guarantees the year filter.
        await _type_field(page,
            'input[placeholder="Year"], input[placeholder*="Birth Year" i], '
            'input[id*="birthYear" i], input[id*="birthLikeDate" i], '
            'input[aria-label*="Year" i]',
            year, "Birth Year", log)

    for sel in ['button:has-text("SEARCH")', 'button:has-text("Search")',
                'button[type="submit"]']:
        try:
            el = page.locator(sel).first
            if await el.count() and await el.is_visible():
                await el.click(timeout=6000)
                log("  SEARCH нажат")
                break
        except Exception:
            continue

    try:
        await page.wait_for_url(lambda u: "discovery/results" in u, timeout=20000)
    except Exception:
        pass
    await asyncio.sleep(5)  # обязательно 5 секунд

    # GUARANTEE the birth-year filter: the form field is unreliable, so if the
    # results URL doesn't carry the birth date, add it (q.birthLikeDate.from/to)
    # and reload. Without it the search isn't limited to e.g. 1897 and the right
    # person never reaches the result list.
    u = page.url
    if "discovery/results" in u:
        new_u, changed = u, False
        if year and "birthLikeDate" not in u:
            lo, hi = _year_range(year, 0 if exact.get("year") else 2)
            if lo:
                sep = "&" if "?" in new_u else "?"
                new_u = f"{new_u}{sep}q.birthLikeDate.from={lo}&q.birthLikeDate.to={hi}"
                changed = True
                tag = "точно" if exact.get("year") else "±2"
                log(f"  → год рождения {year} ({tag}) в URL: {lo}–{hi}")
        u2 = _apply_exact(new_u, exact)
        if u2 != new_u:
            new_u = u2; changed = True
            log(f"  → точное совпадение: {[k for k,v in exact.items() if v]}")
        if changed:
            try:
                await page.goto(new_u, wait_until="domcontentloaded", timeout=25000)
                await asyncio.sleep(5)
            except Exception as _e:
                log(f"  !! не вышло уточнить URL: {type(_e).__name__}")
    log(f"  Результаты: {page.url}")


async def _ensure_year_in_url(page, year, log, exact=None):
    """Re-add the birth-year filter (q.birthLikeDate.from/to) to the current
    results / records URL. The HR tab and the Advanced-Search reload rebuild the
    URL and can drop it, so we call this right before reading rows — otherwise the
    search isn't limited to the wanted year and the right person never shows up."""
    exact = exact or {}
    if not year:
        return
    u = page.url
    lo, hi = _year_range(year, 0 if exact.get("year") else 2)
    if not lo or ("tab=records" not in u and "discovery/results" not in u) \
            or "birthLikeDate" in u:
        return
    sep = "&" if "?" in u else "?"
    new_u = _apply_exact(
        f"{u}{sep}q.birthLikeDate.from={lo}&q.birthLikeDate.to={hi}", exact)
    try:
        await page.goto(new_u, wait_until="domcontentloaded", timeout=25000)
        try:
            await page.wait_for_selector("tbody tr", timeout=15000)
            log(f"  → год {year} закреплён в URL результатов")
        except Exception:
            # no rows with the year filter (wrong param or zero matches) →
            # fall back to the original URL so we never lose all results.
            log("  !! с фильтром года строк нет — возвращаюсь без него")
            await page.goto(u, wait_until="domcontentloaded", timeout=20000)
            await asyncio.sleep(3)
    except Exception as _e:
        log(f"  !! год в URL не закрепился: {type(_e).__name__}")


# ── 2. ЛОГИН ЧЕРЕЗ NAV-КНОПКУ (до HR tab) ────────────────────────────────── #

async def _sign_in_if_needed(page, email: str, password: str,
                             logged_in_ref: list, log) -> bool:
    """Sign in ONCE up front (via the nav «Sign In» button) so records open
    already logged-in and Advanced Search runs on a logged-in results page. With
    the persistent profile the session is usually already there → the button is
    absent → skip. Sets logged_in_ref[0]=True only on an ACTUAL successful login,
    so a false «already logged in» still lets _scrape_page log in at record #1."""
    log("  Проверяю авторизацию...")
    results_url = page.url
    try:
        btn = page.locator('[data-testid="no-loggedin-sign-in-button"]').first
        await btn.wait_for(state="visible", timeout=6000)
    except Exception:
        log("  Уже авторизованы (кнопка Sign In не видна) — вход не нужен")
        return True

    log("  Не авторизованы — вхожу ОДИН раз до открытия записей...")
    try:
        await btn.click(timeout=5000)
        try:
            await page.wait_for_url(
                lambda u: "login" in u or "ident.familysearch" in u, timeout=12000)
        except Exception:
            pass
        await asyncio.sleep(2)
        ok = await _login(page, email, password, log)
        if ok:
            logged_in_ref[0] = True
            if "discovery/results" not in page.url:
                log("  Возвращаюсь на страницу результатов...")
                await page.goto(results_url,
                                wait_until="domcontentloaded", timeout=20000)
                await asyncio.sleep(3)
        return ok
    except Exception as e:
        log(f"  !! Sign in: {e}")
        return False


# ── 3. ТАБ HISTORICAL RECORDS ─────────────────────────────────────────────── #

async def _click_hr(page, log):
    """
    Кликнуть [data-testid='hr-tab'].
    URL меняется на ?tab=records.
    После этого ждать появления строк таблицы (tbody tr) — до 10 секунд.
    """
    clicked = False
    try:
        el = page.locator('[data-testid="hr-tab"]').first
        if await el.count() and await el.is_visible():
            await el.click(timeout=5000)
            clicked = True
    except Exception:
        pass

    if not clicked:
        try:
            el = page.get_by_role("tab", name=re.compile(r"Historical Records", re.I)).first
            if await el.count():
                await el.click(timeout=5000)
                clicked = True
        except Exception:
            pass

    if not clicked:
        log("  (HR tab не найден)")
        return

    # Ждём URL с tab=records
    try:
        await page.wait_for_url(lambda u: "tab=records" in u, timeout=8000)
    except Exception:
        await asyncio.sleep(2)

    # Ждём СТРОКИ ТАБЛИЦЫ (таблица грузится через XHR после смены URL)
    try:
        await page.wait_for_selector("tbody tr", timeout=20000)
    except Exception:
        await asyncio.sleep(5)  # если не дождались — всё равно идём дальше

    log(f"  HR tab → {page.url}")


# ── 3. ЛОГИН — ОДИН РАЗ ──────────────────────────────────────────────────── #

async def _login(page, username: str, password: str, log) -> bool:
    """
    Заполнить #userName и #password через page.fill().
    Нажать #login. Ждать редиректа обратно на familysearch.org.
    Вызывать ТОЛЬКО ОДИН РАЗ за сессию.
    """
    log(f"  Страница логина: {page.url[:80]}")

    # Ждём поле #userName
    try:
        await page.locator("#userName").wait_for(state="visible", timeout=20000)
    except Exception:
        log("  !! #userName не появился")
        return False

    await asyncio.sleep(2)  # дать форме полностью отрисоваться

    # Заполняем с retry: иногда React-форма сбрасывает значение после fill
    for attempt in range(3):
        try:
            await page.fill("#userName", username)
            await asyncio.sleep(0.5)
            u_val = (await page.locator("#userName").input_value()).strip()
            if u_val:
                log(f"  #userName заполнен")
                break
            log(f"  !! #userName пусто после fill (попытка {attempt+1})")
            await asyncio.sleep(1)
        except Exception as e:
            log(f"  !! fill #userName: {e}")
            if attempt == 2:
                return False
            await asyncio.sleep(1)
    else:
        log("  !! #userName не удалось заполнить")
        return False

    for attempt in range(3):
        try:
            await page.fill("#password", password)
            await asyncio.sleep(0.5)
            p_val = (await page.locator("#password").input_value()).strip()
            if p_val:
                log(f"  #password заполнен")
                break
            log(f"  !! #password пусто после fill (попытка {attempt+1})")
            await asyncio.sleep(1)
        except Exception as e:
            log(f"  !! fill #password: {e}")
            if attempt == 2:
                return False
            await asyncio.sleep(1)
    else:
        log("  !! #password не удалось заполнить")
        return False

    # Финальная проверка
    u_val = await page.locator("#userName").input_value()
    p_val = await page.locator("#password").input_value()
    log(f"  Проверка: user={'OK' if u_val else '!ПУСТО'}, "
        f"pass={'OK' if p_val else '!ПУСТО'}")

    if not u_val.strip() or not p_val.strip():
        log("  !! Поля пустые — логин провалится")
        return False

    # Нажать ТОЛЬКО #login — ничего другого
    try:
        await page.click("#login", timeout=5000)
        log("  #login нажат")
    except Exception as e:
        log(f"  !! #login: {e}")
        return False

    # Ждём редирект
    try:
        await page.wait_for_url(
            lambda u: "familysearch.org" in u and "login" not in u,
            timeout=30000)
    except Exception:
        await asyncio.sleep(6)

    ok = "familysearch.org" in page.url and "login" not in page.url
    log(f"  Логин {'OK ✓' if ok else 'ПРОВАЛИЛСЯ ✗'}  URL: {page.url[:80]}")
    return ok


# ── 4. СОБРАТЬ СТРОКИ ТАБЛИЦЫ ────────────────────────────────────────────── #

async def _collect(page, qname: str, log) -> list:
    """
    Структура таблицы (подтверждена):
      cells[0] = ссылка "More"  href=/ark:... ← URL записи, НЕ имя!
      cells[1] = картинка
      cells[2] = <strong>Имя</strong> + коллекция 2-й строкой
      cells[3] = Events
      cells[4] = Relationships
    """
    await asyncio.sleep(2)
    results, seen = [], set()
    # The result table loads via XHR — wait until a row has a REAL record link, not
    # just skeleton <tr>s. Otherwise we read N empty rows → 0 candidates (the «Строк:
    # 12 → Кандидатов: 0» bug after a login redirect resets the table).
    rows = []
    for _ in range(25):                              # up to ~25s
        rows = await page.query_selector_all("tbody tr")
        ready = False
        for row in rows:
            for a in await row.query_selector_all("a[href]"):
                if _is_record(await a.get_attribute("href") or ""):
                    ready = True
                    break
            if ready:
                break
        if ready:
            break
        await asyncio.sleep(1)
    log(f"  Строк: {len(rows)}")

    for row in rows:
        try:
            cells = await row.query_selector_all("td")
            if not cells:
                continue
            # URL из ark-ссылки в cells[0]
            url = ""
            for a in await row.query_selector_all("a[href]"):
                h = (await a.get_attribute("href") or "").strip()
                if _is_record(h):
                    url = _abs(h)
                    break
            if not url or url in seen:
                continue
            seen.add(url)

            # Имя из cells[2] <strong>
            idx  = 2 if len(cells) > 3 else max(0, len(cells) - 3)
            name = ""
            coll = ""
            if len(cells) > idx:
                nc = cells[idx]
                try:
                    b = await nc.query_selector("strong, b")
                    if b:
                        name = (await b.text_content() or "").strip()
                except Exception:
                    pass
                if not name:
                    lines = [ln.strip()
                             for ln in (await nc.text_content() or "").splitlines()
                             if ln.strip()]
                    name = lines[0] if lines else ""
                    coll = lines[1] if len(lines) > 1 else ""
                else:
                    try:
                        lines = [ln.strip()
                                 for ln in (await nc.text_content() or "").splitlines()
                                 if ln.strip()]
                        coll = lines[1] if len(lines) > 1 else ""
                    except Exception:
                        pass

            if not name:
                continue
            # innerText (NOT text_content) so the role and the name aren't glued
            # («SpouseRebecca M Sanders») — innerText keeps FS's line breaks.
            async def _celltext(c):
                try:
                    return (await c.evaluate("e => e.innerText") or "").strip()
                except Exception:
                    return (await c.text_content() or "").strip()
            evts = await _celltext(cells[idx+1]) if len(cells) > idx+1 else ""
            rels = await _celltext(cells[idx+2]) if len(cells) > idx+2 else ""
            score = round(_sim(qname, name), 1)
            results.append({"url": url, "name": name, "coll": coll,
                            "evts": evts, "rels": rels, "score": score})
            log(f"    {score:5.1f}%  {name}")
        except Exception:
            continue

    log(f"  Кандидатов: {len(results)}")
    return results


# ── 5. УСТАНОВИТЬ 60 НА СТРАНИЦУ ─────────────────────────────────────────── #

async def _set_60(page, log):
    for sel in ['select[aria-label*="result" i]',
                'select[name*="result" i]',
                'select[id*="result" i]']:
        try:
            el = page.locator(sel).last
            if not await el.count():
                continue
            opts = await el.evaluate("e => Array.from(e.options).map(o=>o.value)")
            if "60" in opts:
                await el.select_option(value="60")
                await asyncio.sleep(1)
                log("  60 результатов на странице")
                return
        except Exception:
            continue


# ── 6. РАСШИРЕННЫЙ ПОИСК ─────────────────────────────────────────────────── #

async def _advanced(page, adv: dict, log):
    log("  Открываю Advanced Search...")
    try:
        btn = page.locator('[data-testid="advanced-search-form-button"]').first
        await btn.wait_for(state="visible", timeout=10000)
        await btn.scroll_into_view_if_needed(timeout=3000)
        await btn.click(timeout=5000)
        await asyncio.sleep(1.5)
        log("  Модальное окно открыто")
    except Exception as e:
        log(f"  !! Advanced Search: {e}")
        return

    event_map = {
        "birth_place":     ("BIRTH",     'input[placeholder*="City, County, State, Province"]'),
        "birth_year":      ("BIRTH",     'input[placeholder="Year"]'),
        "death_place":     ("DEATH",     'input[placeholder*="City, County, State, Province"]'),
        "death_year":      ("DEATH",     'input[placeholder="Year"]'),
        "marriage_place":  ("MARRIAGE",  'input[placeholder*="City, County, State, Province"]'),
        "marriage_year":   ("MARRIAGE",  'input[placeholder="Year"]'),
        "residence_place": ("RESIDENCE", 'input[placeholder*="City, County, State, Province"]'),
        "residence_year":  ("RESIDENCE", 'input[placeholder="Year"]'),
        "any_place":       ("ANY",       'input[placeholder*="City, County, State, Province"]'),
        "any_year":        ("ANY",       'input[placeholder="Year"]'),
    }
    opened: set = set()
    for key, (tab, sel) in event_map.items():
        val = adv.get(key, "")
        if not val:
            continue
        if tab not in opened:
            try:
                t = page.get_by_text(re.compile(rf"^{re.escape(tab)}$", re.I)).first
                if await t.count():
                    await t.click(timeout=3000)
                    await asyncio.sleep(0.7)
                    opened.add(tab)
            except Exception:
                pass
        await _type_field(page, sel, val, key, log)

    # Семейные члены — через data-testid (точные имена из HTML FamilySearch)
    # (btn_testid, given_field_testid, given_exact_testid, surname_field_testid, surname_exact_testid)
    fam_testids = {
        "spouse": ("spouse-fieldGroupButton",
                   "spouseGivenName0-field",  "q_spouseGivenName_exact",
                   "spouseSurname0-field",    "q_spouseSurname_exact"),
        "father": ("father-fieldGroupButton",
                   "fatherGivenName0-field",  "q_fatherGivenName_exact",
                   "fatherSurname0-field",    "q_fatherSurname_exact"),
        "mother": ("mother-fieldGroupButton",
                   "motherGivenName0-field",  "q_motherGivenName_exact",
                   "motherSurname0-field",    "q_motherSurname_exact"),
        "other":  ("otherPerson-fieldGroupButton",
                   "otherGivenName0-field",   "q_otherGivenName_exact",
                   "otherSurname0-field",     "q_otherSurname_exact"),
    }
    for key, (btn_tid, fn_tid, fn_exact_tid, ln_tid, ln_exact_tid) in fam_testids.items():
        fv          = adv.get(f"{key}_first", "")
        lv          = adv.get(f"{key}_last",  "")
        fv_exact    = adv.get(f"{key}_first_exact", False)
        lv_exact    = adv.get(f"{key}_last_exact",  False)
        if not fv and not lv:
            continue

        # Нажать кнопку раскрытия секции (Spouse / Father / Mother / Other Person)
        try:
            btn = page.locator(f'[data-testid="{btn_tid}"]').first
            if await btn.count():
                await btn.click(timeout=3000)
                await asyncio.sleep(1)
                log(f"  {key} section expanded")
            else:
                log(f"  !! {key}: кнопка {btn_tid} не найдена")
        except Exception as e:
            log(f"  !! {key} expand: {e}")

        for tid, exact_tid, val, exact, label in [
            (fn_tid, fn_exact_tid, fv, fv_exact, f"{key}_first"),
            (ln_tid, ln_exact_tid, lv, lv_exact, f"{key}_last"),
        ]:
            if not val:
                continue
            try:
                el = page.locator(f'[data-testid="{tid}"]').first
                if not await el.count():
                    log(f"  !! {label}: поле {tid} не найдено")
                    continue
                await el.scroll_into_view_if_needed(timeout=3000)
                await el.click(timeout=3000)
                await page.keyboard.press("Control+a")
                await page.keyboard.press("Delete")
                await page.keyboard.type(val, delay=40)
                await asyncio.sleep(0.2)
                log(f"  OK  {label} = {val!r}")
                if exact:
                    try:
                        cb = page.locator(f'[data-testid="{exact_tid}"]').first
                        if await cb.count():
                            if not await cb.is_checked():
                                await cb.click(timeout=2000)
                            log(f"  OK  {label} exact ✓")
                        else:
                            log(f"  !! {label} exact: {exact_tid} не найден")
                    except Exception as e:
                        log(f"  !! {label} exact: {e}")
            except Exception as e:
                log(f"  !! {label}: {e}")

    if adv.get("country"):
        try:
            t = page.get_by_text(re.compile(r"^LOCATION$", re.I)).first
            if await t.count():
                await t.click(timeout=3000)
                await asyncio.sleep(0.5)
        except Exception:
            pass
        await _type_field(page, 'input[placeholder="Country or Location"]',
                         adv["country"], "country", log)
    if adv.get("state"):
        await _type_field(page, 'input[placeholder="State or Province"]',
                         adv["state"], "state", log)
    if adv.get("keywords"):
        await _type_field(page, 'input[placeholder*="keyword" i]',
                         adv["keywords"], "keywords", log)

    for sel in ['[data-testid="search-button"]',
                'button:has-text("SEARCH")', 'button:has-text("Search")']:
        try:
            el = page.locator(sel).last
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                log(f"  Search нажат ({sel})")
                break
        except Exception:
            continue

    try:
        await page.wait_for_url(lambda u: "tab=records" in u, timeout=15000)
    except Exception:
        await asyncio.sleep(3)
    log("  Advanced Search отправлен")


# ── 7. СКАЧАТЬ БАЙТЫ КАРТИНКИ ─────────────────────────────────────────────── #

async def _fetch_bytes(ctx, src: str) -> bytes | None:
    if not src or not src.startswith("http"):
        return None
    pg = await ctx.new_page()
    try:
        r = await pg.goto(src, timeout=15000)
        if r and r.ok:
            body = await r.body()
            if len(body) > 5000:
                return body
    except Exception:
        pass
    finally:
        await pg.close()
    return None


# ── 8. НАЙТИ ЛУЧШУЮ КАРТИНКУ НА СТРАНИЦЕ ─────────────────────────────────── #

async def _best_img(page) -> str:
    """Largest document image on the page. Scrolls first (FS loads the viewer image
    lazily) and falls back to a viewer-looking image even when its size is still
    unknown — so view=index viewer pages aren't wrongly reported as «no image»."""
    try:
        await page.evaluate("window.scrollTo(0, document.body.scrollHeight)")
        await asyncio.sleep(0.6)
        await page.evaluate("window.scrollTo(0, 0)")
        await asyncio.sleep(0.3)
    except Exception:
        pass
    best, area, fallback = "", 0, ""
    for el in await page.query_selector_all("img"):
        try:
            src = (await el.evaluate("e => e.currentSrc || e.src || ''") or "").strip()
            if not src.startswith("http"):
                continue
            if any(b.lower() in src.lower() for b in _IMG_SKIP):
                continue
            w = int(await el.evaluate("e => e.naturalWidth")  or 0)
            h = int(await el.evaluate("e => e.naturalHeight") or 0)
            if w * h > area:
                area = w * h
                best = src
            if not fallback and any(k in src.lower() for k in
                                    ("dz/v1", "apiv2", "/dz/", "sg.familysearch",
                                     "/records/image", "/ark:")):
                fallback = src
        except Exception:
            continue
    return best or fallback


# ── 9. СКАЧАТЬ ПОЛНОФОРМАТНУЮ JPG ЧЕРЕЗ ВЬЮЕР ────────────────────────────── #

async def _download_jpg(ctx, page, dest_dir: Path, title: str, log) -> str | None:
    """
    Кликнуть картинку → вьюер → download → JPG Only →
    [data-testid="full-text-confirm-download"] → сохранить файл.
    """
    dest_dir.mkdir(parents=True, exist_ok=True)
    fname  = safe_fn(title) + ".jpg"
    dest   = dest_dir / fname
    before = set(DOWNLOADS_DIR.glob("*.jpg")) | set(DOWNLOADS_DIR.glob("*.jpeg"))
    tabs_before = set(ctx.pages)

    # Найти и кликнуть лучшую картинку
    img_src = await _best_img(page)
    clicked = False
    if img_src:
        try:
            el = page.locator(f'img[src="{img_src}"]').first
            if await el.count() and await el.is_visible():
                await el.click(timeout=5000)
                await asyncio.sleep(3)
                clicked = True
                log("    Картинка кликнута")
        except Exception:
            pass

    if not clicked:
        for sel in ['img[src*="dz/v1"]', 'img[src*="apiv2"]',
                    'img[alt="Thumbnail"]', 'main img']:
            try:
                el = page.locator(sel).first
                if not await el.count():
                    continue
                src = await el.get_attribute("src") or ""
                if any(b.lower() in src.lower() for b in _IMG_SKIP):
                    continue
                if await el.is_visible():
                    await el.click(timeout=5000)
                    await asyncio.sleep(3)
                    img_src = src
                    clicked = True
                    log(f"    Картинка кликнута ({sel})")
                    break
            except Exception:
                continue

    if not clicked:
        log("    Картинка на странице не найдена")
        return None

    # Вьюер может открыться в новой вкладке
    viewer = page
    await asyncio.sleep(1)
    new_tabs = set(ctx.pages) - tabs_before
    if new_tabs:
        viewer = list(new_tabs)[0]
        try:
            await viewer.wait_for_load_state("domcontentloaded", timeout=10000)
        except Exception:
            pass
        await asyncio.sleep(2)
        log("    Вьюер в новой вкладке")

    # Ждём появления основной кнопки Download (после логина viewer грузится дольше)
    try:
        await viewer.wait_for_selector(
            'button[aria-label*="Download" i]', timeout=8000)
    except Exception:
        pass  # если не появилась — пробуем всё равно

    # Весь блок download обёрнут в expect_download чтобы не пропустить событие
    downloaded = None
    try:
        async with viewer.expect_download(timeout=45000) as dl_info:
            # Кнопка download (стрелка вниз)
            dl_ok = False
            for sel in [
                'button[aria-label*="Download" i]',
                'button[title*="Download" i]',
                '[data-testid*="download" i]:not([data-testid="full-text-confirm-download"])',
                '[class*="toolbar"] button:nth-last-child(3)',
                '[class*="toolbar"] button:nth-last-child(2)',
                '[class*="tools"] button:nth-last-child(2)',
            ]:
                try:
                    el = viewer.locator(sel).first
                    if await el.count() and await el.is_visible():
                        await el.click(timeout=4000)
                        await asyncio.sleep(1.5)
                        dl_ok = True
                        log(f"    Download нажат ({sel})")
                        break
                except Exception:
                    continue

            if not dl_ok:
                raise RuntimeError("кнопка download не найдена")

            # JPG Only
            for lbl in ("JPG Only", "JPG only"):
                try:
                    el = viewer.get_by_text(lbl, exact=True).first
                    if await el.count():
                        await el.click(timeout=3000)
                        await asyncio.sleep(0.5)
                        log("    JPG Only выбран")
                        break
                except Exception:
                    continue

            # Confirm download
            for sel in ['[data-testid="full-text-confirm-download"]',
                        'button:has-text("Download")',
                        'button:has-text("DOWNLOAD")']:
                try:
                    btn = viewer.locator(sel).first
                    if await btn.count() and await btn.is_visible():
                        await btn.click(timeout=5000)
                        log(f"    Download нажат ({sel})")
                        break
                except Exception:
                    continue

        dl = await dl_info.value
        await dl.save_as(str(dest))
        downloaded = str(dest)
        log(f"    Сохранено: {fname} ({dest.stat().st_size//1024}KB)")
    except Exception as exc:
        log(f"    expect_download: {exc} → жду в Downloads...")
        for _ in range(15):
            await asyncio.sleep(1)
            after = set(DOWNLOADS_DIR.glob("*.jpg")) | set(DOWNLOADS_DIR.glob("*.jpeg"))
            nw = after - before
            if nw:
                src_f = max(nw, key=lambda p: p.stat().st_mtime)
                shutil.move(str(src_f), str(dest))
                downloaded = str(dest)
                log(f"    Перемещено: {fname}")
                break

    if not downloaded and not img_src:
        # Кнопка не найдена, нет картинки
        if viewer is not page:
            try: await viewer.close()
            except Exception: pass
        return None

    # Закрыть лишние вкладки
    await asyncio.sleep(0.5)
    for pg in list(ctx.pages):
        if pg not in (page, viewer) and "familysearch" not in pg.url:
            try: await pg.close()
            except Exception: pass
    if viewer is not page:
        try: await viewer.close()
        except Exception: pass

    if downloaded:
        return downloaded
    # Fallback: сохранить превьюшку
    if img_src:
        body = await _fetch_bytes(ctx, img_src)
        if body:
            fp = dest_dir / (safe_fn(title) + "_preview.jpg")
            fp.write_bytes(body)
            log(f"    Fallback превьюшка: {fp.name}")
            return str(fp)
    return None


# ── 10. СКРАПИНГ СТРАНИЦЫ ЗАПИСИ ─────────────────────────────────────────── #

# Clean «label | value» extractor for the FS «person details» panel — labels and
# values are adjacent elements, so text_content() glues them («NameRuby…SexFemale»).
# Require EXACTLY two single-line text children; skip anything wrapping a table.
_FS_FIELDS_JS = r"""() => {
    const norm = s => (s || '').replace(/\s+/g, ' ').trim();
    const out = [], seen = new Set();
    const BAD = /^(sign in|search|menu|save|print|share|view|home|family ?tree|memories|get involved|help|settings|tree|overview|sources|details|person details|tools|edit|add|learn more|collection information|cite this|attach|report|feedback|about)/i;
    // junk values: the «Learn more … FamilySearch Wiki» line, account chrome, etc.
    const BADV = /(learn more|familysearch wiki|sign in|log ?in|cite this|see all|view all)/i;
    for (const el of document.querySelectorAll('li, div, dl, tr, section')) {
        // Skip page chrome (the «A  Alla Khananashvili» account menu lives in the
        // header/nav, the Wiki link in a footer/aside) — never genealogical fields.
        if (el.closest('nav, header, footer, aside, [role=navigation], [role=banner], [role=menu], [role=menubar], [role=contentinfo]'))
            continue;
        let kids;
        if (el.tagName === 'TR')
            kids = [...el.querySelectorAll(':scope > td, :scope > th')];
        else
            kids = [...el.children].filter(c => norm(c.innerText));
        if (kids.length !== 2) continue;
        if (kids.some(c => c.querySelector && c.querySelector('table,td,th,tr,li,dl')))
            continue;
        const k = norm(kids[0].innerText), v = norm(kids[1].innerText);
        if (!k || !v || k === v) continue;
        if (k.includes('\n') || k.length < 2 || k.length >= 45 || v.length >= 300) continue;
        if (!/[A-Za-z]/.test(k) || BAD.test(k) || BADV.test(v) || BADV.test(k)) continue;
        const key = k.toLowerCase();
        if (seen.has(key)) continue;
        seen.add(key); out.push([k, v]);
    }
    return out;
}"""


async def _scrape_page(ctx, page, url: str, name_hint: str,
                       images_root: Path, logged_in_ref: list,
                       email: str, password: str, log) -> dict:
    """
    Навигировать main page на url.
    Если редирект на логин — войти (только если logged_in_ref[0] == False).
    Скрапить данные, скачать картинку.
    НЕ закрывает page — вызывающий делает go_back().
    """
    rec = {"url": url, "title": name_hint, "name": name_hint,
           "table_data": {}, "images": [], "thumb_bytes": None,
           "events": "", "relationships": "", "collection": ""}

    for bad in BAD_PATHS:
        if bad in url:
            return rec

    await page.goto(url, wait_until="domcontentloaded", timeout=30000)
    await asyncio.sleep(3)

    # Если редирект на логин — залогиниться ОДИН РАЗ
    # Проверяем и по URL и по заголовку страницы (JS-редирект может ещё не завершиться)
    page_title = (await page.title()).lower()
    is_login = "login" in page.url or "sign in" in page_title or "sign-in" in page_title
    if is_login:
        if logged_in_ref[0]:
            log(f"  !! Повторный логин-редирект (уже был вход). URL: {page.url[:60]}")
            return rec
        if not email or not password:
            log("  !! Нет credentials для логина")
            return rec
        log("  → Форма логина, входим...")
        ok = await _login(page, email, password, log)
        if not ok:
            log("  !! Вход провалился")
            return rec
        logged_in_ref[0] = True
        # После логина state= редиректит на нужную запись,
        # но если нет — навигируем вручную
        if url.split("?")[0] not in page.url:
            await page.goto(url, wait_until="domcontentloaded", timeout=30000)
            await asyncio.sleep(3)

    # Заголовок — ТОЛЬКО h1/h2 (НЕ [class*="title"] — он хватал всю панель «person
    # details» ОДНИМ слипшимся блоком: «…NameRuby…SexFemale…»). Первая строка,
    # с капом длины: имя не бывает 250 символов, так слипшийся блок отсекается.
    for sel in ["h1", "h2"]:
        try:
            t = (await page.locator(sel).first.text_content(timeout=3000) or "").strip()
            t = t.split("\n")[0].strip()
            if 2 < len(t) < 120:
                rec["title"] = t
                break
        except Exception:
            pass

    # Данные: dl/dt/dd (с защитой от слипания — длинный «label» = не поле, а блок)
    td: dict = {}
    try:
        dts = await page.query_selector_all("dl dt")
        dds = await page.query_selector_all("dl dd")
        for dt, dd in zip(dts, dds):
            k = (await dt.text_content() or "").strip().rstrip(":")
            v = (await dd.text_content() or "").strip()
            if k and v and "\n" not in k and len(k) < 45 and len(v) < 300:
                td[k] = v
    except Exception:
        pass
    # «person details» панель + любые таблицы — КЛАСС-АГНОСТИЧНО через _FS_FIELDS_JS
    # (ровно 2 однострочных ребёнка, label<45/value<300, без вложенных таблиц). НЕ
    # свой loose-обход <table tr>: он хватал ВЕСЬ блок «Isidor Sitron person
    # details…NARA)» одной ячейкой (text_content склеивает) → слипшаяся строка.
    if not td:
        try:
            for k, v in (await page.evaluate(_FS_FIELDS_JS)):
                td.setdefault(k, v)
        except Exception:
            pass
    rec["table_data"] = td

    # Метка для имени файла
    img_label = name_hint
    if td:
        parts = [name_hint]
        for k in ("Event Type", "Type"):
            if td.get(k):
                parts.append(td[k]); break
        for k in ("Event Date", "Date", "Birth Date", "Death Date", "Marriage Date"):
            if td.get(k):
                parts.append(td[k]); break
        if len(parts) > 1:
            img_label = " — ".join(parts)

    img_dir = images_root / safe_fn(img_label)

    # Превьюшка (маленькая) для Word
    img_src = await _best_img(page)
    if img_src:
        rec["thumb_bytes"] = await _fetch_bytes(ctx, img_src)
        if rec["thumb_bytes"]:
            log(f"    Превьюшка: {len(rec['thumb_bytes'])//1024}KB")
        else:
            log("    Превьюшка: не получена")
    else:
        log("    На странице нет картинки документа")

    # Полноформатная JPG через вьюер. view=index тоже пробуем — некоторые
    # индексные записи ИМЕЮТ скачиваемое изображение документа (пользователь
    # подтвердила: «вот же оно, доступно»). _download_jpg сам кликает картинку /
    # открывает вьюер и имеет свои fallback-селекторы, если _best_img пуст.
    is_index = "view=index" in url
    jp = None
    if img_src or is_index:
        jp = await _download_jpg(ctx, page, img_dir, img_label, log)
    if jp:
        rec["images"] = [jp]
    elif rec["thumb_bytes"]:
        fp = img_dir / (safe_fn(img_label) + "_preview.jpg")
        fp.parent.mkdir(parents=True, exist_ok=True)
        fp.write_bytes(rec["thumb_bytes"])
        rec["images"] = [str(fp)]
        log(f"    Сохранена превьюшка: {fp.name}")
    else:
        rec["images"] = []
        log("    Изображение не найдено")

    return rec


# ── Word ──────────────────────────────────────────────────────────────────── #

def _add_link(para, text, url):
    rid = para.part.relate_to(url, HYPERLINK_REL, is_external=True)
    hl  = OxmlElement("w:hyperlink"); hl.set(qn("r:id"), rid)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    c   = OxmlElement("w:color"); c.set(qn("w:val"), "0563C1")
    u   = OxmlElement("w:u");     u.set(qn("w:val"), "single")
    rPr.append(c); rPr.append(u); run.append(rPr)
    t   = OxmlElement("w:t"); t.text = text or url
    t.set(qn("xml:space"), "preserve")
    run.append(t); hl.append(run); para._p.append(hl)


def _docx_add_record(doc, i, rec):
    """Render ONE record into an open Document (shared by fresh write + append)."""
    title = rec.get("title") or rec.get("name", "—")
    doc.add_heading(f"{i}. {title}", level=2)

    if rec.get("url"):
        pp = doc.add_paragraph()
        pp.add_run("Источник (FamilySearch): ").bold = True
        _add_link(pp, "Открыть запись", rec["url"])

    p = doc.add_paragraph()
    p.add_run("Совпадение: ").bold = True
    p.add_run(f"{rec.get('score','?')}%")

    rows_data = []
    if rec.get("collection"):
        rows_data.append(("Коллекция", rec["collection"]))
    if rec.get("events"):
        rows_data.append(("События", rec["events"]))
    if rec.get("relationships"):
        rows_data.append(("Родственники", rec["relationships"]))
    for f, v in rec.get("table_data", {}).items():
        rows_data.append((str(f), str(v)))

    if rows_data:
        tbl = doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Поле"; hdr[1].text = "Значение"
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for f, v in rows_data:
            r = tbl.add_row().cells
            r[0].text = f; set_cell_lines(r[1], v)

    doc.add_paragraph("")

    imgs = rec.get("images", [])
    tb   = rec.get("thumb_bytes")
    if imgs and Path(imgs[0]).exists():
        doc.add_paragraph("Изображение документа:").runs[0].bold = True
        try:
            doc.add_picture(imgs[0], width=Inches(4))
        except Exception:
            doc.add_paragraph(f"  [{Path(imgs[0]).name}]")
        p = doc.add_paragraph(); p.add_run("Файл: ").bold = True
        p.add_run(str(Path(imgs[0]).resolve()))      # exact path where it was saved
    elif tb:
        doc.add_paragraph("Превью документа:").runs[0].bold = True
        try:
            doc.add_picture(io.BytesIO(tb), width=Inches(4))
        except Exception:
            doc.add_paragraph("  [не удалось вставить]")
    else:
        doc.add_paragraph("  [изображение недоступно]")
    doc.add_paragraph("")


def write_docx(path: Path, records: list, qlines: list, append: bool = False):
    if not _DOCX_OK:
        raise RuntimeError("python-docx не установлен")
    existing = append and Path(path).exists()
    if existing:
        # Open the existing document and append a clearly-marked new batch.
        doc = Document(str(path))
        doc.add_page_break()
        sep = doc.add_heading(
            f"➕ Добавлено ещё {len(records)} (совпадение ≥{MIN_MATCH}%) — "
            f"{time.strftime('%Y-%m-%d %H:%M')}", level=1)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Параметры:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph("")
    else:
        doc = Document()
        s = doc.sections[0]
        s.page_width  = Mm(297); s.page_height = Mm(210)
        s.left_margin = s.right_margin = Mm(15)
        s.top_margin  = s.bottom_margin = Mm(15)
        h = doc.add_heading("FamilySearch — Результаты", 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Параметры:")
        for ln in qlines:
            doc.add_paragraph(ln, style="List Bullet")
        doc.add_paragraph(f"Найдено: {len(records)} (совпадение ≥{MIN_MATCH}%)")
        doc.add_paragraph("")

    for i, rec in enumerate(records, 1):
        _docx_add_record(doc, i, rec)
    doc.save(path)


# ── Excel ─────────────────────────────────────────────────────────────────── #

def write_xlsx(path: Path, records: list, qlines: list, append: bool = False):
    if not _OPENPYXL_OK:
        raise RuntimeError("openpyxl не установлен")
    HF = PatternFill("solid", fgColor="006B6B")
    HN = Font(bold=True, color="FFFFFF", size=11)
    LINKF = Font(color="0563C1", underline="single")
    TS = Side(style="thin", color="B0C8C8")
    T  = Border(left=TS, right=TS, top=TS, bottom=TS)

    aff: list = []
    for rec in records:
        for k in rec.get("table_data", {}):
            if k not in aff: aff.append(k)
    # «База» = название сайта (FamilySearch) — чтобы источник был виден в таблице.
    base_cols = ["#", "База", "Имя", "Совп. %", "Коллекция", "События",
                 "Родственники", "Файл JPG", "URL"]

    existing = append and Path(path).exists()
    if existing:
        wb = load_workbook(str(path)); ws = wb.active
        header = [ws.cell(row=1, column=c).value
                  for c in range(1, ws.max_column + 1)]
        for name in base_cols + aff:
            if name not in header:
                header.append(name)
                c = ws.cell(row=1, column=len(header), value=name)
                c.font = HN; c.fill = HF; c.border = T
                c.alignment = Alignment(horizontal="center",
                                        vertical="center", wrap_text=True)
        start_row = ws.max_row + 1
        start_num = ws.max_row - 1            # records already in the sheet
    else:
        wb = Workbook(); ws = wb.active; ws.title = "FamilySearch"
        header = base_cols + aff
        for ci, cn in enumerate(header, 1):
            c = ws.cell(row=1, column=ci, value=cn)
            c.font = HN; c.fill = HF; c.border = T
            c.alignment = Alignment(horizontal="center",
                                    vertical="center", wrap_text=True)
        start_row = 2
        start_num = 0
    col_idx = {name: i + 1 for i, name in enumerate(header)}

    for n, rec in enumerate(records):
        ri   = start_row + n
        td   = rec.get("table_data", {})
        imgs = "\n".join(str(Path(p).resolve()) for p in rec.get("images", []))
        row  = {"#": start_num + n + 1, "База": "FamilySearch",
                "Имя": rec.get("title", rec.get("name", "")),
                "Совп. %": rec.get("score", ""),
                "Коллекция": rec.get("collection", ""),
                "События": rec.get("events", ""),
                "Родственники": rec.get("relationships", ""),
                "Файл JPG": imgs, "URL": rec.get("url", "")}
        for f in aff:
            row[f] = td.get(f, "")
        for name, val in row.items():
            ci = col_idx.get(name)
            if not ci:
                continue
            c = ws.cell(row=ri, column=ci)
            if name == "URL" and val:            # hidden hyperlink, not raw URL
                c.value = "Открыть"; c.hyperlink = val; c.font = LINKF
            else:
                c.value = val
            c.border = T
            c.alignment = Alignment(wrap_text=True, vertical="top")

    for ci in range(1, len(header) + 1):
        ltr = get_column_letter(ci)
        mw  = max(len(str(header[ci-1] or "")),
                  *(len(str(ws.cell(row=r, column=ci).value or "").split("\n")[0])
                    for r in range(2, ws.max_row + 1)), 8)
        ws.column_dimensions[ltr].width = min(mw + 4, 60)
    wb.save(path)


# ── ГЛАВНАЯ ФУНКЦИЯ ───────────────────────────────────────────────────────── #

async def run_scraper(
    *,
    first_names:   str       = "",
    last_names:    str       = "",
    place_lived:   str       = "",
    birth_year:    str       = "",
    year_range:    int       = 0,    # ± years on the birth year (0 = exact)
    tab:           str       = "Historical Records",
    advanced:      dict|None = None,
    exact:         dict|None = None,
    output_format: str       = "both",
    output_folder            = Path("."),
    email:         str|None  = None,
    password:      str|None  = None,
    log                      = print,
    progress                 = None,
    cancel_event             = None,
    ask_file_conflict        = None,  # callable(list[str]) → "overwrite"/"append"/"skip"
) -> dict:

    def _prog(pct, txt):
        log(txt)
        if progress: progress(pct, txt)

    def _done():
        return bool(cancel_event and cancel_event.is_set())

    adv       = advanced or {}
    exact     = exact or {}
    has_adv   = any(v for v in adv.values() if v and v not in (False, "Unspecified"))
    want_docx = output_format in ("docx", "both")
    want_xlsx = output_format in ("xlsx", "both")
    output_folder = Path(output_folder)
    output_folder.mkdir(parents=True, exist_ok=True)
    images_root = output_folder / "images"

    qname     = " ".join(p for p in (first_names, last_names) if p)
    qlines    = [ln for ln in [
        f"First Names: {first_names}", f"Last Names: {last_names}",
        f"Place Lived: {place_lived}", f"Birth Year: {birth_year}",
    ] if not ln.endswith(": ")]
    summary   = {"ok": False}
    file_base = safe_fn(f"familysearch_{qname}") if qname else "familysearch_results"

    # logged_in_ref[0] == True после первого успешного входа
    # Передаём как список чтобы _scrape_page мог изменить флаг
    logged_in_ref = [False]

    _prog(0, "Запускаю браузер...")

    async with async_playwright() as pw:
        # PERSISTENT profile → the FamilySearch login (cookies/session) is kept
        # between runs, so we don't sign in every time («запоминай куки»).
        for _lk in ("SingletonLock", "SingletonCookie", "SingletonSocket"):
            try: (FS_PROFILE_DIR / _lk).unlink()
            except Exception: pass
        FS_PROFILE_DIR.mkdir(parents=True, exist_ok=True)
        ctx = await pw.chromium.launch_persistent_context(
            str(FS_PROFILE_DIR),
            headless=False,
            no_viewport=True,
            accept_downloads=True,
            args=["--start-maximized",
                  "--disable-blink-features=AutomationControlled"],
        )
        browser = ctx                       # so the rest (ctx/close) stays valid
        page = ctx.pages[0] if ctx.pages else await ctx.new_page()
        # persistent profile may restore old tabs — keep just one
        for _p in list(ctx.pages)[1:]:
            try: await _p.close()
            except Exception: pass

        try:
            # ── 1. SEARCH via the full results URL (all q.*/f.* params) ──── #
            # Build the FamilySearch results URL straight from the form (name,
            # birth-year range, life events, family members ×3, record country /
            # subcountry, record types, batch / film / principal). Reliable and
            # matches the site URL — no fragile advanced-popup form filling.
            _prog(5, "Поиск...")
            url = _fs_url(first_names, last_names, place_lived, birth_year,
                          year_range, exact, adv)
            log(f"  Открываю результаты: {url}")
            await page.goto(url, wait_until="domcontentloaded", timeout=40000)
            await asyncio.sleep(5)
            if _done(): return summary

            # ── 2. ЛОГИН ОДИН РАЗ — ДО открытия записей ──────────────── #
            # HR rows show without login; sign in now so opening a record later
            # never lands on the login page. The URL already carries tab=records
            # + every filter, so we must NOT click the HR tab (it rebuilds the
            # URL and would drop the filters — known FS behaviour).
            _prog(16, "Sign in...")
            await _sign_in_if_needed(page, email or "", password or "",
                                     logged_in_ref, log)
            if _done(): return summary
            if "discovery/results" not in page.url:   # login navigated away → back
                await page.goto(url, wait_until="domcontentloaded", timeout=40000)
                await asyncio.sleep(4)

            # ── 5. 60 НА СТРАНИЦУ + СБОР РЕЗУЛЬТАТОВ ─────────────────── #
            _prog(28, "Сбор результатов...")
            await _set_60(page, log)
            raw       = await _collect(page, qname, log)

            qualified = [r for r in raw if r["score"] >= MIN_MATCH]
            log(f"  Подходящих (≥{MIN_MATCH}%): {len(qualified)}")

            if not qualified:
                _prog(100, f"Нет записей с совпадением ≥{MIN_MATCH}%.")
                summary.update({"ok": True, "n_records": 0,
                                "message": f"Нет записей ≥{MIN_MATCH}%."})
                return summary

            # ── 6. СКРАПИНГ ОТФИЛЬТРОВАННЫХ ЗАПИСЕЙ (одна сессия) ─────── #
            # All records open on the same `page` — the session stays one.
            results_url = page.url
            records: list = []
            for i, r in enumerate(qualified, 1):
                if _done(): break
                _prog(30 + int(64 * i / len(qualified)),
                      f"[{i}/{len(qualified)}] {r['name'][:60]}...")
                det = await _scrape_page(ctx, page, r["url"], r["name"],
                                         images_root, logged_in_ref,
                                         email or "", password or "", log)
                det["score"]         = r["score"]
                det["collection"]    = r.get("coll", "")
                det["events"]        = r.get("evts", "")
                det["relationships"] = r.get("rels", "")
                records.append(det)
                log(f"  ✓  {det['title'][:70]}  ({r['score']}%)")
                await page.goto(results_url,
                                wait_until="domcontentloaded", timeout=20000)
                try:
                    await page.wait_for_selector("tbody tr", timeout=8000)
                except Exception:
                    await asyncio.sleep(3)

            # ── 7. СОХРАНЕНИЕ ──────────────────────────────────────── #
            _prog(96, "Сохранение файлов...")
            docx_p = output_folder / f"{file_base}.docx"
            xlsx_p = output_folder / f"{file_base}.xlsx"

            # Если файлы уже есть — ВСЕГДА спросить: перезаписать / дополнить /
            # пропустить. Без callback или без конфликта — обычная перезапись.
            existing_names = [p.name for p, want in
                              ((docx_p, want_docx), (xlsx_p, want_xlsx))
                              if want and records and p.exists()]
            decision = "overwrite"
            if existing_names and ask_file_conflict:
                try:
                    decision = (ask_file_conflict(existing_names)
                                or "overwrite").lower()
                except Exception as _e:
                    log(f"  !! диалог конфликта файлов: {_e}")
                    decision = "overwrite"
                log(f"  → Файл(ы) уже существуют {existing_names} → выбор: {decision}")
            append = (decision == "append")

            sd = sx = False
            if decision == "skip":
                log("  → Сохранение пропущено по выбору пользователя "
                    "(существующие файлы не тронуты).")
            else:
                if want_docx and records:
                    write_docx(docx_p, records, qlines, append=append)
                    sd = True
                    log(f"  Word: {docx_p}"
                        f"{' (дополнен)' if append and docx_p.name in existing_names else ''}")
                if want_xlsx and records:
                    write_xlsx(xlsx_p, records, qlines, append=append)
                    sx = True
                    log(f"  Excel: {xlsx_p}"
                        f"{' (дополнен)' if append and xlsx_p.name in existing_names else ''}")
            _prog(100, f"Готово — {len(records)} записей.")
            summary.update({
                "ok":            True,
                "docx_count":    1 if sd else 0,
                "xlsx_path":     str(xlsx_p) if sx else None,
                "output_folder": str(output_folder),
                "n_records":     len(records),
            })

        except Exception as exc:
            summary.update({"error": "exception",
                            "message": f"{type(exc).__name__}: {exc}"})
            log(f"  !! {exc}")
        finally:
            try: await ctx.close()
            except Exception: pass
            try: await browser.close()
            except Exception: pass

    return summary
